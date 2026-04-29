from __future__ import annotations

import io
import html
import inspect
import json
import base64
import hmac
import os
import re
import hashlib
import secrets
import sqlite3
from collections import Counter
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime
from functools import lru_cache
from pathlib import Path
from typing import Any
from urllib.parse import parse_qsl, urlencode, urljoin, urlparse, urlunparse

import numpy as np
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import requests
import streamlit as st

try:
    from rapidfuzz import fuzz
except Exception:  # pragma: no cover - optional acceleration
    fuzz = None

try:
    import jieba
except Exception:  # pragma: no cover - optional Chinese tokenizer
    jieba = None

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
except Exception:  # pragma: no cover - optional semantic scorer
    TfidfVectorizer = None
    cosine_similarity = None


APP_TITLE = "CareerPilot 全职业岗位分析器"
DB_PATH = Path(__file__).with_name("careerpilot.db")
JD_EXPORT_DIR = Path.home() / "Downloads" / "CareerPilot_JD"
JD_EXPORT_DIRS = [JD_EXPORT_DIR]

APP_NAVIGATION = {
    "main_tabs": ["岗位工作台", "简历工作台", "求职决策", "面试与报告"],
    "jd_modes": ["单条JD分析", "批量JD筛选", "行业招聘监测"],
    "resume_modes": ["简历解析与匹配", "定制简历", "不足与努力方向"],
    "decision_modes": ["Offer预测", "实习评估", "投递管理"],
    "report_modes": ["面经分析", "可视化与报告"],
}

MAIN_WORKSPACE_LABELS = {
    "jd": "岗位工作台",
    "resume": "简历工作台",
    "decision": "求职决策",
    "report": "面试与报告",
}

RUNTIME_CACHE_SCHEMA = "2026-04-29-quality-match-v3"
DB_SCHEMA_VERSION = "auth-user-v1"
PASSWORD_HASH_ITERATIONS = 260_000
AUTH_SESSION_KEY = "auth_user"

DEFAULT_TARGET_INTENTION_NAME = "目标意向"
CLEAN_DEFAULT_PROFILE = ""

CLEAN_DEFAULT_PROFILES = [
    {
        "name": DEFAULT_TARGET_INTENTION_NAME,
        "content": "",
    }
]

LEGACY_AUTO_TARGET_CITIES = ["上海", "北京", "深圳", "广州", "杭州"]
LEGACY_AUTO_TARGET_INDUSTRIES = ["互联网", "制造业", "咨询", "金融", "新能源", "医疗", "消费", "专业服务"]
LEGACY_AUTO_AVOID_KEYWORDS = "职责模糊、纯杂务、长期无转正路径、销售指标与岗位名称不一致"
LEGACY_AUTO_NOTES = "优先真实业务场景、清晰产出、可量化成果和成长路径。"

DEFAULT_TARGET_PREFERENCES: dict[str, Any] = {
    "target_roles": [],
    "target_cities": [],
    "extra_cities": "",
    "accept_remote": True,
    "accept_nationwide": True,
    "min_monthly_salary": 10000,
    "min_daily_salary": 120,
    "preferred_industries": [],
    "extra_industries": "",
    "job_keywords": [],
    "avoid_keywords": "",
    "notes": "",
}

LEGACY_TEMPLATE_RESUME_MARKERS = [
    "示例简历",
    "模板简历",
    "请替换为你的",
    "请在此填写",
    "张三",
    "李四",
    "某大学",
    "某公司",
    "某项目",
    "负责XX",
    "量化成果XX",
]

CHINA_285_CITY_OPTIONS_BY_PROVINCE: dict[str, list[str]] = {
    "北京市": ["北京"],
    "上海市": ["上海"],
    "天津市": ["天津"],
    "重庆市": ["重庆"],
    "河北省": ["石家庄", "唐山", "秦皇岛", "邯郸", "邢台", "保定", "张家口", "承德", "沧州", "廊坊", "衡水"],
    "山西省": ["太原", "大同", "阳泉", "长治", "晋城", "朔州", "晋中", "运城", "忻州", "临汾", "吕梁"],
    "辽宁省": ["沈阳", "大连", "鞍山", "抚顺", "本溪", "丹东", "锦州", "营口", "阜新", "辽阳", "盘锦", "铁岭", "朝阳", "葫芦岛"],
    "吉林省": ["长春", "吉林", "四平", "辽源", "通化", "白山", "松原", "白城"],
    "黑龙江省": ["哈尔滨", "齐齐哈尔", "鸡西", "鹤岗", "双鸭山", "大庆", "伊春", "佳木斯", "七台河", "牡丹江", "黑河", "绥化"],
    "江苏省": ["南京", "无锡", "徐州", "常州", "苏州", "南通", "连云港", "淮安", "盐城", "扬州", "镇江", "泰州", "宿迁"],
    "浙江省": ["杭州", "宁波", "温州", "嘉兴", "湖州", "绍兴", "金华", "衢州", "舟山", "台州", "丽水"],
    "安徽省": ["合肥", "芜湖", "蚌埠", "淮南", "马鞍山", "淮北", "铜陵", "安庆", "黄山", "滁州", "阜阳", "宿州", "六安", "亳州", "池州", "宣城"],
    "福建省": ["福州", "厦门", "莆田", "三明", "泉州", "漳州", "南平", "龙岩", "宁德"],
    "江西省": ["南昌", "景德镇", "萍乡", "九江", "新余", "鹰潭", "赣州", "吉安", "宜春", "抚州", "上饶"],
    "山东省": ["济南", "青岛", "淄博", "枣庄", "东营", "烟台", "潍坊", "济宁", "泰安", "威海", "日照", "临沂", "德州", "聊城", "滨州", "菏泽"],
    "河南省": ["郑州", "开封", "洛阳", "平顶山", "安阳", "鹤壁", "新乡", "焦作", "濮阳", "许昌", "漯河", "三门峡", "南阳", "商丘", "信阳", "周口", "驻马店"],
    "湖北省": ["武汉", "黄石", "十堰", "宜昌", "襄阳", "鄂州", "荆门", "孝感", "荆州", "黄冈", "咸宁", "随州"],
    "湖南省": ["长沙", "株洲", "湘潭", "衡阳", "邵阳", "岳阳", "常德", "张家界", "益阳", "郴州", "永州", "怀化", "娄底"],
    "广东省": ["广州", "韶关", "深圳", "珠海", "汕头", "佛山", "江门", "湛江", "茂名", "肇庆", "惠州", "梅州", "汕尾", "河源", "阳江", "清远", "东莞", "中山", "潮州", "揭阳", "云浮"],
    "海南省": ["海口", "三亚"],
    "四川省": ["成都", "自贡", "攀枝花", "泸州", "德阳", "绵阳", "广元", "遂宁", "内江", "乐山", "南充", "眉山", "宜宾", "广安", "达州", "雅安", "巴中", "资阳"],
    "贵州省": ["贵阳", "六盘水", "遵义", "安顺", "毕节", "铜仁"],
    "云南省": ["昆明", "曲靖", "玉溪", "保山", "昭通", "丽江", "普洱", "临沧"],
    "陕西省": ["西安", "铜川", "宝鸡", "咸阳", "渭南", "延安", "汉中", "榆林", "安康", "商洛"],
    "甘肃省": ["兰州", "嘉峪关", "金昌", "白银", "天水", "武威", "张掖", "平凉", "酒泉", "庆阳", "定西", "陇南"],
    "内蒙古自治区": ["呼和浩特", "包头", "乌海", "赤峰", "通辽", "鄂尔多斯", "呼伦贝尔", "巴彦淖尔", "乌兰察布"],
    "广西壮族自治区": ["南宁", "柳州", "桂林", "梧州", "北海", "防城港", "钦州", "贵港", "玉林", "百色", "贺州", "河池", "来宾", "崇左"],
    "宁夏回族自治区": ["银川", "石嘴山", "吴忠", "固原", "中卫"],
    "新疆维吾尔自治区": ["乌鲁木齐", "克拉玛依"],
}

CHINA_CITY_OPTIONS = [
    city
    for province_cities in CHINA_285_CITY_OPTIONS_BY_PROVINCE.values()
    for city in province_cities
]

RECRUITMENT_INDUSTRY_OPTIONS = [
    "互联网/电子商务",
    "移动互联网",
    "计算机软件",
    "计算机硬件",
    "IT服务/系统集成",
    "云计算/大数据",
    "人工智能",
    "游戏",
    "电子/半导体/集成电路",
    "通信/网络设备",
    "金融",
    "银行",
    "证券/期货",
    "基金",
    "保险",
    "投资/融资",
    "新能源",
    "汽车/新能源汽车",
    "机械/设备/重工",
    "仪器仪表/工业自动化",
    "原材料及加工",
    "化工",
    "环保",
    "能源/矿产/电力",
    "贸易/进出口",
    "批发/零售",
    "食品/饮料/酒水",
    "服装/纺织/皮革",
    "家居/家具/家电",
    "快消品",
    "物流/仓储",
    "交通/运输",
    "供应链/采购",
    "房地产/建筑",
    "物业管理",
    "咨询",
    "专业服务",
    "检测/认证",
    "会计/审计",
    "法律",
    "人力资源服务",
    "广告/公关/会展",
    "媒体/出版/影视",
    "文化/体育/娱乐",
    "教育/培训",
    "学术/科研",
    "医疗服务",
    "医药/生物工程",
    "医疗器械",
    "养老/健康服务",
    "餐饮",
    "酒店/旅游",
    "生活服务",
    "政府/公共事业",
    "非营利组织",
    "农林牧渔",
    "可持续发展/ESG",
    "跨境电商",
]

TARGET_ROLE_OPTIONS = [
    "数据分析",
    "商业分析",
    "产品经理",
    "产品运营",
    "用户运营",
    "内容运营",
    "增长运营",
    "项目管理",
    "咨询顾问",
    "行业研究",
    "市场营销",
    "品牌公关",
    "商务拓展",
    "客户成功",
    "供应链管理",
    "采购",
    "财务分析",
    "法务合规",
    "人力资源",
    "前端开发",
    "后端开发",
    "算法工程师",
    "数据工程师",
    "测试工程师",
    "UI/UX设计",
    "ESG/可持续发展",
    "碳管理/LCA",
]

INDUSTRY_ALIAS_MAP: dict[str, list[str]] = {
    "互联网/电子商务": ["互联网", "电商", "电子商务", "平台", "社区", "本地生活", "O2O", "SaaS"],
    "移动互联网": ["移动互联网", "App", "小程序", "移动端", "客户端"],
    "计算机软件": ["软件", "SaaS", "企业服务", "ERP", "CRM", "低代码", "数据库", "操作系统"],
    "计算机硬件": ["硬件", "服务器", "存储", "芯片", "嵌入式", "智能终端"],
    "IT服务/系统集成": ["IT服务", "系统集成", "实施", "运维", "信息化", "数字化转型"],
    "云计算/大数据": ["云计算", "大数据", "数据仓库", "数据平台", "数据湖", "数仓", "BI"],
    "人工智能": ["人工智能", "AI", "AIGC", "大模型", "LLM", "NLP", "机器学习", "深度学习", "推荐算法", "计算机视觉", "智能驾驶"],
    "游戏": ["游戏", "手游", "端游", "页游", "电竞", "Unity", "Unreal"],
    "电子/半导体/集成电路": ["电子", "半导体", "集成电路", "芯片", "晶圆", "封测", "EDA", "功率器件"],
    "通信/网络设备": ["通信", "5G", "网络设备", "基站", "光通信", "运营商", "物联网"],
    "金融": ["金融", "资管", "财富管理", "金融科技", "FinTech", "信贷", "风控", "支付"],
    "银行": ["银行", "商业银行", "零售银行", "对公", "信贷", "柜面"],
    "证券/期货": ["证券", "券商", "投行", "研究所", "期货", "交易", "量化"],
    "基金": ["基金", "公募", "私募", "资管", "投资研究", "FOF"],
    "保险": ["保险", "寿险", "财险", "再保险", "精算"],
    "投资/融资": ["投资", "融资", "VC", "PE", "FA", "并购", "投融资"],
    "新能源": ["新能源", "光伏", "储能", "锂电", "动力电池", "风电", "充电桩", "氢能", "碳中和", "绿电"],
    "汽车/新能源汽车": ["汽车", "新能源汽车", "整车", "主机厂", "智能驾驶", "车联网", "自动驾驶", "座舱", "三电"],
    "机械/设备/重工": ["机械", "设备", "重工", "装备制造", "机电", "数控", "工程机械"],
    "仪器仪表/工业自动化": ["仪器仪表", "工业自动化", "PLC", "机器人", "传感器", "智能制造", "工控"],
    "原材料及加工": ["原材料", "钢铁", "有色", "加工", "新材料", "复合材料"],
    "化工": ["化工", "精细化工", "材料化学", "化学品", "石化"],
    "环保": ["环保", "环境", "水处理", "固废", "废气", "环评", "污染治理"],
    "能源/矿产/电力": ["能源", "矿产", "电力", "电网", "发电", "煤炭", "油气", "天然气"],
    "贸易/进出口": ["贸易", "进出口", "外贸", "报关", "国际贸易"],
    "批发/零售": ["批发", "零售", "门店", "商超", "新零售", "连锁"],
    "食品/饮料/酒水": ["食品", "饮料", "酒水", "乳制品", "预制菜", "快消食品"],
    "服装/纺织/皮革": ["服装", "纺织", "皮革", "鞋服", "时尚"],
    "家居/家具/家电": ["家居", "家具", "家电", "厨电", "智能家居"],
    "快消品": ["快消", "FMCG", "日化", "美妆", "个护", "食品饮料", "消费品"],
    "物流/仓储": ["物流", "仓储", "快递", "货运", "配送", "仓配", "运输"],
    "交通/运输": ["交通", "运输", "航空", "铁路", "航运", "港口", "地铁"],
    "供应链/采购": ["供应链", "采购", "计划", "S&OP", "供应商管理", "库存", "履约"],
    "房地产/建筑": ["房地产", "建筑", "工程", "土建", "地产", "施工", "设计院"],
    "物业管理": ["物业", "园区", "楼宇", "设施管理"],
    "咨询": ["咨询", "管理咨询", "战略咨询", "业务咨询", "数字化咨询", "咨询顾问"],
    "专业服务": ["专业服务", "事务所", "审计", "税务", "认证", "检测", "咨询", "律所"],
    "检测/认证": ["检测", "认证", "TIC", "检验", "审核", "ISO", "SGS", "TÜV", "BV"],
    "会计/审计": ["会计", "审计", "税务", "四大", "财务咨询", "内控"],
    "法律": ["法律", "法务", "律所", "合规", "合同", "知识产权"],
    "人力资源服务": ["人力资源服务", "猎头", "招聘", "RPO", "薪酬", "培训"],
    "广告/公关/会展": ["广告", "公关", "会展", "品牌传播", "媒介", "活动策划"],
    "媒体/出版/影视": ["媒体", "出版", "影视", "内容", "短视频", "直播", "新媒体"],
    "文化/体育/娱乐": ["文化", "体育", "娱乐", "演出", "赛事", "艺人"],
    "教育/培训": ["教育", "培训", "教培", "在线教育", "课程", "留学"],
    "学术/科研": ["科研", "研究院", "实验室", "课题", "高校", "学术"],
    "医疗服务": ["医疗", "医院", "诊所", "互联网医疗", "健康管理"],
    "医药/生物工程": ["医药", "制药", "生物", "CRO", "CDMO", "临床", "药企"],
    "医疗器械": ["医疗器械", "器械", "IVD", "影像设备", "耗材"],
    "养老/健康服务": ["养老", "健康服务", "康养", "护理", "体检"],
    "餐饮": ["餐饮", "餐厅", "连锁餐饮", "茶饮", "咖啡"],
    "酒店/旅游": ["酒店", "旅游", "旅行", "OTA", "景区", "民宿"],
    "生活服务": ["生活服务", "家政", "维修", "到家服务", "本地服务"],
    "政府/公共事业": ["政府", "公共事业", "事业单位", "政务", "公共服务"],
    "非营利组织": ["非营利", "NGO", "公益", "基金会", "社会组织"],
    "农林牧渔": ["农业", "林业", "牧业", "渔业", "种植", "养殖"],
    "可持续发展/ESG": ["ESG", "可持续", "可持续发展", "双碳", "碳管理", "碳核算", "LCA", "碳足迹", "CBAM", "气候"],
    "跨境电商": ["跨境电商", "亚马逊", "Temu", "Shein", "TikTok Shop", "独立站", "出海"],
}

CRAWLER_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0 Safari/537.36"
    ),
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
}

HIGH_VALUE_KEYWORDS = [
    "核心项目",
    "项目交付",
    "独立负责",
    "业务分析",
    "数据分析",
    "指标体系",
    "用户研究",
    "需求分析",
    "产品设计",
    "项目管理",
    "增长",
    "策略",
    "咨询",
    "行业研究",
    "客户成功",
    "商业分析",
    "SQL",
    "Python",
    "Tableau",
    "Power BI",
    "Java",
    "C++",
    "前端",
    "后端",
    "算法",
    "机器学习",
    "财务分析",
    "法务合规",
    "供应链管理",
    "采购",
    "英文",
    "转正",
    "校招",
    "管培",
    "LCA",
    "生命周期评价",
    "生命周期评估",
    "CFP",
    "产品碳足迹",
    "ISO14067",
    "ISO 14067",
    "EPD",
    "CBAM",
    "碳边境调节机制",
    "碳核查",
    "碳盘查",
    "供应链碳",
    "供应链碳数据",
    "SimaPro",
    "GaBi",
    "openLCA",
    "PEF",
    "GHG Protocol",
    "温室气体核算",
    "出海合规",
]

LOW_VALUE_KEYWORDS = [
    "打杂",
    "纯执行",
    "无明确职责",
    "长期实习无转正",
    "电话销售",
    "地推",
    "陌拜",
    "销售指标",
    "行政杂务",
    "材料搬运",
    "只做排版",
    "只做整理",
    "无培训",
    "无成长",
    "ESG报告",
    "ESG 报告",
    "信息披露",
    "CSR",
    "CSR传播",
    "公众号",
    "活动策划",
    "行政",
    "行政协同",
    "品牌传播",
    "可持续传播",
    "材料整理",
    "会议组织",
]

LOW_VALUE_RISK_RULES: dict[str, list[str]] = {
    "职责模糊": ["职责不清", "职责模糊", "工作内容不限", "服从安排", "临时安排", "领导交办", "其他事项", "协助完成", "支持部门"],
    "纯执行/杂务": ["打杂", "纯执行", "行政杂务", "材料整理", "会议组织", "只做排版", "只做整理", "资料归档", "跑腿", "材料搬运"],
    "销售指标伪装": ["电话销售", "地推", "陌拜", "销售指标", "邀约客户", "拉新指标", "业绩考核", "转化指标", "客户开发"],
    "无转正/培养弱": ["长期实习无转正", "无转正", "不提供转正", "无培训", "无成长", "无导师", "自学为主", "短期实习"],
    "标题内容不一致": ["岗位名称仅供参考", "实际工作以销售为主", "实际工作以行政为主", "挂名", "储备干部", "综合岗"],
    "实习生当全职用": ["独立负责", "可独立承担", "高强度", "抗压", "加班", "全职实习", "每周5天", "长期稳定实习", "人手紧张"],
    "泛宣传/披露": ["ESG报告", "ESG 报告", "信息披露", "CSR", "CSR传播", "公众号", "活动策划", "品牌传播", "可持续传播", "宣传"],
}

SKILL_ALIASES: dict[str, list[str]] = {
    "产品能力": ["产品经理", "产品助理", "需求分析", "PRD", "原型", "Axure", "用户故事", "竞品分析", "产品设计"],
    "运营": ["运营", "用户运营", "内容运营", "活动运营", "社群运营", "增长运营", "转化", "留存", "拉新"],
    "用户研究": ["用户研究", "用户访谈", "问卷", "可用性测试", "用户画像", "需求调研"],
    "项目管理": ["项目管理", "项目推进", "项目交付", "PMO", "进度管理", "跨部门", "里程碑"],
    "沟通协作": ["沟通", "协作", "跨部门", "客户沟通", "汇报", "presentation", "PPT"],
    "业务分析": ["业务分析", "商业分析", "经营分析", "策略分析", "指标体系", "漏斗分析", "归因分析"],
    "市场营销": ["市场", "营销", "品牌", "投放", "增长", "SEO", "SEM", "渠道", "Campaign"],
    "销售/商务": ["销售", "商务", "BD", "客户开发", "商机", "续约", "客情", "招投标"],
    "财务分析": ["财务分析", "财务", "预算", "成本", "审计", "报表", "现金流", "Excel模型"],
    "法务合规": ["法务", "合规", "合同", "风控", "监管", "隐私", "数据合规"],
    "人力资源": ["HR", "人力资源", "招聘", "培训", "薪酬", "绩效", "组织发展", "员工关系"],
    "供应链管理": ["供应链", "采购", "计划", "物流", "库存", "供应商管理", "S&OP"],
    "设计": ["UI", "UX", "Figma", "视觉设计", "交互设计", "用户体验", "设计规范"],
    "前端": ["前端", "React", "Vue", "JavaScript", "TypeScript", "HTML", "CSS"],
    "后端": ["后端", "Java", "Spring", "Go", "PHP", "Node.js", "服务端", "微服务", "API"],
    "机器学习/AI": ["机器学习", "深度学习", "算法", "模型训练", "NLP", "LLM", "AIGC", "推荐系统"],
    "LCA": ["LCA", "生命周期评价", "生命周期评估"],
    "ISO14067": ["ISO14067", "ISO 14067", "产品碳足迹标准"],
    "GHG Protocol": ["GHG Protocol", "GHG", "温室气体核算"],
    "PEF": ["PEF", "Product Environmental Footprint"],
    "CBAM": ["CBAM", "碳边境调节机制", "欧盟碳边境"],
    "EPD": ["EPD", "环境产品声明"],
    "SimaPro": ["SimaPro"],
    "GaBi": ["GaBi"],
    "openLCA": ["openLCA", "open LCA"],
    "供应链碳管理": ["供应链碳", "供应链碳管理", "供应链碳数据", "供应商碳"],
    "英文能力": ["英文", "英语", "CET-6", "六级", "雅思", "托福", "口语"],
    "数据分析": ["数据分析", "计量", "建模", "统计", "可视化"],
    "SQL": ["SQL", "数据库"],
    "Python": ["Python", "pandas", "numpy", "爬虫"],
    "Excel": ["Excel", "VLOOKUP", "Power Query", "数据透视"],
    "Power BI": ["Power BI", "Tableau", "BI"],
}

SEMANTIC_CAPABILITY_GROUPS: dict[str, list[str]] = {
    "数据分析": [
        "数据分析", "数据处理", "数据清洗", "数据建模", "统计分析", "指标体系", "指标拆解", "经营分析", "商业分析",
        "业务分析", "报表", "看板", "BI", "可视化", "SQL", "Python", "pandas", "Excel", "Tableau", "Power BI",
        "漏斗分析", "归因分析", "用户数据", "A/B测试", "AB测试",
    ],
    "产品能力": [
        "产品经理", "产品助理", "产品运营", "需求分析", "需求调研", "需求梳理", "PRD", "原型", "Axure",
        "用户故事", "用户研究", "用户访谈", "用户画像", "竞品分析", "功能设计", "产品设计", "体验优化",
        "可用性测试", "产品方案",
    ],
    "运营增长": [
        "运营", "用户运营", "内容运营", "活动运营", "社群运营", "增长运营", "策略运营", "增长", "转化",
        "留存", "拉新", "促活", "复购", "渠道运营", "活动策划", "私域", "社群", "增长策略",
    ],
    "项目/咨询": [
        "项目管理", "项目推进", "项目交付", "项目协调", "PMO", "进度管理", "里程碑", "跨部门", "交付",
        "解决方案", "咨询", "尽调", "研究报告", "客户访谈", "方案设计", "项目复盘", "交付物",
    ],
    "市场商务": [
        "市场", "营销", "品牌", "投放", "渠道", "SEO", "SEM", "Campaign", "销售", "商务", "BD",
        "客户开发", "商机", "线索", "续约", "招投标", "大客户", "客户管理", "商业化",
    ],
    "研发工程": [
        "前端", "后端", "全栈", "Java", "Spring", "Go", "PHP", "C++", "React", "Vue", "JavaScript",
        "TypeScript", "服务端", "微服务", "API", "算法", "机器学习", "深度学习", "NLP", "LLM",
        "AIGC", "推荐系统", "模型训练",
    ],
    "财务法务人力": [
        "财务", "财务分析", "预算", "成本", "审计", "报表", "现金流", "金融", "投研", "估值",
        "法务", "合规", "合同", "风控", "监管", "隐私", "内控", "人力资源", "招聘", "培训",
        "薪酬", "绩效", "组织发展",
    ],
    "供应链": [
        "供应链", "采购", "物流", "库存", "供应商管理", "供应商", "计划", "S&OP", "仓储",
        "交付计划", "产销协同", "需求计划",
    ],
    "LCA/碳足迹": [
        "LCA", "生命周期评价", "生命周期评估", "清单分析", "系统边界", "功能单位", "SimaPro", "GaBi",
        "openLCA", "eFootprint", "产品碳足迹", "碳足迹", "CFP", "ISO14067", "ISO 14067", "PCR",
        "EPD", "环境产品声明", "产品环境声明",
    ],
    "碳核算/ESG": [
        "碳核算", "碳盘查", "碳核查", "GHG", "温室气体", "排放因子", "ISO14064", "Scope 1",
        "Scope 2", "Scope 3", "ESG", "可持续发展", "信息披露", "可持续发展报告", "评级",
    ],
    "出海合规": [
        "CBAM", "碳边境", "碳边境调节机制", "欧盟碳边境", "出海", "出海合规", "欧盟", "海外合规",
        "电池法", "EUDR", "供应链合规", "贸易合规",
    ],
    "英文沟通": [
        "英文", "英语", "English", "CET-6", "六级", "雅思", "托福", "口语", "presentation",
        "英语汇报", "英文报告", "跨文化沟通", "海外客户",
    ],
}

CATEGORY_SEMANTIC_GROUPS: dict[str, list[str]] = {
    "数据/商业分析岗": ["数据分析"],
    "产品岗": ["产品能力"],
    "运营岗": ["运营增长"],
    "市场/销售岗": ["市场商务"],
    "咨询/项目岗": ["项目/咨询"],
    "研发/工程岗": ["研发工程"],
    "财务/金融岗": ["财务法务人力"],
    "法务/合规岗": ["财务法务人力", "出海合规"],
    "人力/行政岗": ["财务法务人力"],
    "供应链/采购岗": ["供应链"],
    "LCA技术岗": ["LCA/碳足迹"],
    "产品碳足迹岗": ["LCA/碳足迹"],
    "ESG岗": ["碳核算/ESG"],
    "碳核算岗": ["碳核算/ESG"],
    "CBAM/出海合规岗": ["出海合规"],
    "咨询岗": ["项目/咨询"],
    "泛运营岗": ["运营增长"],
}

SKILL_SEMANTIC_GROUPS: dict[str, list[str]] = {
    "产品能力": ["产品能力"],
    "运营": ["运营增长"],
    "用户研究": ["产品能力"],
    "项目管理": ["项目/咨询"],
    "沟通协作": ["项目/咨询", "英文沟通"],
    "业务分析": ["数据分析"],
    "市场营销": ["市场商务"],
    "销售/商务": ["市场商务"],
    "财务分析": ["财务法务人力"],
    "法务合规": ["财务法务人力", "出海合规"],
    "人力资源": ["财务法务人力"],
    "供应链管理": ["供应链"],
    "前端": ["研发工程"],
    "后端": ["研发工程"],
    "机器学习/AI": ["研发工程"],
    "LCA": ["LCA/碳足迹"],
    "ISO14067": ["LCA/碳足迹"],
    "GHG Protocol": ["碳核算/ESG"],
    "PEF": ["LCA/碳足迹"],
    "CBAM": ["出海合规"],
    "EPD": ["LCA/碳足迹"],
    "SimaPro": ["LCA/碳足迹"],
    "GaBi": ["LCA/碳足迹"],
    "openLCA": ["LCA/碳足迹"],
    "供应链碳管理": ["供应链", "碳核算/ESG"],
    "英文能力": ["英文沟通"],
    "数据分析": ["数据分析"],
    "SQL": ["数据分析"],
    "Python": ["数据分析", "研发工程"],
    "Excel": ["数据分析"],
    "Power BI": ["数据分析"],
}

CATEGORY_RULES: dict[str, list[str]] = {
    "数据/商业分析岗": ["数据分析", "商业分析", "经营分析", "SQL", "Python", "Tableau", "Power BI", "指标", "看板", "漏斗", "归因"],
    "产品岗": ["产品经理", "产品助理", "需求分析", "PRD", "原型", "用户故事", "竞品", "产品设计", "功能设计"],
    "运营岗": ["用户运营", "内容运营", "活动运营", "社群运营", "增长运营", "转化", "留存", "拉新", "运营策略"],
    "市场/销售岗": ["市场", "营销", "品牌", "投放", "渠道", "销售", "商务", "BD", "客户开发", "线索", "商机"],
    "咨询/项目岗": ["咨询", "客户", "解决方案", "项目交付", "尽调", "研究报告", "项目管理", "PMO", "交付"],
    "研发/工程岗": ["后端", "前端", "全栈", "Java", "Python", "Go", "PHP", "C++", "React", "Vue", "架构", "服务端", "算法", "开发工程师"],
    "设计岗": ["UI", "UX", "交互", "视觉", "Figma", "用户体验", "设计规范", "原型设计"],
    "财务/金融岗": ["财务", "金融", "投研", "审计", "预算", "成本", "风控", "估值", "报表"],
    "法务/合规岗": ["法务", "合规", "合同", "风控", "监管", "隐私", "数据合规", "内控"],
    "人力/行政岗": ["人力资源", "招聘", "培训", "薪酬", "绩效", "组织发展", "行政"],
    "供应链/采购岗": ["供应链", "采购", "物流", "库存", "供应商", "计划", "S&OP", "仓储"],
    "LCA技术岗": ["LCA", "生命周期", "SimaPro", "GaBi", "openLCA", "清单分析", "边界"],
    "产品碳足迹岗": ["产品碳足迹", "CFP", "ISO14067", "ISO 14067", "EPD", "PCR"],
    "ESG岗": ["ESG", "信息披露", "可持续发展报告", "评级", "CSR"],
    "碳核算岗": ["碳核算", "碳盘查", "碳核查", "GHG", "温室气体", "排放因子"],
    "CBAM/出海合规岗": ["CBAM", "碳边境", "出海", "欧盟", "合规", "EUDR", "电池法"],
    "咨询岗": ["咨询", "客户", "解决方案", "项目交付", "尽调", "研究报告"],
    "泛运营岗": ["运营", "公众号", "活动", "行政", "宣传", "品牌", "会议"],
}

INTERVIEW_CATEGORIES: dict[str, list[str]] = {
    "技术问题": [
        "LCA",
        "边界",
        "功能单位",
        "清单",
        "排放因子",
        "ISO14067",
        "CBAM",
        "EPD",
        "碳核算",
        "碳足迹",
        "数据库",
    ],
    "行为面问题": [
        "介绍",
        "项目经历",
        "困难",
        "冲突",
        "优势",
        "缺点",
        "为什么",
        "离职",
        "职业规划",
    ],
    "英文面试问题": [
        "英文",
        "英语",
        "English",
        "introduce",
        "presentation",
        "self-introduction",
    ],
    "案例分析题": ["case", "案例", "测算", "估算", "方案", "情景", "分析一个", "如何设计"],
}

CERT_KEYWORDS = ["CFA", "CPA", "FRM", "PMP", "法律职业资格", "教师资格", "人力资源证书", "SCR", "碳核查员", "ISO14064"]
PROJECT_KEYWORDS = ["数据分析", "产品", "运营", "研发", "咨询", "市场", "销售", "财务", "法务", "人力", "供应链", "制造业", "新能源", "金融", "医疗", "CBAM", "EPD"]


FULLWIDTH_TRANSLATION = str.maketrans(
    {
        **{chr(0xFF10 + index): str(index) for index in range(10)},
        **{chr(0xFF21 + index): chr(0x41 + index) for index in range(26)},
        **{chr(0xFF41 + index): chr(0x61 + index) for index in range(26)},
        "（": "(",
        "）": ")",
        "【": "[",
        "】": "]",
        "：": ":",
        "；": ";",
        "，": ",",
        "。": ".",
        "、": "/",
        "｜": "|",
        "＋": "+",
        "－": "-",
        "～": "~",
        "—": "-",
        "–": "-",
        "　": " ",
    }
)


def normalize_unicode_text(text: str | None) -> str:
    if text is None:
        return ""
    value = html.unescape(str(text))
    value = repair_private_use_digits(value)
    value = value.translate(FULLWIDTH_TRANSLATION)
    value = re.sub(r"[\u200b-\u200f\u202a-\u202e\ufeff]", "", value)
    return value


def normalize_text(text: str | None) -> str:
    if not text:
        return ""
    text = normalize_unicode_text(text)
    return re.sub(r"\s+", " ", text).strip()


def normalize_multiline_text(text: str | None) -> str:
    if not text:
        return ""
    value = normalize_unicode_text(text)
    lines = []
    for line in value.splitlines():
        clean = re.sub(r"[ \t\r\f\v]+", " ", line).strip()
        if clean:
            lines.append(clean)
    output = []
    for line in lines:
        if not output or output[-1] != line:
            output.append(line)
    return "\n".join(output).strip()


def repair_private_use_digits(text: str | None) -> str:
    if not text:
        return ""
    value = str(text)
    boss_digits = {chr(0xE031 + index): str(index) for index in range(10)}
    return value.translate(str.maketrans(boss_digits))


def content_fingerprint(text: str | None) -> str:
    normalized = normalize_text(text)
    if not normalized:
        return ""
    return hashlib.sha1(normalized.encode("utf-8", errors="ignore")).hexdigest()


def decode_bytes_best_effort(data: bytes) -> str:
    if not data:
        return ""
    candidates: list[tuple[float, str]] = []
    for encoding in ["utf-8-sig", "utf-8", "gb18030", "gbk", "cp936", "big5"]:
        try:
            value = data.decode(encoding)
        except Exception:
            continue
        normalized = normalize_multiline_text(value)
        replacement_penalty = value.count("\ufffd") * 10
        mojibake_penalty = 35 if looks_mojibake(value) else 0
        chinese_count = sum("\u4e00" <= char <= "\u9fff" for char in normalized)
        ascii_count = sum(char.isascii() and char.isalnum() for char in normalized)
        score = chinese_count * 2 + ascii_count * 0.35 + len(normalized) * 0.02 - replacement_penalty - mojibake_penalty
        candidates.append((score, value))
    if candidates:
        return max(candidates, key=lambda item: item[0])[1]
    return data.decode("utf-8", errors="ignore")


def read_text_file_best_effort(path: Path) -> str:
    return decode_bytes_best_effort(path.read_bytes())


def read_json_file_best_effort(path: Path) -> Any:
    try:
        return json.loads(path.read_text(encoding="utf-8-sig"))
    except Exception:
        return json.loads(read_text_file_best_effort(path))


def looks_mojibake(text: str | None) -> bool:
    if not text:
        return False
    marker_codes = [0x9359, 0x6D93, 0x9422, 0x5B80, 0x7EC9, 0x7039, 0x93B6, 0x93C4, 0x9496, 0x95BE]
    markers = [chr(code) for code in marker_codes]
    return sum(1 for marker in markers if marker in text) >= 2


def repair_mojibake_text(text: str | None) -> str:
    if text is None:
        return ""
    value = str(text)
    if not looks_mojibake(value):
        return value
    candidates = [value]
    for encoding in ["gb18030", "gbk", "cp936", "latin1"]:
        try:
            candidates.append(value.encode(encoding).decode("utf-8"))
        except Exception:
            continue
    readable = [candidate for candidate in candidates if candidate and not looks_mojibake(candidate)]
    if readable:
        return max(readable, key=lambda item: sum("\u4e00" <= char <= "\u9fff" for char in item))
    return value


def repair_dataframe_text(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return df
    clean = df.copy()
    for column in clean.columns:
        if not (pd.api.types.is_object_dtype(clean[column]) or pd.api.types.is_string_dtype(clean[column])):
            continue
        clean[column] = clean[column].apply(repair_mojibake_text)
    return clean


def effective_profile_text(profile_text: str | None) -> str:
    profile_text = repair_mojibake_text(profile_text)
    if not profile_text or looks_mojibake(profile_text):
        return ""
    return profile_text


def text_contains(text: str, keyword: str) -> bool:
    return keyword.lower() in text.lower()


def text_without_urls(text: str) -> str:
    return normalize_text(re.sub(r"https?://\S+", " ", str(text or ""), flags=re.I))


def count_alias_hits(text: str, aliases: list[str]) -> int:
    return sum(1 for alias in aliases if text_contains(text, alias))


@lru_cache(maxsize=2048)
def semantic_groups_for_terms_cached(label: str, aliases: tuple[str, ...] = ()) -> tuple[str, ...]:
    seeds = split_preference_items([label] + list(aliases))
    groups: list[str] = []
    for group, terms in SEMANTIC_CAPABILITY_GROUPS.items():
        group_terms = split_preference_items([group] + terms)
        for seed in seeds:
            if not seed:
                continue
            if seed == group or seed in group_terms:
                groups.append(group)
                break
            if any(seed in term or term in seed for term in group_terms if len(seed) >= 3 and len(term) >= 3):
                groups.append(group)
                break
    return tuple(dict.fromkeys(groups))


def semantic_groups_for_terms(label: str, aliases: list[str] | None = None) -> list[str]:
    return list(semantic_groups_for_terms_cached(label, tuple(aliases or ())))


@lru_cache(maxsize=2048)
def expanded_aliases_cached(label: str, aliases: tuple[str, ...] = ()) -> tuple[str, ...]:
    items = split_preference_items([label] + list(aliases))
    group_names = list(dict.fromkeys(SKILL_SEMANTIC_GROUPS.get(label, []) + list(semantic_groups_for_terms_cached(label, aliases))))
    for group in group_names:
        items.extend(SEMANTIC_CAPABILITY_GROUPS.get(group, []))
    return tuple(split_preference_items(items))


def expanded_aliases(label: str, aliases: list[str] | None = None) -> list[str]:
    return list(expanded_aliases_cached(label, tuple(aliases or ())))


def related_alias_hits(text: str, label: str, aliases: list[str] | None = None) -> list[str]:
    direct = set(alias_hits(text, aliases or [label]))
    related = []
    for alias in expanded_aliases(label, aliases):
        if alias in direct:
            continue
        if text_contains(text, alias):
            related.append(alias)
    return list(dict.fromkeys(related))


@lru_cache(maxsize=4096)
def semantic_expansion_terms(text: str) -> tuple[str, ...]:
    clean = normalize_text(text)
    expansions: list[str] = []
    for group, terms in SEMANTIC_CAPABILITY_GROUPS.items():
        hits = [term for term in terms if text_contains(clean, term)]
        if hits:
            expansions.extend([group] * min(3, len(hits)))
            expansions.extend(hits[:8])
    return tuple(split_preference_items(expansions))


def enrich_text_for_semantic_similarity(text: str) -> str:
    clean = normalize_text(text)
    expansions = semantic_expansion_terms(clean)
    if not expansions:
        return clean
    return clean + " " + " ".join(expansions)


@lru_cache(maxsize=4096)
def capability_group_hits(text: str) -> dict[str, tuple[str, ...]]:
    clean = normalize_text(text)
    hits: dict[str, tuple[str, ...]] = {}
    if not clean:
        return hits
    for group, terms in SEMANTIC_CAPABILITY_GROUPS.items():
        group_hits = [term for term in terms if text_contains(clean, term)]
        if group_hits:
            hits[group] = tuple(split_preference_items(group_hits))
    return hits


def capability_overlap_similarity(text_a: str, text_b: str) -> float:
    hits_a = capability_group_hits(text_a)
    hits_b = capability_group_hits(text_b)
    if not hits_a or not hits_b:
        return 0.0

    shared_groups = set(hits_a) & set(hits_b)
    all_groups = set(hits_a) | set(hits_b)
    if not shared_groups:
        return 0.0

    group_score = len(shared_groups) / max(len(all_groups), 1)
    evidence_score_parts = []
    for group in shared_groups:
        terms_a = set(hits_a[group])
        terms_b = set(hits_b[group])
        shared_terms = terms_a & terms_b
        if shared_terms:
            evidence_score_parts.append(1.0)
        else:
            evidence_score_parts.append(min(len(terms_a), len(terms_b)) / max(len(terms_a), len(terms_b), 1) * 0.72)
    evidence_score = sum(evidence_score_parts) / max(len(evidence_score_parts), 1)
    return float(np.clip(group_score * 0.42 + evidence_score * 0.58, 0, 1))


STOPWORDS = {
    "岗位",
    "职位",
    "职责",
    "要求",
    "负责",
    "相关",
    "公司",
    "工作",
    "能力",
    "优先",
    "具备",
    "熟悉",
    "进行",
    "以及",
    "或者",
    "我们",
    "需要",
    "以上",
    "以下",
    "包括",
}


@lru_cache(maxsize=4096)
def tokenize_for_similarity(text: str) -> str:
    text = normalize_text(text)
    if not text:
        return ""
    if jieba:
        tokens = jieba.lcut(text)
    else:
        tokens = re.findall(r"[A-Za-z][A-Za-z0-9+#.\-]{1,}|[0-9]+(?:\.[0-9]+)?|[\u4e00-\u9fa5]{2,}", text)
    cleaned = []
    for token in tokens:
        token = token.strip().lower()
        if len(token) < 2 or token in STOPWORDS:
            continue
        cleaned.append(token)
    return " ".join(cleaned)


def semantic_similarity(text_a: str, text_b: str) -> float:
    """Return a local 0-1 similarity score. Uses sklearn/rapidfuzz when available."""
    text_a = normalize_text(text_a)[:3200]
    text_b = normalize_text(text_b)[:3200]
    if not text_a or not text_b:
        return 0.0

    token_a = tokenize_for_similarity(enrich_text_for_semantic_similarity(text_a))
    token_b = tokenize_for_similarity(enrich_text_for_semantic_similarity(text_b))
    base_score = 0.0
    if TfidfVectorizer and cosine_similarity and token_a and token_b:
        try:
            matrix = TfidfVectorizer(token_pattern=r"(?u)\b\w+\b", ngram_range=(1, 2)).fit_transform([token_a, token_b])
            base_score = max(base_score, float(np.clip(cosine_similarity(matrix[0], matrix[1])[0][0], 0, 1)))
            char_matrix = TfidfVectorizer(analyzer="char", ngram_range=(2, 4)).fit_transform([text_a, text_b])
            base_score = max(base_score, float(np.clip(cosine_similarity(char_matrix[0], char_matrix[1])[0][0], 0, 1)) * 0.86)
        except Exception:
            pass
    if fuzz:
        base_score = max(base_score, float(np.clip(fuzz.token_set_ratio(text_a, text_b) / 100, 0, 1)) * 0.72)

    if not base_score:
        tokens_a = set(token_a.split())
        tokens_b = set(token_b.split())
        if tokens_a and tokens_b:
            base_score = len(tokens_a & tokens_b) / max(len(tokens_a | tokens_b), 1)

    capability_score = capability_overlap_similarity(text_a, text_b)
    if capability_score:
        blended = base_score * 0.58 + capability_score * 0.42
        capability_floor = capability_score * (0.78 if base_score >= 0.08 else 0.62)
        base_score = max(base_score, blended, capability_floor)
    return float(np.clip(base_score, 0, 1))


@lru_cache(maxsize=8192)
def semantic_similarity_fast(text_a: str, text_b: str) -> float:
    text_a = normalize_text(text_a)[:2200]
    text_b = normalize_text(text_b)[:2200]
    if not text_a or not text_b:
        return 0.0
    token_a = tokenize_for_similarity(enrich_text_for_semantic_similarity(text_a))
    token_b = tokenize_for_similarity(enrich_text_for_semantic_similarity(text_b))
    tokens_a = set(token_a.split())
    tokens_b = set(token_b.split())
    base_score = 0.0
    if tokens_a and tokens_b:
        overlap = len(tokens_a & tokens_b) / max(len(tokens_a | tokens_b), 1)
        coverage = len(tokens_a & tokens_b) / max(min(len(tokens_a), len(tokens_b)), 1)
        base_score = max(base_score, overlap * 0.72 + coverage * 0.28)
    if fuzz:
        base_score = max(base_score, float(np.clip(fuzz.token_set_ratio(text_a, text_b) / 100, 0, 1)) * 0.62)
    capability_score = capability_overlap_similarity(text_a, text_b)
    if capability_score:
        base_score = max(base_score, base_score * 0.58 + capability_score * 0.42, capability_score * 0.68)
    return float(np.clip(base_score, 0, 1))


def alias_hits(text: str, aliases: list[str]) -> list[str]:
    return [alias for alias in aliases if text_contains(text, alias)]


def keyword_table(text: str) -> pd.DataFrame:
    rows = []
    for skill, aliases in SKILL_ALIASES.items():
        hits = count_alias_hits(text, aliases)
        related = related_alias_hits(text, skill, aliases)
        if hits or related:
            hit_words = [a for a in aliases if text_contains(text, a)]
            hit_words.extend(item for item in related if item not in hit_words)
            rows.append({"技能": skill, "命中次数": hits + len(related), "命中词": " / ".join(hit_words[:10])})
    if not rows:
        return pd.DataFrame(columns=["技能", "命中次数", "命中词"])
    return pd.DataFrame(rows).sort_values(["命中次数", "技能"], ascending=[False, True]).reset_index(drop=True)


def score_keywords(text: str, keywords: list[str]) -> int:
    return sum(1 for keyword in keywords if text_contains(text, keyword))


def extract_by_patterns(text: str, patterns: list[str], default: str = "未识别") -> str:
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.I)
        if match:
            value = match.group(1).strip(" ：:，,;；|")
            if value:
                return value[:80]
    return default


COMPANY_SUFFIX_PATTERN = (
    r"(?:有限公司|有限责任公司|股份有限公司|集团|科技|咨询|检测|认证|事务所|研究院|设计院|"
    r"中心|公司|工厂|厂|商行|合作社|医院|学校|学院|米业|家具|酒店|餐饮|贸易|"
    r"实业|制造|传媒|教育|医药|生物|物流|电子)"
)

COMPANY_LABEL_PATTERN = (
    r"(?:公司名称|公司名|企业名称|企业名|单位名称|单位名|招聘单位|用人单位|雇主|"
    r"所属公司|发布公司|Employer|Company|Organization)"
)


def clean_company_candidate(value: str) -> str:
    value = normalize_text(value)
    value = re.sub(rf"^(?:{COMPANY_LABEL_PATTERN})\s*[:：]?\s*", "", value, flags=re.I)
    value = value.strip(" ：:，,;；|（）()[]【】")
    for boundary in [
        " 岗位",
        " 职位",
        " 招聘",
        " 薪资",
        " 月薪",
        " 地点",
        " 城市",
        " 地址",
        " 学历",
        " 经验",
        " 福利",
        " 发布时间",
        " 链接",
        " URL",
    ]:
        position = value.find(boundary)
        if position > 0:
            value = value[:position]
    return value.strip(" ：:，,;；|")[:90]


def is_probable_company_name(value: str) -> bool:
    clean = clean_company_candidate(value)
    if not (2 <= len(clean) <= 90):
        return False
    if re.search(r"https?://|www\.|@|^\d+[\d,.\-~ 至到]*[kKwW万千元]|^[\d\s\-~至到]+$", clean, flags=re.I):
        return False
    bad_words = ["岗位", "职位", "薪资", "月薪", "学历", "经验", "职责", "要求", "福利", "登录", "注册", "首页", "筛选", "排序", "详情", "查看", "招聘人数"]
    if any(word in clean for word in bad_words):
        return False
    geographic_only = {"北京", "上海", "广州", "深圳", "杭州", "苏州", "浙江省", "黑龙江省", "安徽省", "广东省", "江苏省"}
    if clean in geographic_only:
        return False
    return bool(re.search(COMPANY_SUFFIX_PATTERN, clean))


def company_candidate_score(value: str, context: str = "") -> int:
    clean = clean_company_candidate(value)
    if not is_probable_company_name(clean):
        return 0
    score = 10
    if re.search(r"(?:有限公司|有限责任公司|股份有限公司|集团|公司)$", clean):
        score += 12
    elif re.search(COMPANY_SUFFIX_PATTERN + r"$", clean):
        score += 8
    if re.search(COMPANY_LABEL_PATTERN, context, flags=re.I):
        score += 10
    if 4 <= len(clean) <= 40:
        score += 4
    return score


def extract_company_candidates(text: str, title: str = "") -> list[str]:
    candidates: list[tuple[int, str]] = []
    multiline = normalize_multiline_text(text)
    compact = normalize_text(text)

    for pattern in [
        rf"(?:{COMPANY_LABEL_PATTERN}|公司|企业|单位)\s*[:：]\s*([^\n，。；;|]{{2,90}})",
        rf"{COMPANY_LABEL_PATTERN}\s+([^\n，。；;|]{{2,90}})",
    ]:
        for match in re.finditer(pattern, multiline, flags=re.I):
            candidate = clean_company_candidate(match.group(1))
            score = company_candidate_score(candidate, match.group(0))
            if score:
                candidates.append((score + 8, candidate))

    lines = [line.strip() for line in multiline.splitlines() if line.strip()]
    title_clean = clean_company_candidate(title)
    title_indexes = [i for i, item in enumerate(lines) if title_clean and title_clean in item]
    for index, line in enumerate(lines):
        candidate = clean_company_candidate(line)
        if title_clean and candidate == title_clean:
            continue
        score = company_candidate_score(candidate, line)
        if score:
            if index <= 4:
                score += 3
            if title_indexes and min(abs(index - item) for item in title_indexes) <= 3:
                score += 5
            candidates.append((score, candidate))

    suffix_pattern = rf"([\u4e00-\u9fa5A-Za-z0-9（）()·&\-]{{2,60}}{COMPANY_SUFFIX_PATTERN})"
    for match in re.finditer(suffix_pattern, compact):
        candidate = clean_company_candidate(match.group(1))
        score = company_candidate_score(candidate, match.group(0))
        if score:
            candidates.append((score, candidate))

    best_by_value: dict[str, int] = {}
    for score, candidate in candidates:
        best_by_value[candidate] = max(best_by_value.get(candidate, 0), score)
    return [candidate for candidate, _score in sorted(best_by_value.items(), key=lambda item: item[1], reverse=True)]


def extract_salary(text: str) -> str:
    clean = normalize_text(text)
    patterns = [
        r"(\d+(?:\.\d+)?\s*[kK千万wW]?\s*[-~—至到]\s*\d+(?:\.\d+)?\s*(?:[kK千万wW]|元\s*/?\s*(?:天|日|月)|元/天|元/日|元/月)(?:\s*[·*/xX]\s*\d+\s*薪)?)",
        r"(\d+(?:\.\d+)?\s*(?:元\s*/?\s*(?:天|日|月)|元/天|元/日|元/月)\s*[-~—至到]\s*\d+(?:\.\d+)?\s*(?:元\s*/?\s*(?:天|日|月)|元/天|元/日|元/月)?)",
        r"((?:薪资|薪酬|月薪|日薪|工资|待遇)[:：]?\s*[^，。；;\n]{2,30})",
        r"(面议|薪资面议|薪酬面议)",
        r"(\d+(?:\.\d+)?\s*/\s*(?:天|日|月))",
    ]
    for pattern in patterns:
        match = re.search(pattern, clean)
        if not match:
            continue
        candidate = normalize_text(match.group(1))
        if not candidate:
            continue
        if "面议" in candidate:
            return candidate
        if re.search(r"\d+\s*[-~—至到]\s*\d+\s*年", candidate):
            continue
        if re.search(r"(经验|学历|本科|硕士|大专|天/周|个月)", candidate):
            continue
        if re.search(r"(?:[kK千万wW]|\d+\s*/\s*(?:天|日|月)|元\s*/?\s*(?:天|日|月)|元/天|元/日|元/月|[·xX*]\s*\d+\s*薪|\d+\s*薪)", candidate):
            return candidate
    return "未识别"


def extract_location(text: str) -> str:
    explicit = extract_by_patterns(
        text,
        [
            r"(?:工作地点|地点|办公地点|城市)[:：]\s*([^，。；;\n]{2,30})",
            r"(?:base|Base|BASE)[:：]?\s*([^，。；;\n]{2,30})",
        ],
    )
    if explicit != "未识别":
        return explicit
    cities = list(CITY_REGION_MAP)
    found = [city for city in cities if any(text_contains(text, alias) for alias in city_aliases(city))]
    return " / ".join(found[:3]) if found else "未识别"


CITY_PROVINCE_MAP = {
    "上海": "上海市",
    "北京": "北京市",
    "天津": "天津市",
    "重庆": "重庆市",
    "苏州": "江苏省",
    "南京": "江苏省",
    "无锡": "江苏省",
    "常州": "江苏省",
    "南通": "江苏省",
    "扬州": "江苏省",
    "镇江": "江苏省",
    "泰州": "江苏省",
    "盐城": "江苏省",
    "徐州": "江苏省",
    "淮安": "江苏省",
    "连云港": "江苏省",
    "宿迁": "江苏省",
    "杭州": "浙江省",
    "宁波": "浙江省",
    "嘉兴": "浙江省",
    "湖州": "浙江省",
    "绍兴": "浙江省",
    "金华": "浙江省",
    "义乌": "浙江省",
    "温州": "浙江省",
    "台州": "浙江省",
    "舟山": "浙江省",
    "衢州": "浙江省",
    "丽水": "浙江省",
    "合肥": "安徽省",
    "芜湖": "安徽省",
    "马鞍山": "安徽省",
    "铜陵": "安徽省",
    "安庆": "安徽省",
    "滁州": "安徽省",
    "蚌埠": "安徽省",
    "阜阳": "安徽省",
    "厦门": "福建省",
    "福州": "福建省",
    "泉州": "福建省",
    "漳州": "福建省",
    "莆田": "福建省",
    "宁德": "福建省",
    "龙岩": "福建省",
    "青岛": "山东省",
    "济南": "山东省",
    "烟台": "山东省",
    "潍坊": "山东省",
    "淄博": "山东省",
    "威海": "山东省",
    "临沂": "山东省",
    "济宁": "山东省",
    "东营": "山东省",
    "深圳": "广东省",
    "广州": "广东省",
    "佛山": "广东省",
    "东莞": "广东省",
    "珠海": "广东省",
    "中山": "广东省",
    "惠州": "广东省",
    "江门": "广东省",
    "肇庆": "广东省",
    "汕头": "广东省",
    "武汉": "湖北省",
    "宜昌": "湖北省",
    "襄阳": "湖北省",
    "荆州": "湖北省",
    "黄石": "湖北省",
    "长沙": "湖南省",
    "株洲": "湖南省",
    "湘潭": "湖南省",
    "岳阳": "湖南省",
    "衡阳": "湖南省",
    "郑州": "河南省",
    "洛阳": "河南省",
    "开封": "河南省",
    "新乡": "河南省",
    "许昌": "河南省",
    "成都": "四川省",
    "绵阳": "四川省",
    "德阳": "四川省",
    "宜宾": "四川省",
    "泸州": "四川省",
    "西安": "陕西省",
    "咸阳": "陕西省",
    "宝鸡": "陕西省",
    "石家庄": "河北省",
    "唐山": "河北省",
    "保定": "河北省",
    "廊坊": "河北省",
    "邯郸": "河北省",
    "秦皇岛": "河北省",
    "沈阳": "辽宁省",
    "大连": "辽宁省",
    "鞍山": "辽宁省",
    "长春": "吉林省",
    "吉林": "吉林省",
    "哈尔滨": "黑龙江省",
    "大庆": "黑龙江省",
    "南昌": "江西省",
    "赣州": "江西省",
    "九江": "江西省",
    "南宁": "广西壮族自治区",
    "柳州": "广西壮族自治区",
    "桂林": "广西壮族自治区",
    "海口": "海南省",
    "三亚": "海南省",
    "贵阳": "贵州省",
    "遵义": "贵州省",
    "昆明": "云南省",
    "曲靖": "云南省",
    "兰州": "甘肃省",
    "乌鲁木齐": "新疆维吾尔自治区",
    "呼和浩特": "内蒙古自治区",
    "包头": "内蒙古自治区",
    "鄂尔多斯": "内蒙古自治区",
    "太原": "山西省",
    "大同": "山西省",
    "南阳": "河南省",
    "香港": "港澳台",
    "澳门": "港澳台",
    "台北": "港澳台",
}
CITY_PROVINCE_MAP.update(
    {
        city: province
        for province, province_cities in CHINA_285_CITY_OPTIONS_BY_PROVINCE.items()
        for city in province_cities
    }
)

PROVINCE_REGION_MAP = {
    "上海市": "上海",
    "江苏省": "长三角",
    "浙江省": "长三角",
    "安徽省": "长三角",
    "北京市": "华北",
    "天津市": "华北",
    "河北省": "华北",
    "山东省": "华北",
    "广东省": "华南",
    "福建省": "华东其他",
    "湖北省": "华中",
    "湖南省": "华中",
    "河南省": "华中",
    "四川省": "西南",
    "重庆市": "西南",
    "陕西省": "西北",
    "辽宁省": "东北",
    "吉林省": "东北",
    "黑龙江省": "东北",
    "江西省": "华东其他",
    "广西壮族自治区": "华南",
    "海南省": "华南",
    "贵州省": "西南",
    "云南省": "西南",
    "甘肃省": "西北",
    "宁夏回族自治区": "西北",
    "青海省": "西北",
    "西藏自治区": "西南",
    "新疆维吾尔自治区": "西北",
    "内蒙古自治区": "华北",
    "山西省": "华北",
    "港澳台": "港澳台",
}

CITY_REGION_MAP = {
    city: PROVINCE_REGION_MAP.get(province, "其他地区")
    for city, province in CITY_PROVINCE_MAP.items()
}
CITY_REGION_MAP["上海"] = "上海"

CITY_ALIAS_MAP: dict[str, list[str]] = {
    "北京": ["京", "帝都", "北京市"],
    "上海": ["沪", "魔都", "上海市"],
    "广州": ["穗", "羊城", "广州市"],
    "深圳": ["深", "鹏城", "深圳市"],
    "杭州": ["杭", "杭州市"],
    "南京": ["宁", "金陵", "南京市"],
    "苏州": ["苏", "苏州市"],
    "成都": ["蓉", "成都市"],
    "重庆": ["渝", "重庆市"],
    "武汉": ["汉", "武汉市"],
    "西安": ["镐", "西安市"],
    "天津": ["津", "天津市"],
    "厦门": ["鹭", "厦门市"],
    "长沙": ["长株潭", "长沙市"],
    "青岛": ["青", "青岛市"],
    "宁波": ["甬", "宁波市"],
}

CITY_GROUP_MAP: dict[str, list[str]] = {
    "一线城市": ["北京", "上海", "广州", "深圳"],
    "北上广深": ["北京", "上海", "广州", "深圳"],
    "新一线": ["成都", "杭州", "重庆", "武汉", "苏州", "西安", "南京", "长沙", "天津", "郑州", "东莞", "青岛", "昆明", "宁波", "合肥"],
    "长三角": ["上海", "南京", "苏州", "无锡", "常州", "南通", "杭州", "宁波", "嘉兴", "湖州", "绍兴", "合肥", "芜湖"],
    "江浙沪": ["上海", "南京", "苏州", "无锡", "常州", "南通", "杭州", "宁波", "嘉兴", "湖州", "绍兴"],
    "珠三角": ["广州", "深圳", "佛山", "东莞", "珠海", "中山", "惠州", "江门", "肇庆"],
    "大湾区": ["广州", "深圳", "佛山", "东莞", "珠海", "中山", "惠州", "江门", "肇庆", "香港", "澳门"],
    "粤港澳大湾区": ["广州", "深圳", "佛山", "东莞", "珠海", "中山", "惠州", "江门", "肇庆", "香港", "澳门"],
    "京津冀": ["北京", "天津", "石家庄", "唐山", "保定", "廊坊", "邯郸", "秦皇岛", "张家口", "承德", "沧州", "衡水", "邢台"],
    "成渝": ["成都", "重庆", "绵阳", "德阳", "眉山", "资阳"],
    "华东": ["上海", "南京", "苏州", "杭州", "宁波", "合肥", "福州", "厦门", "南昌", "济南", "青岛"],
    "华南": ["广州", "深圳", "佛山", "东莞", "珠海", "中山", "南宁", "海口", "三亚", "厦门"],
    "华北": ["北京", "天津", "石家庄", "太原", "济南", "青岛", "呼和浩特"],
    "华中": ["武汉", "长沙", "郑州", "南昌", "合肥"],
    "西南": ["成都", "重庆", "昆明", "贵阳", "南宁"],
    "西北": ["西安", "兰州", "银川", "乌鲁木齐"],
    "东北": ["沈阳", "大连", "长春", "哈尔滨", "大庆"],
}

PROVINCE_BOXES = {
    "北京市": (115.7, 39.4, 117.4, 41.1),
    "天津市": (116.7, 38.6, 118.1, 40.3),
    "河北省": (113.4, 36.0, 119.9, 42.6),
    "山东省": (114.5, 34.3, 122.8, 38.4),
    "上海市": (120.8, 30.7, 122.1, 31.9),
    "江苏省": (116.3, 30.7, 121.9, 35.2),
    "浙江省": (118.0, 27.0, 123.0, 31.3),
    "安徽省": (114.8, 29.4, 119.7, 34.7),
    "福建省": (115.8, 23.5, 120.7, 28.6),
    "广东省": (109.6, 20.1, 117.3, 25.6),
    "湖北省": (108.3, 29.0, 116.2, 33.4),
    "湖南省": (108.8, 24.6, 114.3, 30.1),
    "河南省": (110.3, 31.4, 116.7, 36.4),
    "四川省": (97.3, 26.0, 108.5, 34.3),
    "重庆市": (105.3, 28.1, 110.2, 32.2),
    "陕西省": (105.5, 31.7, 111.3, 39.6),
    "辽宁省": (118.8, 38.7, 125.8, 43.5),
    "吉林省": (121.6, 40.8, 131.3, 46.3),
    "黑龙江省": (121.2, 43.4, 135.1, 53.6),
    "江西省": (113.5, 24.5, 118.5, 30.1),
    "广西壮族自治区": (104.4, 20.8, 112.1, 26.4),
    "海南省": (108.6, 18.0, 111.2, 20.2),
    "贵州省": (103.6, 24.6, 109.6, 29.3),
    "云南省": (97.5, 21.1, 106.2, 29.3),
    "甘肃省": (92.2, 32.1, 108.7, 42.8),
    "宁夏回族自治区": (104.0, 35.0, 107.7, 39.4),
    "青海省": (89.4, 31.6, 103.1, 39.2),
    "西藏自治区": (78.4, 26.8, 99.2, 36.5),
    "新疆维吾尔自治区": (73.5, 34.0, 96.4, 49.2),
    "内蒙古自治区": (97.2, 37.4, 126.1, 53.3),
    "山西省": (110.2, 34.5, 114.6, 40.7),
}
ONLINE_CHINA_GEOJSON_URL = "https://geo.datav.aliyun.com/areas_v3/bound/100000_full.json"


@lru_cache(maxsize=512)
def city_aliases(city: str) -> tuple[str, ...]:
    clean = normalize_text(city)
    aliases = [clean]
    aliases.extend(CITY_ALIAS_MAP.get(clean, []))
    if clean.endswith("市"):
        aliases.append(clean[:-1])
    return tuple(split_preference_items(aliases))


def expand_city_preference(city_or_group: str) -> list[str]:
    clean = normalize_text(city_or_group)
    if clean in CITY_GROUP_MAP:
        return CITY_GROUP_MAP[clean]
    if clean in PROVINCE_REGION_MAP:
        return [city for city, province in CITY_PROVINCE_MAP.items() if province == clean]
    if clean.endswith("省") or clean.endswith("自治区") or clean.endswith("市"):
        province_cities = [city for city, province in CITY_PROVINCE_MAP.items() if province == clean]
        if province_cities:
            return province_cities
    return [clean]


def city_match_label(location_pool: str, target_cities: list[str]) -> str:
    for item in split_preference_items(target_cities):
        expanded_cities = expand_city_preference(item)
        for city in expanded_cities:
            aliases = city_aliases(city)
            if any(text_contains(location_pool, alias) for alias in aliases):
                return item if item == city else f"{item}({city})"
        if item in CITY_GROUP_MAP and text_contains(location_pool, item):
            return item
    return ""


@lru_cache(maxsize=4096)
def extract_standard_cities_cached(location: str, text: str = "") -> tuple[str, ...]:
    combined = normalize_text(f"{location} {text}")
    cities: list[str] = []
    for city in CITY_REGION_MAP:
        if any(text_contains(combined, alias) for alias in city_aliases(city)):
            cities.append(city)
    for group, group_cities in CITY_GROUP_MAP.items():
        if text_contains(combined, group):
            cities.extend(group_cities)
    return tuple(dict.fromkeys(cities))


def extract_standard_cities(location: str, text: str = "") -> list[str]:
    return list(extract_standard_cities_cached(location, text))


def extract_provinces(location: str, text: str = "") -> list[str]:
    """Map recognized cities to provinces; avoid guessing province from province words alone."""
    provinces = []
    for city in extract_standard_cities(location, text):
        province = CITY_PROVINCE_MAP.get(city)
        if province:
            provinces.append(province)
    return list(dict.fromkeys(provinces))


def classify_region(location: str, text: str = "") -> dict[str, str]:
    combined = normalize_text(f"{location} {text}")
    if any(text_contains(combined, keyword) for keyword in ["远程", "居家", "remote", "Remote"]):
        return {"标准城市": "远程", "省份": "远程"}
    has_city_group = any(text_contains(combined, group) for group in CITY_GROUP_MAP)
    if not has_city_group and any(text_contains(combined, keyword) for keyword in ["全国", "多地", "不限城市", "多个城市", "异地办公"]):
        return {"标准城市": "多地/全国", "省份": "多地/全国"}

    cities = extract_standard_cities(location, text)
    provinces = extract_provinces(location, text)
    if not cities and not provinces:
        location_only = normalize_text(location)
        overseas_text = location_only if location_only and location_only != "未识别" else combined
        if any(text_contains(overseas_text, keyword) for keyword in ["海外", "新加坡", "香港", "澳门", "台湾"]):
            return {"标准城市": "海外/港澳台", "省份": "海外/港澳台"}
        return {"标准城市": "未识别", "省份": "未识别"}

    if not provinces:
        provinces = [CITY_PROVINCE_MAP[city] for city in cities if city in CITY_PROVINCE_MAP]
    province_label = " / ".join(provinces[:4]) if provinces else "未识别"

    return {
        "标准城市": " / ".join(cities[:4]),
        "省份": province_label,
    }


def split_location_values(value: Any) -> list[str]:
    if value is None:
        return []
    values = []
    for item in re.split(r"\s*/\s*|,|，|、", str(value)):
        item = item.strip()
        if item:
            values.append(item)
    return list(dict.fromkeys(values))


def province_filter_options(df: pd.DataFrame) -> list[str]:
    if df is None or df.empty or "省份" not in df.columns:
        return ["全部"]
    values: list[str] = []
    for item in df["省份"].dropna().astype(str):
        values.extend(split_location_values(item))
    provinces = sorted(value for value in set(values) if value in PROVINCE_BOXES)
    special = [value for value in ["远程", "多地/全国", "海外/港澳台", "未识别"] if value in values]
    return ["全部"] + provinces + special


def filter_by_province(df: pd.DataFrame, selected: str) -> pd.DataFrame:
    if selected == "全部" or df is None or df.empty or "省份" not in df.columns:
        return df
    mask = df["省份"].astype(str).apply(lambda value: selected in split_location_values(value))
    return df[mask]


def recognized_province_total(df: pd.DataFrame) -> int:
    if df is None or df.empty or "省份" not in df.columns:
        return 0
    values: set[str] = set()
    for item in df["省份"].dropna().astype(str):
        values.update(value for value in split_location_values(item) if value in PROVINCE_BOXES)
    return len(values)


def province_counts_from_df(df: pd.DataFrame) -> pd.DataFrame:
    if df is None or df.empty or "省份" not in df.columns:
        return pd.DataFrame(columns=["省份", "岗位数"])
    counts: Counter[str] = Counter()
    for value in df["省份"].dropna().astype(str):
        for province in split_location_values(value):
            if province in PROVINCE_BOXES:
                counts[province] += 1
    if not counts:
        return pd.DataFrame(columns=["省份", "岗位数"])
    return pd.DataFrame([{"省份": key, "岗位数": value} for key, value in counts.items()]).sort_values("岗位数", ascending=False)


def province_map_points_from_counts(province_df: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for _, row in province_df.iterrows():
        province = str(row["省份"])
        bbox = PROVINCE_BOXES.get(province)
        if not bbox:
            continue
        min_lon, min_lat, max_lon, max_lat = bbox
        count = int(row["岗位数"])
        rows.append(
            {
                "省份": province,
                "显示名": province.replace("省", "").replace("市", "").replace("壮族自治区", "").replace("维吾尔自治区", "").replace("自治区", ""),
                "岗位数": count,
                "lon": (min_lon + max_lon) / 2,
                "lat": (min_lat + max_lat) / 2,
                "marker_size": int(np.clip(18 + np.sqrt(count) * 12, 20, 64)),
            }
        )
    return pd.DataFrame(rows)


@st.cache_data(ttl=24 * 60 * 60, show_spinner=False)
def load_online_china_geojson() -> dict[str, Any] | None:
    try:
        response = requests.get(ONLINE_CHINA_GEOJSON_URL, timeout=8)
        response.raise_for_status()
        data = response.json()
        if data.get("type") == "FeatureCollection":
            return data
    except Exception:
        return None
    return None


def render_china_province_map(df: pd.DataFrame, title: str = "省份岗位分布") -> None:
    province_df = province_counts_from_df(df)
    if province_df.empty:
        st.caption("暂无可绘制的省份数据。")
        return
    geojson = load_online_china_geojson()
    if geojson:
        try:
            fig = px.choropleth_mapbox(
                province_df,
                geojson=geojson,
                locations="省份",
                color="岗位数",
                featureidkey="properties.name",
                color_continuous_scale="YlGnBu",
                mapbox_style="open-street-map",
                center={"lat": 35.5, "lon": 104.0},
                zoom=3,
                opacity=0.72,
                hover_name="省份",
                hover_data={"岗位数": True, "省份": False},
                title=title,
            )
            fig.update_layout(height=560, margin=dict(l=0, r=0, t=45, b=0))
            st.plotly_chart(fig, width="stretch")
            st.caption("地图使用在线省级边界 GeoJSON 与 OpenStreetMap 底图；颜色深浅表示各省岗位数量。")
            return
        except Exception:
            pass
    map_df = province_map_points_from_counts(province_df)
    if map_df.empty:
        st.caption("暂无可在地图上定位的省份数据。")
        return
    fig = go.Figure()
    fig.add_trace(
        go.Densitymapbox(
            lat=map_df["lat"],
            lon=map_df["lon"],
            z=map_df["岗位数"],
            radius=45,
            colorscale="YlGnBu",
            opacity=0.45,
            showscale=False,
            hoverinfo="skip",
        )
    )
    fig.add_trace(
        go.Scattermapbox(
            lat=map_df["lat"],
            lon=map_df["lon"],
            mode="markers",
            marker=dict(
                size=map_df["marker_size"] + 5,
                color="white",
                opacity=0.88,
            ),
            hoverinfo="skip",
        )
    )
    fig.add_trace(
        go.Scattermapbox(
            lat=map_df["lat"],
            lon=map_df["lon"],
            mode="markers+text",
            text=map_df["显示名"],
            textposition="top center",
            customdata=map_df[["省份", "岗位数"]].to_numpy(),
            marker=dict(
                size=map_df["marker_size"],
                color=map_df["岗位数"],
                colorscale="YlGnBu",
                opacity=0.88,
                colorbar=dict(title="岗位数"),
            ),
            hovertemplate="%{customdata[0]}<br>岗位数：%{customdata[1]}<extra></extra>",
        )
    )
    fig.update_layout(
        title=title,
        height=560,
        margin=dict(l=0, r=0, t=45, b=0),
        mapbox=dict(
            style="open-street-map",
            center=dict(lat=35.5, lon=104.0),
            zoom=3,
        ),
        showlegend=False,
    )
    st.plotly_chart(fig, width="stretch")
    st.caption("地图底图使用 OpenStreetMap 在线瓦片；气泡和热力强度表示各省岗位数量。")


def extract_company(text: str) -> str:
    candidates = extract_company_candidates(text)
    if candidates:
        return candidates[0]
    explicit = extract_by_patterns(
        text,
        [
            rf"(?:{COMPANY_LABEL_PATTERN}|公司|企业|单位)\s*[:：]\s*([^，。；;\n]{{2,80}}?)(?=\s+(?:岗位名称|职位名称|招聘职位|岗位|职位|地点|城市|薪资|学历|经验|链接|URL|JD|职责|要求)[:：]|[，。；;\n]|$)",
            rf"([^，。；;\n]{{2,45}}{COMPANY_SUFFIX_PATTERN})",
        ],
    )
    return explicit


def extract_job_title(text: str) -> str:
    explicit = extract_by_patterns(
        text,
        [
            r"(?:岗位名称|职位名称|招聘职位|岗位|职位|Job Title|Position)[:：]\s*([^，。；;\n]{2,80}?)(?=\s+(?:公司名称|企业名称|公司|企业|地点|城市|薪资|学历|经验|链接|URL|JD|职责|要求)[:：]|[，。；;\n]|$)",
            r"招聘\s*([^，。；;\n]{2,40}(?:工程师|顾问|专员|分析师|经理|实习生))",
            r"([^，。；;\n]{2,40}(?:产品|运营|数据|研发|算法|前端|后端|市场|销售|财务|法务|人力|供应链|咨询|项目|合规|LCA|ESG|CBAM)[^，。；;\n]{0,20}(?:工程师|顾问|专员|分析师|经理|实习生|助理|管培生)?)",
        ],
    )
    return explicit


def extract_education(text: str) -> str:
    levels = []
    if re.search(r"博士|PhD", text, flags=re.I):
        levels.append("博士")
    if re.search(r"硕士|研究生|master", text, flags=re.I):
        levels.append("硕士")
    if re.search(r"本科|学士|bachelor", text, flags=re.I):
        levels.append("本科")
    if re.search(r"大专|专科", text, flags=re.I):
        levels.append("大专")
    return " / ".join(levels) if levels else "未识别"


def extract_experience(text: str) -> str:
    return extract_by_patterns(
        text,
        [
            r"((?:\d+\s*[-~至到]\s*)?\d+\s*年(?:以上)?(?:工作)?经验)",
            r"(经验不限)",
            r"(应届生|校招|秋招|实习)",
        ],
    )


def classify_job(text: str) -> tuple[str, dict[str, int]]:
    scores = {}
    for category, words in CATEGORY_RULES.items():
        aliases = list(words)
        for group in CATEGORY_SEMANTIC_GROUPS.get(category, []):
            aliases.extend(SEMANTIC_CAPABILITY_GROUPS.get(group, []))
        scores[category] = score_keywords(text, split_preference_items(aliases))
    if not any(scores.values()):
        return "待人工判断", scores
    category = max(scores, key=scores.get)
    return category, scores


def low_value_risk_labels(text: str) -> tuple[list[str], list[str]]:
    tags: list[str] = []
    details: list[str] = []
    for tag, keywords in LOW_VALUE_RISK_RULES.items():
        hits = [word for word in keywords if text_contains(text, word)]
        if hits:
            tags.append(tag)
            details.append(f"{tag}：" + " / ".join(hits[:4]))
    return tags, details


def value_verdict(text: str, category: str) -> dict[str, Any]:
    high_score = score_keywords(text, HIGH_VALUE_KEYWORDS)
    low_score = score_keywords(text, LOW_VALUE_KEYWORDS)
    risk_tags, risk_details = low_value_risk_labels(text)
    high_value_categories = [
        "数据/商业分析岗",
        "产品岗",
        "咨询/项目岗",
        "研发/工程岗",
        "财务/金融岗",
        "法务/合规岗",
        "供应链/采购岗",
        "LCA技术岗",
        "产品碳足迹岗",
        "CBAM/出海合规岗",
        "碳核算岗",
    ]
    high_value = high_score >= 3 or (high_score >= 2 and category in high_value_categories)
    generic_esg = (low_score >= 2 or len(risk_tags) >= 2) and high_score <= 2
    if category == "泛运营岗":
        generic_esg = True
    if high_value and not generic_esg:
        label = "岗位职责相对清晰，能沉淀可量化成果，建议优先深看"
    elif generic_esg:
        label = "疑似低价值/杂务型岗位，建议先核实真实职责和产出"
    else:
        label = "中等价值岗位，需要结合薪资与项目内容判断"
    return {
        "high_score": high_score,
        "low_score": low_score,
        "is_high_value": high_value,
        "is_generic_esg": generic_esg,
        "is_low_value_risk": generic_esg,
        "risk_tags": risk_tags,
        "risk_details": risk_details,
        "label": label,
    }


def analyze_jd(text: str) -> dict[str, Any]:
    clean = text_without_urls(normalize_text(text))
    category, category_scores = classify_job(clean)
    value = value_verdict(clean, category)
    skills = keyword_table(clean)
    location = extract_location(clean)
    region_info = classify_region(location, clean)
    return {
        "raw_text": clean,
        "basic": {
            "公司名": extract_company(clean),
            "岗位名": extract_job_title(clean),
            "地点": location,
            **region_info,
            "薪资": extract_salary(clean),
            "学历要求": extract_education(clean),
            "经验要求": extract_experience(clean),
        },
        "skills": skills,
        "category": category,
        "category_scores": category_scores,
        "value": value,
    }


def clear_jd_dependent_results() -> None:
    for key in [
        "resume_match",
        "resume_text",
        "custom_resume",
        "gap_analysis",
        "offer_prediction",
    ]:
        st.session_state.pop(key, None)


def clear_resume_dependent_results() -> None:
    for key in [
        "resume_match",
        "resume_text",
        "custom_resume",
        "gap_analysis",
        "offer_prediction",
    ]:
        st.session_state.pop(key, None)


def clear_legacy_runtime_state() -> None:
    if st.session_state.get("_careerpilot_runtime_cache_schema") == RUNTIME_CACHE_SCHEMA:
        return
    try:
        st.cache_data.clear()
    except Exception:
        pass
    for key in [
        "main_workspace",
        "batch_export_records",
        "batch_export_table",
        "batch_export_deleted_ids",
        "batch_export_selected_ids",
        "batch_export_info_set_selected_ids",
        "batch_url_records",
        "batch_url_deleted_ids",
        "batch_url_selected_ids",
        "batch_url_crawl_summary",
        "batch_jd_analysis",
        "jd_exported_text",
        "jd_exported_table",
        "jd_crawl_results",
        "crawled_jd_text",
        "recruitment_monitor",
    ]:
        st.session_state.pop(key, None)
    st.session_state["_careerpilot_runtime_cache_schema"] = RUNTIME_CACHE_SCHEMA


def set_current_target_jd(text: str, source: str = "单条JD分析", title: str = "") -> dict[str, Any] | None:
    clean = normalize_text(text)
    if not clean:
        return None
    fingerprint = content_fingerprint(clean)
    if st.session_state.get("target_jd_fingerprint") and st.session_state.get("target_jd_fingerprint") != fingerprint:
        clear_jd_dependent_results()
    jd_analysis = analyze_jd(clean)
    st.session_state.jd_analysis = jd_analysis
    st.session_state.target_jd_fingerprint = fingerprint
    st.session_state.target_jd_meta = {
        "source": source,
        "title": title or jd_analysis.get("basic", {}).get("岗位名", "") or jd_analysis.get("category", "目标JD"),
        "updated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }
    return jd_analysis


def current_jd_fingerprint() -> str:
    if st.session_state.get("target_jd_fingerprint"):
        return st.session_state["target_jd_fingerprint"]
    jd_analysis = st.session_state.get("jd_analysis")
    if not jd_analysis:
        return ""
    return content_fingerprint(jd_analysis.get("raw_text", ""))


def resume_strengths(resume_text: str, profile_text: str | None = None) -> list[str]:
    profile_text = normalize_text(effective_profile_text(profile_text) + "\n" + resume_text)
    strengths = []
    checks = [
        ("数据分析与指标拆解能力", ["数据分析", "SQL", "Python", "Excel", "指标", "建模"]),
        ("项目推进与交付能力", ["项目", "交付", "推进", "跨部门", "协调", "复盘"]),
        ("产品/运营/业务理解能力", ["产品", "运营", "用户", "需求", "业务", "增长"]),
        ("研究与结构化表达能力", ["研究", "报告", "论文", "调研", "分析", "PPT"]),
        ("Python / Pandas 数据处理潜力", ["Python", "pandas", "numpy"]),
        ("英文或跨文化沟通能力", ["英文", "英语", "English", "presentation"]),
    ]
    for label, words in checks:
        if any(text_contains(profile_text, word) for word in words):
            strengths.append(label)
    return strengths


def match_resume_to_jd(
    jd_analysis: dict[str, Any],
    resume_text: str,
    profile_text: str | None = None,
    preferences: dict[str, Any] | None = None,
    fast: bool = False,
) -> dict[str, Any]:
    profile_text = effective_profile_text(profile_text)
    resume_only = normalize_text(resume_text)
    candidate_text = normalize_text(profile_text + "\n" + resume_only)
    jd_text = jd_analysis.get("raw_text", "")
    preferences = preferences or load_target_preferences()

    required_skills = jd_skill_list(jd_analysis)
    if not required_skills:
        required_skills = [skill for skill, aliases in SKILL_ALIASES.items() if count_alias_hits(jd_text, expanded_aliases(skill, aliases))]

    matched: list[str] = []
    missing: list[str] = []
    evidence: list[str] = []
    total_weight = 0
    matched_weight = 0
    for skill in required_skills:
        aliases = SKILL_ALIASES.get(skill, [skill])
        weight = 10
        if skill in ["产品能力", "业务分析", "机器学习/AI", "后端", "前端", "LCA", "ISO14067", "EPD", "CBAM", "SimaPro", "GaBi", "openLCA"]:
            weight = 14
        elif skill in ["英文能力", "Python", "SQL", "数据分析", "项目管理", "运营", "供应链管理", "供应链碳管理"]:
            weight = 9
        weight, _weight_reason = contextual_profile_group_weight(skill, aliases, weight, profile_text, resume_only, preferences)
        total_weight += weight
        hits = alias_hits(candidate_text, aliases)
        related_hits = related_alias_hits(candidate_text, skill, aliases)
        if hits:
            matched.append(skill)
            matched_weight += weight
            evidence.append(f"{skill}：{', '.join(hits[:3])}")
        elif related_hits:
            matched.append(skill)
            matched_weight += weight * 0.68
            evidence.append(f"{skill}相关：{', '.join(related_hits[:3])}")
        else:
            missing.append(skill)

    coverage_score = matched_weight / max(total_weight, 1) * 100 if total_weight else 45
    semantic_fn = semantic_similarity_fast if fast else semantic_similarity
    semantic_score = semantic_fn(jd_text, candidate_text) * 100
    score = 28 + coverage_score * 0.54 + semantic_score * 0.22
    if jd_analysis["category"] in ["数据/商业分析岗", "产品岗", "咨询/项目岗", "研发/工程岗", "财务/金融岗", "法务/合规岗", "供应链/采购岗", "LCA技术岗", "产品碳足迹岗", "ESG岗", "碳核算岗", "CBAM/出海合规岗"]:
        score += 7
    if jd_analysis["value"]["is_generic_esg"]:
        score -= 10
    if any(skill in missing for skill in ["SimaPro", "GaBi", "openLCA"]):
        score -= 4
    if any(skill in matched for skill in ["LCA", "ISO14067", "PEF", "Python", "数据分析"]):
        score += 4
    if jd_has_engineering_requirement(jd_text, jd_analysis, skills=required_skills) and resume_only and not resume_has_engineering_evidence(resume_only):
        score -= 14
    score = int(np.clip(round(score), 0, 100))

    gap_examples = []
    if any(skill in missing for skill in ["SQL", "Python", "数据分析", "业务分析"]):
        gap_examples.append("数据处理、指标拆解或业务分析证据不足")
    if "产品能力" in missing:
        gap_examples.append("需求分析、用户研究或产品方案证据不足")
    if "运营" in missing:
        gap_examples.append("运营动作、指标结果或复盘证据不足")
    if "项目管理" in missing:
        gap_examples.append("项目推进、跨部门协作或交付结果证据不足")
    if any(skill in missing for skill in ["前端", "后端", "机器学习/AI"]):
        gap_examples.append("工程实现、模型/系统落地或代码项目证据不足")
    if any(skill in missing for skill in ["财务分析", "法务合规", "人力资源", "供应链管理"]):
        gap_examples.append("专业职能方法、案例或行业语境证据不足")
    if any(skill in missing for skill in ["SimaPro", "GaBi"]):
        gap_examples.append("SimaPro / GaBi 工具实操证据不足")
    if "英文能力" in missing:
        gap_examples.append("英文口语或英文汇报证据不足")
    if "EPD" in missing:
        gap_examples.append("EPD / PCR 项目经验不足")
    if "CBAM" in missing:
        gap_examples.append("CBAM 行业案例分析不足")
    if "供应链碳管理" in missing:
        gap_examples.append("供应链碳数据管理案例不足")
    if not gap_examples:
        gap_examples.append("需把已有经历改写得更贴近岗位关键词")

    return {
        "score": score,
        "matched_skills": matched,
        "missing_skills": missing,
        "strengths": resume_strengths(resume_text, profile_text),
        "gap_examples": gap_examples,
        "semantic_score": int(round(semantic_score)),
        "coverage_score": int(round(coverage_score)),
        "evidence": evidence[:8],
    }


def sentence_split(text: str) -> list[str]:
    parts = re.split(r"[。！？!?；;\n\r]+", text)
    return [part.strip(" -:：\t") for part in parts if len(part.strip()) >= 4]


def analyze_interview(text: str, jd_analysis: dict[str, Any] | None = None) -> dict[str, Any]:
    clean = normalize_text(text)
    sentences = sentence_split(text)
    buckets: dict[str, list[str]] = {category: [] for category in INTERVIEW_CATEGORIES}
    for sentence in sentences:
        for category, keywords in INTERVIEW_CATEGORIES.items():
            if any(text_contains(sentence, keyword) for keyword in keywords):
                buckets[category].append(sentence[:160])

    jd_skills = []
    if jd_analysis and not jd_analysis.get("skills", pd.DataFrame()).empty:
        jd_skills = jd_analysis["skills"]["技能"].head(8).tolist()
    generated = generate_interview_questions(jd_skills)
    return {
        "raw_text": clean,
        "buckets": buckets,
        "generated_questions": generated,
        "answer_templates": generate_interview_answer_templates(jd_skills),
    }


def generate_interview_answer_templates(skills: list[str]) -> list[str]:
    focus = "、".join(skills[:4]) if skills else "岗位核心技能、项目经历、数据/业务分析"
    return [
        f"中文自我介绍：我目前的核心经历集中在{focus}。我适合这个岗位，是因为我能把业务问题拆成具体任务，明确数据、流程、协作对象和交付结果，并且能把结果复盘成下一步动作。",
        "项目回答结构：这个项目的目标是___；我负责___；我先确认问题和评价指标，再拆解任务、推进协作、产出___；最终结果是___；如果重做，我会加强___。",
        "英文项目开头：In this project, my role was to clarify the business problem, break it down into executable tasks, coordinate with stakeholders, and deliver measurable outputs such as reports, dashboards, product changes, or process improvements.",
        "反问招聘方：这个岗位前三个月最重要的产出是什么？会用哪些指标衡量表现？新人能接触哪些真实项目或业务问题？试用期结束前可以沉淀什么可复盘成果？",
    ]


def generate_interview_questions(skills: list[str]) -> dict[str, list[str]]:
    technical = []
    for skill in skills[:6]:
        if skill == "LCA":
            technical.append("用你做过的项目说明：评价目标是什么、功能单位怎么定、系统边界包括哪些环节、清单数据从哪里来。")
        elif skill == "CBAM":
            technical.append("如果一家制造企业要应对 CBAM，你会先要哪些活动数据、排放因子、产品编码和供应商资料？")
        elif skill == "ISO14067":
            technical.append("请用 1 分钟讲清 ISO14067 与 PEF 在边界、分配、数据质量和结果解释上的差异。")
        elif skill == "EPD":
            technical.append("EPD 项目里 PCR 为什么重要？如果没有合适 PCR，你会怎么判断报告边界和数据口径？")
        elif skill in ["SimaPro", "GaBi", "openLCA"]:
            technical.append(f"请描述一次 {skill} 建模流程：建过程、选数据库、设分配规则、看贡献分析、导出结果。")
        elif skill == "供应链碳管理":
            technical.append("供应商数据质量不一致时，你会怎么设计模板、校验规则、缺失值处理和追踪机制？")
        elif skill in ["产品能力", "用户研究"]:
            technical.append("请讲一个你从用户问题出发，完成需求判断、方案设计和结果验证的例子。")
        elif skill in ["业务分析", "数据分析", "SQL"]:
            technical.append("如果一个核心指标下降，你会如何拆解口径、定位原因、验证假设并给出行动建议？")
        elif skill in ["运营", "市场营销"]:
            technical.append("请讲一个你提升拉新、转化、留存或复购指标的方案，重点说动作、数据和复盘。")
        elif skill in ["项目管理", "沟通协作"]:
            technical.append("跨部门项目进度失控时，你会如何重排优先级、同步风险并推动交付？")
        elif skill in ["前端", "后端", "机器学习/AI"]:
            technical.append("请讲一个你做过的工程/模型项目：问题、技术选型、关键实现、测试方式和上线结果。")
        elif skill == "Python":
            technical.append("请举例说明你如何用 Python 清洗数据、检查异常值、生成分析结果，并把结论交付给业务或项目方。")

    return {
        "技术问题": technical or ["请结合一个项目说明你如何定义问题、拆解任务、推进执行，并用结果验证价值。"],
        "行为面问题": ["用 STAR 结构讲一个你把模糊任务拆成可执行步骤的项目。", "为什么选择这个岗位方向？请结合经历、能力证据和长期目标回答。"],
        "英文面试问题": ["Please introduce one project that best demonstrates your fit for this role.", "How would you explain your project impact to a non-technical stakeholder?"],
        "案例分析题": ["如果业务负责人给你一个模糊目标，请按问题定义、信息收集、方案设计、执行协作、结果复盘五步拆解。"],
    }


def gap_action_for_item(item: str, jd_text: str = "") -> dict[str, str]:
    item_text = normalize_text(item)
    if any(text_contains(item_text, keyword) for keyword in ["业务分析", "商业分析", "经营分析", "指标"]):
        return {
            "短板": "业务分析证据不够落地",
            "优先级": "P1",
            "今天就做": "选一个业务问题做拆解：目标指标、影响因素、数据口径、验证路径、结论和行动建议各写 1 句。",
            "交付物": "一张“指标拆解-原因假设-验证数据-行动建议”表。",
            "简历可写": "围绕核心业务指标拆解问题，完成数据口径确认、原因假设验证和行动建议输出。",
            "面试说法": "我做分析会先确认指标定义，再拆影响因素和数据口径，最后把结论落到能执行的业务动作上。",
        }
    if any(text_contains(item_text, keyword) for keyword in ["产品", "需求", "用户研究"]):
        return {
            "短板": "产品/用户问题证据不足",
            "优先级": "P1",
            "今天就做": "选一个熟悉产品做 1 页需求拆解：目标用户、使用场景、痛点、现有方案、优化方案、验证指标。",
            "交付物": "1 页产品需求拆解 + 1 个可讲的用户故事。",
            "简历可写": "围绕用户场景拆解需求，输出产品优化方案、验证指标和迭代建议。",
            "面试说法": "我会先确认用户、场景和指标，再判断需求优先级，最后用数据或反馈验证方案是否有效。",
        }
    if any(text_contains(item_text, keyword) for keyword in ["运营", "增长", "市场"]):
        return {
            "短板": "运营动作与指标结果不够具体",
            "优先级": "P1",
            "今天就做": "把一段运营经历按目标、用户分层、触达动作、转化路径、数据结果和复盘改写。",
            "交付物": "一张运营漏斗表 + 2 条可直接放进简历的 bullet。",
            "简历可写": "围绕拉新、转化、留存或复购目标设计运营动作，并根据数据复盘优化策略。",
            "面试说法": "我讲运营不会只讲做了活动，而会说明目标指标、触达对象、动作设计、数据结果和复盘。",
        }
    if any(text_contains(item_text, keyword) for keyword in ["销售", "商务", "BD", "客户开发"]):
        return {
            "短板": "客户开发/商务转化证据不足",
            "优先级": "P1",
            "今天就做": "整理一个客户推进案例：客户类型、需求判断、触达话术、异议处理、推进结果和复盘。",
            "交付物": "一段客户推进 STAR 案例 + 3 句可直接使用的沟通话术。",
            "简历可写": "参与客户需求判断与商机推进，完成客户信息梳理、沟通跟进和转化复盘。",
            "面试说法": "我会先判断客户真实需求和决策链，再设计触达节奏，推进中记录异议并及时复盘调整。",
        }
    if any(text_contains(item_text, keyword) for keyword in ["项目管理", "沟通协作", "交付"]):
        return {
            "短板": "项目推进和跨部门协作证据不足",
            "优先级": "P1",
            "今天就做": "整理一个项目推进案例，写清里程碑、协作对象、风险点、你推动的动作和最终交付。",
            "交付物": "一份 STAR 项目复盘 + 风险同步话术。",
            "简历可写": "推动跨部门项目按期交付，拆解任务、同步风险并沉淀流程复盘。",
            "面试说法": "遇到项目卡点时，我会先明确责任人和时间线，再同步风险、给出备选方案并推动闭环。",
        }
    if any(text_contains(item_text, keyword) for keyword in ["客户沟通", "跨部门", "协作", "咨询式表达"]):
        return {
            "短板": "沟通协作场景不够具体",
            "优先级": "P2",
            "今天就做": "整理一个沟通推进案例：对方是谁、分歧是什么、你如何同步信息、推动决策和闭环结果。",
            "交付物": "一段沟通协作 STAR 案例 + 3 句面试回答话术。",
            "简历可写": "在项目推进中协调多方信息，拆解分歧点并推动任务闭环，保障交付物按时完成。",
            "面试说法": "我会先对齐目标和信息差，再把争议拆成可决策的问题，最后明确责任人、时间点和下一步。",
        }
    if any(text_contains(item_text, keyword) for keyword in ["财务", "预算", "审计", "成本"]):
        return {
            "短板": "财务分析场景证据不足",
            "优先级": "P1",
            "今天就做": "把一个财务/成本问题拆成收入、成本、费用、现金流或预算差异，写出分析口径和结论。",
            "交付物": "一张财务指标拆解表 + 2 条简历 bullet。",
            "简历可写": "基于财务数据完成预算、成本或经营指标分析，识别异常项并输出改进建议。",
            "面试说法": "我会先确认财务口径，再看趋势、结构和异常项，最后把分析结论转成管理动作。",
        }
    if any(text_contains(item_text, keyword) for keyword in ["法务", "合规", "合同", "风控"]):
        return {
            "短板": "法务/合规案例表达不足",
            "优先级": "P1",
            "今天就做": "选一类合同或合规问题，整理风险点、适用规则、处理建议和业务沟通话术。",
            "交付物": "一页风险识别清单 + 一段业务沟通说明。",
            "简历可写": "协助梳理合同或合规风险，提炼关键条款、风险点和业务侧处理建议。",
            "面试说法": "我不会只背规则，会先判断业务场景和风险等级，再给出可执行、能沟通的处理建议。",
        }
    if any(text_contains(item_text, keyword) for keyword in ["人力", "招聘", "培训", "绩效"]):
        return {
            "短板": "HR 场景成果不够清晰",
            "优先级": "P2",
            "今天就做": "整理一个 HR 场景：招聘漏斗、培训反馈、绩效流程或员工沟通，补上对象、动作和结果。",
            "交付物": "一张 HR 流程/漏斗表 + 2 条可投递表达。",
            "简历可写": "参与招聘、培训或绩效流程支持，完成信息整理、过程跟进和结果复盘。",
            "面试说法": "HR 工作的重点是把人、流程和数据连起来，我会说明对象、规则、动作和结果。",
        }
    if any(text_contains(item_text, keyword) for keyword in ["设计", "UI", "UX", "交互", "Figma"]):
        return {
            "短板": "设计决策与作品集证据不足",
            "优先级": "P1",
            "今天就做": "选一个界面或流程改版案例，补齐问题、用户路径、方案对比、设计取舍和验证指标。",
            "交付物": "一页作品集案例说明 + 改版前后对比图。",
            "简历可写": "基于用户路径和业务目标完成界面/流程优化，输出设计方案并沉淀组件或规范。",
            "面试说法": "我会从用户目标和业务指标出发解释设计取舍，而不是只讲视觉效果。",
        }
    if any(text_contains(item_text, keyword) for keyword in ["前端", "后端", "Java", "C++", "机器学习", "AI"]):
        return {
            "短板": "工程/技术项目证据不足",
            "优先级": "P1",
            "今天就做": "选一个代码项目补齐 README：问题、技术栈、核心模块、测试方式、部署方式和结果截图。",
            "交付物": "项目 README + 代码仓库链接/截图 + 2 条技术 bullet。",
            "简历可写": "独立实现核心模块，完成接口/组件/模型开发、测试验证和结果复盘。",
            "面试说法": "我会从问题定义、技术选型、核心实现、测试验证和上线结果五步讲这个项目。",
        }
    if any(text_contains(item_text, keyword) for keyword in ["SimaPro", "GaBi", "openLCA"]):
        return {
            "短板": "LCA 软件建模证据不足",
            "优先级": "P1",
            "今天就做": "用 openLCA 或任一公开案例跑通一个简单产品系统，记录功能单位、边界、数据库和贡献分析截图。",
            "交付物": "1 页项目复盘 + 3 张截图：模型结构、贡献分析、结果表。",
            "简历可写": "使用 openLCA/SimaPro 建立产品生命周期模型，完成系统边界、数据库选择、贡献分析和结果解释。",
            "面试说法": "我不只看软件按钮，而是按功能单位、边界、清单、数据库、分配规则、结果解释这条线复盘建模过程。",
        }
    if any(text_contains(item_text, keyword) for keyword in ["ISO14067", "PEF", "EPD", "PCR"]):
        return {
            "短板": "标准/报告方法学表达不够硬",
            "优先级": "P1",
            "今天就做": "整理 ISO14067、PEF、EPD/PCR 的差异表，各写 3 条边界、数据质量和结果解释差异。",
            "交付物": "一张方法学对比表 + 一段 90 秒口述稿。",
            "简历可写": "梳理 ISO14067、PEF 与 EPD/PCR 方法要求，归纳边界设定、数据质量和报告解释差异。",
            "面试说法": "我会先确认报告目的和适用标准，再反推边界、数据口径、PCR/方法学依据和结果呈现方式。",
        }
    if text_contains(item_text, "CBAM"):
        return {
            "短板": "CBAM 案例和数据清单证据不足",
            "优先级": "P1",
            "今天就做": "选钢铁、铝或电池材料做一个 CBAM 数据收集模板，列出产品、产量、能源、工艺排放、供应商证明和默认值风险。",
            "交付物": "一份 CBAM 数据收集表 + 5 条申报风险提示。",
            "简历可写": "跟踪 CBAM 出海合规要求，拆解企业活动数据、排放因子、默认值使用和申报风险。",
            "面试说法": "我会从产品范围、活动数据、直接/间接排放、默认值、供应商证明和申报节奏六块拆解项目。",
        }
    if any(text_contains(item_text, keyword) for keyword in ["英文", "英语", "English", "汇报"]):
        return {
            "短板": "英文汇报证据不足",
            "优先级": "P1",
            "今天就做": "写一版 90 秒英文项目介绍，包含 background、task、method、result、learning，并录音复盘。",
            "交付物": "英文自我介绍 + 英文项目介绍各一版。",
            "简历可写": "阅读英文标准/政策材料，提炼项目要求并完成英文项目说明。",
            "面试说法": "I can explain a project by covering the business problem, my role, key actions, measurable results, and what I learned.",
        }
    if any(text_contains(item_text, keyword) for keyword in ["Python", "SQL", "数据"]):
        return {
            "短板": "数据处理成果不够可展示",
            "优先级": "P2",
            "今天就做": "用一份公开数据或历史项目数据做清洗、口径统一、指标计算、异常值检查和结论输出。",
            "交付物": "一个 Excel/Python/SQL 分析模板 + 结果截图 + 3 条结论。",
            "简历可写": "使用 Python/SQL/Excel 完成数据清洗、指标计算、异常检查和业务结论输出。",
            "面试说法": "我的数据处理习惯是先统一口径和指标定义，再清洗异常值，最后把结论转成业务动作。",
        }
    if text_contains(item_text, "供应链"):
        if not text_contains(item_text, "碳"):
            return {
                "短板": "供应链/采购场景证据不足",
                "优先级": "P2",
                "今天就做": "设计一份供应商评估表，字段包括价格、交期、质量、风险、历史表现和备选方案。",
                "交付物": "供应商评估模板 + 一段采购/供应链面试案例。",
                "简历可写": "参与供应商信息整理与评估，围绕交期、质量、成本和风险建立对比模板。",
                "面试说法": "供应链工作不只是催进度，我会把供应商、交期、质量、成本和风险放在同一张表里管理。",
            }
        return {
            "短板": "供应链碳数据管理案例不足",
            "优先级": "P2",
            "今天就做": "设计供应商碳数据收集模板，字段包括活动数据、单位、时间范围、数据来源、证明材料和置信度。",
            "交付物": "一份供应商数据模板 + 校验规则。",
            "简历可写": "设计供应链碳数据收集与校验模板，支持供应商活动数据追踪、质量分级和缺失项补充。",
            "面试说法": "供应链数据的关键不是只催数，而是把字段、单位、证明材料和质量分级规则提前设计清楚。",
        }
    return {
        "短板": item_text or "岗位证据表达不足",
        "优先级": "P2",
        "今天就做": "把当前简历里最相关的一段经历改写成“背景-任务-动作-结果”，并补上 JD 关键词。",
        "交付物": "一版可直接替换到简历里的项目 bullet。",
        "简历可写": "围绕岗位关键词改写项目经历，突出数据、方法、交付物和结果。",
        "面试说法": "我会先解释项目目标，再讲我负责的数据、方法和最终产出，而不是只描述参与过。",
    }


def concrete_gap_row(row: dict[str, str], category: str, jd_text: str = "") -> dict[str, str]:
    shortcoming = normalize_text(row.get("短板", ""))
    category = normalize_text(category)
    concrete = dict(row)
    concrete["不要写"] = "不要写“负责/参与/协助/学习能力强”。没有真实材料、数字或交付物的内容不要硬塞。"

    if any(text_contains(shortcoming, key) for key in ["数据", "业务分析", "指标", "Python", "SQL"]):
        concrete.update(
            {
                "今天就做": "找一段真实经历，补齐 5 个字段：数据来源、样本量/时间范围、处理工具、发现的问题、最后交付物。",
                "交付物": "一张表：字段=指标口径、数据来源、处理步骤、异常发现、结论、建议动作。",
                "简历可写": "使用【Excel/Python/SQL】处理【数据来源/样本量】数据，统一【指标口径】，定位【具体问题】，输出【表格/看板/报告】并提出【1-2条建议】。",
                "面试说法": "这个分析我先确认【指标定义】，再用【工具】处理【数据】，发现【问题】，最后给出【建议动作】。",
            }
        )
    elif any(text_contains(shortcoming, key) for key in ["LCA", "碳", "ISO14067", "EPD", "PCR", "SimaPro", "GaBi", "openLCA"]):
        concrete.update(
            {
                "今天就做": "补一个可展示的 LCA/碳核算案例：产品、功能单位、系统边界、数据来源、排放因子/数据库、结果热点。",
                "交付物": "1 页案例说明 + 3 张证据图：边界图、清单数据表、热点/贡献分析结果。",
                "简历可写": "围绕【产品/材料】建立【LCA/碳核算】案例，设定功能单位【填写】，梳理【原料/能源/运输】清单数据，使用【工具/数据库】完成热点分析并形成【报告/截图】。",
                "面试说法": "这个案例我会按目标、功能单位、边界、清单数据、数据库/因子、热点结果六步讲清楚。",
            }
        )
    elif text_contains(shortcoming, "CBAM"):
        concrete.update(
            {
                "今天就做": "做一份 CBAM 数据清单，不要只看政策摘要；必须列产品编码、产量、直接排放、间接排放、供应商证明、默认值风险。",
                "交付物": "CBAM 数据收集表 + 5 条申报风险提示。",
                "简历可写": "拆解【钢铁/铝/电池等产品】CBAM 申报数据，整理产量、能源消耗、直接/间接排放、供应商证明和默认值使用风险，形成【数据清单/风险表】。",
                "面试说法": "我会从产品范围、活动数据、排放边界、供应商证明、默认值和申报节奏拆这个问题。",
            }
        )
    elif any(text_contains(shortcoming, key) for key in ["产品", "用户", "需求"]):
        concrete.update(
            {
                "今天就做": "选一个真实产品或功能，写清目标用户、场景、痛点、现方案、你的优化方案、验证指标。",
                "交付物": "1 页需求拆解：用户-场景-痛点-方案-指标-优先级。",
                "简历可写": "针对【用户群体】在【使用场景】中的【痛点】，梳理需求优先级，输出【功能/流程】优化方案，并用【指标】验证方案效果。",
                "面试说法": "我会先讲用户和场景，再讲为什么这个需求优先，最后讲用什么指标判断方案有没有用。",
            }
        )
    elif any(text_contains(shortcoming, key) for key in ["运营", "增长", "市场"]):
        concrete.update(
            {
                "今天就做": "把一个活动或运营任务补成漏斗：目标、触达人群、渠道、转化节点、数据结果、复盘动作。",
                "交付物": "运营漏斗表 + 2 条带数字的简历 bullet。",
                "简历可写": "面向【人群】设计【运营动作/活动】，通过【渠道】触达【人数】，跟踪【转化/留存/点击】指标，复盘后调整【动作】。",
                "面试说法": "这段经历我会按目标、用户、动作、数据、复盘讲，不只说办了活动。",
            }
        )
    elif any(text_contains(shortcoming, key) for key in ["项目", "协作", "跨部门", "沟通"]):
        concrete.update(
            {
                "今天就做": "整理一个推进卡点：目标、参与方、卡点、你推动的动作、时间线、最终交付。",
                "交付物": "项目推进 STAR 案例 + 一张责任人/时间线表。",
                "简历可写": "在【项目名称】中对接【协作方】，拆解【任务/风险】，推动【关键动作】，在【时间】前交付【文档/报告/系统/方案】。",
                "面试说法": "我会说明卡点是什么、我找了谁、做了什么取舍、最后怎么闭环。",
            }
        )
    elif any(text_contains(shortcoming, key) for key in ["英文", "英语", "English"]):
        concrete.update(
            {
                "今天就做": "写 90 秒英文项目介绍并录音，内容必须包含 project background、my role、method、result、lesson。",
                "交付物": "英文自我介绍 + 英文项目介绍录音稿。",
                "简历可写": "阅读并整理【英文标准/论文/政策/客户材料】，提炼【关键要求】，输出【英文摘要/汇报材料/项目说明】。",
                "面试说法": "In this project, I was responsible for 【role】. I used 【method/tool】 to solve 【problem】 and delivered 【result】.",
            }
        )
    elif any(text_contains(shortcoming, key) for key in ["工程", "技术", "前端", "后端", "机器学习", "AI"]):
        concrete.update(
            {
                "今天就做": "补项目 README：问题、技术栈、核心模块、你写的代码、测试方式、运行截图。",
                "交付物": "README + 代码/截图 + 2 条技术简历 bullet。",
                "简历可写": "基于【技术栈】实现【模块/功能】，完成【接口/组件/模型】开发与【测试方式】验证，解决【具体问题】。",
                "面试说法": "我会从需求、技术选型、核心实现、测试验证、结果五步讲这个项目。",
            }
        )
    elif category:
        concrete.update(
            {
                "今天就做": f"从当前简历里选一段最接近“{category}”的经历，补齐：任务背景、你负责的动作、交付物、可验证结果。",
                "交付物": "一段可替换简历 bullet + 一段 60 秒面试说明。",
                "简历可写": f"围绕【{category}相关项目】完成【具体任务】，使用【工具/方法】输出【交付物】，结果体现为【数字/报告/方案/截图】。",
                "面试说法": "我会用真实项目讲，不讲抽象能力：背景是什么、我做了什么、交付了什么、结果如何验证。",
            }
        )
    return concrete


def build_gap_analysis(
    jd_analysis: dict[str, Any] | None,
    resume_match: dict[str, Any] | None,
    interview_analysis: dict[str, Any] | None,
) -> dict[str, Any]:
    jd_text = jd_analysis.get("raw_text", "") if jd_analysis else ""
    missing = resume_match.get("missing_skills", []) if resume_match else []
    score = resume_match.get("score", 0) if resume_match else 0
    category = jd_analysis.get("category", "目标岗位") if jd_analysis else "目标岗位"

    skill_order = [
        "业务分析",
        "数据分析",
        "SQL",
        "Python",
        "产品能力",
        "用户研究",
        "运营",
        "项目管理",
        "沟通协作",
        "市场营销",
        "销售/商务",
        "财务分析",
        "法务合规",
        "人力资源",
        "供应链管理",
        "设计",
        "前端",
        "后端",
        "机器学习/AI",
        "英文能力",
        "SimaPro",
        "GaBi",
        "openLCA",
        "ISO14067",
        "EPD",
        "CBAM",
        "GHG Protocol",
    ]
    hard = [skill for skill in skill_order if skill in missing]
    for skill in ["SimaPro", "GaBi", "openLCA", "ISO14067", "EPD", "CBAM"]:
        if skill not in hard and text_contains(jd_text, skill):
            hard.append(skill)

    soft = []
    if "英文能力" in missing or any(text_contains(jd_text, word) for word in ["英文", "英语", "English"]):
        soft.append("英文口语 / 英文汇报")
    if any(text_contains(jd_text, word) for word in ["客户", "咨询", "汇报", "presentation", "跨部门", "协作"]):
        soft.append("客户沟通 / 跨部门协作")
    if score < 70:
        soft.append("面试表达结构化程度")

    certs = [cert for cert in CERT_KEYWORDS if text_contains(jd_text, cert)]

    category_project_map = {
        "数据/商业分析岗": "业务分析项目案例",
        "产品岗": "产品需求拆解案例",
        "运营岗": "运营增长复盘案例",
        "市场/销售岗": "客户开发或营销转化案例",
        "咨询/项目岗": "咨询交付或项目推进案例",
        "研发/工程岗": "工程项目或代码作品",
        "设计岗": "设计作品集案例",
        "财务/金融岗": "财务/经营分析案例",
        "法务/合规岗": "法务/合规风险案例",
        "人力/行政岗": "招聘、培训或绩效流程案例",
        "供应链/采购岗": "供应链/采购评估案例",
        "LCA技术岗": "LCA建模或碳核算项目案例",
        "产品碳足迹岗": "产品碳足迹项目案例",
        "碳核算岗": "碳核算项目案例",
        "CBAM/出海合规岗": "CBAM/合规数据清单案例",
    }
    projects = []
    project_seed = category_project_map.get(category)
    if project_seed:
        projects.append(project_seed)
    for keyword in PROJECT_KEYWORDS:
        if text_contains(jd_text, keyword):
            candidate = f"{keyword}相关项目案例"
            if candidate not in projects:
                projects.append(candidate)
    projects = projects[:3]

    if not hard and not projects:
        projects.append(f"{category}可量化项目案例")

    raw_gap_items = list(dict.fromkeys(hard[:6] + projects[:3] + soft[:2] + certs[:2]))
    action_rows = []
    seen_shortcomings: set[str] = set()
    for raw_item in raw_gap_items:
        action = gap_action_for_item(raw_item, jd_text)
        if action["短板"] not in seen_shortcomings:
            action_rows.append(action)
            seen_shortcomings.add(action["短板"])
    action_rows = [concrete_gap_row(row, category, jd_text) for row in action_rows]
    action_rows = sorted(action_rows, key=lambda row: row["优先级"])[:8]
    immediate_actions = [row["今天就做"] for row in action_rows[:4]]
    resume_patches = [row["简历可写"] for row in action_rows[:5]]
    interview_scripts = [row["面试说法"] for row in action_rows[:5]]
    portfolio_tasks = [row["交付物"] for row in action_rows[:5]]
    priorities = [(row["优先级"], row["今天就做"]) for row in action_rows[:5]]
    if not any(priority == "P3" for priority, _ in priorities):
        priorities.append(("P3", "准备 3 个真实项目故事：每个都写清背景、你的动作、交付物、结果。没有数字就写证据形态，比如报告、截图、表格。"))
    weekly_plan = [
        {"时间": "今天", "任务": immediate_actions[0] if immediate_actions else "选择一段真实经历，补齐背景、动作、交付物、结果四项。", "完成标准": "产出 1 段带【工具/数据/交付物/结果】的简历 bullet。"},
        {"时间": "明天", "任务": immediate_actions[1] if len(immediate_actions) > 1 else "整理目标 JD 的 8 个关键词，并逐一对应当前简历里的证据。", "完成标准": "形成关键词-证据对应表，空缺项标出补证据方式。"},
        {"时间": "本周内", "任务": "完成一个可展示的小作品：" + (portfolio_tasks[0] if portfolio_tasks else "1 页项目复盘或数据表。"), "完成标准": "至少有截图、表格、报告目录或代码链接之一。"},
        {"时间": "投递前", "任务": "检查每条简历改写句：是否有真实项目名、工具/方法、交付物、结果。缺一项就不要投递版使用。", "完成标准": "删掉“负责、参与、协助、学习能力强”等空泛表述。"},
    ]

    if score >= 80:
        summary = f"当前对 {category} 已有投递基础，重点不是再加形容词，而是把项目证据写实：工具、数据、交付物、结果至少出现 3 项。"
    elif score >= 65:
        summary = f"当前可以投递 {category}，但投前先补 2-3 个硬证据；缺证据的关键词不要硬写，改成可验证的项目句。"
    else:
        summary = f"当前与 {category} 的证据距离偏大，先补一个可展示作品或项目复盘，再投要求高的岗位；否则容易像套模板。"

    interview_focus = []
    if interview_analysis:
        for category, questions in interview_analysis.get("generated_questions", {}).items():
            interview_focus.extend([f"{category}: {question}" for question in questions[:2]])
    if not interview_focus:
        interview_focus = interview_scripts[:4]

    return {
        "summary": summary,
        "current_gaps": {
            "技能证据不足": hard or ["未发现明显硬技能缺口，重点优化表达和证据呈现"],
            "表达/协作不足": soft or ["暂无明显软技能缺口"],
            "证书/资格要求": certs,
            "项目案例不足": projects,
        },
        "priorities": priorities,
        "action_rows": action_rows,
        "immediate_actions": immediate_actions,
        "resume_patches": resume_patches,
        "portfolio_tasks": portfolio_tasks,
        "interview_scripts": interview_scripts,
        "weekly_plan": weekly_plan,
        "interview_focus": interview_focus,
    }


def jd_skill_list(jd_analysis: dict[str, Any] | None) -> list[str]:
    if not jd_analysis:
        return []
    skills_df = jd_analysis.get("skills", pd.DataFrame())
    if skills_df is None or skills_df.empty:
        return []
    for column in ["技能"]:
        if column in skills_df.columns:
            return [str(item) for item in skills_df[column].dropna().tolist()]
    return [str(item) for item in skills_df.iloc[:, 0].dropna().tolist()]


def clean_extracted_value(value: str) -> str:
    value = normalize_text(value)
    boundaries = [
        " 岗位名称",
        " 职位名称",
        " 招聘职位",
        " 岗位",
        " 职位",
        " 公司名称",
        " 企业名称",
        " 地点",
        " 城市",
        " 薪资",
        " 月薪",
        " 学历",
        " 经验",
        " 链接",
        " URL",
        " JD",
        " 要求",
        " 职责",
        " 任职",
        " 应届",
        " 校招",
        " 全职",
        " 实习",
        " 工作内容",
    ]
    for boundary in boundaries:
        position = value.find(boundary)
        if position > 0:
            value = value[:position]
    return value.strip(" ：:，,;；|")[:80]


def jd_basic_value(jd_analysis: dict[str, Any] | None, *keys: str) -> str:
    if not jd_analysis:
        return ""
    basic = jd_analysis.get("basic", {})
    for key in keys:
        if key in basic and basic[key] != "未识别":
            return clean_extracted_value(str(basic[key]))
    return ""


def pick_resume_evidence_lines(resume_text: str, keywords: list[str], limit: int = 6) -> list[str]:
    lines = normalize_resume_lines(resume_text)
    hits = []
    for line in lines:
        if any(text_contains(line, keyword) for keyword in keywords):
            clean_line = re.sub(r"^\s*[-*•·]\s*", "", line).strip()
            if 12 <= len(clean_line) <= 160:
                hits.append(clean_line)
        if len(hits) >= limit:
            break
    return list(dict.fromkeys(hits))


def format_skill_group(title: str, skills: list[str]) -> str:
    clean = [skill for skill in skills if skill]
    return f"{title}：" + (" / ".join(clean) if clean else "按目标岗位补充")


def build_custom_resume(master_resume: str, jd_analysis: dict[str, Any] | None, profile_text: str | None = None) -> dict[str, Any]:
    profile_text = profile_text or ""
    master_resume = normalize_text(master_resume)
    jd_text = jd_analysis.get("raw_text", "") if jd_analysis else ""
    category = jd_analysis.get("category", "目标岗位") if jd_analysis else "目标岗位"
    skills = jd_skill_list(jd_analysis)
    matched_skills = []
    missing_skills = []

    combined_candidate_text = normalize_text(profile_text + "\n" + master_resume)
    for skill in skills:
        aliases = SKILL_ALIASES.get(skill, [skill])
        if count_alias_hits(combined_candidate_text, aliases):
            matched_skills.append(skill)
        else:
            missing_skills.append(skill)

    company = jd_basic_value(jd_analysis, "公司名")
    job_title = jd_basic_value(jd_analysis, "岗位名") or category
    keyword_line = " / ".join((matched_skills + missing_skills)[:12]) or "数据分析 / 项目管理 / 产品运营 / 沟通协作"

    role_terms = [
        skill
        for skill in [
            "产品能力",
            "运营",
            "用户研究",
            "业务分析",
            "项目管理",
            "市场营销",
            "销售/商务",
            "财务分析",
            "法务合规",
            "人力资源",
            "供应链管理",
            "设计",
            "前端",
            "后端",
            "机器学习/AI",
            "LCA",
            "产品碳足迹",
            "碳核算",
        ]
        if skill in matched_skills or skill in skills or text_contains(combined_candidate_text, skill)
    ]
    tool_terms = [
        skill
        for skill in ["Python", "SQL", "Excel", "Power BI", "Tableau", "pandas", "Stata", "MATLAB", "数据分析"]
        if text_contains(combined_candidate_text, skill) or skill in matched_skills or skill in skills
    ]
    delivery_terms = [
        skill
        for skill in ["英文能力", "沟通协作", "项目交付", "报告写作", "客户沟通", "行业研究", "复盘"]
        if skill in matched_skills or text_contains(combined_candidate_text, skill) or text_contains(jd_text, skill)
    ]
    if not role_terms:
        role_terms = skills[:5] or [category]
    if not tool_terms:
        tool_terms = ["信息整理", "数据校验", "结果复盘"]
    if not delivery_terms:
        delivery_terms = ["结构化表达", "跨部门协作", "项目复盘"]

    evidence_keywords = list(dict.fromkeys(["数据", "项目", "产品", "运营", "用户", "业务", "客户", "分析", "Python", "SQL", "Excel", "英文"] + skills))
    evidence_lines = pick_resume_evidence_lines(master_resume + "\n" + profile_text, evidence_keywords, limit=8)

    target_direction = job_title if job_title and job_title != "未识别" else category
    summary_lines = [
        f"目标方向为{target_direction}，已积累 {' / '.join(role_terms[:5])} 相关经历或能力证据。",
        f"具备 {format_skill_group('', tool_terms).lstrip('：')} 能力，可支持信息整理、数据处理、指标判断和结果输出。",
    ]
    summary_lines.append(f"能够围绕岗位目标推进任务拆解、过程协作、交付物沉淀和复盘改进，重点匹配 {' / '.join((skills or role_terms)[:5])}。")

    skills_section = [
        format_skill_group("岗位核心能力", role_terms[:6]),
        format_skill_group("数据与工具", tool_terms[:6]),
        format_skill_group("沟通与交付", delivery_terms[:6]),
    ]

    bullets = []
    if any(skill in skills for skill in ["产品能力", "用户研究"]) or category == "产品岗":
        bullets.append("围绕用户场景和业务目标拆解需求，梳理痛点、优先级、方案路径和验证指标，输出可执行的产品优化建议。")
    if any(skill in skills for skill in ["运营", "市场营销"]) or category in ["运营岗", "市场/销售岗"]:
        bullets.append("基于目标用户和转化路径设计运营动作，跟踪拉新、转化、留存等指标，并根据数据复盘优化策略。")
    if any(skill in skills for skill in ["业务分析", "数据分析", "SQL", "Python"]) or category == "数据/商业分析岗":
        bullets.append("围绕核心业务指标完成数据口径确认、清洗整理、异常检查和原因分析，输出可落地的业务建议。")
    if any(skill in skills for skill in ["项目管理", "沟通协作"]) or category in ["咨询/项目岗"]:
        bullets.append("拆解项目目标、里程碑和协作分工，跟进风险与进度，推动交付物按时输出并沉淀复盘。")
    if category == "研发/工程岗" or any(skill in skills for skill in ["前端", "后端", "机器学习/AI"]):
        bullets.append("参与功能模块、接口、组件或模型开发，完成关键逻辑实现、测试验证和问题复盘。")
    if tool_terms:
        bullets.append(f"使用 {' / '.join(tool_terms[:4])} 处理项目或业务数据，完成清洗、匹配、指标构建和可视化输出。")
    if "ISO14067" in skills or "PEF" in skills or text_contains(combined_candidate_text, "ISO14067") or text_contains(combined_candidate_text, "PEF"):
        bullets.append("围绕 ISO14067 与 PEF 方法差异整理研究材料，对边界设定、分配规则、数据质量和结果解释进行对比。")
    if "CBAM" in skills or text_contains(jd_text, "CBAM"):
        bullets.append("跟踪欧盟 CBAM 要求，拆解企业活动数据、排放因子、默认值使用和申报风险，沉淀资料清单。")
    if "EPD" in skills or text_contains(jd_text, "EPD"):
        bullets.append("整理 EPD/PCR 与产品碳足迹报告相关材料，归纳方法学依据、数据质量说明和报告核心章节。")
    if "英文能力" in skills or any(text_contains(jd_text, word) for word in ["英文", "英语", "English"]):
        bullets.append("阅读英文资料、客户材料或项目文档，提炼关键要求，并准备中英文项目说明。")
    bullets = list(dict.fromkeys(bullets))[:7] or generate_resume_bullets(jd_analysis, {"matched_skills": matched_skills})

    project_rewrite = [
        f"项目名称：{target_direction}相关项目 / 业务分析与交付复盘",
        "项目职责：明确项目目标、核心问题、评价指标和协作对象，拆解个人负责模块和交付标准。",
        "关键动作：完成信息收集、数据/材料整理、方案判断、过程推进和结果复盘，确保产出可追踪、可解释。",
        "项目产出：形成分析表、方案文档、结果图表、流程模板或复盘结论，为业务判断或后续执行提供依据。",
    ]

    application_pitch = (
        f"我对{target_direction}方向比较匹配，核心优势是{' / '.join(role_terms[:4])}和{' / '.join(tool_terms[:3])}。"
        f"如果岗位涉及{' / '.join(skills[:4]) or '业务拆解、数据处理和项目交付'}，可以较快进入资料整理、执行推进、结果输出和复盘改进工作。"
    )

    ready_resume_text = "\n".join(
        [
            "求职摘要",
            *[f"- {line}" for line in summary_lines],
            "",
            "核心能力",
            *[f"- {line}" for line in skills_section],
            "",
            "经历改写",
            *[f"- {line}" for line in bullets],
            "",
            "项目经历可替换版本",
            *[f"- {line}" for line in project_rewrite],
        ]
    )

    if missing_skills:
        risk_notes = [
            f"JD 明确提到但当前简历证据不足：{' / '.join(missing_skills[:8])}。",
            "没有实际做过的工具、项目或标准不要写成熟练掌握；可以先补课程项目、独立案例或作品截图。"
        ]
    else:
        risk_notes = ["关键词覆盖较好，重点是把经历写成项目交付语言，而不是课程/实习流水账。"]

    return {
        "company": company,
        "job_title": job_title,
        "category": category,
        "keyword_line": keyword_line,
        "matched_skills": matched_skills,
        "missing_skills": missing_skills,
        "summary": "\n".join(summary_lines),
        "summary_lines": summary_lines,
        "skills_section": skills_section,
        "bullets": bullets,
        "experience_bullets": bullets,
        "project_rewrite": project_rewrite,
        "evidence_lines": evidence_lines,
        "application_pitch": application_pitch,
        "ready_resume_text": ready_resume_text,
        "risk_notes": risk_notes,
    }


INTERNSHIP_DIMENSIONS = {
    "岗位相关度": {
        "weight": 28,
        "keywords": [
            "数据分析",
            "产品",
            "运营",
            "用户研究",
            "项目管理",
            "咨询",
            "研发",
            "前端",
            "后端",
            "算法",
            "市场",
            "销售",
            "财务",
            "法务",
            "人力",
            "供应链",
            "LCA",
            "生命周期",
            "产品碳足迹",
            "CFP",
            "ISO14067",
            "ISO 14067",
            "PEF",
            "EPD",
            "PCR",
            "CBAM",
            "碳核算",
            "碳盘查",
            "碳核查",
            "GHG",
            "排放因子",
            "SimaPro",
            "GaBi",
            "openLCA",
        ],
    },
    "项目含金量": {
        "weight": 24,
        "keywords": [
            "项目交付",
            "客户",
            "建模",
            "指标",
            "需求",
            "用户",
            "转化",
            "留存",
            "代码",
            "数据收集",
            "供应商",
            "报告撰写",
            "方案",
            "数据库",
            "方法学",
            "现场",
            "调研",
        ],
    },
    "目标行业匹配": {
        "weight": 18,
        "keywords": [
            "互联网",
            "制造",
            "金融",
            "咨询",
            "新能源",
            "医疗",
            "消费",
            "专业服务",
            "汽车",
            "电池",
            "钢铁",
            "化工",
            "纺织",
            "咨询",
            "出海",
            "合规",
            "供应链",
            "第三方",
        ],
    },
    "可迁移技能": {
        "weight": 16,
        "keywords": [
            "Python",
            "SQL",
            "Excel",
            "Power BI",
            "Tableau",
            "数据分析",
            "英文",
            "英语",
            "汇报",
            "PPT",
            "访谈",
            "研究",
            "沟通",
            "协作",
            "复盘",
        ],
    },
    "简历信号": {
        "weight": 14,
        "keywords": [
            "独立负责",
            "参与",
            "主导",
            "输出",
            "成果",
            "案例",
            "报告",
            "模型",
            "模板",
            "工具",
            "客户认可",
            "量化",
        ],
    },
}

INTERNSHIP_RISK_KEYWORDS = [
    "ESG报告",
    "ESG 报告",
    "CSR",
    "公众号",
    "推文",
    "活动策划",
    "行政",
    "会议组织",
    "材料整理",
    "翻译",
    "打杂",
    "品牌传播",
    "披露",
    "评级问卷",
]


def analyze_internship(
    company: str,
    role: str,
    industry: str,
    duration: str,
    internship_text: str,
    profile_text: str | None = None,
) -> dict[str, Any]:
    profile_text = profile_text or ""
    text = normalize_text("\n".join([company, role, industry, duration, internship_text, profile_text]))
    internship_only_text = normalize_text("\n".join([company, role, industry, duration, internship_text]))

    dimension_rows = []
    total_score = 0
    for dimension, config in INTERNSHIP_DIMENSIONS.items():
        keywords = config["keywords"]
        hits = [keyword for keyword in keywords if text_contains(internship_only_text, keyword)]
        dimension_score = int(np.clip(round(35 + len(hits) * 13), 0, 100)) if hits else 35
        weighted_score = dimension_score * config["weight"] / 100
        total_score += weighted_score
        dimension_rows.append(
            {
                "维度": dimension,
                "得分": dimension_score,
                "权重": config["weight"],
                "命中关键词": " / ".join(hits[:10]) if hits else "暂无明显证据",
            }
        )

    high_hits = [keyword for keyword in HIGH_VALUE_KEYWORDS if text_contains(internship_only_text, keyword)]
    low_hits = [keyword for keyword in INTERNSHIP_RISK_KEYWORDS if text_contains(internship_only_text, keyword)]
    target_hits = []
    for keyword in ["技术", "产品", "运营", "数据", "研发", "咨询", "金融", "制造", "供应链", "合规", "项目", "客户", "LCA", "CBAM", "EPD"]:
        if text_contains(text, keyword):
            target_hits.append(keyword)

    bonus = min(len(high_hits) * 2, 10) + min(len(target_hits), 6)
    penalty = min(len(low_hits) * 4, 18)
    final_score = int(np.clip(round(total_score + bonus - penalty), 0, 100))

    if final_score >= 82:
        verdict = "非常有利，建议优先争取"
        decision = "如果薪资、导师和时间安排可接受，可以作为后续全职投递的重要跳板。"
    elif final_score >= 68:
        verdict = "比较有利，值得考虑"
        decision = "适合去，但入职前要确认能接触真实项目、数据和可写进简历的产出。"
    elif final_score >= 52:
        verdict = "中等价值，需要谈清楚工作内容"
        decision = "不要只看公司名，重点确认是否能接触真实项目、核心任务和可量化产出。"
    else:
        verdict = "帮助有限，谨慎选择"
        decision = "如果主要是行政、宣传、披露材料整理，除非没有更好选择，否则不建议作为核心实习。"

    reasons = []
    if high_hits:
        reasons.append("命中了高价值岗位关键词：" + " / ".join(high_hits[:8]))
    if any(keyword in internship_only_text for keyword in ["SimaPro", "GaBi", "openLCA"]):
        reasons.append("有机会补齐 LCA 软件建模证据，这是后续投递技术岗的强信号。")
    if any(text_contains(internship_only_text, keyword) for keyword in ["客户", "项目交付", "报告", "现场", "供应商"]):
        reasons.append("包含真实项目交付、客户沟通或供应链场景，简历可讲性较强。")
    if any(text_contains(internship_only_text, keyword) for keyword in ["产品", "运营", "数据", "研发", "财务", "法务", "人力", "供应链"]):
        reasons.append("岗位职能较明确，便于沉淀可迁移的项目证据。")
    if any(text_contains(internship_only_text, keyword) for keyword in ["CBAM", "出海", "欧盟", "合规"]):
        reasons.append("出海合规属性明确，符合 CBAM/海外合规方向的差异化定位。")
    if not reasons:
        reasons.append("目前描述里的岗位证据不够明确，需要补充具体任务后再判断。")

    risks = []
    if low_hits:
        risks.append("存在低价值/杂务风险词：" + " / ".join(low_hits[:8]))
    if not any(text_contains(internship_only_text, keyword) for keyword in ["数据", "指标", "需求", "开发", "运营", "客户", "报告", "项目"]):
        risks.append("没有看到清晰的数据、需求、开发、运营或项目交付任务，可能难以转化为简历卖点。")
    if text_contains(internship_only_text, "ESG") and not high_hits:
        risks.append("如果只做披露、问卷或材料整理，对后续求职帮助有限。")
    if not risks:
        risks.append("主要风险是实习中拿不到可量化成果，需要主动争取项目产出。")

    recommended_outputs = [
        "至少沉淀 1 个可讲的 STAR 项目：背景、任务、你的动作、量化结果。",
        "保留可脱敏展示的表格、看板、代码截图、需求文档、流程模板或报告目录。",
        "把实习内容改写成岗位关键词和业务结果，而不是泛泛写“协助”“支持”。",
    ]
    if any(text_contains(internship_only_text, keyword) for keyword in ["SimaPro", "GaBi", "openLCA", "LCA"]):
        recommended_outputs.append("争取完整跑通一次建模流程：功能单位、系统边界、数据库选择、分配规则、结果解释。")
    if any(text_contains(internship_only_text, keyword) for keyword in ["CBAM", "出海", "欧盟"]):
        recommended_outputs.append("争取完成一个行业 CBAM 数据收集表或申报逻辑拆解案例。")
    if any(text_contains(internship_only_text, keyword) for keyword in ["英文", "英语", "海外", "欧盟"]):
        recommended_outputs.append("争取做一次英文材料整理或英文汇报，补强英文面试证据。")

    questions_to_ask = [
        "我是否能参与真实客户项目，还是主要做资料整理？",
        "是否能接触数据、需求、客户、开发、运营或报告核心章节？",
        "实习结束前能否产出一个可脱敏复盘的项目案例？",
        "导师是否能给我反馈方法、交付物和沟通上的问题？",
    ]
    negotiation_script = [
        "我希望实习期间能参与至少一个真实项目的数据、需求、执行或报告核心环节，这样可以更快理解业务流程。",
        "如果可以，我想在入职前确认是否能接触关键任务和交付物，而不只是做材料排版或事务性整理。",
        "实习结束时我希望能沉淀一个可脱敏复盘的项目案例，方便后续总结和持续改进。",
    ]
    first_week_plan = [
        "第 1 天：确认项目类型、交付物、导师要求和可接触的数据范围。",
        "第 2-3 天：整理项目资料目录，建立数据、需求、任务、证明材料和问题清单。",
        "第 4-5 天：主动争取一个可独立负责的小任务，如数据校验、竞品分析、结果图表、代码模块或报告章节初稿。",
    ]
    resume_bullets = [
        "参与真实业务项目的信息整理、数据校验或任务推进，梳理关键资料、问题清单和交付要求。",
        "协助整理项目报告、结果图表、需求文档或流程模板，提升交付物的可复核性和可复用性。",
    ]

    return {
        "score": final_score,
        "verdict": verdict,
        "decision": decision,
        "dimension_scores": pd.DataFrame(dimension_rows),
        "high_hits": high_hits,
        "low_hits": low_hits,
        "target_hits": target_hits,
        "reasons": reasons,
        "risks": risks,
        "recommended_outputs": recommended_outputs,
        "questions_to_ask": questions_to_ask,
        "negotiation_script": negotiation_script,
        "first_week_plan": first_week_plan,
        "resume_bullets": resume_bullets,
    }


def recruitment_fingerprint(record: dict[str, Any]) -> str:
    raw = "|".join(
        [
            str(record.get("公司", "")),
            str(record.get("岗位", "")),
            str(record.get("地点", "")),
            str(record.get("薪资", "")),
            str(record.get("类型", "")),
        ]
    ).lower()
    return hashlib.sha1(raw.encode("utf-8", errors="ignore")).hexdigest()


def infer_company_tier(company: str, text: str = "") -> str:
    combined = normalize_text(company + "\n" + text)
    top_keywords = ["SGS", "TÜV", "TUV", "Intertek", "必维", "BV", "UL", "DNV", "德勤", "PwC", "普华永道", "安永", "EY", "KPMG", "毕马威"]
    strong_keywords = ["上市", "集团", "外资", "世界500强", "头部", "龙头", "检测", "认证", "咨询", "汽车", "电池", "新能源"]
    if any(text_contains(combined, keyword) for keyword in top_keywords):
        return "头部/知名机构"
    if any(text_contains(combined, keyword) for keyword in strong_keywords):
        return "中高层级"
    if company and company != "未识别":
        return "普通公司"
    return "未知"


def parse_recruitment_record(text: str, source: str = "") -> dict[str, Any]:
    jd = analyze_jd(text)
    skills = jd_skill_list(jd)
    fields = extract_card_fields_v2({"source": source, "text": text}, text, jd)
    company = fields["company"] or "未识别"
    job_title = fields["title"] or "未识别"
    salary = fields["salary"] or "未识别"
    location = fields["location"] or "未识别"
    region_info = classify_region(location, text)
    education = fields["education"] or "未识别"
    experience = fields["experience"] or "未识别"
    record = {
        "公司": company,
        "岗位": job_title,
        "类型": detect_job_type_safe(text),
        "是否招实习": detect_internship_opening_safe(text),
        "是否招应届生": detect_fresh_graduate_safe(text),
        "薪资": salary,
        "地点": location,
        **region_info,
        "学历要求": education,
        "经验要求": experience,
        "公司层级": infer_company_tier(company, text),
        "岗位分类": jd.get("category", "待判断"),
        "高价值岗位": "是" if jd.get("value", {}).get("is_high_value") else "否",
        "低价值风险": "是" if jd.get("value", {}).get("is_generic_esg") else "否",
        "技能关键词": " / ".join(skills[:10]),
        "来源": source,
        "链接": source if re.match(r"^https?://", source, flags=re.I) else "",
        "JD原文": text,
        "原文片段": text[:260],
    }
    record["fingerprint"] = recruitment_fingerprint(record)
    return record


def analyze_recruitment_sources(text: str, crawl_results: list[dict[str, Any]] | None = None) -> pd.DataFrame:
    records = []
    for item in dedupe_jd_records(split_batch_jd_text(text)):
        records.append(parse_recruitment_record(item["text"], item.get("source", "粘贴文本")))
    for item in crawl_results or []:
        for record in records_from_crawl_result(item):
            records.append(parse_recruitment_record(record["text"], record.get("url") or item.get("url", "链接抓取")))
    if not records:
        return pd.DataFrame()
    return pd.DataFrame(records).drop_duplicates(subset=["fingerprint"], keep="first").reset_index(drop=True)


def register_recruitment_records(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return df
    init_db()
    user_id = require_user_id()
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    output = df.copy()
    output["是否新发现"] = "否"
    with sqlite3.connect(DB_PATH) as conn:
        existing = {
            row[0]
            for row in conn.execute("SELECT fingerprint FROM recruitment_posts WHERE user_id = ?", (user_id,)).fetchall()
        }
        for index, row in output.iterrows():
            fingerprint = row.get("fingerprint") or recruitment_fingerprint(row.to_dict())
            is_new = fingerprint not in existing
            output.at[index, "是否新发现"] = "是" if is_new else "否"
            values = (
                user_id,
                fingerprint,
                row.get("公司", ""),
                row.get("岗位", ""),
                row.get("类型", ""),
                row.get("是否招应届生", ""),
                row.get("薪资", ""),
                row.get("地点", ""),
                row.get("标准城市", ""),
                row.get("省份", ""),
                row.get("区域", ""),
                row.get("地域优先级", ""),
                row.get("学历要求", ""),
                row.get("经验要求", ""),
                row.get("公司层级", ""),
                row.get("岗位分类", ""),
                row.get("高价值岗位", ""),
                row.get("低价值风险", row.get("泛ESG风险", "")),
                row.get("技能关键词", ""),
                row.get("来源", ""),
                row.get("原文片段", ""),
                now,
                now,
            )
            if is_new:
                conn.execute(
                    """
                    INSERT INTO recruitment_posts (
                        user_id, fingerprint, company, job_title, job_type, fresh_graduate,
                        salary, location, standard_city, province, region, region_priority,
                        education, experience, company_tier,
                        category, high_value, generic_esg, skills, source,
                        snippet, first_seen, last_seen
                    )
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    values,
                )
            else:
                conn.execute(
                    """
                    UPDATE recruitment_posts
                    SET last_seen = ?, salary = ?, location = ?, standard_city = ?, province = ?,
                        region = ?, region_priority = ?, source = ?, snippet = ?
                    WHERE user_id = ? AND fingerprint = ?
                    """,
                    (
                        now,
                        row.get("薪资", ""),
                        row.get("地点", ""),
                        row.get("标准城市", ""),
                        row.get("省份", ""),
                        row.get("区域", ""),
                        row.get("地域优先级", ""),
                        row.get("来源", ""),
                        row.get("原文片段", ""),
                        user_id,
                        fingerprint,
                    ),
                )
            existing.add(fingerprint)
    return output


def load_recruitment_posts() -> pd.DataFrame:
    init_db()
    user_id = require_user_id()
    with sqlite3.connect(DB_PATH) as conn:
        return pd.read_sql_query("SELECT * FROM recruitment_posts WHERE user_id = ? ORDER BY last_seen DESC", conn, params=(user_id,))


def count_shortcomings(resume_match: dict[str, Any] | None, gap_analysis: dict[str, Any] | None) -> int:
    count = 0
    if resume_match:
        count += len(resume_match.get("missing_skills", []))
        count += len(resume_match.get("gap_examples", []))
    if gap_analysis:
        for items in gap_analysis.get("current_gaps", {}).values():
            count += len(items)
    return min(count, 12)


def historical_rates() -> dict[str, float]:
    df = load_applications()
    if df.empty:
        return {"pass_rate": 0.35, "interview_rate": 0.22, "offer_rate": 0.06, "sample": 0}
    sample = len(df)
    interview_count = 0
    offer_count = 0
    for _, row in df.iterrows():
        interview_status = str(row.get("interview_status", ""))
        offer_status = str(row.get("offer_status", ""))
        if interview_status and not any(keyword in interview_status for keyword in ["未开始", "已投递", "无"]):
            interview_count += 1
        if "Offer" in offer_status or "offer" in offer_status.lower():
            offer_count += 1
    return {
        "pass_rate": max(interview_count / sample, 0.08),
        "interview_rate": max(interview_count / sample, 0.08),
        "offer_rate": max(offer_count / sample, 0.02),
        "sample": sample,
    }


def predict_offer_probabilities(
    jd_analysis: dict[str, Any] | None,
    resume_match: dict[str, Any] | None,
    gap_analysis: dict[str, Any] | None,
    company_tier: str,
    education_fit: str,
    english_level: str,
) -> dict[str, Any]:
    match_score = int(resume_match.get("score", 55) if resume_match else 55)
    jd_text = jd_analysis.get("raw_text", "") if jd_analysis else ""
    english_required = any(text_contains(jd_text, keyword) for keyword in ["英文", "英语", "English", "口语", "CET-6", "六级"])
    shortcomings = count_shortcomings(resume_match, gap_analysis)
    history = historical_rates()

    pass_rate = 18 + match_score * 0.62
    interview_rate = 8 + match_score * 0.42
    offer_rate = 2 + match_score * 0.22

    tier_adjust = {
        "头部/知名机构": -10,
        "中高层级": -4,
        "普通公司": 4,
        "小公司/冷门岗位": 10,
        "未知": 0,
    }.get(company_tier, 0)
    pass_rate += tier_adjust
    interview_rate += tier_adjust * 0.7
    offer_rate += tier_adjust * 0.35

    if education_fit == "高于或满足要求":
        pass_rate += 6
        interview_rate += 4
        offer_rate += 2
    elif education_fit == "勉强满足":
        pass_rate -= 4
        interview_rate -= 3
        offer_rate -= 2
    elif education_fit == "低于要求":
        pass_rate -= 14
        interview_rate -= 10
        offer_rate -= 6

    if english_required:
        if english_level == "强":
            pass_rate += 5
            interview_rate += 5
            offer_rate += 4
        elif english_level == "一般":
            pass_rate -= 3
            interview_rate -= 4
            offer_rate -= 3
        else:
            pass_rate -= 10
            interview_rate -= 9
            offer_rate -= 7

    pass_rate -= shortcomings * 2.0
    interview_rate -= shortcomings * 1.6
    offer_rate -= shortcomings * 0.9

    if history["sample"] >= 5:
        pass_rate = pass_rate * 0.75 + history["pass_rate"] * 100 * 0.25
        interview_rate = interview_rate * 0.75 + history["interview_rate"] * 100 * 0.25
        offer_rate = offer_rate * 0.8 + history["offer_rate"] * 100 * 0.2

    pass_rate = float(np.clip(round(pass_rate, 1), 3, 95))
    interview_rate = float(np.clip(round(min(interview_rate, pass_rate * 0.92), 1), 2, 85))
    offer_rate = float(np.clip(round(min(offer_rate, interview_rate * 0.65), 1), 1, 60))

    drivers = [
        f"岗位匹配度：{match_score}/100",
        f"公司层级：{company_tier}",
        f"短板数量：{shortcomings}",
        f"历史投递样本：{history['sample']} 条",
    ]
    if english_required:
        drivers.append(f"JD 有英文要求，当前英文判断：{english_level}")
    if jd_analysis:
        drivers.append(f"岗位分类：{jd_analysis.get('category', '待判断')}")

    actions = []
    if resume_match and resume_match.get("missing_skills"):
        missing_line = " / ".join(resume_match["missing_skills"][:6])
        actions.append(f"投递前把简历摘要和最近一段项目经历补上这些词：{missing_line}。每个词至少对应一句真实证据。")
    if english_required and english_level != "强":
        actions.append("今天准备两段英文材料：45 秒自我介绍 + 90 秒目标岗位相关项目复盘，明天录音复盘一次。")
    if company_tier == "头部/知名机构":
        actions.append("头部公司投递版简历必须出现 3 类证据：项目交付物、工具/数据处理、可复核的结果或报告产出。")
    actions.append("投递后第 3 天记录状态；第 5 天若无反馈，补投 3 个同方向但竞争稍低的岗位。")
    application_steps = [
        "投递前：用定制简历页生成投递版，把摘要、核心能力和项目 bullet 替换进简历。",
        "投递当天：保存 JD、投递版本和投递渠道，避免后续面试时找不到原岗位。",
        "投递后 3-5 天：若无反馈，做一次轻量跟进，同时复盘这个岗位还能投哪些相邻公司。",
    ]

    return {
        "简历通过率": pass_rate,
        "进入面试概率": interview_rate,
        "拿 offer 概率": offer_rate,
        "drivers": drivers,
        "actions": actions,
        "application_steps": application_steps,
        "shortcomings": shortcomings,
        "history": history,
    }


def generate_resume_bullets(jd_analysis: dict[str, Any] | None, resume_match: dict[str, Any] | None) -> list[str]:
    if not jd_analysis:
        return []
    skills = []
    if resume_match:
        skills = resume_match.get("matched_skills", [])[:5]
    if not skills and not jd_analysis.get("skills", pd.DataFrame()).empty:
        skills = jd_analysis["skills"]["技能"].head(5).tolist()

    bullets = []
    category = jd_analysis.get("category", "")
    if "产品能力" in skills or category == "产品岗":
        bullets.append("围绕用户场景和业务目标拆解需求，梳理痛点、优先级、方案路径和验证指标，输出可执行的产品优化建议。")
    if "运营" in skills or category == "运营岗":
        bullets.append("基于目标用户和转化路径设计运营动作，跟踪拉新、转化、留存等指标，并根据数据复盘优化策略。")
    if "数据分析" in skills or "Python" in skills or "SQL" in skills or category == "数据/商业分析岗":
        bullets.append("基于 Python / SQL / Excel 处理项目或业务数据，完成数据清洗、指标构建、异常检查和结论可视化。")
    if "项目管理" in skills or category == "咨询/项目岗":
        bullets.append("拆解项目目标、里程碑和协作分工，跟进风险与进度，推动交付物按时输出并沉淀复盘。")
    if category == "研发/工程岗" or any(skill in skills for skill in ["前端", "后端", "机器学习/AI"]):
        bullets.append("参与功能模块、接口、组件或模型开发，完成关键逻辑实现、测试验证和问题复盘。")
    if category in ["LCA技术岗", "产品碳足迹岗"] or "LCA" in skills:
        bullets.append("围绕 LCA/产品碳足迹项目梳理目标、边界、清单数据和结果解释，支持报告或业务判断。")
    if "ISO14067" in skills or "PEF" in skills:
        bullets.append(
            "围绕 ISO14067 与 PEF 方法差异开展论文研究，能够比较边界设定、分配规则、数据质量和结果解释要求。"
        )
    if "CBAM" in skills:
        bullets.append(
            "关注欧盟 CBAM 与出海合规要求，能够拆解企业活动数据、排放因子、默认值使用和申报风险点。"
        )
    if not bullets:
        bullets.append("围绕目标岗位拆解项目任务，完成信息整理、数据处理、过程推进和结果复盘，形成可直接交付的文档或表格。")
    return bullets


def ocr_pdf_bytes(data: bytes, max_pages: int = 6) -> str:
    try:
        import pypdfium2 as pdfium
        import pytesseract

        pdf = pdfium.PdfDocument(data)
        chunks = []
        for index in range(min(len(pdf), max_pages)):
            page = pdf[index]
            bitmap = page.render(scale=2).to_pil()
            text = pytesseract.image_to_string(bitmap, lang="chi_sim+eng")
            if text.strip():
                chunks.append(text.strip())
        return "\n\n".join(chunks)
    except Exception as exc:
        return f"[PDF OCR 失败：请确认已安装 pypdfium2、pytesseract 和本机 Tesseract OCR，或上传文字版 PDF。错误：{exc}]"


def extract_text_from_upload(uploaded_file: Any) -> str:
    if uploaded_file is None:
        return ""
    suffix = Path(uploaded_file.name).suffix.lower()
    data = uploaded_file.getvalue()

    if suffix in [".txt", ".md", ".csv"]:
        for encoding in ["utf-8-sig", "utf-8", "gbk", "gb18030"]:
            try:
                return data.decode(encoding)
            except UnicodeDecodeError:
                continue
        return data.decode("utf-8", errors="ignore")

    if suffix in [".xlsx", ".xls"]:
        try:
            sheets = pd.read_excel(io.BytesIO(data), sheet_name=None)
            chunks = []
            for name, df in sheets.items():
                chunks.append(f"Sheet: {name}")
                chunks.append(df.astype(str).replace("nan", "").to_string(index=False))
            return "\n".join(chunks)
        except Exception as exc:
            return f"[Excel 解析失败：{exc}]"

    if suffix == ".pdf":
        try:
            import pdfplumber

            with pdfplumber.open(io.BytesIO(data)) as pdf:
                page_chunks = []
                table_chunks = []
                for page in pdf.pages:
                    page_text = page.extract_text(x_tolerance=1.5, y_tolerance=3, layout=True) or page.extract_text() or ""
                    if page_text.strip():
                        page_chunks.append(page_text)
                    for table in page.extract_tables() or []:
                        rows = [" | ".join((cell or "").strip() for cell in row if cell is not None) for row in table]
                        table_text = "\n".join(row for row in rows if row.strip())
                        if table_text:
                            table_chunks.append(table_text)
                text = "\n\n".join(page_chunks + table_chunks)
            if len(text.strip()) >= 50:
                return text
        except Exception:
            pass
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(data))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            if len(text.strip()) >= 50:
                return text
        except Exception as exc:
            text = f"[PDF 文字层解析失败：{exc}]"
        ocr_text = ocr_pdf_bytes(data)
        if len(normalize_text(ocr_text)) >= 50 and not ocr_text.startswith("[PDF OCR 失败"):
            return ocr_text
        if ocr_text.startswith("[PDF OCR 失败"):
            return ocr_text
        return "[PDF 解析失败：文字层过少且 OCR 未能读取有效内容，请上传文字版 PDF / DOCX / TXT。]"

    if suffix in [".docx", ".doc"]:
        if suffix == ".doc":
            return "[DOC 解析暂不支持，请另存为 DOCX / PDF / TXT 后上传。]"
        try:
            from docx import Document

            doc = Document(io.BytesIO(data))
            paragraphs = [paragraph.text for paragraph in doc.paragraphs]
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    table_text.append(" | ".join(cell.text for cell in row.cells))
            return "\n".join(paragraphs + table_text)
        except Exception as exc:
            return f"[Word 解析失败：请安装 python-docx 或改用文本粘贴。错误：{exc}]"

    if suffix in [".html", ".htm"]:
        raw = data.decode("utf-8", errors="ignore")
        try:
            from bs4 import BeautifulSoup

            return BeautifulSoup(raw, "html.parser").get_text("\n")
        except Exception:
            return re.sub(r"<[^>]+>", " ", raw)

    if suffix in [".png", ".jpg", ".jpeg", ".bmp", ".webp", ".tif", ".tiff"]:
        try:
            from PIL import Image
            import pytesseract

            image = Image.open(io.BytesIO(data))
            return pytesseract.image_to_string(image, lang="chi_sim+eng")
        except Exception as exc:
            return f"[图片 OCR 失败：请安装 pillow、pytesseract 和本机 Tesseract OCR，或改用文本粘贴。错误：{exc}]"

    return "[暂不支持该文件类型，请上传 PDF / DOCX / TXT / HTML / Excel / 图片。]"


RESUME_SECTION_RULES: dict[str, list[str]] = {
    "基本信息": ["基本信息", "个人信息", "联系方式", "求职意向", "个人简介", "profile"],
    "教育背景": ["教育背景", "教育经历", "教育", "education"],
    "实习/工作经历": ["实习经历", "工作经历", "实践经历", "工作经验", "internship", "work experience", "professional experience"],
    "项目经历": ["项目经历", "项目经验", "项目", "project experience", "projects"],
    "科研/论文": ["科研经历", "研究经历", "论文", "论文发表", "科研项目", "research", "publication"],
    "技能": ["专业技能", "技能", "技能证书", "软件技能", "skills", "software"],
    "证书/语言": ["证书", "语言", "英语", "cet", "ielts", "toefl", "certification", "language"],
    "获奖/校园": ["获奖", "荣誉", "竞赛", "校园经历", "社团经历", "awards", "honors"],
}
RESUME_SECTION_ORDER = list(RESUME_SECTION_RULES)


def normalize_resume_lines(text: str) -> list[str]:
    normalized = (text or "").replace("\r\n", "\n").replace("\r", "\n").replace("\u3000", " ")
    lines = []
    for raw_line in normalized.split("\n"):
        line = re.sub(r"[ \t]+", " ", raw_line).strip()
        line = line.strip("•·●◆◇■□")
        if not line or re.fullmatch(r"\d{1,3}", line):
            continue
        if line.lower() in {"page", "resume", "curriculum vitae"}:
            continue
        lines.append(line)
    return lines


def _clean_heading_token(text: str) -> str:
    return re.sub(r"[\s:：|｜/\-—_·•●◆◇■□（）()\[\]【】.]", "", text or "").lower()


def detect_resume_section(line: str) -> str | None:
    clean_line = _clean_heading_token(line)
    if not clean_line:
        return None
    for section, aliases in RESUME_SECTION_RULES.items():
        for alias in aliases:
            if re.match(rf"^\s*{re.escape(alias)}\s*[:：|｜/\-—]", line, flags=re.I):
                return section
            clean_alias = _clean_heading_token(alias)
            if clean_line == clean_alias:
                return section
            if clean_line.startswith(clean_alias) and len(clean_line) <= len(clean_alias) + 12:
                return section
            if len(clean_alias) >= 4 and clean_alias in clean_line and len(clean_line) <= 28:
                return section
    return None


def split_resume_section_heading(line: str) -> tuple[str | None, str]:
    section = detect_resume_section(line)
    if not section:
        return None, line
    aliases = sorted(RESUME_SECTION_RULES[section], key=len, reverse=True)
    for alias in aliases:
        match = re.match(rf"^\s*{re.escape(alias)}\s*[:：|｜/\-—]?\s*(.*)$", line, flags=re.I)
        if match:
            return section, match.group(1).strip()
    return section, ""


def parse_resume_sections(text: str) -> dict[str, list[str]]:
    sections: dict[str, list[str]] = {section: [] for section in RESUME_SECTION_ORDER}
    current = "基本信息"
    for line in normalize_resume_lines(text):
        section, remainder = split_resume_section_heading(line)
        if section:
            current = section
            if remainder:
                sections[current].append(remainder)
            continue
        sections.setdefault(current, []).append(line)
    return {section: lines for section, lines in sections.items() if lines}


def resume_skill_hits(text: str) -> list[str]:
    skills = []
    for skill, aliases in SKILL_ALIASES.items():
        if count_alias_hits(text, expanded_aliases(skill, aliases)):
            skills.append(skill)
    return skills


def resume_section_table(sections: dict[str, list[str]]) -> pd.DataFrame:
    rows = []
    for section in RESUME_SECTION_ORDER:
        lines = sections.get(section, [])
        if not lines:
            continue
        rows.append(
            {
                "板块": section,
                "读取行数": len(lines),
                "内容预览": "；".join(lines[:3])[:220],
            }
        )
    return pd.DataFrame(rows, columns=["板块", "读取行数", "内容预览"])


def _first_matching_lines(lines: list[str], keywords: list[str], limit: int = 4) -> list[str]:
    hits = []
    for line in lines:
        if any(text_contains(line, keyword) for keyword in keywords):
            hits.append(line)
        if len(hits) >= limit:
            break
    return hits


def extract_resume_highlights(text: str, sections: dict[str, list[str]]) -> dict[str, list[str]]:
    lines = normalize_resume_lines(text)
    skill_hits = resume_skill_hits(text)
    education_lines = sections.get("教育背景", [])[:4] or _first_matching_lines(lines, ["硕士", "研究生", "本科", "学士", "博士", "大学", "学院"])
    experience_lines = sections.get("实习/工作经历", [])[:5] or _first_matching_lines(lines, ["实习", "工作", "负责", "参与", "SAES", "咨询", "项目"])
    project_lines = sections.get("项目经历", [])[:5] or _first_matching_lines(lines, ["LCA", "碳足迹", "ISO14067", "PEF", "CBAM", "EPD", "数据分析"])
    language_lines = _first_matching_lines(lines, ["英语", "英文", "CET", "六级", "雅思", "托福", "IELTS", "TOEFL"])
    carbon_lines = _first_matching_lines(lines, HIGH_VALUE_KEYWORDS + ["固废", "资源循环", "减排", "排放因子"])
    return {
        "学历线索": education_lines,
        "经历线索": experience_lines,
        "项目线索": project_lines,
        "语言线索": language_lines,
        "碳/LCA证据": carbon_lines,
        "技能命中": skill_hits,
    }


def build_structured_resume_text(sections: dict[str, list[str]], fallback: str = "") -> str:
    parts = []
    for section in RESUME_SECTION_ORDER:
        lines = sections.get(section, [])
        if lines:
            parts.append(f"【{section}】\n" + "\n".join(lines))
    structured = "\n\n".join(parts).strip()
    return structured or "\n".join(normalize_resume_lines(fallback)).strip()


def parse_resume_content(text: str) -> dict[str, Any]:
    raw_text = text or ""
    lines = normalize_resume_lines(raw_text)
    clean_text = "\n".join(lines)
    sections = parse_resume_sections(raw_text)
    structured_text = build_structured_resume_text(sections, raw_text)
    highlights = extract_resume_highlights(raw_text, sections)
    return {
        "raw_text": raw_text,
        "clean_text": clean_text,
        "lines": lines,
        "sections": sections,
        "table": resume_section_table(sections),
        "skills": highlights["技能命中"],
        "highlights": highlights,
        "structured_text": structured_text,
    }


def parsed_resume_from_upload(uploaded_file: Any) -> dict[str, Any] | None:
    if uploaded_file is None:
        return None
    raw_text = extract_text_from_upload(uploaded_file)
    return parse_resume_content(raw_text)


def resume_text_from_parsed(parsed: dict[str, Any] | None) -> str:
    if not parsed:
        return ""
    return str(parsed.get("structured_text") or parsed.get("clean_text") or parsed.get("raw_text") or "")


def uploaded_file_signature(uploaded_file: Any) -> str:
    if uploaded_file is None:
        return ""
    return "|".join(
        [
            str(getattr(uploaded_file, "name", "")),
            str(getattr(uploaded_file, "size", "")),
            str(getattr(uploaded_file, "type", "")),
        ]
    )


def ensure_widget_text(key: str, default: str = "") -> None:
    if key not in st.session_state:
        st.session_state[key] = default or ""


def sync_uploaded_resume_to_widget(
    uploaded_file: Any,
    parsed: dict[str, Any] | None,
    text_key: str,
    signature_key: str,
) -> str:
    upload_text = resume_text_from_parsed(parsed)
    signature = uploaded_file_signature(uploaded_file)
    if upload_text and signature and st.session_state.get(signature_key) != signature:
        st.session_state[text_key] = upload_text
        st.session_state[signature_key] = signature
    return upload_text


def resume_parse_quality(parsed: dict[str, Any] | None) -> dict[str, Any]:
    if not parsed:
        return {"label": "未设置", "detail": "暂无简历内容", "score": 0}
    raw_text = str(parsed.get("raw_text", "") or "")
    clean_text = str(parsed.get("clean_text", "") or "")
    sections = parsed.get("sections", {}) or {}
    skills = parsed.get("skills", []) or []
    if raw_text.startswith("[") and "失败" in raw_text[:80]:
        return {"label": "解析失败", "detail": raw_text[:120], "score": 0}
    score = 0
    score += min(len(clean_text) // 400, 4) * 15
    score += min(len(sections), 6) * 6
    score += min(len(skills), 8) * 4
    score = int(np.clip(score, 0, 100))
    if len(clean_text) < 80:
        label = "文本过少"
        detail = "可读文字太少，可能是扫描版 PDF 或版式提取失败。"
    elif len(sections) < 2:
        label = "可用但结构弱"
        detail = "已读到文字，但板块识别较少，建议检查教育、经历、项目、技能是否完整。"
    elif score >= 70:
        label = "良好"
        detail = "简历文字和板块结构较完整，可直接用于匹配。"
    else:
        label = "可用"
        detail = "可用于分析，但建议补齐关键经历和技能板块。"
    return {"label": label, "detail": detail, "score": score}


def render_global_resume_status(active_resume: dict[str, Any]) -> None:
    content = str(active_resume.get("content") or "")
    if not content.strip():
        st.caption("当前简历未设置。上传或粘贴后，所有功能都会调用这一份。")
        return
    parsed = parse_resume_content(content)
    quality = resume_parse_quality(parsed)
    st.caption(f"当前简历：{active_resume.get('name', '未命名')}｜更新：{active_resume.get('updated_at', '') or '未知'}")
    cols = st.columns(3)
    cols[0].metric("有效行数", len(parsed.get("lines", [])))
    cols[1].metric("识别板块", len(parsed.get("sections", {})))
    cols[2].metric("质量", quality["label"])
    if parsed.get("skills"):
        st.caption("技能命中：" + " / ".join(parsed["skills"][:8]))
    st.caption(str(quality["detail"]))


def unique_resume_name(base_name: str) -> str:
    base_name = (base_name or "简历").strip()
    existing = set(load_user_resumes()["name"].astype(str).tolist())
    if base_name not in existing:
        return base_name
    suffix = datetime.now().strftime("%m%d%H%M")
    candidate = f"{base_name}-{suffix}"
    counter = 2
    while candidate in existing:
        candidate = f"{base_name}-{suffix}-{counter}"
        counter += 1
    return candidate


def normalize_url(url: str) -> str:
    url = url.strip()
    if not url:
        return ""
    if not re.match(r"^https?://", url, flags=re.I):
        url = "https://" + url
    return url


def canonical_job_url(url: str) -> str:
    clean = normalize_url(str(url or ""))
    if not clean:
        return ""
    try:
        parsed = urlparse(clean)
    except Exception:
        return clean.split("#", 1)[0].rstrip("/")
    if not parsed.netloc:
        return ""
    host = parsed.netloc.lower()
    if "." not in host and host not in {"localhost"} and not re.match(r"^\d{1,3}(?:\.\d{1,3}){3}(?::\d+)?$", host):
        return ""
    ignored_params = {
        "",
        "_",
        "ka",
        "lid",
        "page",
        "pageindex",
        "page_index",
        "securityid",
        "sessionid",
        "source",
        "src",
        "from",
        "utm_source",
        "utm_medium",
        "utm_campaign",
        "utm_term",
        "utm_content",
    }
    query_items = []
    for key, value in parse_qsl(parsed.query, keep_blank_values=False):
        key_lower = key.lower()
        if key_lower in ignored_params or key_lower.startswith("utm_"):
            continue
        query_items.append((key_lower, value))
    path = re.sub(r"/+$", "", parsed.path or "/")
    return urlunparse(
        (
            parsed.scheme.lower() or "https",
            host,
            path,
            "",
            urlencode(sorted(query_items), doseq=True),
            "",
        )
    )


def compact_identity_value(value: Any) -> str:
    text = normalize_text(str(value or ""))
    if not text or text == "未识别":
        return ""
    text = re.sub(r"\s+", "", text)
    text = re.sub(r"[|｜,，;；:：。]+", "", text)
    return text.lower()


def record_identity_fields(record: dict[str, str], text: str) -> dict[str, str]:
    title = normalize_text(str(record.get("title") or record.get("job_title") or ""))
    company = normalize_text(str(record.get("company") or ""))
    salary = normalize_text(str(record.get("salary") or ""))
    location = normalize_text(str(record.get("location") or ""))

    if not title or title == "未识别":
        title = extract_job_title(text)
    if not company or company == "未识别":
        company = extract_company(text)
    if not salary or salary == "未识别":
        salary = extract_salary(text)
    if not location or location == "未识别":
        location = extract_location(text)

    return {
        "title": compact_identity_value(title),
        "company": compact_identity_value(company),
        "salary": compact_identity_value(salary),
        "location": compact_identity_value(location),
    }


def jd_record_dedupe_keys(record: dict[str, str], text: str) -> list[str]:
    keys: list[str] = []
    record_url = str(record.get("url") or "")
    source_url = str(record.get("source") or "")
    canonical_url = canonical_job_url(record_url)
    if canonical_url and detail_url_is_usable(record_url, source_url):
        keys.append("url:" + canonical_url)

    fields = record_identity_fields(record, text)
    if fields["company"] and fields["title"] and (fields["salary"] or fields["location"]):
        keys.append("|".join(["job", fields["company"], fields["title"], fields["salary"], fields["location"]]))
    elif fields["title"] and fields["salary"] and fields["location"]:
        keys.append("|".join(["role", fields["title"], fields["salary"], fields["location"]]))
    elif fields["title"] and fields["salary"]:
        keys.append("|".join(["role_salary", fields["title"], fields["salary"]]))

    compact_text = re.sub(r"\s+", "", text)
    if compact_text:
        keys.append("text:" + content_fingerprint(compact_text[:900]))
    return keys


def valid_jd_record(record: dict[str, str], text: str) -> bool:
    company = normalize_text(str(record.get("company") or ""))
    title = normalize_text(str(record.get("title") or record.get("job_title") or ""))
    location = normalize_text(str(record.get("location") or ""))
    if company and re.search(r"(?:先生|女士|老师|HR|hr|在线)$", company):
        return False
    if company and invalid_plugin_company_line(company):
        return False
    if title and re.search(r"(?:能力|职责|要求|工作地点|任职资格|职位描述|项目经验|开发经验|工作经验|任职经历)$", title):
        return False
    if location and re.search(r"^\d+\)|工作地点[:：]?", location):
        return False
    return True


def extract_urls(text: str) -> list[str]:
    urls = []
    for line in text.splitlines():
        clean = normalize_text(line)
        for match in re.findall(r"https?://[^\s<>'\"，。；;、）)]+", clean, flags=re.I):
            candidate = normalize_url(match)
            parsed = urlparse(candidate)
            if parsed.netloc:
                urls.append(candidate)
        if re.match(r"^[A-Za-z0-9.-]+\.[A-Za-z]{2,}(?:/.*)?$", clean):
            candidate = normalize_url(clean)
            parsed = urlparse(candidate)
            if parsed.netloc:
                urls.append(candidate)
    return list(dict.fromkeys(urls))


def job_card_score(text: str) -> int:
    clean = normalize_text(text)
    if not clean:
        return 0
    keywords = [
        "岗位",
        "职位",
        "招聘",
        "薪资",
        "月薪",
        "日薪",
        "经验",
        "学历",
        "本科",
        "硕士",
        "实习",
        "全职",
        "公司",
        "工作地点",
        "职责",
        "要求",
        "会计",
        "财务",
        "专员",
        "助理",
        "顾问",
        "分析师",
        "LCA",
        "ESG",
        "CBAM",
        "EPD",
        "ISO",
        "Python",
    ]
    score = sum(1 for keyword in keywords if text_contains(clean, keyword))
    if extract_salary(clean) != "未识别":
        score += 5
    if re.search(r"(有限公司|集团|科技|咨询|检测|认证|股份|事务所|公司)", clean):
        score += 2
    if re.search(r"(上海|北京|深圳|广州|苏州|杭州|南京|成都|武汉|重庆|宁波|无锡)", clean):
        score += 2
    if 40 <= len(clean) <= 2200:
        score += 2
    elif len(clean) > 3500:
        score -= 4
    return score


def jsonld_items(value: Any) -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []
    if isinstance(value, list):
        for item in value:
            items.extend(jsonld_items(item))
    elif isinstance(value, dict):
        items.append(value)
        graph = value.get("@graph")
        if isinstance(graph, list):
            items.extend(jsonld_items(graph))
    return items


def jsonld_type_matches(item: dict[str, Any], type_name: str) -> bool:
    value = item.get("@type") or item.get("type")
    if isinstance(value, list):
        return any(str(part).lower() == type_name.lower() for part in value)
    return str(value).lower() == type_name.lower()


def first_jsonld_text(value: Any, *keys: str) -> str:
    if isinstance(value, dict):
        for key in keys:
            item = value.get(key)
            if isinstance(item, str) and item.strip():
                return normalize_text(item)
            if isinstance(item, dict):
                nested = first_jsonld_text(item, "name", "addressLocality", "addressRegion", "streetAddress")
                if nested:
                    return nested
            if isinstance(item, list):
                nested_values = [first_jsonld_text(part, "name", "addressLocality", "addressRegion") for part in item]
                nested_values = [part for part in nested_values if part]
                if nested_values:
                    return " / ".join(nested_values)
    elif isinstance(value, list):
        nested_values = [first_jsonld_text(part, *keys) for part in value]
        nested_values = [part for part in nested_values if part]
        if nested_values:
            return " / ".join(nested_values)
    return ""


def salary_from_jsonld(value: Any) -> str:
    if not isinstance(value, dict):
        return ""
    base_salary = value.get("baseSalary")
    if not isinstance(base_salary, dict):
        return ""
    salary_value = base_salary.get("value")
    currency = str(base_salary.get("currency") or "")
    unit = str(base_salary.get("unitText") or "")
    if isinstance(salary_value, dict):
        min_value = salary_value.get("minValue")
        max_value = salary_value.get("maxValue")
        raw_value = salary_value.get("value")
        if min_value and max_value:
            return normalize_text(f"{min_value}-{max_value} {currency} {unit}")
        if raw_value:
            return normalize_text(f"{raw_value} {currency} {unit}")
    return ""


GENERIC_JOB_TITLE_KEYS = [
    "title",
    "name",
    "jobName",
    "job_name",
    "jobTitle",
    "positionName",
    "position",
    "postName",
    "post",
    "recruitPost",
    "zwmc",
    "岗位名称",
    "职位名称",
    "岗位",
    "职位",
]

GENERIC_JOB_COMPANY_KEYS = [
    "company",
    "companyName",
    "companyFullName",
    "companyShortName",
    "company_name",
    "corpName",
    "corp_name",
    "brandName",
    "enterpriseName",
    "enterprise_name",
    "orgName",
    "org_name",
    "unitName",
    "unit_name",
    "recruitUnit",
    "employerName",
    "employer",
    "organization",
    "aab004",
    "单位名称",
    "公司名称",
    "企业名称",
]

GENERIC_JOB_SALARY_KEYS = [
    "salary",
    "salaryText",
    "salaryDesc",
    "salaryRange",
    "pay",
    "wage",
    "monthSalary",
    "薪资",
    "薪酬",
    "月薪",
]

GENERIC_JOB_LOCATION_KEYS = [
    "location",
    "city",
    "cityName",
    "area",
    "areaName",
    "district",
    "workPlace",
    "workAddress",
    "address",
    "aab302",
    "area_",
    "地点",
    "城市",
    "工作地点",
]

GENERIC_JOB_DESCRIPTION_KEYS = [
    "description",
    "jobDesc",
    "jobDescription",
    "requirement",
    "requirements",
    "responsibility",
    "responsibilities",
    "duties",
    "content",
    "summary",
    "acb22a",
    "岗位描述",
    "职位描述",
    "岗位职责",
    "任职要求",
]

GENERIC_JOB_URL_KEYS = ["url", "link", "href", "jobUrl", "detailUrl", "applyUrl", "ace760", "链接"]
GENERIC_JOB_ID_KEYS = ["id", "jobId", "positionId", "postId", "recruitId", "acb200"]


def first_value_by_keys(item: dict[str, Any], keys: list[str]) -> str:
    lower_map = {str(key).lower(): key for key in item.keys()}
    for key in keys:
        actual = lower_map.get(key.lower())
        if actual is not None:
            value = clean_json_job_value(item.get(actual))
            if value:
                return value
    return ""


def first_nested_value_by_keys(value: Any, keys: list[str], depth: int = 0) -> str:
    if depth > 2:
        return ""
    if isinstance(value, dict):
        direct = first_value_by_keys(value, keys)
        if direct:
            return direct
        for nested in value.values():
            found = first_nested_value_by_keys(nested, keys, depth + 1)
            if found:
                return found
    elif isinstance(value, list):
        for item in value[:8]:
            found = first_nested_value_by_keys(item, keys, depth + 1)
            if found:
                return found
    return ""


def generic_salary_from_item(item: dict[str, Any]) -> str:
    direct = first_nested_value_by_keys(item, GENERIC_JOB_SALARY_KEYS)
    if direct:
        return direct
    return public_job_salary(item)


def generic_url_from_item(item: dict[str, Any], page_url: str) -> str:
    direct = first_nested_value_by_keys(item, GENERIC_JOB_URL_KEYS)
    if direct:
        return normalize_url(urljoin(page_url, html.unescape(direct)))
    item_id = first_value_by_keys(item, GENERIC_JOB_ID_KEYS)
    if item_id:
        return normalize_url(urljoin(page_url, str(item_id)))
    return page_url


def record_from_generic_job_item(item: dict[str, Any], page_url: str) -> dict[str, str] | None:
    title = first_nested_value_by_keys(item, GENERIC_JOB_TITLE_KEYS)
    company = first_nested_value_by_keys(item, GENERIC_JOB_COMPANY_KEYS)
    location = first_nested_value_by_keys(item, GENERIC_JOB_LOCATION_KEYS)
    salary = generic_salary_from_item(item)
    description = first_nested_value_by_keys(item, GENERIC_JOB_DESCRIPTION_KEYS)
    if not title:
        return None
    key_text = "".join(map(str, item.keys())).lower()
    if title == company:
        return None
    if "name" in {str(key).lower() for key in item.keys()} and not any(marker in key_text for marker in ["job", "position", "post", "recruit", "岗位", "职位"]):
        return None
    if len(title) > 120 or re.search(r"^\d+$|https?://|登录|注册|首页|筛选|排序", title, flags=re.I):
        return None
    url = generic_url_from_item(item, page_url)
    lines = [
        f"岗位：{title}",
        f"公司：{company}" if company else "",
        f"薪资：{salary}" if salary else "",
        f"地点：{location}" if location else "",
        f"链接：{url}" if url else "",
        f"岗位描述：{description}" if description else "",
    ]
    text = normalize_multiline_text("\n".join(line for line in lines if line))
    if job_card_score(text) < 5 and not any(key in key_text for key in ["job", "position", "post", "recruit", "岗位", "职位"]):
        return None
    return {
        "source": page_url,
        "title": title,
        "url": url,
        "text": text,
        "company": company,
        "salary": salary,
        "location": location,
    }


def walk_json_dicts(value: Any, limit: int = 1200) -> list[dict[str, Any]]:
    output: list[dict[str, Any]] = []

    def visit(node: Any) -> None:
        if len(output) >= limit:
            return
        if isinstance(node, dict):
            output.append(node)
            for child in node.values():
                visit(child)
        elif isinstance(node, list):
            for child in node[:500]:
                visit(child)

    visit(value)
    return output


def parse_json_candidate(text: str) -> Any | None:
    text = html.unescape(text or "").strip()
    if not text:
        return None
    try:
        return json.loads(text)
    except Exception:
        pass
    start_positions = [pos for pos in [text.find("{"), text.find("[")] if pos >= 0]
    if not start_positions:
        return None
    start = min(start_positions)
    end_char = "}" if text[start] == "{" else "]"
    end = text.rfind(end_char)
    if end <= start:
        return None
    try:
        return json.loads(text[start:end + 1])
    except Exception:
        return None


def extract_generic_json_job_records_from_html(raw_html: str, page_url: str) -> list[dict[str, str]]:
    try:
        from bs4 import BeautifulSoup
    except Exception:
        return []
    soup = BeautifulSoup(raw_html, "html.parser")
    blobs: list[str] = []
    for script in soup.find_all("script"):
        script_text = script.string or script.get_text("", strip=True)
        if script_text and any(marker in script_text.lower() for marker in ["job", "position", "recruit", "岗位", "职位"]):
            blobs.append(script_text)
    for node in soup.find_all(["input", "textarea"]):
        value = node.get("value") or node.get_text("", strip=True)
        if value and len(value) >= 30:
            blobs.append(str(value))

    records: list[dict[str, str]] = []
    seen_keys: set[str] = set()
    for blob in blobs[:80]:
        data = parse_json_candidate(blob)
        if data is None:
            continue
        for item in walk_json_dicts(data):
            record = record_from_public_job_item(item, page_url) or record_from_generic_job_item(item, page_url)
            if not record:
                continue
            key = content_fingerprint(record.get("url", "") + "|" + record.get("title", "") + "|" + record.get("text", "")[:260])
            if key in seen_keys:
                continue
            seen_keys.add(key)
            records.append(record)
            if len(records) >= 200:
                break
    return dedupe_jd_records(records)


def extract_jobposting_records_from_html(raw_html: str, page_url: str) -> list[dict[str, str]]:
    try:
        from bs4 import BeautifulSoup
    except Exception:
        return []
    soup = BeautifulSoup(raw_html, "html.parser")
    records: list[dict[str, str]] = []
    for script in soup.find_all("script", attrs={"type": re.compile("ld\\+json", re.I)}):
        script_text = script.string or script.get_text("", strip=True)
        if not script_text:
            continue
        try:
            data = json.loads(script_text)
        except Exception:
            continue
        for item in jsonld_items(data):
            if not jsonld_type_matches(item, "JobPosting"):
                continue
            title = first_jsonld_text(item, "title", "name")
            company = first_jsonld_text(item.get("hiringOrganization"), "name")
            location = first_jsonld_text(item.get("jobLocation"), "addressLocality", "addressRegion", "name")
            salary = salary_from_jsonld(item)
            link = normalize_url(str(item.get("url") or page_url))
            description = normalize_multiline_text(html_to_text(str(item.get("description") or "")))
            lines = [
                f"岗位：{title}" if title else "",
                f"公司：{company}" if company else "",
                f"地点：{location}" if location else "",
                f"薪资：{salary}" if salary else "",
                f"链接：{link}" if link else "",
                description,
            ]
            text = normalize_multiline_text("\n".join(part for part in lines if part))
            if len(normalize_text(text)) >= 30:
                records.append(
                    {
                        "source": page_url,
                        "title": title or "结构化岗位",
                        "url": link,
                        "text": text,
                        "company": company,
                        "salary": salary,
                        "location": location,
                    }
                )
    return records


def looks_like_company_name(value: str) -> bool:
    return is_probable_company_name(value)


def clean_json_job_value(value: Any) -> str:
    if value is None:
        return ""
    text = normalize_text(str(value))
    if text.lower() in {"none", "null", "nan"}:
        return ""
    return text


def public_job_salary(item: dict[str, Any]) -> str:
    low = clean_json_job_value(item.get("acb241"))
    high = clean_json_job_value(item.get("acb242"))
    if low and high and low != "0" and high != "0":
        return f"{low}-{high}元/月"
    if low and low != "0":
        return f"{low}元以上/月"
    return ""


def public_job_detail_url(item: dict[str, Any], page_url: str) -> str:
    external = clean_json_job_value(item.get("ace760"))
    if external:
        return normalize_url(html.unescape(external))
    job_id = clean_json_job_value(item.get("acb200"))
    if job_id:
        return normalize_url(urljoin(page_url, f"../jobinfolist/cb21/showgw?id={job_id}"))
    return page_url


def record_from_public_job_item(item: dict[str, Any], page_url: str) -> dict[str, str] | None:
    title = clean_json_job_value(item.get("aca112") or item.get("title") or item.get("name"))
    if not title:
        return None
    company = clean_json_job_value(item.get("aab004") or item.get("company"))
    location_detail = clean_json_job_value(item.get("aab302") or item.get("area_"))
    address = clean_json_job_value(item.get("acb202") or item.get("aae006"))
    if not company and looks_like_company_name(address):
        company = address
    salary = public_job_salary(item)
    category = clean_json_job_value(item.get("aca111_") or item.get("aca111_Local_"))
    description = clean_json_job_value(item.get("acb22a") or item.get("description"))
    recruit_count = clean_json_job_value(item.get("acb240"))
    publish_date = clean_json_job_value(item.get("s_aae397") or item.get("s_ctime"))
    expire_date = clean_json_job_value(item.get("s_aae398"))
    url = public_job_detail_url(item, page_url)

    lines = [
        f"岗位：{title}",
        f"公司：{company}" if company else "",
        f"薪资：{salary}" if salary else "",
        f"地点：{location_detail}" if location_detail else "",
        f"地址：{address}" if address and address != company else "",
        f"岗位类别：{category}" if category else "",
        f"招聘人数：{recruit_count}" if recruit_count and recruit_count != "0" else "",
        f"发布日期：{publish_date}" if publish_date else "",
        f"截止日期：{expire_date}" if expire_date else "",
        f"链接：{url}" if url else "",
        f"岗位描述：{description}" if description else "",
    ]
    text = normalize_multiline_text("\n".join(line for line in lines if line))
    if len(normalize_text(text)) < 20:
        return None
    return {
        "source": page_url,
        "title": title,
        "url": url,
        "text": text,
        "company": company,
        "salary": salary,
        "location": location_detail,
        "experience": extract_experience(description) if description else "",
    }


def extract_hidden_json_job_records_from_html(raw_html: str, page_url: str) -> list[dict[str, str]]:
    try:
        from bs4 import BeautifulSoup
    except Exception:
        return []
    soup = BeautifulSoup(raw_html, "html.parser")
    records: list[dict[str, str]] = []
    for input_node in soup.find_all("input"):
        value = input_node.get("value")
        if not value:
            continue
        decoded = html.unescape(str(value)).strip()
        if not decoded.startswith("[") or "aca112" not in decoded or "acb200" not in decoded:
            continue
        try:
            data = json.loads(decoded)
        except Exception:
            continue
        if not isinstance(data, list):
            continue
        for item in data:
            if not isinstance(item, dict):
                continue
            record = record_from_public_job_item(item, page_url)
            if record:
                records.append(record)
    return dedupe_jd_records(records)


def field_from_node(node: Any, selectors: list[str]) -> str:
    for selector in selectors:
        try:
            target = node.select_one(selector)
        except Exception:
            target = None
        if target:
            value = normalize_text(target.get_text(" ", strip=True))
            if value:
                return value
            for attr in ["title", "aria-label", "data-company", "data-name", "alt"]:
                attr_value = normalize_text(str(target.get(attr) or ""))
                if attr_value:
                    return attr_value
    return ""


def records_from_html_tables(soup: Any, page_url: str) -> list[dict[str, str]]:
    records: list[dict[str, str]] = []
    for table in soup.find_all("table"):
        rows = table.find_all("tr")
        if len(rows) < 2:
            continue
        header_cells = rows[0].find_all(["th", "td"])
        headers = [normalize_text(cell.get_text(" ", strip=True)) for cell in header_cells]
        header_text = " ".join(headers)
        if not any(word in header_text for word in ["岗位", "职位", "公司", "单位", "薪资", "地点", "Job", "Position", "Company"]):
            continue
        for row in rows[1:]:
            cells = row.find_all(["td", "th"])
            values = [normalize_text(cell.get_text(" ", strip=True)) for cell in cells]
            if len([value for value in values if value]) < 2:
                continue
            row_dict = {
                headers[index] if index < len(headers) and headers[index] else f"列{index + 1}": value
                for index, value in enumerate(values)
            }
            df = pd.DataFrame([row_dict])
            table_records = records_from_dataframe(df, page_url)
            if table_records:
                record = table_records[0]
            else:
                record = {"source": page_url, "title": likely_title_from_row_values(values), "url": page_url, "text": "\n".join(values)}
            anchor = row.select_one("a[href]")
            if anchor and anchor.get("href"):
                record["url"] = normalize_url(urljoin(page_url, str(anchor.get("href"))))
            if job_card_score(record.get("text", "")) >= 5:
                records.append(record)
    return dedupe_jd_records(records)


def extract_job_card_records_from_html(raw_html: str, page_url: str) -> list[dict[str, str]]:
    try:
        from bs4 import BeautifulSoup
    except Exception:
        return []
    soup = BeautifulSoup(raw_html, "html.parser")
    for tag in soup(["script", "style", "noscript", "svg", "canvas", "iframe"]):
        tag.decompose()

    records = extract_hidden_json_job_records_from_html(raw_html, page_url)
    records.extend(extract_generic_json_job_records_from_html(raw_html, page_url))
    records.extend(extract_jobposting_records_from_html(raw_html, page_url))
    records.extend(records_from_html_tables(soup, page_url))
    selector = ",".join(
        [
            "[data-jobid]",
            "[data-jid]",
            "[data-job-id]",
            "[class*='job']",
            "[class*='Job']",
            "[class*='position']",
            "[class*='Position']",
            "[class*='recruit']",
            "[class*='vacancy']",
            "article",
            "li",
        ]
    )
    seen = {normalize_text(record.get("text", ""))[:260] for record in records}
    candidates = []
    for node in soup.select(selector):
        text = normalize_multiline_text(node.get_text("\n", strip=True))
        compact = normalize_text(text)
        if not (18 <= len(compact) <= 2400):
            continue
        score = job_card_score(text)
        if score < 7:
            continue
        candidates.append((score, len(compact), node, text))

    for _score, _length, node, text in sorted(candidates, key=lambda item: (item[0], -item[1]), reverse=True):
        compact_key = normalize_text(text)[:260]
        if compact_key in seen or any(compact_key in existing or existing in compact_key for existing in seen):
            continue
        title = field_from_node(
            node,
            [
                "[class*='title']",
                "[class*='name']",
                "[class*='job-name']",
                "[class*='position-name']",
                "h1",
                "h2",
                "h3",
                "a",
            ],
        )
        company = field_from_node(
            node,
            [
                "[class*='company']",
                "[class*='Company']",
                "[class*='corp']",
                "[class*='Corp']",
                "[class*='brand']",
                "[class*='employer']",
                "[class*='Employer']",
                "[class*='enterprise']",
                "[class*='Enterprise']",
                "[class*='unit']",
                "[class*='org']",
                "[data-company]",
                "[data-employer]",
            ],
        )
        if not company:
            candidates = extract_company_candidates(text, title)
            company = candidates[0] if candidates else ""
        link = ""
        anchor = node.select_one("a[href]")
        if anchor and anchor.get("href"):
            link = normalize_url(urljoin(page_url, str(anchor.get("href"))))
        if not link:
            link = page_url
        record = {
            "source": page_url,
            "title": title[:100] if title else "网页岗位",
            "url": link,
            "text": text,
            "company": company,
        }
        records.append(record)
        seen.add(compact_key)
        if len(records) >= 120:
            break
    return dedupe_jd_records(records)


def html_to_text(html: str) -> str:
    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "noscript", "svg", "canvas", "iframe"]):
            tag.decompose()

        title = soup.title.get_text(" ", strip=True) if soup.title else ""
        meta_parts = []
        for attrs in [
            {"name": "description"},
            {"property": "og:description"},
            {"name": "keywords"},
        ]:
            tag = soup.find("meta", attrs=attrs)
            if tag and tag.get("content"):
                meta_parts.append(tag["content"])

        main_candidates = soup.find_all(["main", "article", "section", "div", "p", "li"])
        pieces = [title] + meta_parts
        if main_candidates:
            scored = []
            for node in main_candidates:
                text = node.get_text("\n", strip=True)
                if len(text) >= 80:
                    scored.append((len(text), text))
            for _, text in sorted(scored, reverse=True)[:8]:
                pieces.append(text)
        else:
            pieces.append(soup.get_text("\n", strip=True))

        return normalize_multiline_text("\n\n".join(pieces))
    except Exception:
        no_tags = re.sub(r"<[^>]+>", " ", html)
        return normalize_multiline_text(no_tags)


def build_crawler_headers(url: str) -> dict[str, str]:
    headers = dict(CRAWLER_HEADERS)
    if url:
        headers["Referer"] = url
    return headers


def fetch_jd_url_with_playwright(url: str, timeout_ms: int = 18000) -> dict[str, Any]:
    normalized = normalize_url(url)
    result = {"url": normalized, "ok": False, "text": "", "error": "", "records": []}
    if not normalized:
        result["error"] = "空链接"
        return result
    try:
        from playwright.sync_api import sync_playwright

        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page(
                user_agent=CRAWLER_HEADERS["User-Agent"],
                locale="zh-CN",
                viewport={"width": 1440, "height": 1000},
            )
            page.goto(normalized, wait_until="domcontentloaded", timeout=timeout_ms)
            try:
                page.wait_for_load_state("networkidle", timeout=5000)
            except Exception:
                pass
            page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
            page.wait_for_timeout(900)
            html = page.content()
            browser.close()
        result["records"] = extract_job_card_records_from_html(html, normalized)
        text = html_to_text(html)
        if len(text) < 80:
            result["error"] = "动态抓取文本仍然过短，页面可能需要登录或验证码"
            result["text"] = text
            return result
        result["ok"] = True
        result["text"] = f"来源：{normalized}\n{text}"
        return result
    except Exception as exc:
        result["error"] = f"Playwright 动态抓取失败：{exc}"
        return result


def fetch_jd_url(url: str, timeout: int = 12, use_dynamic: bool = False) -> dict[str, Any]:
    normalized = normalize_url(url)
    result = {"url": normalized, "ok": False, "text": "", "error": "", "records": []}
    if not normalized:
        result["error"] = "空链接"
        return result

    try:
        response = requests.get(
            normalized,
            headers=build_crawler_headers(normalized),
            timeout=timeout,
        )
        response.raise_for_status()
        response.encoding = response.apparent_encoding or response.encoding
        content_type = response.headers.get("content-type", "").lower()
        if "pdf" in content_type or normalized.lower().endswith(".pdf"):
            try:
                import pdfplumber

                with pdfplumber.open(io.BytesIO(response.content)) as pdf:
                    text = "\n".join(page.extract_text() or "" for page in pdf.pages)
            except Exception:
                try:
                    from pypdf import PdfReader

                    reader = PdfReader(io.BytesIO(response.content))
                    text = "\n".join(page.extract_text() or "" for page in reader.pages)
                except Exception as exc:
                    result["error"] = f"PDF 提取失败：{exc}"
                    return result
        else:
            result["records"] = extract_job_card_records_from_html(response.text, normalized)
            text = html_to_text(response.text)

        if len(text) < 80:
            if use_dynamic:
                dynamic_result = fetch_jd_url_with_playwright(normalized)
                if dynamic_result.get("ok"):
                    return dynamic_result
            result["error"] = "抓取文本过短，页面可能需要登录、验证码或前端动态渲染"
            result["text"] = text
            return result

        result["ok"] = True
        result["text"] = f"来源：{normalized}\n{text}"
        return result
    except Exception as exc:
        if use_dynamic:
            dynamic_result = fetch_jd_url_with_playwright(normalized)
            if dynamic_result.get("ok"):
                return dynamic_result
            result["error"] = f"{exc}；{dynamic_result.get('error', '')}"
            return result
        result["error"] = str(exc)
        return result


def crawl_jd_urls(urls: list[str], max_workers: int = 5, use_dynamic: bool = False) -> list[dict[str, Any]]:
    clean_urls = [url for url in urls if url]
    if not clean_urls:
        return []

    workers = max(1, min(max_workers, len(clean_urls), 3 if use_dynamic else 8))
    results = []
    with ThreadPoolExecutor(max_workers=workers) as executor:
        futures = {executor.submit(fetch_jd_url, url, 12, use_dynamic): url for url in clean_urls}
        for future in as_completed(futures):
            try:
                results.append(future.result())
            except Exception as exc:
                results.append({"url": futures[future], "ok": False, "text": "", "error": str(exc)})
    return sorted(results, key=lambda item: clean_urls.index(item["url"]) if item["url"] in clean_urls else 999)


def detail_url_is_usable(url: str, source_url: str = "") -> bool:
    clean = normalize_url(str(url or ""))
    if not clean:
        return False
    if clean.lower().startswith(("javascript:", "mailto:", "tel:")):
        return False
    parsed = urlparse(clean)
    if parsed.scheme not in {"http", "https"} or not parsed.netloc:
        return False
    if source_url and normalize_url(source_url).split("#", 1)[0] == clean.split("#", 1)[0]:
        return False
    if re.search(r"/search|/jobs\?|/interns\?|/list|/index", parsed.path + ("?" + parsed.query if parsed.query else ""), re.I):
        return False
    return True


def fetch_detail_for_enrichment(url: str) -> dict[str, str]:
    result = fetch_jd_url(url, timeout=10, use_dynamic=False)
    if not result.get("ok"):
        return {"ok": "", "error": str(result.get("error") or ""), "url": normalize_url(url)}
    text = normalize_multiline_text(str(result.get("text") or ""))
    fields = extract_detail_fields_from_text(url, text)
    if not fields:
        detail_records = records_from_crawl_result(result)
        if detail_records:
            best = detail_records[0]
            fields = {
                "company": str(best.get("company") or ""),
                "title": str(best.get("title") or ""),
                "salary": str(best.get("salary") or ""),
                "location": str(best.get("location") or ""),
                "education": str(best.get("education") or ""),
                "experience": str(best.get("experience") or ""),
            }
    return {"ok": "1", "url": normalize_url(url), "text": text, **{key: value for key, value in fields.items() if value}}


def merge_record_with_detail(record: dict[str, str], detail: dict[str, str]) -> dict[str, str]:
    if not detail.get("ok"):
        return record
    merged = dict(record)
    detail_text = normalize_multiline_text(str(detail.get("text") or ""))
    for key in ["company", "title", "salary", "location", "education", "experience"]:
        value = normalize_text(str(detail.get(key) or ""))
        if value:
            merged[key] = value
    header_lines = [
        f"岗位：{merged.get('title', '')}" if merged.get("title") else "",
        f"公司：{merged.get('company', '')}" if merged.get("company") else "",
        f"薪资：{merged.get('salary', '')}" if merged.get("salary") else "",
        f"地点：{merged.get('location', '')}" if merged.get("location") else "",
        f"学历：{merged.get('education', '')}" if merged.get("education") else "",
        f"经验：{merged.get('experience', '')}" if merged.get("experience") else "",
        f"链接：{detail.get('url') or merged.get('url', '')}",
    ]
    original_text = normalize_multiline_text(str(record.get("text") or ""))
    merged["text"] = normalize_multiline_text("\n".join(line for line in header_lines if line) + "\n\n详情页原文：\n" + (detail_text or original_text))
    merged["url"] = str(detail.get("url") or merged.get("url") or "")
    merged["detail_enriched"] = "1"
    return merged


def enrich_records_with_detail_pages(records: list[dict[str, str]], limit: int = 60) -> list[dict[str, str]]:
    if not records:
        return []
    enriched: list[dict[str, str]] = []
    cache: dict[str, dict[str, str]] = {}
    fetched = 0
    for record in records:
        url = str(record.get("url") or "")
        source = str(record.get("source") or "")
        if fetched < limit and detail_url_is_usable(url, source):
            normalized = normalize_url(url)
            if normalized not in cache:
                cache[normalized] = fetch_detail_for_enrichment(normalized)
                fetched += 1
            enriched.append(merge_record_with_detail(record, cache[normalized]))
        else:
            enriched.append(record)
    return dedupe_jd_records(enriched)


def read_exported_jd_file(path: Path) -> dict[str, str]:
    suffix = path.suffix.lower()
    title = path.stem
    url = ""
    text = ""
    is_batch = False

    if suffix == ".json":
        data = read_json_file_best_effort(path)
        title = str(data.get("title") or title)
        url = str(data.get("url") or "")
        jobs = data.get("jobs") or []
        if isinstance(jobs, list) and jobs:
            is_batch = True
            chunks = []
            for index, job in enumerate(jobs, start=1):
                if not isinstance(job, dict):
                    continue
                job_title = str(job.get("title") or f"岗位{index}")
                job_url = str(job.get("url") or "")
                job_text = str(job.get("text") or "")
                chunks.append(normalize_text(f"岗位{index}：{job_title} 链接：{job_url} {job_text}"))
            text = "\n\n".join(chunks)
        else:
            text = str(data.get("text") or data.get("content") or data.get("body") or "")
    elif suffix in [".txt", ".md"]:
        text = read_text_file_best_effort(path)
    elif suffix in [".html", ".htm"]:
        text = html_to_text(read_text_file_best_effort(path))
    elif suffix == ".pdf":
        try:
            import pdfplumber

            with pdfplumber.open(str(path)) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        except Exception:
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
    elif suffix == ".docx":
        from docx import Document

        doc = Document(str(path))
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

    if not is_batch:
        text = normalize_text(text)
    return {
        "title": title,
        "url": url,
        "text": text,
        "job_count": len(jobs) if suffix == ".json" and isinstance(jobs, list) else 1,
        "path": str(path),
        "modified": datetime.fromtimestamp(path.stat().st_mtime).strftime("%Y-%m-%d %H:%M:%S"),
    }


def existing_jd_export_dirs() -> list[Path]:
    return [path for path in JD_EXPORT_DIRS if path.exists()]


def export_date_label(path: Path) -> str:
    for root in JD_EXPORT_DIRS:
        try:
            relative = path.relative_to(root)
        except ValueError:
            continue
        return path.parent.name if relative.parent != Path(".") else "未归档"
    return path.parent.name


def scan_browser_export_folder(limit: int = 80) -> tuple[str, pd.DataFrame]:
    export_dirs = existing_jd_export_dirs()
    if not export_dirs:
        return "", pd.DataFrame(columns=["文件", "日期", "标题", "链接", "岗位数", "字数", "修改时间"])

    supported = {".json", ".txt", ".md", ".html", ".htm", ".pdf", ".docx"}
    files = [
        path
        for export_dir in export_dirs
        for path in export_dir.rglob("*")
        if path.is_file() and path.suffix.lower() in supported
    ]
    files = sorted(files, key=lambda item: item.stat().st_mtime, reverse=True)[:limit]

    rows = []
    chunks = []
    for path in files:
        try:
            item = read_exported_jd_file(path)
            if len(item["text"]) < 30:
                continue
            rows.append(
                {
                    "文件": path.name,
                    "日期": export_date_label(path),
                    "标题": item["title"],
                    "链接": item["url"],
                    "岗位数": item.get("job_count", 1),
                    "字数": len(item["text"]),
                    "修改时间": item["modified"],
                }
            )
            source = item["url"] or item["path"]
            chunks.append(f"来源：{source}\n标题：{item['title']}\n{item['text']}")
        except Exception:
            continue
    return "\n\n".join(chunks), pd.DataFrame(rows)


def line_looks_like_job_title(line: str) -> bool:
    value = normalize_text(line)
    if not (3 <= len(value) <= 90):
        return False
    if re.search(r"https?://|www\.|@|登录|注册|筛选|排序|首页|公司主页|收藏|分享", value, flags=re.I):
        return False
    if extract_salary(value) != "未识别":
        return False
    title_words = [
        "实习",
        "工程师",
        "顾问",
        "专员",
        "分析师",
        "经理",
        "助理",
        "管培生",
        "运营",
        "产品",
        "数据",
        "市场",
        "销售",
        "财务",
        "法务",
        "合规",
        "供应链",
        "LCA",
        "ESG",
        "CBAM",
        "Job",
        "Position",
        "Consultant",
        "Analyst",
        "Intern",
    ]
    return any(text_contains(value, word) for word in title_words)


def split_listing_lines_into_blocks(text: str) -> list[str]:
    lines = [line.strip() for line in normalize_multiline_text(text).splitlines() if line.strip()]
    if len(lines) < 6:
        return []
    starts = []
    for index, line in enumerate(lines):
        window = "\n".join(lines[index:index + 10])
        if line_looks_like_job_title(line) and job_card_score(window) >= 7:
            starts.append(index)
    starts = sorted(dict.fromkeys(starts))
    if len(starts) <= 1:
        return []
    blocks = []
    for index, start in enumerate(starts):
        end = starts[index + 1] if index + 1 < len(starts) else min(len(lines), start + 24)
        block = "\n".join(lines[start:end]).strip()
        compact = normalize_text(block)
        if 30 <= len(compact) <= 2600 and job_card_score(block) >= 7:
            blocks.append(block)
    return blocks


def split_batch_jd_text(text: str) -> list[dict[str, str]]:
    text = text.strip()
    if not text:
        return []

    blocks = [block.strip() for block in re.split(r"\n\s*\n+", text) if len(block.strip()) >= 20]
    if len(blocks) > 1:
        scored_blocks = [block for block in blocks if job_card_score(block) >= 6]
        if len(scored_blocks) >= 2:
            blocks = scored_blocks
    if len(blocks) <= 1:
        starts = list(re.finditer(r"(?m)^(?:公司|企业|岗位|职位|招聘|Job|Position)[:：]", text))
        if len(starts) > 1:
            blocks = []
            for index, match in enumerate(starts):
                end = starts[index + 1].start() if index + 1 < len(starts) else len(text)
                chunk = text[match.start():end].strip()
                if len(chunk) >= 20:
                    blocks.append(chunk)
    if len(blocks) <= 1:
        inline_markers = list(re.finditer(r"(?=(?:岗位|职位|招聘职位|Job|Position)[:：])", text, flags=re.I))
        if len(inline_markers) > 1:
            inline_blocks = []
            for index, match in enumerate(inline_markers):
                end = inline_markers[index + 1].start() if index + 1 < len(inline_markers) else len(text)
                chunk = text[match.start():end].strip(" \n\t|；;")
                if len(chunk) >= 20 and job_card_score(chunk) >= 5:
                    inline_blocks.append(chunk)
            if len(inline_blocks) > 1:
                blocks = inline_blocks
    if len(blocks) <= 1:
        line_blocks = split_listing_lines_into_blocks(text)
        if len(line_blocks) > 1:
            blocks = line_blocks
    if not blocks:
        blocks = [text]

    return [
        {
            "source": "粘贴文本",
            "title": f"文本岗位{index}",
            "url": "",
            "text": block,
        }
        for index, block in enumerate(blocks, start=1)
    ]


def boss_salary_line(value: str) -> bool:
    text = normalize_text(value)
    return bool(re.search(r"\d+(?:\.\d+)?\s*[-~至到]\s*\d+(?:\.\d+)?\s*(?:K|k|元/天|元/日|元/月|万)", text))


def boss_education_line(value: str) -> bool:
    return normalize_text(value) in {"学历不限", "初中及以下", "中专/中技", "高中", "大专", "本科", "硕士", "博士"}


def boss_meta_line(value: str) -> bool:
    text = normalize_text(value)
    return bool(
        boss_education_line(text)
        or re.search(r"经验不限|无经验|\d+\s*[-~至到]\s*\d+\s*年|\d+\s*年以上|\d+\s*年以下|1年以下|在校/应届|应届|应届毕业生|\d+\s*天/周|\d+\s*个月", text)
    )


def boss_location_line(value: str) -> bool:
    text = normalize_text(value)
    if not (2 <= len(text) <= 40):
        return False
    if boss_salary_line(text) or is_probable_company_name(text):
        return False
    city_prefix = r"(?:上海|北京|深圳|广州|杭州|苏州|南京|武汉|成都|重庆|宁波|无锡|常州|嘉兴)"
    return bool(re.search(r"·", text) or re.match(rf"^{city_prefix}(?:$|·)", text) or (len(text) <= 20 and re.search(r"[区县市]$", text)))


def boss_title_line(value: str) -> bool:
    text = normalize_text(value)
    if not (2 <= len(text) <= 80):
        return False
    if ("：" in text or "，" in text) and len(text) > 18:
        return False
    if boss_salary_line(text) or boss_meta_line(text) or boss_location_line(text):
        return False
    bad = {"首页", "职位", "公司", "校园", "海归", "推荐", "行业", "地图", "搜索", "工作区域", "职位类型", "求职类型", "薪资待遇", "工作经验", "学历要求", "公司行业", "公司规模", "融资阶段", "清空", "消息", "简历"}
    if text in bad or "添加求职期望" in text:
        return False
    return any(word in text for word in ["ESG", "碳", "审计", "咨询", "分析师", "顾问", "专员", "助理", "管培生", "销售", "会计", "财务", "可持续", "低碳", "绿色", "市场", "项目", "经理"])


def boss_company_line(value: str) -> bool:
    text = normalize_text(value)
    if not (2 <= len(text) <= 80):
        return False
    if bad_generic_card_line(text):
        return False
    if re.match(r"^\d+\s*[.、]", text) or any(mark in text for mark in ["。", "；", ";"]):
        return False
    if boss_salary_line(text) or boss_meta_line(text):
        return False
    bad_exact = {
        "收藏",
        "立即沟通",
        "举报",
        "微信扫码分享",
        "职位描述",
        "任职要求",
        "去App",
        "前往App",
        "查看更多信息",
        "求职工具",
        "热门职位",
        "热门城市",
        "附近城市",
        "满意",
        "不满意",
        "一般",
        "提交",
        "在线",
        "实习",
        "校园",
        "校招",
        "全职",
        "正式工",
        "兼职",
        "去聊聊",
        "投递",
        "立即投递",
    }
    if text in bad_exact:
        return False
    if any(re.search(pattern, text) for pattern in GENERIC_STATUS_PATTERNS):
        return False
    if "：" in text and len(text) > 18:
        return False
    if any(word in text for word in ["岗位", "职位", "职责", "薪资", "经验", "学历", "招聘", "首页", "搜索", "扫码", "沟通"]):
        return False
    if is_probable_company_name(text):
        return True
    if boss_location_line(text):
        return False
    return True


def records_from_boss_listing_text(text: str, source: str = "公开链接", url: str = "", page_index: str = "") -> list[dict[str, str]]:
    lines = [line.strip() for line in normalize_multiline_text(text).splitlines() if line.strip()]
    records: list[dict[str, str]] = []
    index = 0
    while index < len(lines) - 4:
        title = lines[index]
        if not boss_title_line(title):
            index += 1
            continue
        cursor = index + 1
        salary = ""
        if cursor < len(lines) and boss_salary_line(lines[cursor]):
            salary = lines[cursor]
            cursor += 1
        elif cursor + 1 < len(lines) and boss_salary_line(lines[cursor + 1]):
            # Some crawlers insert one short status/label line between title and salary.
            salary = lines[cursor + 1]
            cursor += 2
        location_before_meta = ""
        if cursor < len(lines) and boss_location_line(lines[cursor]):
            location_before_meta = lines[cursor]
            cursor += 1
        meta_lines: list[str] = []
        meta_start = cursor
        while cursor < len(lines) and cursor < meta_start + 6 and boss_meta_line(lines[cursor]):
            meta_lines.append(lines[cursor])
            cursor += 1
        if cursor >= len(lines) or not boss_company_line(lines[cursor]):
            index += 1
            continue
        company = lines[cursor]
        location = lines[cursor + 1] if cursor + 1 < len(lines) and boss_location_line(lines[cursor + 1]) else location_before_meta
        if not salary and not (meta_lines and location):
            index += 1
            continue
        education = next((item for item in meta_lines if boss_education_line(item)), "")
        experience = " / ".join(item for item in meta_lines if not boss_education_line(item))
        block_lines = [
            f"岗位：{title}",
            f"公司：{company}",
            f"薪资：{salary}" if salary else "",
            f"经验：{experience}" if experience else "",
            f"学历：{education}" if education else "",
            f"地点：{location}" if location else "",
        ]
        records.append(
            {
                "source": source,
                "title": title,
                "url": url,
                "text": "\n".join(line for line in block_lines if line),
                "company": company,
                "salary": salary,
                "location": location,
                "education": education,
                "experience": experience,
                "page_index": page_index,
            }
        )
        index = cursor + 2
    return dedupe_jd_records(records)


def zhilian_location_line(value: str) -> bool:
    text = strip_listing_marks(value)
    return bool(re.fullmatch(r"[「【]?(?:全国|海外|[一-龥]{2,10})(?:[·\-][一-龥A-Za-z0-9]{1,18})?[」】]?", text))


def zhilian_scale_line(value: str) -> bool:
    text = strip_listing_marks(value)
    return bool(re.fullmatch(r"(?:\d+\s*-\s*\d+人|\d+人以上|\d+人以下|少于\d+人)", text))


def zhilian_nature_line(value: str) -> bool:
    text = strip_listing_marks(value)
    return text in {"民营", "国企", "外资", "合资", "上市公司", "事业单位", "其它", "股份制企业", "社会团体", "代表处"}


def clean_zhilian_location(value: str) -> str:
    return strip_listing_marks(value).strip("「」【】")


def records_from_zhilian_listing_text(text: str, source: str = "公开链接", url: str = "", page_index: str = "") -> list[dict[str, str]]:
    lines = split_listing_lines(text)
    records: list[dict[str, str]] = []
    index = 0
    while index < len(lines) - 8:
        title = clean_generic_title(lines[index])
        if not generic_title_line(title):
            index += 1
            continue
        if not zhilian_location_line(lines[index + 1]):
            index += 1
            continue
        salary = strip_listing_marks(lines[index + 2])
        if not generic_salary_line(salary):
            index += 1
            continue
        job_type = strip_listing_marks(lines[index + 3])
        if job_type not in GENERIC_JOB_TYPE_LINES:
            index += 1
            continue
        education = strip_listing_marks(lines[index + 4])
        experience = strip_listing_marks(lines[index + 5])
        if not boss_education_line(education) or not generic_meta_line(experience):
            index += 1
            continue

        window_end = min(len(lines), index + 32)
        scale_pos = -1
        for pos in range(index + 6, window_end):
            if zhilian_scale_line(lines[pos]) and (pos + 1 >= len(lines) or zhilian_nature_line(lines[pos + 1]) or strip_listing_marks(lines[pos + 1]) == "立即投递"):
                scale_pos = pos
                break
        if scale_pos < 0:
            index += 1
            continue

        company_candidates: list[tuple[int, int, str]] = []
        for pos in range(index + 6, scale_pos):
            candidate = strip_listing_marks(lines[pos])
            score = generic_company_score(candidate, scale_pos - pos)
            if pos == scale_pos - 2:
                score += 12
            if score >= 18:
                company_candidates.append((score, pos, candidate))
        if not company_candidates:
            index += 1
            continue
        _company_score, company_pos, company = max(company_candidates, key=lambda item: item[0])
        industry = strip_listing_marks(lines[company_pos + 1]) if company_pos + 1 < scale_pos else ""
        if generic_company_score(company) <= 0:
            index += 1
            continue

        location = clean_zhilian_location(lines[index + 1])
        block_lines = [
            f"岗位：{title}",
            f"公司：{company}",
            f"薪资：{salary}",
            f"地点：{location}",
            f"类型：{job_type}",
            f"经验：{experience}",
            f"学历：{education}",
            f"行业：{industry}" if industry else "",
        ]
        records.append(
            {
                "source": source,
                "title": title,
                "url": url,
                "text": "\n".join(line for line in block_lines if line),
                "company": company,
                "salary": salary,
                "location": location,
                "education": education,
                "experience": experience,
                "page_index": page_index,
            }
        )
        index = scale_pos + 2
    return dedupe_jd_records(records)


def liepin_bracket_location(lines: list[str], start: int, end: int) -> tuple[str, int]:
    for pos in range(start, min(end, len(lines))):
        text = strip_listing_marks(lines[pos])
        if text in {"【", "】"}:
            continue
        if pos > start and strip_listing_marks(lines[pos - 1]) == "【":
            return text, pos
        if generic_location_line(text):
            return text, pos
    return "", -1


def records_from_liepin_listing_text(text: str, source: str = "公开链接", url: str = "", page_index: str = "") -> list[dict[str, str]]:
    lines = split_listing_lines(text)
    records: list[dict[str, str]] = []
    index = 0
    while index < len(lines) - 6:
        title = clean_generic_title(lines[index])
        if not generic_title_line(title):
            index += 1
            continue
        window_end = min(len(lines), index + 16)
        salary_pos = next((pos for pos in range(index + 1, window_end) if generic_salary_line(lines[pos])), -1)
        if salary_pos < 0:
            index += 1
            continue
        salary = strip_listing_marks(lines[salary_pos])
        location, _location_pos = liepin_bracket_location(lines, index + 1, salary_pos)
        meta_lines: list[str] = []
        company = ""
        company_pos = -1
        for pos in range(salary_pos + 1, window_end):
            line = strip_listing_marks(lines[pos])
            if not line or line in {"实习", "全职", "兼职", "提供转正"} or generic_meta_line(line) or boss_education_line(line):
                if line:
                    meta_lines.append(line)
                continue
            if "·" in line and re.search(r"(招聘|HR|专员|经理|顾问|在线)", line):
                break
            if generic_company_score(line) >= 20:
                company = line
                company_pos = pos
                break
        if not company:
            index += 1
            continue
        education = next((item for item in meta_lines if boss_education_line(item)), "")
        experience = " / ".join(item for item in meta_lines if item != education and item not in GENERIC_JOB_TYPE_LINES)
        block_lines = [
            f"岗位：{title}",
            f"公司：{company}",
            f"薪资：{salary}",
            f"地点：{location}" if location else "",
            f"经验：{experience}" if experience else "",
            f"学历：{education}" if education else "",
        ]
        records.append(
            {
                "source": source,
                "title": title,
                "url": url,
                "text": "\n".join(line for line in block_lines if line),
                "company": company,
                "salary": salary,
                "location": location,
                "education": education,
                "experience": experience,
                "page_index": page_index,
            }
        )
        index = max(company_pos + 1, index + 1)
    return dedupe_jd_records(records)


def job51_company_profile_line(value: str) -> bool:
    text = strip_listing_marks(value)
    if not (4 <= len(text) <= 100):
        return False
    return bool(
        re.search(r"(?:\d+\s*-\s*\d+人|\d+人以上|\d+人以下|少于\d+人)", text)
        or re.search(r"(?:已上市|上市公司|外资|民营|国企|合资|事业单位|股份制|创业公司)", text)
        or re.search(r"(?:专业服务|咨询|人力资源|财会|检测|认证|互联网|电子|半导体|制造|新能源|环保|金融|贸易|集团公司)", text)
    )


def records_from_51job_listing_text(text: str, source: str = "公开链接", url: str = "", page_index: str = "") -> list[dict[str, str]]:
    lines = split_listing_lines(text)
    if len(lines) < 6:
        return []
    title = clean_generic_title(lines[0])
    if not generic_title_line(title):
        return []
    salary = strip_listing_marks(lines[1]) if len(lines) > 1 and generic_salary_line(lines[1]) else ""
    location = strip_listing_marks(lines[2]) if len(lines) > 2 and generic_location_line(lines[2]) else ""
    company = ""
    profile = ""
    stop_words = {"去聊聊", "投递", "申请职位", "立即投递"}
    for pos in range(len(lines) - 1, max(2, len(lines) - 28), -1):
        line = strip_listing_marks(lines[pos])
        if line in stop_words:
            continue
        if job51_company_profile_line(line) and pos > 0:
            candidate = strip_listing_marks(lines[pos - 1])
            if generic_company_score(candidate) >= 20:
                company = candidate
                profile = line
                break
    if not company:
        for pos in range(3, min(len(lines), 28)):
            candidate = strip_listing_marks(lines[pos])
            if generic_company_score(candidate) >= 45:
                company = candidate
                profile = strip_listing_marks(lines[pos + 1]) if pos + 1 < len(lines) else ""
                break
    if not company:
        return []
    block_lines = [
        f"岗位：{title}",
        f"公司：{company}",
        f"薪资：{salary}" if salary else "",
        f"地点：{location}" if location else "",
        f"公司信息：{profile}" if profile else "",
    ]
    return [
        {
            "source": source,
            "title": title,
            "url": url,
            "text": "\n".join(line for line in block_lines if line),
            "company": company,
            "salary": salary,
            "location": location,
            "education": "",
            "experience": "",
            "page_index": page_index,
        }
    ]


def shixiseng_meta_line(value: str) -> bool:
    text = strip_listing_marks(value)
    return "|" in text and any(city in text for city in ["上海", "北京", "广州", "深圳", "杭州", "苏州", "南京"])


def shixiseng_company_line(value: str) -> bool:
    text = strip_listing_marks(value)
    if not (2 <= len(text) <= 50):
        return False
    if "/" in text and re.search(r"(互联网|咨询|服务|建筑|物业|金融|经济|投资|财会|游戏|软件|房产|家居)", text):
        return False
    if bad_generic_card_line(text) or generic_salary_line(text) or shixiseng_meta_line(text):
        return False
    if re.search(r"(实习|薪资|周末|导师|地铁|津贴|免费|双休|不加班|可转正|氛围)", text):
        return False
    return generic_company_score(text) >= 8 or bool(re.search(r"[A-Za-z]{2,}", text))


def records_from_shixiseng_listing_text(text: str, source: str = "公开链接", url: str = "", page_index: str = "") -> list[dict[str, str]]:
    lines = [normalize_text(line) for line in text.splitlines() if normalize_text(line)]
    records: list[dict[str, str]] = []
    for index in range(1, len(lines) - 1):
        meta_line = lines[index]
        company_line = lines[index + 1]
        if not shixiseng_meta_line(meta_line) or not shixiseng_company_line(company_line):
            continue
        title_line = lines[index - 1]
        meta_parts = [strip_listing_marks(part) for part in meta_line.split("|")]
        location = meta_parts[0] if meta_parts else ""
        title = clean_generic_title(title_line)
        salary = extract_salary(title_line)
        if salary != "未识别":
            title = normalize_text(title.replace(salary, "")).strip(" -_｜|，,")
        else:
            salary = "薪资面议" if "薪资面议" in title_line else ""
        company = strip_listing_marks(company_line)
        industry = lines[index + 2] if index + 2 < len(lines) else ""
        block_lines = [
            f"岗位：{title}",
            f"公司：{company}",
            f"薪资：{salary}" if salary else "",
            f"地点：{location}" if location else "",
            f"要求：{' / '.join(meta_parts[1:])}" if len(meta_parts) > 1 else "",
            f"公司信息：{industry}" if industry else "",
        ]
        records.append(
            {
                "source": source,
                "title": title,
                "url": url,
                "text": "\n".join(line for line in block_lines if line),
                "company": company,
                "salary": salary,
                "location": location,
                "education": "",
                "experience": " / ".join(meta_parts[1:]) if len(meta_parts) > 1 else "",
                "page_index": page_index,
            }
        )
    return dedupe_jd_records(records)


def site_specific_listing_records(text: str, source: str, url: str = "", page_index: str = "", site_hint: str = "") -> list[dict[str, str]]:
    hint = normalize_text("\n".join([source, url, site_hint]))
    parsers: list[Any] = []
    if "BOSS" in hint or "zhipin" in hint:
        parsers.append(records_from_boss_listing_text)
    if "智联" in hint or "zhaopin" in hint:
        parsers.append(records_from_zhilian_listing_text)
    if "实习僧" in hint or "shixiseng" in hint:
        parsers.append(records_from_shixiseng_listing_text)
    if "猎聘" in hint or "liepin" in hint:
        parsers.append(records_from_liepin_listing_text)
    if "前程无忧" in hint or "51job" in hint:
        parsers.append(records_from_51job_listing_text)
    if not parsers:
        parsers = [
            records_from_boss_listing_text,
            records_from_zhilian_listing_text,
            records_from_shixiseng_listing_text,
            records_from_liepin_listing_text,
            records_from_51job_listing_text,
        ]
    for parser in parsers:
        records = parser(text, source=source, url=url, page_index=page_index)
        if records:
            return records
    return []


GENERIC_JOB_TYPE_LINES = {"实习", "校园", "校招", "全职", "正式工", "兼职", "兼职/临时", "社招"}
GENERIC_BAD_CARD_LINES = {
    "收藏",
    "立即沟通",
    "举报",
    "微信扫码分享",
    "职位描述",
    "任职要求",
    "去App",
    "前往App",
    "查看更多信息",
    "求职工具",
    "热门职位",
    "热门城市",
    "附近城市",
    "满意",
    "不满意",
    "一般",
    "提交",
    "在线",
    "去聊聊",
    "投递",
    "立即投递",
    "申请职位",
    "查看详情",
    "与BOSS随时沟通",
    "领导NICE",
    "领导nice",
    "新能源",
    "战略咨询",
    "合同签订",
    "cpa",
    "ESG",
    "运营管理",
    "生物医药",
    "补充医疗保险",
    "审核",
    "统招本科",
    "数据开发",
    "电子信息工程",
    "物流管理",
    "市场营销",
    "职位推荐",
    "职位搜索",
    "首页",
    "搜索",
}
GENERIC_STATUS_PATTERNS = [
    r"\d+\s*(?:天|分钟|小时)内?处理简历",
    r"\d+\s*分钟前处理简历",
    r"回复率高|简历处理快|喜欢聊天|活跃|在线|远程办公",
]
GENERIC_CITY_PREFIX = (
    "上海|北京|深圳|广州|苏州|杭州|南京|宁波|无锡|常州|嘉兴|成都|武汉|重庆|"
    "厦门|郑州|合肥|天津|芜湖|东莞|泉州|鄂尔多斯|昆明|宁德|西安|长沙|福州|"
    "青岛|济南|大连|沈阳|长春|哈尔滨|石家庄|佛山|南昌|贵阳"
)


def strip_listing_marks(value: str) -> str:
    return normalize_text(value).strip("「」【】[]()（）| ")


def split_listing_lines(text: str) -> list[str]:
    normalized = normalize_multiline_text(text)
    rough_lines: list[str] = []
    for line in normalized.splitlines():
        parts = re.split(r"\s+\|\s+|\s{2,}", line)
        rough_lines.extend(parts if len(parts) > 1 else [line])
    lines: list[str] = []
    for line in rough_lines:
        clean = normalize_text(line).strip()
        if not clean or clean in {"|", "【", "】", "「", "」"}:
            continue
        lines.append(clean)
    return lines


def generic_salary_line(value: str) -> bool:
    text = strip_listing_marks(value)
    return bool(
        extract_salary(text) != "未识别"
        or text in {"面议", "薪资面议", "薪酬面议"}
        or re.fullmatch(r"\d+(?:\.\d+)?\s*[-~至到]\s*\d+(?:\.\d+)?\s*元(?:[·xX]\s*\d+\s*薪)?", text)
        or ("/天" in text and len(text) <= 80)
    )


def generic_location_line(value: str) -> bool:
    text = strip_listing_marks(value)
    if not (2 <= len(text) <= 36):
        return False
    if generic_salary_line(text) or generic_meta_line(text):
        return False
    if boss_location_line(text):
        return True
    return bool(re.match(rf"^(?:{GENERIC_CITY_PREFIX})(?:$|[·\-\s].{{1,18}}$)", text))


def generic_meta_line(value: str) -> bool:
    text = strip_listing_marks(value)
    if text in GENERIC_JOB_TYPE_LINES:
        return True
    if boss_meta_line(text):
        return True
    return bool(
        re.fullmatch(r"(?:面议|提供转正|不提供转正|.{1,6}个月|.{1,6}天/周|在校/应届|应届毕业生|经验不限|无经验|1年以下|学历不限)", text)
        or re.fullmatch(r"\d+\s*年以内", text)
        or re.fullmatch(r"\d+\s*[-~至到]\s*\d+\s*年(?:以上)?", text)
        or re.fullmatch(r"\d+\s*年以上", text)
        or re.fullmatch(r"\d+\s*年以下", text)
    )


def bad_generic_card_line(value: str) -> bool:
    text = strip_listing_marks(value)
    if not text:
        return True
    skill_tags = {
        "Python", "SQL", "Sql", "SPSS", "R语言", "Power BI", "SmartBI", "FineBI", "BI工具",
        "数据仓库", "数据建模", "数据治理", "数据挖掘", "报表开发", "水质", "水质监测",
        "废气", "废水", "海洋", "检测", "分析", "环境工程", "工程设计", "工艺设计",
        "设备运维", "云计算", "数字基建", "商业数据分析", "游戏数据分析", "电商数据分析",
        "行业数据分析", "数据分析", "市场数据分析", "用户数据分析", "数据看板",
        "FPGA开发", "射频开发", "PCB设计", "示波器", "HFSS", "ADS", "reach", "REACH",
    }
    if text in skill_tags:
        return True
    if text in GENERIC_BAD_CARD_LINES:
        return True
    if text in {"上市公司", "大型集团", "学生可投", "ESG报告", "认证", "运营", "电子信息", "餐饮补贴", "通信专业监理", "ccaa审核员", "cdp"}:
        return True
    if any(re.search(pattern, text) for pattern in GENERIC_STATUS_PATTERNS):
        return True
    if any(word in text for word in ["筛选", "不限", "清空", "下载APP", "求职助手", "搜索", "登录", "注册"]):
        return True
    if re.search(r"(周边|双休|导师|津贴|免费|班车|氛围|福利|五险|奖金|年假|培训|晋升)", text) and len(text) <= 40:
        return True
    return False


def generic_title_line(value: str) -> bool:
    text = strip_listing_marks(value)
    if not (2 <= len(text) <= 120):
        return False
    if ("：" in text or "，" in text) and len(text) > 20:
        return False
    if bad_generic_card_line(text) or generic_salary_line(text) or generic_location_line(text) or generic_meta_line(text):
        return False
    if re.search(r"^[\d,.\-~至到]+$|https?://|www\.", text, flags=re.I):
        return False
    title_keywords = [
        "实习", "校招", "管培", "经理", "顾问", "专员", "助理", "工程师", "分析", "运营", "产品",
        "销售", "审核", "认证", "研究", "项目", "数据", "咨询", "ESG", "LCA", "碳", "会计", "财务",
        "市场", "管理", "培训生", "岗", "主管", "总监", "Specialist", "Analyst", "Engineer", "Consultant",
    ]
    return any(keyword.lower() in text.lower() for keyword in title_keywords)


def generic_company_score(value: str, distance: int = 0) -> int:
    text = strip_listing_marks(value)
    if not (2 <= len(text) <= 90):
        return 0
    strong_company_shape = bool(is_probable_company_name(text) or re.search(COMPANY_SUFFIX_PATTERN, text))
    if bad_generic_card_line(text) or generic_salary_line(text) or generic_location_line(text) or generic_meta_line(text):
        return 0
    if re.match(r"^\d+\s*[.、]", text) or any(mark in text for mark in ["。", "；", ";"]):
        return 0
    if any(word in text for word in ["岗位", "职位", "职责", "薪资", "学历", "经验", "招聘", "福利", "投递"]):
        return 0
    if generic_meta_line(text):
        return 0
    if ("：" in text or "，" in text or "。" in text) and len(text) > 18:
        return 0
    if not strong_company_shape and re.search(r"(五险|奖金|年假|培训|晋升|双休|不打卡|班车|导师|津贴|Python|SQL|R语言|数据分析|市场开拓)", text, flags=re.I):
        return 0
    if not strong_company_shape and "/" in text and re.search(r"(互联网|咨询|服务|建筑|物业|金融|经济|投资|财会|游戏|软件|房产|家居|检测|认证|技术|商贸)", text):
        return 0
    if not strong_company_shape and re.search(r"(人$|\d+-\d+人|人以上|已上市|融资|民营|国企|外资|事业单位|互联网|咨询服务|技术服务|综合商贸|专业服务|检测，认证)", text):
        return 0
    score = max(0, 10 - distance)
    if is_probable_company_name(text):
        score += 50
    if re.search(COMPANY_SUFFIX_PATTERN, text):
        score += 35
    if re.search(r"(集团|公司|科技|咨询|认证|检测|事务所|银行|学校|医院|研究院|中心|股份|有限|PDD|SHEIN|KPMG|Deloitte|PwC|EY)", text, re.I):
        score += 28
    if re.search(r"[A-Za-z]", text) and len(text) <= 35:
        score += 12
    if "·" in text or "/" in text:
        score -= 20
    if len(text) <= 32:
        score += 6
    return max(score, 0)


def clean_generic_title(value: str, salary: str = "") -> str:
    title = strip_listing_marks(value)
    salary_value = strip_listing_marks(salary)
    if salary_value:
        title = normalize_text(title.replace(salary_value, ""))
    extracted_salary = extract_salary(title)
    if extracted_salary != "未识别":
        title = normalize_text(title.replace(extracted_salary, ""))
    return title.strip(" -_｜|，,")


def records_from_generic_listing_text(text: str, source: str = "公开链接", url: str = "", page_index: str = "") -> list[dict[str, str]]:
    lines = split_listing_lines(text)
    records: list[dict[str, str]] = []
    index = 0
    while index < len(lines) - 3:
        raw_title = lines[index]
        if not generic_title_line(raw_title):
            index += 1
            continue

        window_end = min(len(lines), index + 26)
        salary = ""
        salary_pos = -1
        location = ""
        location_pos = -1
        meta_lines: list[tuple[int, str]] = []
        for pos in range(index, window_end):
            line = lines[pos]
            if not salary and generic_salary_line(line):
                salary = extract_salary(line)
                if salary == "未识别":
                    salary = strip_listing_marks(line)
                salary_pos = pos
            if not location and generic_location_line(line):
                location = strip_listing_marks(line)
                location_pos = pos
            if pos > index and generic_meta_line(line):
                meta_lines.append((pos, strip_listing_marks(line)))

        if salary_pos < 0 and location_pos < 0 and not meta_lines:
            index += 1
            continue

        anchor_positions = [pos for pos in [salary_pos, location_pos] if pos >= 0] + [pos for pos, _ in meta_lines]
        preferred_company_start = max(anchor_positions) + 1 if anchor_positions else index + 1
        company_candidates: list[tuple[int, int, str]] = []
        for pos in range(max(index + 1, preferred_company_start), window_end):
            score = generic_company_score(lines[pos], pos - preferred_company_start)
            if score:
                company_candidates.append((score, pos, strip_listing_marks(lines[pos])))
        if not company_candidates and preferred_company_start > index + 2:
            for pos in range(index + 1, preferred_company_start):
                score = generic_company_score(lines[pos], pos - index)
                if score:
                    company_candidates.append((score - 12, pos, strip_listing_marks(lines[pos])))
        if not company_candidates:
            index += 1
            continue

        if meta_lines:
            nearby = [item for item in company_candidates if item[1] <= preferred_company_start + 2 and item[0] >= 8]
            company_score, company_pos, company = nearby[0] if nearby else max(company_candidates, key=lambda item: item[0])
        else:
            company_score, company_pos, company = max(company_candidates, key=lambda item: item[0])
        if company_score < (8 if meta_lines else 25):
            index += 1
            continue
        title = clean_generic_title(raw_title, salary)
        if not title or generic_location_line(title) or generic_meta_line(title):
            index += 1
            continue
        education = next((item for _pos, item in meta_lines if boss_education_line(item) or item in {"学历不限", "本科", "硕士", "博士", "大专"}), "")
        experience = " / ".join(
            item
            for _pos, item in meta_lines
            if item != education and item not in GENERIC_JOB_TYPE_LINES
        )
        job_type = next((item for _pos, item in meta_lines if item in GENERIC_JOB_TYPE_LINES), "")
        block_lines = [
            f"岗位：{title}",
            f"公司：{company}",
            f"薪资：{salary}" if salary else "",
            f"地点：{location}" if location else "",
            f"类型：{job_type}" if job_type else "",
            f"经验：{experience}" if experience else "",
            f"学历：{education}" if education else "",
        ]
        records.append(
            {
                "source": source,
                "title": title,
                "url": url,
                "text": "\n".join(line for line in block_lines if line),
                "company": company,
                "salary": salary,
                "location": location,
                "education": education,
                "experience": experience,
                "page_index": page_index,
            }
        )
        index = max(company_pos + 1, index + 1)
    return dedupe_jd_records(records)


def valid_plugin_job_title(title: str) -> bool:
    clean = normalize_text(title)
    if re.search(r"(?:能力|职责|要求|工作地点|任职资格|职位描述|项目经验|开发经验|工作经验|任职经历)$", clean):
        return False
    return bool(clean and generic_title_line(clean) and clean not in {"清空", "推荐", "搜索", "职位", "岗位"})


def invalid_plugin_company_line(company: str) -> bool:
    clean = normalize_text(company)
    if not clean:
        return True
    if len(clean) > 42:
        return True
    if generic_meta_line(clean) or extract_experience(clean) != "未识别":
        return True
    if re.match(r"^\d+[/、).．]\s*(?:负责|参与|协助|主导|完成|建立|制定|维护|对接|支持)", clean):
        return True
    action_pattern = r"(?:负责|完成|推动|建立|建⽴|对接|掌握|提供|寻求|开拓|维护|匹配|销售任务|客户|业务需求)"
    if re.search(r"[:：]", clean) and re.search(action_pattern, clean):
        return True
    if re.search(r"(?:岗位职责|职位描述|任职要求|工作内容|项目经验|开发经验|工作经验|任职经历|市场洞察|渠道建设|商机开拓|方案匹配)", clean):
        return True
    return False


def infer_plugin_job_title(text: str, salary: str, company: str) -> str:
    lines = [line.strip() for line in normalize_multiline_text(text).splitlines() if line.strip()]
    salary_clean = normalize_text(salary)
    company_clean = normalize_text(company)
    if salary_clean:
        for pos, line in enumerate(lines):
            if normalize_text(line) != salary_clean:
                continue
            if company_clean and company_clean not in [normalize_text(item) for item in lines[pos + 1:pos + 7]]:
                continue
            for title_pos in range(pos - 1, max(-1, pos - 5), -1):
                candidate = normalize_text(lines[title_pos])
                if valid_plugin_job_title(candidate):
                    return candidate
    if company_clean:
        for pos, line in enumerate(lines):
            if normalize_text(line) != company_clean:
                continue
            for title_pos in range(pos - 1, max(-1, pos - 8), -1):
                candidate = normalize_text(lines[title_pos])
                if valid_plugin_job_title(candidate):
                    return candidate
    return ""


def valid_plugin_job_url(value: Any) -> str:
    raw = str(value or "")
    canonical = canonical_job_url(raw)
    if not canonical or not detail_url_is_usable(canonical, ""):
        return ""
    parsed = urlparse(canonical)
    host = parsed.netloc.lower()
    path = parsed.path.rstrip("/").lower()
    if host.endswith("zhipin.com") and "/job_detail/" not in path:
        return ""
    return canonical


def plugin_job_detail_url(job: dict[str, Any], base_url: str = "") -> str:
    for key in ["detailUrl", "detail_url", "jobUrl", "job_url", "href", "link", "url"]:
        url = valid_plugin_job_url(job.get(key))
        if url:
            return url
    for url in extract_urls(str(job.get("detailText") or "") + "\n" + str(job.get("text") or "")):
        valid_url = valid_plugin_job_url(url)
        if valid_url:
            return valid_url
    return valid_plugin_job_url(base_url)


def ordered_plugin_detail_urls(jobs: list[Any], base_url: str = "") -> list[str]:
    urls: list[str] = []
    seen: set[str] = set()
    for job in jobs:
        if not isinstance(job, dict):
            continue
        url = plugin_job_detail_url(job, base_url)
        if url and url not in seen:
            urls.append(url)
            seen.add(url)
    return urls


def attach_detail_urls_by_order(records: list[dict[str, str]], urls: list[str]) -> list[dict[str, str]]:
    if not records or not urls:
        return records
    output = [dict(record) for record in records]
    used: set[str] = {preferred_jd_record_url(record) for record in output if preferred_jd_record_url(record)}
    cursor = 0
    for record in output:
        if preferred_jd_record_url(record):
            continue
        while cursor < len(urls) and urls[cursor] in used:
            cursor += 1
        if cursor >= len(urls):
            break
        record["url"] = urls[cursor]
        text = normalize_multiline_text(str(record.get("text") or ""))
        if urls[cursor] not in text:
            record["text"] = normalize_multiline_text(text + f"\n链接：{urls[cursor]}")
        used.add(urls[cursor])
        cursor += 1
    return output


def record_from_plugin_job(job: dict[str, Any], source: str, base_title: str, base_url: str, index: int) -> dict[str, str] | None:
    text = normalize_multiline_text(str(job.get("text") or ""))
    company = normalize_text(str(job.get("company") or ""))
    salary = normalize_text(str(job.get("salary") or ""))
    location = normalize_text(str(job.get("location") or ""))
    education = normalize_text(str(job.get("education") or ""))
    experience = normalize_text(str(job.get("experience") or ""))
    title = normalize_text(str(job.get("title") or ""))
    if not valid_plugin_job_title(title):
        title = infer_plugin_job_title(text, salary, company)

    if invalid_plugin_company_line(company) or generic_location_line(company) or boss_location_line(company) or bad_generic_card_line(company):
        return None
    if re.search(r"(?:先生|女士|老师|HR|hr|在线)$", company):
        return None
    if not title or not company:
        return None
    if not salary:
        return None

    block_lines = [
        f"岗位：{title}",
        f"公司：{company}" if company else "",
        f"薪资：{salary}" if salary else "",
        f"地点：{location}" if location else "",
        f"经验：{experience}" if experience else "",
        f"学历：{education}" if education else "",
    ]
    record_url = plugin_job_detail_url(job, base_url)
    return {
        "source": source,
        "title": title or f"{base_title}-{index}",
        "url": record_url,
        "text": "\n".join(line for line in block_lines if line),
        "company": company,
        "salary": salary,
        "location": location,
        "education": education,
        "experience": experience,
        "page_index": str(job.get("pageIndex") or job.get("page_index") or ""),
    }


def records_from_exported_jd_file(path: Path, enrich_detail_pages: bool = False, detail_limit: int = 60) -> list[dict[str, str]]:
    suffix = path.suffix.lower()
    records = []
    if suffix == ".json":
        data = read_json_file_best_effort(path)
        base_title = str(data.get("title") or path.stem)
        base_url = str(data.get("url") or "")
        jobs = data.get("jobs") or []
        if isinstance(jobs, list) and jobs:
            detail_urls = ordered_plugin_detail_urls(jobs, base_url)
            parsed_records: list[dict[str, str]] = []
            fallback_records: list[dict[str, str]] = []
            for index, job in enumerate(jobs, start=1):
                if not isinstance(job, dict):
                    continue
                text = normalize_multiline_text(str(job.get("text") or ""))
                if len(text) < 20:
                    continue
                record_page_index = str(job.get("pageIndex") or job.get("page_index") or "")
                has_structured_fields = any(normalize_text(str(job.get(key) or "")) for key in ["company", "salary", "location", "education", "experience", "title"])
                plugin_record = record_from_plugin_job(job, path.name, base_title, base_url, index)
                if has_structured_fields and plugin_record:
                    parsed_records.append(plugin_record)
                    continue
                parsed_listing_records = site_specific_listing_records(
                    text,
                    source=path.name,
                    url="",
                    page_index=record_page_index,
                    site_hint=base_title,
                )
                if not parsed_listing_records:
                    parsed_listing_records = records_from_generic_listing_text(
                        text,
                        source=path.name,
                        url="",
                        page_index=record_page_index,
                    )
                if len(parsed_listing_records) >= 2:
                    for parsed_record in parsed_listing_records:
                        parsed_record["url"] = ""
                    parsed_records.extend(parsed_listing_records)
                    if plugin_record:
                        parsed_records.append(plugin_record)
                    continue
                if plugin_record:
                    parsed_records.append(plugin_record)
                    continue
                if parsed_listing_records:
                    parsed_records.extend(parsed_listing_records)
                    continue
                fallback_records.append(
                    {
                        "source": path.name,
                        "title": str(job.get("title") or f"{base_title}-{index}"),
                        "url": plugin_job_detail_url(job, base_url),
                        "text": text,
                        "company": str(job.get("company") or ""),
                        "salary": str(job.get("salary") or ""),
                        "location": str(job.get("location") or ""),
                        "education": str(job.get("education") or ""),
                        "experience": str(job.get("experience") or ""),
                        "page_index": str(job.get("pageIndex") or job.get("page_index") or ""),
                    }
                )
            records.extend(attach_detail_urls_by_order(parsed_records + fallback_records, detail_urls))
        else:
            text = normalize_text(str(data.get("text") or data.get("content") or data.get("body") or ""))
            if len(text) >= 20:
                records.append({"source": path.name, "title": base_title, "url": base_url, "text": text})
        records = dedupe_jd_records(records)
        if isinstance(jobs, list) and jobs:
            records = attach_detail_urls_by_order(records, ordered_plugin_detail_urls(jobs, base_url))
            records = [record for record in records if preferred_jd_record_url(record)]
            records = dedupe_jd_records(records)
        return enrich_records_with_detail_pages(records, detail_limit) if enrich_detail_pages else records

    item = read_exported_jd_file(path)
    if len(item["text"]) >= 20:
        records.append({"source": path.name, "title": item["title"], "url": item["url"], "text": item["text"]})
    records = dedupe_jd_records(records)
    return enrich_records_with_detail_pages(records, detail_limit) if enrich_detail_pages else records


@st.cache_data(show_spinner=False, ttl=5)
def available_export_dates() -> list[str]:
    export_dirs = existing_jd_export_dirs()
    if not export_dirs:
        return []
    supported = {".json", ".txt", ".md", ".html", ".htm", ".pdf", ".docx"}
    dates = set()
    for export_dir in export_dirs:
        for path in export_dir.rglob("*"):
            if path.is_file() and path.suffix.lower() in supported:
                dates.add(export_date_label(path))
    return sorted(dates, reverse=True)


def scan_exported_jd_records(
    limit: int = 120,
    date_filter: str | None = None,
    enrich_detail_pages: bool = False,
    detail_limit: int = 60,
) -> tuple[list[dict[str, str]], pd.DataFrame]:
    export_dirs = existing_jd_export_dirs()
    if not export_dirs:
        return [], pd.DataFrame(columns=["文件", "日期", "岗位数", "修改时间"])
    supported = {".json", ".txt", ".md", ".html", ".htm", ".pdf", ".docx"}
    files = [
        path
        for export_dir in export_dirs
        for path in export_dir.rglob("*")
        if path.is_file() and path.suffix.lower() in supported
    ]
    if date_filter and date_filter != "全部":
        files = [
            path
            for path in files
            if export_date_label(path) == date_filter
        ]
    files = sorted(files, key=lambda item: item.stat().st_mtime, reverse=True)[:limit]

    records = []
    rows = []
    remaining_detail_limit = max(0, detail_limit)
    for path in files:
        try:
            file_records = records_from_exported_jd_file(
                path,
                enrich_detail_pages=enrich_detail_pages and remaining_detail_limit > 0,
                detail_limit=remaining_detail_limit,
            )
            if enrich_detail_pages:
                remaining_detail_limit = max(0, remaining_detail_limit - sum(1 for record in file_records if record.get("detail_enriched")))
        except Exception:
            continue
        if not file_records:
            continue
        file_id = str(path)
        file_records = [
            {
                **record,
                "_export_file_id": file_id,
                "_export_file_name": path.name,
            }
            for record in file_records
        ]
        records.extend(file_records)
        rows.append(
            {
                "信息集ID": file_id,
                "文件": path.name,
                "日期": export_date_label(path),
                "岗位数": len(file_records),
                "修改时间": datetime.fromtimestamp(path.stat().st_mtime).strftime("%Y-%m-%d %H:%M:%S"),
            }
        )
    records = dedupe_jd_records(records)
    return records, pd.DataFrame(rows)


TABLE_FIELD_ALIASES = {
    "company": [
        "公司",
        "公司名",
        "公司名称",
        "企业",
        "企业名",
        "企业名称",
        "雇主",
        "单位",
        "招聘单位",
        "用人单位",
        "company",
        "company name",
        "employer",
        "organization",
        "organisation",
    ],
    "title": [
        "岗位",
        "岗位名",
        "岗位名称",
        "职位",
        "职位名",
        "职位名称",
        "招聘职位",
        "职位标题",
        "岗位标题",
        "job",
        "job title",
        "position",
        "position name",
        "title",
        "role",
    ],
    "salary": ["薪资", "薪酬", "工资", "月薪", "日薪", "待遇", "salary", "pay", "compensation"],
    "location": ["地点", "城市", "工作地点", "工作城市", "办公地点", "地区", "location", "city", "workplace"],
    "education": ["学历", "学历要求", "教育", "教育要求", "education", "degree"],
    "experience": ["经验", "经验要求", "工作经验", "年限", "experience"],
    "url": ["链接", "职位链接", "岗位链接", "网址", "url", "link", "jd链接"],
}

TABLE_FIELD_NEGATIVE_HINTS = {
    "company": ["规模", "性质", "行业", "介绍", "详情", "福利", "地址", "logo"],
    "title": ["分类", "类别", "详情", "描述", "职责", "要求", "公司", "链接"],
    "salary": ["判断", "备注"],
    "location": ["偏好", "备注"],
}


def normalize_table_label(label: Any) -> str:
    value = normalize_text(str(label or "")).lower()
    return re.sub(r"[\s_\-./\\()（）【】\[\]{}:：,，;；|]+", "", value)


def clean_table_cell(value: Any) -> str:
    if value is None:
        return ""
    try:
        if pd.isna(value):
            return ""
    except Exception:
        pass
    text = normalize_text(str(value))
    if text.lower() in {"nan", "none", "null", "nat", "未识别"}:
        return ""
    return text


def table_column_score(column: Any, aliases: list[str], negative_hints: list[str] | None = None) -> int:
    label = normalize_table_label(column)
    if not label or label.startswith("unnamed"):
        return 0
    negative_hints = negative_hints or []
    if any(normalize_table_label(hint) in label for hint in negative_hints):
        return 0

    best = 0
    for alias in aliases:
        alias_label = normalize_table_label(alias)
        if not alias_label:
            continue
        if label == alias_label:
            best = max(best, 100 + len(alias_label))
        elif label.endswith(alias_label):
            best = max(best, 75 + len(alias_label))
        elif alias_label in label:
            best = max(best, 55 + len(alias_label))
    return best


def pick_table_column(df: pd.DataFrame, field: str) -> Any | None:
    aliases = TABLE_FIELD_ALIASES.get(field, [])
    negative_hints = TABLE_FIELD_NEGATIVE_HINTS.get(field, [])
    scored = [
        (table_column_score(column, aliases, negative_hints), column)
        for column in df.columns
    ]
    scored = [item for item in scored if item[0] > 0]
    if not scored:
        return None
    return max(scored, key=lambda item: item[0])[1]


def records_from_crawl_result(item: dict[str, Any]) -> list[dict[str, str]]:
    if not item.get("ok"):
        return []
    page_url = str(item.get("url") or "")
    structured_records = item.get("records") or []
    output: list[dict[str, str]] = []
    if isinstance(structured_records, list):
        for index, record in enumerate(structured_records, start=1):
            if not isinstance(record, dict):
                continue
            text = normalize_multiline_text(str(record.get("text") or ""))
            if len(normalize_text(text)) < 20:
                continue
            title = str(record.get("title") or f"链接岗位{index}")
            company = str(record.get("company") or "")
            if not company:
                candidates = extract_company_candidates(text, title)
                company = candidates[0] if candidates else ""
            output.append(
                {
                    "source": page_url or str(record.get("source") or "公开链接"),
                    "title": title,
                    "url": str(record.get("url") or page_url),
                    "text": text,
                    "company": company,
                    "salary": str(record.get("salary") or ""),
                    "location": str(record.get("location") or ""),
                    "education": str(record.get("education") or ""),
                    "experience": str(record.get("experience") or ""),
                }
            )
    if len(output) >= 2:
        return dedupe_jd_records(output)

    text = normalize_multiline_text(str(item.get("text") or ""))
    fallback = []
    for index, record in enumerate(split_batch_jd_text(text), start=1):
        fallback.append(
            {
                "source": page_url or "公开链接",
                "title": record.get("title") or f"链接岗位{index}",
                "url": page_url,
                "text": record.get("text", ""),
            }
        )
    if output and len(fallback) <= 1:
        return dedupe_jd_records(output)
    return dedupe_jd_records(output + fallback)


def likely_title_from_row_values(values: list[str]) -> str:
    for value in values:
        if not value or len(value) > 90:
            continue
        if re.search(r"https?://|www\.|@|有限公司|集团|公司|^\d+[\d,.\-~ 至到]*[kKwW万千元]", value):
            continue
        if any(word in value for word in ["实习", "工程师", "顾问", "专员", "分析师", "经理", "助理", "管培生", "运营", "产品", "数据", "市场", "销售", "财务", "法务", "合规", "LCA", "ESG"]):
            return value[:80]
    return values[0][:80] if values else ""


def records_from_dataframe(df: pd.DataFrame, source_prefix: str) -> list[dict[str, str]]:
    if df is None or df.empty:
        return []
    df = repair_dataframe_text(df).fillna("")
    column_map = {field: pick_table_column(df, field) for field in TABLE_FIELD_ALIASES}
    records: list[dict[str, str]] = []

    for index, row in df.iterrows():
        row_values = {str(column): clean_table_cell(row.get(column, "")) for column in df.columns}
        values = [value for value in row_values.values() if value]
        if not values:
            continue

        fields = {
            field: clean_table_cell(row.get(column, "")) if column else ""
            for field, column in column_map.items()
        }
        title = fields.get("title") or likely_title_from_row_values(values)
        company = fields.get("company", "")
        url = fields.get("url", "")

        structured_lines = []
        label_pairs = [
            ("公司", company),
            ("岗位", title),
            ("薪资", fields.get("salary", "")),
            ("地点", fields.get("location", "")),
            ("学历", fields.get("education", "")),
            ("经验", fields.get("experience", "")),
            ("链接", url),
        ]
        for label, value in label_pairs:
            if value:
                structured_lines.append(f"{label}：{value}")
        for column, value in row_values.items():
            if value and f"{column}：{value}" not in structured_lines:
                structured_lines.append(f"{column}：{value}")

        records.append(
            {
                "source": f"{source_prefix}/{index + 1}",
                "title": title,
                "url": url,
                "text": "\n".join(structured_lines),
                "company": company,
                "salary": fields.get("salary", ""),
                "location": fields.get("location", ""),
                "education": fields.get("education", ""),
                "experience": fields.get("experience", ""),
            }
        )
    return records


def read_csv_upload(data: bytes) -> pd.DataFrame:
    for encoding in ["utf-8-sig", "utf-8", "gbk", "gb18030"]:
        try:
            return pd.read_csv(io.BytesIO(data), encoding=encoding)
        except UnicodeDecodeError:
            continue
    return pd.read_csv(io.StringIO(data.decode("utf-8", errors="ignore")))


def records_from_uploaded_file(uploaded_file: Any) -> list[dict[str, str]]:
    if uploaded_file is None:
        return []
    suffix = Path(uploaded_file.name).suffix.lower()
    data = uploaded_file.getvalue()
    records = []

    if suffix in [".xlsx", ".xls"]:
        sheets = pd.read_excel(io.BytesIO(data), sheet_name=None)
        for sheet_name, df in sheets.items():
            records.extend(records_from_dataframe(df, f"{uploaded_file.name}/{sheet_name}"))
        return dedupe_jd_records(records)

    if suffix == ".csv":
        try:
            records.extend(records_from_dataframe(read_csv_upload(data), uploaded_file.name))
        except Exception:
            extracted = extract_text_from_upload(uploaded_file)
            for record in split_batch_jd_text(extracted):
                record["source"] = uploaded_file.name
                records.append(record)
        return dedupe_jd_records(records)

    extracted = extract_text_from_upload(uploaded_file)
    for record in split_batch_jd_text(extracted):
        record["source"] = uploaded_file.name
        records.append(record)
    return records


def preferred_jd_record_url(record: dict[str, str]) -> str:
    return valid_plugin_job_url(record.get("url"))


def merge_deduped_jd_record(existing: dict[str, str], candidate: dict[str, str]) -> dict[str, str]:
    merged = dict(existing)
    existing_url = preferred_jd_record_url(merged)
    candidate_url = preferred_jd_record_url(candidate)
    if candidate_url and not existing_url:
        merged["url"] = candidate_url
    elif existing_url:
        merged["url"] = existing_url

    for key in ["source", "title", "company", "salary", "location", "education", "experience", "page_index"]:
        current = normalize_text(str(merged.get(key) or ""))
        incoming = normalize_text(str(candidate.get(key) or ""))
        if not current and incoming:
            merged[key] = incoming

    current_text = normalize_text(str(merged.get("text") or ""))
    incoming_text = normalize_text(str(candidate.get("text") or ""))
    incoming_has_detail = bool(re.search(r"(?:岗位职责|职位描述|任职要求|工作内容|岗位要求)", incoming_text))
    current_has_detail = bool(re.search(r"(?:岗位职责|职位描述|任职要求|工作内容|岗位要求)", current_text))
    if incoming_text and (not current_text or (incoming_has_detail and not current_has_detail) or len(incoming_text) > len(current_text) * 1.8):
        merged["text"] = incoming_text
    elif current_text:
        merged["text"] = current_text

    if candidate_url and candidate_url not in normalize_text(str(merged.get("text") or "")):
        merged["text"] = normalize_multiline_text(str(merged.get("text") or "") + f"\n链接：{candidate_url}")
    return merged


def dedupe_jd_records(records: list[dict[str, str]]) -> list[dict[str, str]]:
    key_to_index: dict[str, int] = {}
    output: list[dict[str, str]] = []
    for record in records:
        text = normalize_text(record.get("text", ""))
        if len(text) < 20:
            continue
        if not valid_jd_record(record, text):
            continue
        keys = jd_record_dedupe_keys(record, text)
        duplicate_index = next((key_to_index[key] for key in keys if key in key_to_index), None)
        clean_record = dict(record)
        clean_record["text"] = text
        clean_url = preferred_jd_record_url(clean_record)
        if clean_url:
            clean_record["url"] = clean_url
        if duplicate_index is not None:
            output[duplicate_index] = merge_deduped_jd_record(output[duplicate_index], clean_record)
            for key in keys:
                key_to_index.setdefault(key, duplicate_index)
            continue
        output_index = len(output)
        for key in keys:
            key_to_index[key] = output_index
        output.append(clean_record)
    return output


def jd_record_id(record: dict[str, str]) -> str:
    text = normalize_text(str(record.get("text") or ""))
    keys = jd_record_dedupe_keys(record, text)
    return content_fingerprint(keys[0] if keys else text[:500])


def record_preview_fields(record: dict[str, str]) -> dict[str, str]:
    text = str(record.get("text") or "")
    try:
        jd = analyze_jd(text)
        fields = extract_card_fields_v2(record, text, jd)
    except Exception:
        fields = {
            "company": normalize_text(str(record.get("company") or "")) or "未识别",
            "title": normalize_text(str(record.get("title") or "")) or "未识别",
            "salary": normalize_text(str(record.get("salary") or "")) or "未识别",
            "location": normalize_text(str(record.get("location") or "")) or "未识别",
        }
    return {
        "公司": fields.get("company", "") or "未识别",
        "岗位": fields.get("title", "") or "未识别",
        "薪资": fields.get("salary", "") or "未识别",
        "地点": fields.get("location", "") or "未识别",
        "来源": normalize_text(str(record.get("source") or ""))[:80],
        "链接": normalize_text(str(record.get("url") or ""))[:120],
    }


def export_info_set_id(record: dict[str, str]) -> str:
    return str(record.get("_export_file_id") or record.get("_export_file_name") or record.get("source") or "")


def selected_export_info_set_records(records: list[dict[str, str]], table: pd.DataFrame, key_prefix: str) -> list[dict[str, str]]:
    if table is None or table.empty:
        return []
    selected_key = f"{key_prefix}_info_set_selected_ids"
    table = table.copy()
    if "信息集ID" not in table.columns:
        table["信息集ID"] = table["文件"].astype(str) if "文件" in table.columns else table.index.astype(str)
    info_set_ids = set(table["信息集ID"].astype(str).tolist())
    if selected_key not in st.session_state:
        st.session_state[selected_key] = set(info_set_ids)
    selected_ids = set(st.session_state.get(selected_key, set())) & info_set_ids
    return [
        record
        for record in records
        if export_info_set_id(record) in selected_ids
    ]


def selector_table_height(row_count: int, *, max_height: int = 360) -> int:
    return min(max_height, 48 + max(1, min(row_count, 10)) * 42)


def info_set_checkbox_key(key_prefix: str, info_set_id: str) -> str:
    return f"{key_prefix}_info_set_checked_{content_fingerprint(info_set_id)[:16]}"


def render_export_info_set_selector(records: list[dict[str, str]], table: pd.DataFrame, key_prefix: str) -> list[dict[str, str]]:
    if table is None or table.empty:
        return []
    selected_key = f"{key_prefix}_info_set_selected_ids"
    table = table.copy()
    if "信息集ID" not in table.columns:
        table["信息集ID"] = table["文件"].astype(str) if "文件" in table.columns else table.index.astype(str)
    info_set_ids = set(table["信息集ID"].astype(str).tolist())
    if selected_key not in st.session_state:
        st.session_state[selected_key] = set(info_set_ids)
    selected_ids = set(st.session_state.get(selected_key, set())) & info_set_ids

    st.markdown("##### 选择导入信息集")
    st.caption("按插件一次导出的文件导入；取消勾选会跳过该文件中的全部岗位。")

    action_cols = st.columns([1, 1, 3])
    if action_cols[0].button("全选", key=f"{key_prefix}_info_set_select_all"):
        selected_ids = set(info_set_ids)
        st.session_state[selected_key] = selected_ids
        for info_set_id in info_set_ids:
            st.session_state[info_set_checkbox_key(key_prefix, info_set_id)] = True
    if action_cols[1].button("清空", key=f"{key_prefix}_info_set_clear_all"):
        selected_ids = set()
        st.session_state[selected_key] = selected_ids
        for info_set_id in info_set_ids:
            st.session_state[info_set_checkbox_key(key_prefix, info_set_id)] = False

    checked_ids: set[str] = set()
    try:
        list_container = st.container(height=280)
    except TypeError:
        list_container = st.container()
    with list_container:
        for _, row in table.iterrows():
            info_set_id = str(row.get("信息集ID", ""))
            if not info_set_id:
                continue
            checkbox_key = info_set_checkbox_key(key_prefix, info_set_id)
            if checkbox_key not in st.session_state:
                st.session_state[checkbox_key] = info_set_id in selected_ids
            file_name = normalize_text(str(row.get("文件", ""))) or "未命名导出文件"
            date_label = normalize_text(str(row.get("日期", "")))
            job_count = normalize_text(str(row.get("岗位数", "")))
            modified = normalize_text(str(row.get("修改时间", "")))
            label_parts = [file_name]
            meta_parts = [part for part in [date_label, f"{job_count}条岗位" if job_count else "", modified] if part]
            if meta_parts:
                label_parts.append(" | ".join(meta_parts))
            checked = st.checkbox(
                "    ".join(label_parts),
                key=checkbox_key,
                help=info_set_id,
            )
            if checked:
                checked_ids.add(info_set_id)
    selected_ids = checked_ids
    st.session_state[selected_key] = selected_ids

    selected_records = selected_export_info_set_records(records, table, key_prefix)
    st.caption(f"已选择 {len(selected_ids)} 个信息集，包含 {len(selected_records)} 条岗位。")
    return selected_records


if hasattr(st, "dialog"):
    @st.dialog("选择导入文件", width="large")
    def render_export_info_set_dialog(records: list[dict[str, str]], table: pd.DataFrame, key_prefix: str) -> None:
        render_export_info_set_selector(records, table, key_prefix)
        if st.button("完成选择", type="primary", key=f"{key_prefix}_info_set_dialog_done"):
            st.rerun()
else:
    def render_export_info_set_dialog(records: list[dict[str, str]], table: pd.DataFrame, key_prefix: str) -> None:
        render_export_info_set_selector(records, table, key_prefix)


def render_import_record_selector(
    records: list[dict[str, str]],
    key_prefix: str,
    *,
    allow_import_toggle: bool = True,
    title: str = "选择要导入分析的岗位",
    caption: str | None = None,
) -> list[dict[str, str]]:
    if not records:
        return []
    deleted_key = f"{key_prefix}_deleted_ids"
    selected_key = f"{key_prefix}_selected_ids"
    deleted_ids = set(st.session_state.get(deleted_key, set()))
    available = [record for record in dedupe_jd_records(records) if jd_record_id(record) not in deleted_ids]
    available_ids = [jd_record_id(record) for record in available]
    if allow_import_toggle and selected_key not in st.session_state:
        st.session_state[selected_key] = set(available_ids)
    selected_ids = set(st.session_state.get(selected_key, set())) & set(available_ids) if allow_import_toggle else set(available_ids)

    if not allow_import_toggle:
        if caption:
            st.caption(caption)
        st.caption(f"将导入 {len(available)} 条岗位；需要排除单条岗位时再打开下方预览。")
        show_record_preview = st.checkbox("按岗位预览/排除", value=False, key=f"{key_prefix}_show_record_preview")
        if not show_record_preview:
            return available

    rows = []
    for index, record in enumerate(available, start=1):
        record_id = jd_record_id(record)
        preview = record_preview_fields(record)
        row = {
            "记录ID": record_id,
            "删除": False,
            "序号": index,
            **preview,
        }
        if allow_import_toggle:
            row = {"记录ID": record_id, "导入": record_id in selected_ids, **{key: value for key, value in row.items() if key != "记录ID"}}
        rows.append(row)
    if not rows:
        st.info("当前导入候选已清空。")
        return []

    if caption is None:
        caption = "取消“导入”会跳过该岗位；勾选“删除”并点击下方按钮，会从当前候选列表移除，不会删除硬盘里的原始导出文件。"
    column_config = {
        "记录ID": None,
        "删除": st.column_config.CheckboxColumn("排除", help="勾选后从本次候选中排除"),
        "序号": st.column_config.NumberColumn("序号", width="small"),
        "公司": st.column_config.TextColumn("公司", width="medium"),
        "岗位": st.column_config.TextColumn("岗位", width="large"),
        "薪资": st.column_config.TextColumn("薪资", width="small"),
        "地点": st.column_config.TextColumn("地点", width="small"),
        "来源": st.column_config.TextColumn("来源", width="small"),
        "链接": st.column_config.LinkColumn("链接", width="small"),
    }
    if allow_import_toggle:
        column_config["导入"] = st.column_config.CheckboxColumn("导入", help="勾选后进入本次批量分析")
    disabled_columns = ["记录ID", "序号", "公司", "岗位", "薪资", "地点", "来源", "链接"]
    editor_df = pd.DataFrame(rows)
    editor_area = st.container() if allow_import_toggle else st.expander(f"{title}（可选）", expanded=False)
    with editor_area:
        if allow_import_toggle:
            st.markdown(f"##### {title}")
        st.caption(caption)
        edited = st.data_editor(
            editor_df,
            width="stretch",
            hide_index=True,
            height=selector_table_height(len(editor_df)),
            key=f"{key_prefix}_editor",
            column_config=column_config,
            disabled=disabled_columns,
        )
        action_cols = st.columns([1, 1, 3])
        delete_button_label = "排除勾选岗位" if not allow_import_toggle else "删除勾选记录"
        if action_cols[0].button(delete_button_label, key=f"{key_prefix}_delete_selected", disabled=not set(edited.loc[edited["删除"].astype(bool), "记录ID"].astype(str).tolist())):
            delete_ids_now = set(edited.loc[edited["删除"].astype(bool), "记录ID"].astype(str).tolist())
            import_ids_now = (
                set(edited.loc[edited["导入"].astype(bool), "记录ID"].astype(str).tolist())
                if allow_import_toggle
                else set(available_ids)
            )
            st.session_state[deleted_key] = deleted_ids | delete_ids_now
            if allow_import_toggle:
                st.session_state[selected_key] = import_ids_now - delete_ids_now
            st.rerun()
        if action_cols[1].button("恢复已排除", key=f"{key_prefix}_restore_deleted", disabled=not deleted_ids):
            st.session_state[deleted_key] = set()
            if allow_import_toggle:
                st.session_state[selected_key] = set(available_ids)
            st.rerun()
        current_delete_ids = set(edited.loc[edited["删除"].astype(bool), "记录ID"].astype(str).tolist())
        if allow_import_toggle:
            current_import_ids = set(edited.loc[edited["导入"].astype(bool), "记录ID"].astype(str).tolist())
            action_cols[2].caption(f"候选 {len(available)} 条，当前勾选导入 {len(current_import_ids - current_delete_ids)} 条。")
        else:
            action_cols[2].caption(f"{len(available)} 条岗位，已排除 {len(deleted_ids)} 条。")
    import_ids = (
        set(edited.loc[edited["导入"].astype(bool), "记录ID"].astype(str).tolist())
        if allow_import_toggle
        else set(available_ids)
    )
    delete_ids = set(edited.loc[edited["删除"].astype(bool), "记录ID"].astype(str).tolist())

    selected_ids = import_ids - delete_ids
    if allow_import_toggle:
        st.session_state[selected_key] = selected_ids
    selected = [record for record in available if jd_record_id(record) in selected_ids]
    return selected


def batch_recommendation(score: int, jd_analysis: dict[str, Any], job_type: str, fresh: str) -> str:
    generic = jd_analysis.get("value", {}).get("is_generic_esg")
    high_value = jd_analysis.get("value", {}).get("is_high_value")
    if generic and score < 70:
        return "谨慎：疑似低价值/杂务"
    if fresh == "否/偏社招" and score < 82:
        return "谨慎：经验门槛偏高"
    if "实习" in job_type and score >= 78 and high_value:
        return "P1实习优先"
    if score >= 82 and high_value:
        return "P1优先投递"
    if score >= 72:
        return "P2值得投递"
    if score >= 60:
        return "P3可备选"
    if "实习" in job_type and score >= 55:
        return "可作为实习备选"
    return "暂不优先"


def detect_job_type_safe(text: str) -> str:
    title_like = text[:120]
    if any(text_contains(title_like, keyword) for keyword in ["实习", "intern", "Internship", "日常实习", "暑期实习", "元/天", "天/周"]):
        return "实习"
    if any(text_contains(text, keyword) for keyword in ["可转正实习", "实习转正", "实习生"]):
        return "实习"
    if any(text_contains(text, keyword) for keyword in ["全职", "正式", "校招", "社招", "管培", "届生", "15薪", "年薪", "五险一金"]):
        return "正式工"
    if parse_salary_floor(text, "")[1] == "日薪":
        return "实习"
    return "未明确"


def detect_internship_opening_safe(text: str) -> str:
    title_like = text[:140]
    yes_keywords = [
        "实习",
        "实习生",
        "日常实习",
        "暑期实习",
        "寒假实习",
        "intern",
        "Intern",
        "Internship",
        "元/天",
        "天/周",
        "周到岗",
        "在校生",
        "可转正",
        "可转正实习",
        "实习转正",
    ]
    no_keywords = [
        "全职",
        "正式员工",
        "正式岗",
        "社招",
        "5-10年",
        "3-5年",
        "5年以上",
        "资深",
        "专家",
        "负责人",
        "经理",
        "总监",
    ]
    if any(text_contains(title_like, keyword) for keyword in yes_keywords) or parse_salary_floor(text, "")[1] == "日薪":
        return "是"
    if any(text_contains(text, keyword) for keyword in no_keywords):
        return "否"
    return "未明确"


def detect_fresh_graduate_safe(text: str) -> str:
    negative = ["3-5年", "5-10年", "3年以上", "5年以上", "资深", "专家", "负责人", "经理"]
    positive = ["应届", "校招", "秋招", "春招", "毕业生", "经验不限", "0-1年", "1年以内", "2026届", "2027届", "在校生", "实习"]
    if any(text_contains(text, keyword) for keyword in negative):
        return "否/偏社招"
    if any(text_contains(text, keyword) for keyword in positive):
        return "是"
    return "未明确"


PROFILE_MATCH_GROUPS = {
    "数据分析": {"weight": 12, "aliases": ["数据分析", "SQL", "Python", "Excel", "Power BI", "Tableau", "pandas", "指标", "看板"]},
    "业务/商业分析": {"weight": 12, "aliases": ["业务分析", "商业分析", "经营分析", "策略分析", "指标体系", "漏斗分析", "归因分析"]},
    "产品能力": {"weight": 11, "aliases": ["产品经理", "产品助理", "需求分析", "PRD", "原型", "用户研究", "竞品分析"]},
    "运营增长": {"weight": 10, "aliases": ["运营", "增长", "用户运营", "内容运营", "活动运营", "转化", "留存", "拉新"]},
    "项目管理": {"weight": 10, "aliases": ["项目管理", "项目推进", "项目交付", "跨部门", "PMO", "里程碑"]},
    "研发工程": {"weight": 12, "aliases": ["前端", "后端", "Java", "Go", "C++", "React", "Vue", "算法", "机器学习"]},
    "市场/销售": {"weight": 9, "aliases": ["市场", "营销", "投放", "渠道", "销售", "商务", "BD", "客户开发"]},
    "财务/法务/人力": {"weight": 9, "aliases": ["财务", "审计", "预算", "法务", "合规", "合同", "招聘", "培训", "绩效"]},
    "供应链/采购": {"weight": 9, "aliases": ["供应链", "采购", "物流", "库存", "供应商管理", "计划"]},
    "LCA": {"weight": 14, "aliases": ["LCA", "生命周期评价", "生命周期评估"]},
    "产品碳足迹": {"weight": 15, "aliases": ["产品碳足迹", "CFP", "ISO14067", "ISO 14067", "碳足迹"]},
    "EPD/PCR": {"weight": 11, "aliases": ["EPD", "PCR", "环境产品声明", "产品环境声明"]},
    "CBAM/出海合规": {"weight": 13, "aliases": ["CBAM", "碳边境", "出海合规", "欧盟", "海外合规", "电池法", "EUDR"]},
    "碳核算/核查": {"weight": 11, "aliases": ["碳核算", "碳盘查", "碳核查", "GHG", "温室气体", "排放因子", "ISO14064"]},
    "供应链碳管理": {"weight": 9, "aliases": ["供应链碳", "供应链碳数据", "供应商碳", "供应链碳管理", "活动数据"]},
    "LCA软件": {"weight": 9, "aliases": ["SimaPro", "GaBi", "openLCA", "open LCA", "eFootprint"]},
    "数据能力": {"weight": 9, "aliases": ["Python", "SQL", "数据分析", "pandas", "Stata", "MATLAB", "Excel", "Power BI"]},
    "英文能力": {"weight": 7, "aliases": ["英文", "英语", "English", "CET-6", "六级", "口语", "presentation"]},
}

TARGET_INDUSTRY_WORDS = ["互联网", "制造", "金融", "咨询", "新能源", "医疗", "消费", "教育", "物流", "供应链", "专业服务", "研发", "产品", "数据", "运营", "合规", "ESG", "可持续", "双碳", "碳", "LCA", "碳核算", "碳足迹", "绿色"]
GENERIC_ESG_RISK_WORDS = ["打杂", "纯执行", "无明确职责", "行政杂务", "公众号", "活动策划", "行政协同", "品牌传播", "宣传", "会议组织", "材料整理", "只做排版", "只做整理", "评级问卷", "公益", "销售任务"]
ENGINEERING_REQUIRED_GROUPS = {"研发工程", "前端", "后端", "机器学习/AI"}
ENGINEERING_EVIDENCE_TERMS = [
    "前端", "后端", "全栈", "服务端", "Java", "Spring", "Go", "C++", "C#", "PHP",
    "JavaScript", "TypeScript", "React", "Vue", "Node", "数据库", "SQL", "API",
    "接口", "微服务", "代码仓库", "Git", "自动化测试", "单元测试", "部署", "系统开发",
    "开发项目", "算法", "机器学习", "模型训练", "深度学习", "NLP", "LLM",
]


def group_present(text: str, aliases: list[str]) -> bool:
    return any(text_contains(text, alias) for alias in aliases)


def preference_context_text(preferences: dict[str, Any] | None = None) -> str:
    preferences = preferences or {}
    parts = [
        " ".join(split_preference_items(preferences.get("target_roles", []))),
        " ".join(split_preference_items(preferences.get("preferred_industries", []))),
        " ".join(split_preference_items(preferences.get("job_keywords", []))),
        str(preferences.get("notes", "")),
    ]
    return normalize_text("\n".join(parts))


def profile_group_names(label: str, aliases: list[str] | None = None) -> list[str]:
    return list(dict.fromkeys(SKILL_SEMANTIC_GROUPS.get(label, []) + semantic_groups_for_terms(label, aliases)))


def text_has_capability_group(text: str, group_names: list[str]) -> bool:
    if not group_names:
        return False
    hits = capability_group_hits(text)
    return any(group in hits for group in group_names)


def contextual_profile_group_weight(
    label: str,
    aliases: list[str],
    base_weight: float,
    profile_text: str,
    resume_text: str = "",
    preferences: dict[str, Any] | None = None,
) -> tuple[float, str]:
    all_aliases = expanded_aliases(label, aliases)
    target_text = normalize_text(profile_text + "\n" + preference_context_text(preferences))
    resume_clean = normalize_text(resume_text)
    target_hits = alias_hits(target_text, all_aliases) or related_alias_hits(target_text, label, aliases)
    resume_hits = alias_hits(resume_clean, all_aliases) or related_alias_hits(resume_clean, label, aliases)
    group_names = profile_group_names(label, aliases)
    target_group_hit = text_has_capability_group(target_text, group_names)
    resume_group_hit = text_has_capability_group(resume_clean, group_names)
    multiplier = 1.0
    reason = ""
    if (target_hits or target_group_hit) and (resume_hits or resume_group_hit):
        multiplier = 1.65
        reason = f"目标和简历共同强调：{label}"
    elif target_hits or target_group_hit:
        multiplier = 1.35
        reason = f"目标意向加权：{label}"
    elif resume_hits or resume_group_hit:
        multiplier = 1.25
        reason = f"简历证据加权：{label}"
    return base_weight * multiplier, reason


def resume_has_engineering_evidence(resume_text: str) -> bool:
    clean = normalize_text(resume_text)
    if not clean:
        return False
    hits = [term for term in ENGINEERING_EVIDENCE_TERMS if text_contains(clean, term)]
    specific_hits = [
        term for term in hits
        if term not in {"SQL", "数据库", "接口", "部署", "算法", "机器学习", "LLM"}
    ]
    return len(specific_hits) >= 1 or len(hits) >= 3


def jd_has_engineering_requirement(
    jd_text: str,
    jd_analysis: dict[str, Any],
    required_groups: list[str] | None = None,
    skills: list[str] | None = None,
) -> bool:
    clean = text_without_urls(jd_text)
    if jd_analysis.get("category") == "研发/工程岗":
        return True
    if any(name in ENGINEERING_REQUIRED_GROUPS for name in (required_groups or [])):
        return True
    if any(skill in {"前端", "后端", "机器学习/AI"} for skill in (skills or [])):
        return True
    title = normalize_text(jd_analysis.get("basic", {}).get("岗位名", ""))
    engineering_terms = [
        "开发工程师", "软件工程师", "前端工程师", "后端工程师", "测试工程师", "数据工程师",
        "算法工程师", "PHP开发", "Java开发", "C++", "Layout工程师", "架构师",
    ]
    return any(text_contains(title + "\n" + clean[:240], term) for term in engineering_terms)


def industry_aliases(industry: str) -> list[str]:
    clean = normalize_text(industry)
    aliases = [clean]
    aliases.extend(INDUSTRY_ALIAS_MAP.get(clean, []))
    for option, option_aliases in INDUSTRY_ALIAS_MAP.items():
        if clean and (clean in option or option in clean):
            aliases.append(option)
            aliases.extend(option_aliases)
    return split_preference_items(aliases)


def industry_preference_hits(jd_text: str, preferred_industries: list[str]) -> list[str]:
    hits: list[str] = []
    for industry in split_preference_items(preferred_industries):
        aliases = industry_aliases(industry)
        matched_aliases = [alias for alias in aliases if text_contains(jd_text, alias)]
        if matched_aliases:
            alias_preview = "、".join(matched_aliases[:2])
            hits.append(f"{industry}({alias_preview})" if alias_preview != industry else industry)
    return list(dict.fromkeys(hits))


def preference_semantic_hits(jd_text: str, items: list[str], *, conservative: bool = False) -> list[str]:
    hits: list[str] = []
    for item in split_preference_items(items):
        aliases = expanded_aliases(item, [item])
        if conservative:
            aliases = [alias for alias in aliases if alias == item or len(alias) >= 3]
        matched = [alias for alias in aliases if text_contains(jd_text, alias)]
        if matched:
            preview = "、".join(matched[:2])
            hits.append(f"{item}({preview})" if preview != item else item)
    return list(dict.fromkeys(hits))


def parse_salary_floor(text: str, job_type: str = "") -> tuple[float | None, str]:
    text = normalize_text(text)
    day_match = re.search(r"(\d+(?:\.\d+)?)\s*[-~—至到]\s*(\d+(?:\.\d+)?)\s*元\s*/?\s*天", text)
    if day_match:
        return float(day_match.group(1)), "日薪"
    single_day = re.search(r"(\d+(?:\.\d+)?)\s*元\s*/?\s*天", text)
    if single_day:
        return float(single_day.group(1)), "日薪"
    k_match = re.search(r"(\d+(?:\.\d+)?)\s*[-~—至到]\s*(\d+(?:\.\d+)?)\s*[kK](?:\s*[·*xX]\s*(\d+)\s*薪)?", text)
    if k_match:
        base = float(k_match.group(1)) * 1000
        months = int(k_match.group(3)) if k_match.group(3) else 12
        if months > 12:
            return base * months / 12, "月薪折算"
        return base, "月薪"
    single_k = re.search(r"(\d+(?:\.\d+)?)\s*[kK]\s*(?:以上|\+)?", text)
    if single_k:
        return float(single_k.group(1)) * 1000, "月薪"
    annual_match = re.search(r"(\d+(?:\.\d+)?)\s*[-~—至到]\s*(\d+(?:\.\d+)?)\s*[万wW]\s*/?\s*年?", text)
    if annual_match:
        return float(annual_match.group(1)) * 10000 / 12, "年薪折月"
    return None, ""


def salary_attraction_label(text: str, job_type: str, preferences: dict[str, Any] | None = None) -> tuple[str, int]:
    preferences = preferences or DEFAULT_TARGET_PREFERENCES
    floor, unit = parse_salary_floor(text, job_type)
    if floor is None:
        return "薪资未明确，需确认是否达标", 0

    min_daily = int(preferences.get("min_daily_salary") or DEFAULT_TARGET_PREFERENCES["min_daily_salary"])
    min_monthly = int(preferences.get("min_monthly_salary") or DEFAULT_TARGET_PREFERENCES["min_monthly_salary"])
    if unit == "日薪" or job_type == "实习":
        if floor >= min_daily * 1.5:
            return f"实习薪资高于目标日薪 {min_daily} 元", 6
        if floor >= min_daily:
            return f"实习薪资达到目标日薪 {min_daily} 元", 3
        return f"实习薪资低于目标日薪 {min_daily} 元", -6

    if floor >= min_monthly * 1.3:
        return f"薪资明显高于目标月薪 {min_monthly // 1000}k", 8
    if floor >= min_monthly:
        return f"薪资达到目标月薪 {min_monthly // 1000}k", 5
    if floor >= min_monthly * 0.8:
        return f"薪资略低于目标月薪 {min_monthly // 1000}k", -3
    return f"薪资低于目标月薪 {min_monthly // 1000}k", -8


def split_preference_items(value: Any) -> list[str]:
    if isinstance(value, list):
        raw_items = value
    else:
        raw_items = re.split(r"[、,，/\n\s]+", str(value or ""))
    return list(dict.fromkeys(item.strip() for item in raw_items if str(item).strip()))


def free_input_multiselect(
    label: str,
    options: list[str],
    default: list[str],
    *,
    key: str,
    help_text: str,
) -> list[str]:
    default_items = split_preference_items(default)
    option_set = set(options)
    ordered_options = list(dict.fromkeys(default_items + options))
    kwargs: dict[str, Any] = {
        "label": label,
        "options": ordered_options,
        "default": default_items,
        "key": key,
        "help": help_text,
    }
    multiselect_params = inspect.signature(st.multiselect).parameters
    if "accept_new_options" in multiselect_params:
        kwargs["accept_new_options"] = True
    if "placeholder" in multiselect_params:
        kwargs["placeholder"] = "输入关键词搜索，或直接输入后回车确认"
    if "accept_new_options" in multiselect_params:
        return split_preference_items(st.multiselect(**kwargs))

    selected = split_preference_items(st.multiselect(**kwargs))
    custom_default = "、".join(item for item in default_items if item not in option_set)
    custom_items = st.text_input(f"{label}（输入后确认）", value=custom_default, key=f"{key}_custom")
    return split_preference_items(selected + split_preference_items(custom_items))


def target_preference_adjustment(
    jd_text: str,
    jd_analysis: dict[str, Any],
    job_type: str,
    preferences: dict[str, Any] | None = None,
) -> tuple[int, list[str]]:
    preferences = preferences or DEFAULT_TARGET_PREFERENCES
    score = 0
    reasons: list[str] = []

    basic = jd_analysis.get("basic", {}) if jd_analysis else {}
    location_text = normalize_text(str(basic.get("地点", "")))
    region_info = classify_region(location_text, jd_text)
    target_cities = split_preference_items(preferences.get("target_cities", []))
    if target_cities:
        location_pool = normalize_text(" ".join([location_text, region_info.get("标准城市", ""), region_info.get("省份", "")]))
        matched_city = city_match_label(location_pool, target_cities)
        if matched_city:
            score += 7
            reasons.append("城市匹配目标意向：" + matched_city)
        elif region_info.get("省份") == "远程" and preferences.get("accept_remote", True):
            score += 4
            reasons.append("远程岗位符合城市偏好")
        elif region_info.get("省份") == "多地/全国" and preferences.get("accept_nationwide", True):
            score += 3
            reasons.append("多地/全国岗位可接受")
        elif region_info.get("省份") not in {"未识别", "海外/港澳台", "远程", "多地/全国"}:
            score -= 6
            reasons.append("城市不在当前目标意向中")

    preferred_industries = split_preference_items(preferences.get("preferred_industries", []))
    industry_hits = industry_preference_hits(jd_text, preferred_industries)
    if industry_hits:
        score += min(6, len(industry_hits) * 3)
        reasons.append("行业偏好命中：" + " / ".join(industry_hits[:3]))

    target_roles = split_preference_items(preferences.get("target_roles", []))
    role_hits = preference_semantic_hits(jd_text, target_roles)
    if role_hits:
        score += min(8, len(role_hits) * 4)
        reasons.append("求职方向/相关能力命中：" + " / ".join(role_hits[:3]))

    job_keywords = split_preference_items(preferences.get("job_keywords", []))
    keyword_hits = preference_semantic_hits(jd_text, job_keywords)
    if keyword_hits:
        score += min(10, len(keyword_hits) * 3)
        reasons.append("岗位关键词/关联词命中：" + " / ".join(keyword_hits[:4]))

    avoid_hits = preference_semantic_hits(jd_text, split_preference_items(preferences.get("avoid_keywords", "")), conservative=True)
    if avoid_hits:
        score -= min(14, len(avoid_hits) * 5)
        reasons.append("命中规避项/近义风险：" + " / ".join(avoid_hits[:3]))

    return score, reasons


def extract_detail_fields_from_text(url: str, text: str) -> dict[str, str]:
    clean = normalize_multiline_text(text)
    lines = [normalize_text(line) for line in clean.splitlines() if normalize_text(line)]
    joined = "\n".join(lines[:80])
    host = urlparse(normalize_url(url)).netloc.lower() if url else ""
    fields = {"company": "", "title": "", "salary": "", "location": "", "education": "", "experience": ""}

    if "liepin.com" in host:
        for line in lines[:12]:
            match = re.search(r"【([^】]{1,30})\s+(.{2,80}?)招聘】-([^-\n]{2,80}?)招聘信息", line)
            if match:
                fields["location"] = normalize_text(match.group(1))
                fields["title"] = normalize_text(match.group(2))
                fields["company"] = clean_company_candidate(re.sub(rf"(?:{GENERIC_CITY_PREFIX})$", "", normalize_text(match.group(3))))
                break
        if not fields["company"]:
            match = re.search(r"·\s*([^\n·]{2,80}(?:公司|集团|有限责任公司|股份有限公司|有限公司|中心|学校|研究院))", joined)
            if match:
                fields["company"] = clean_company_candidate(match.group(1))

    if "shixiseng.com" in host:
        for line in lines[:12]:
            match = re.search(r"(.{2,80}?)(?:实习|校招)?招聘-([^-\n]{2,60}?)(?:实习生)?招聘-实习僧", line)
            if match:
                fields["title"] = normalize_text(match.group(1))
                fields["company"] = clean_company_candidate(match.group(2))
                break
        if not fields["title"]:
            for line in lines[:30]:
                if generic_title_line(line):
                    fields["title"] = clean_generic_title(line)
                    break

    if "51job.com" in host or "we.51job.com" in host:
        for line in lines[:20]:
            if generic_title_line(line):
                fields["title"] = clean_generic_title(line)
                break
        candidates = extract_company_candidates(joined, fields["title"])
        if candidates:
            fields["company"] = candidates[0]

    salary = extract_salary(joined)
    if salary != "未识别":
        fields["salary"] = salary
    location = extract_location(joined)
    if location != "未识别":
        fields["location"] = fields["location"] or location
    education = extract_education(joined)
    if education != "未识别":
        fields["education"] = education
    experience = extract_experience(joined)
    if experience != "未识别":
        fields["experience"] = experience

    return {key: clean_extracted_value(value) for key, value in fields.items() if value}


def field_is_low_confidence(value: str, field: str = "") -> bool:
    text = normalize_text(value)
    if not text or text == "未识别":
        return True
    if field == "company":
        if text in GENERIC_BAD_CARD_LINES or bad_generic_card_line(text) or generic_location_line(text):
            return True
        if len(text) > 60 and not re.search(COMPANY_SUFFIX_PATTERN, text):
            return True
    if field == "title":
        if bad_generic_card_line(text) or generic_location_line(text) or generic_meta_line(text):
            return True
        if ("：" in text or "，" in text) and len(text) > 26:
            return True
    return False


def extract_card_fields_v2(record: dict[str, str], text: str, jd_analysis: dict[str, Any]) -> dict[str, str]:
    fallback_title = normalize_text(record.get("title", ""))
    salary = normalize_text(record.get("salary", "")) or jd_basic_value(jd_analysis, "薪资") or extract_salary(text)
    location = normalize_text(record.get("location", "")) or jd_basic_value(jd_analysis, "地点") or extract_location(text)
    education = normalize_text(record.get("education", "")) or jd_basic_value(jd_analysis, "学历要求") or extract_education(text)
    experience = normalize_text(record.get("experience", "")) or jd_basic_value(jd_analysis, "经验要求") or extract_experience(text)

    title = fallback_title or jd_basic_value(jd_analysis, "岗位名")
    if not title or title == "未识别":
        title = fallback_title
    if (not title or len(title) > 90) and salary and salary != "未识别":
        before_salary = text.split(salary, 1)[0].strip()
        title = before_salary[-80:].strip(" ：:，,;；|")
    if not title:
        title = text[:60]

    company = clean_company_candidate(normalize_text(record.get("company", ""))) or jd_basic_value(jd_analysis, "公司名")
    if not company or company == "未识别":
        candidates = extract_company_candidates(text, title)
        company = candidates[0] if candidates else ""
    if not company:
        company = "未识别"

    detail_url = str(record.get("url") or "")
    detail_fields = extract_detail_fields_from_text(detail_url, text)
    if detail_fields:
        detail_host = urlparse(normalize_url(detail_url)).netloc.lower() if detail_url else ""
        trusted_detail_domain = any(host in detail_host for host in ["liepin.com", "shixiseng.com", "51job.com", "zhaopin.com"])
        if detail_fields.get("company") and (trusted_detail_domain or field_is_low_confidence(company, "company")):
            company = detail_fields["company"]
        if detail_fields.get("title") and (trusted_detail_domain or field_is_low_confidence(title, "title")):
            title = detail_fields["title"]
        if detail_fields.get("salary") and (trusted_detail_domain or not salary or salary == "未识别"):
            salary = detail_fields["salary"]
        if detail_fields.get("location") and (trusted_detail_domain or not location or location == "未识别"):
            location = detail_fields["location"]
        if detail_fields.get("education") and (trusted_detail_domain or not education or education == "未识别"):
            education = detail_fields["education"]
        if detail_fields.get("experience") and (trusted_detail_domain or not experience or experience == "未识别"):
            experience = detail_fields["experience"]

    return {
        "company": clean_extracted_value(company),
        "title": clean_extracted_value(title),
        "salary": clean_extracted_value(salary),
        "location": clean_extracted_value(location),
        "education": clean_extracted_value(education),
        "experience": clean_extracted_value(experience),
    }


def extraction_confidence(fields: dict[str, str], text: str, jd_analysis: dict[str, Any]) -> str:
    score = extraction_confidence_score(fields, text, jd_analysis)
    if score >= 88:
        return f"高({score})"
    if score >= 68:
        return f"中({score})"
    return f"低({score})：建议打开详情页补全"


def extraction_confidence_score(fields: dict[str, str], text: str, jd_analysis: dict[str, Any]) -> int:
    score = 0
    field_weights = {
        "company": 22,
        "title": 22,
        "salary": 12,
        "location": 12,
        "education": 6,
        "experience": 6,
    }
    for key, weight in field_weights.items():
        value = fields.get(key, "")
        if value and value != "未识别" and not field_is_low_confidence(value, key):
            score += weight
    if len(normalize_text(text)) >= 160:
        score += 8
    if len(normalize_text(text)) >= 420:
        score += 4
    if jd_skill_list(jd_analysis):
        score += 8
    if fields.get("company") and fields.get("title") and fields["company"] != fields["title"]:
        score += 4
    if fields.get("company") and field_is_low_confidence(fields["company"], "company"):
        score -= 16
    if fields.get("title") and field_is_low_confidence(fields["title"], "title"):
        score -= 12
    return int(np.clip(score, 0, 100))


def evaluate_profile_fit_v2(
    jd_analysis: dict[str, Any],
    profile_text: str,
    job_type: str = "",
    fresh: str = "",
    preferences: dict[str, Any] | None = None,
    resume_text: str = "",
    fast: bool = False,
) -> dict[str, Any]:
    profile_text = effective_profile_text(profile_text)
    preferences = preferences or load_target_preferences()
    jd_text = jd_analysis.get("raw_text", "")
    candidate_profile_text = normalize_text(profile_text + "\n" + resume_text)

    required = []
    matched = []
    missing = []
    evidence = []
    related_matched = []
    weight_reasons = []
    total_weight = 0
    matched_weight = 0
    for name, config in PROFILE_MATCH_GROUPS.items():
        aliases = config["aliases"]
        all_aliases = expanded_aliases(name, aliases)
        weight, weight_reason = contextual_profile_group_weight(name, aliases, config["weight"], profile_text, resume_text, preferences)
        if group_present(jd_text, all_aliases):
            required.append(name)
            total_weight += weight
            if weight_reason:
                weight_reasons.append(weight_reason)
            hits = alias_hits(candidate_profile_text, aliases)
            related_hits = related_alias_hits(candidate_profile_text, name, aliases)
            if hits:
                matched.append(name)
                matched_weight += weight
                evidence.append(f"{name}：{', '.join(hits[:3])}")
            elif related_hits:
                matched.append(name)
                related_matched.append(name)
                matched_weight += weight * 0.68
                evidence.append(f"{name}相关：{', '.join(related_hits[:3])}")
            else:
                missing.append(name)

    if total_weight:
        tech_score = int(round(matched_weight / total_weight * 100))
    else:
        tech_score = 45
    semantic_fn = semantic_similarity_fast if fast else semantic_similarity
    semantic_score = int(round(semantic_fn(jd_text, candidate_profile_text) * 100))

    value_score = 0
    reasons = []
    if jd_analysis.get("value", {}).get("is_high_value"):
        value_score += 16
        reasons.append("核心职责和可沉淀成果较明确")
    if any(text_contains(jd_text, word) for word in TARGET_INDUSTRY_WORDS):
        value_score += 9
        reasons.append("行业或职能方向清晰，便于横向比较")
    if fresh == "是":
        value_score += 5
        reasons.append("对应届/在校生友好")
    elif fresh == "否/偏社招":
        reasons.append("应届友好度偏低")

    salary_label, salary_bonus = salary_attraction_label(jd_text, job_type, preferences)
    value_score += salary_bonus
    if salary_label != "未明确":
        reasons.append(salary_label)

    preference_score, preference_reasons = target_preference_adjustment(jd_text, jd_analysis, job_type, preferences)
    if preference_score >= 0:
        value_score += preference_score
    else:
        reasons.append("目标意向扣分：" + str(abs(preference_score)))
    reasons.extend(preference_reasons)
    if weight_reasons:
        reasons.extend(list(dict.fromkeys(weight_reasons))[:4])

    risk_score = max(0, -preference_score)
    if jd_analysis.get("value", {}).get("is_generic_esg"):
        risk_score += 25
    risk_hits = [word for word in GENERIC_ESG_RISK_WORDS if text_contains(jd_text, word)]
    if risk_hits:
        risk_score += min(22, len(risk_hits) * 6)
        reasons.append("低价值/职责不清风险：" + " / ".join(risk_hits[:4]))
    if any(text_contains(jd_text, word) for word in ["3-5年", "5-10年", "5年以上", "资深", "负责人", "总监"]):
        risk_score += 12
        reasons.append("经验门槛偏高")
    if group_present(jd_text, ["销售", "市场拓展", "商务资源"]) and not group_present(jd_text, ["客户管理", "销售分析", "解决方案", "行业", "大客户"]):
        risk_score += 12
        reasons.append("销售/商务指标较重，需确认是否符合个人目标")
    engineering_required = jd_has_engineering_requirement(jd_text, jd_analysis, required, jd_skill_list(jd_analysis))
    if engineering_required and normalize_text(resume_text) and not resume_has_engineering_evidence(resume_text):
        risk_score += 18
        reasons.append("工程/开发岗位需要代码或系统项目证据，当前简历证据不足")
    if total_weight == 0 and jd_analysis.get("category") in ["ESG岗", "泛运营岗", "待人工判断"]:
        risk_score += 8

    semantic_bonus = min(6, len(semantic_expansion_terms(jd_text)) + len(semantic_expansion_terms(candidate_profile_text))) if semantic_score >= 42 else 0
    if semantic_bonus:
        reasons.append(f"语义相关度较高：{semantic_score}/100")

    score = 28 + tech_score * 0.48 + semantic_score * 0.16 + semantic_bonus + value_score - risk_score
    if total_weight == 0:
        score -= 8
    score = int(np.clip(round(score), 0, 100))

    if matched:
        reasons.insert(0, "匹配能力：" + " / ".join(matched[:6]))
    if evidence:
        reasons.insert(1 if matched else 0, "意向证据：" + "；".join(evidence[:4]))
    if not reasons:
        reasons.append("信息不足，建议打开详情页补充JD")

    gaps = missing[:8]
    if not gaps and total_weight == 0:
        gaps = ["缺少岗位核心技能关键词，需要补充更完整 JD 或人工确认职责"]
    if "LCA软件" in gaps and group_present(jd_text, ["建模", "模型", "SimaPro", "GaBi", "openLCA"]):
        gaps = ["LCA软件实操证据"] + [gap for gap in gaps if gap != "LCA软件"]

    return {
        "score": score,
        "tech_score": tech_score,
        "semantic_score": semantic_score,
        "value_score": int(np.clip(value_score, 0, 100)),
        "risk_score": int(np.clip(risk_score, 0, 100)),
        "matched": matched,
        "missing": gaps,
        "reasons": reasons,
        "salary_label": salary_label,
        "required": required,
        "evidence": evidence[:8],
        "related_matched": related_matched,
    }


def next_action_for_job(score: int, jd_analysis: dict[str, Any], job_type: str, internship_opening: str, missing: list[str]) -> str:
    if score >= 85:
        return "立即定制简历并投递"
    if score >= 72:
        if missing:
            return "补强关键词后投递：" + " / ".join(missing[:3])
        return "可直接投递"
    if internship_opening == "是" and score >= 60:
        return "可作为实习积累，先问清项目内容"
    if jd_analysis.get("value", {}).get("is_generic_esg"):
        return "除非缺实习，否则不优先"
    return "暂存观察，优先投更贴合岗位"


def analyze_batch_jd_records(
    records: list[dict[str, str]],
    profile_text: str,
    resume_text: str = "",
) -> pd.DataFrame:
    rows = []
    preferences = load_target_preferences()
    profile_text = effective_profile_text(profile_text)
    resume_text = normalize_text(resume_text)
    has_resume = bool(resume_text.strip())
    deduped_records = dedupe_jd_records(records)

    for index, record in enumerate(deduped_records, start=1):
        text = record["text"]
        jd = analyze_jd(text)
        job_type = detect_job_type_safe(text)
        internship_opening = detect_internship_opening_safe(text)
        fresh = detect_fresh_graduate_safe(text)
        fit = evaluate_profile_fit_v2(jd, profile_text, job_type, fresh, preferences, resume_text=resume_text, fast=True)
        if has_resume:
            resume_match = match_resume_to_jd(jd, resume_text, profile_text, preferences, fast=True)
            combined_score = int(round(fit["score"] * 0.55 + resume_match["score"] * 0.45))
            resume_gaps = resume_match.get("missing_skills", [])[:6]
            resume_hits = resume_match.get("matched_skills", [])[:8]
            resume_score = resume_match.get("score", "")
            resume_semantic = resume_match.get("semantic_score", "")
            resume_evidence = resume_match.get("evidence", [])[:6]
        else:
            combined_score = fit["score"]
            resume_gaps = []
            resume_hits = []
            resume_score = ""
            resume_semantic = ""
            resume_evidence = []

        fields = extract_card_fields_v2(record, text, jd)
        record_url = preferred_jd_record_url(record)
        if not record_url:
            record_url = next((url for url in extract_urls(text) if detail_url_is_usable(url, "")), "")
        company = fields["company"]
        title = fields["title"]
        salary = fields["salary"]
        location = fields["location"]
        education = fields["education"]
        experience = fields["experience"]
        region_info = classify_region(location, text)
        if region_info["省份"] != "未识别":
            fit["reasons"].append(f"省份识别：{region_info['省份']}")
        confidence = extraction_confidence(fields, text, jd)
        confidence_score = extraction_confidence_score(fields, text, jd)
        gaps = list(dict.fromkeys((fit["missing"] or []) + resume_gaps))
        recommendation = batch_recommendation(combined_score, jd, job_type, fresh)
        next_action = next_action_for_job(combined_score, jd, job_type, internship_opening, gaps)

        rows.append(
            {
                "序号": index,
                "意向匹配度": combined_score,
                "技术匹配分": fit.get("tech_score", ""),
                "语义相似分": fit.get("semantic_score", ""),
                "简历匹配分": resume_score,
                "JD简历相似分": resume_semantic,
                "岗位价值分": fit.get("value_score", ""),
                "风险分": fit.get("risk_score", ""),
                "置信度": confidence,
                "读取质量": confidence_score,
                "投递建议": recommendation,
                "下一步动作": next_action,
                "公司": company,
                "岗位": title,
                "类型": job_type,
                "是否招实习": internship_opening,
                "应届生": fresh,
                "薪资": salary,
                "地点": location,
                **region_info,
                "学历": education,
                "经验": experience,
                "岗位分类": jd.get("category", ""),
                "高价值": "是" if jd.get("value", {}).get("is_high_value") else "否",
                "低价值风险": "是" if jd.get("value", {}).get("is_generic_esg") else "否",
                "风险标签": " / ".join(jd.get("value", {}).get("risk_tags", [])),
                "技能关键词": " / ".join(jd_skill_list(jd)[:10]),
                "匹配原因": "；".join(fit["reasons"]),
                "简历命中": " / ".join(resume_hits),
                "简历证据": "；".join(resume_evidence),
                "主要缺口": " / ".join(gaps[:8]),
                "薪资判断": fit.get("salary_label", ""),
                "来源": record.get("source", ""),
                "页面": record.get("page_index", ""),
                "链接": record_url,
                "JD原文": text,
                "原文片段": text[:300],
            }
        )

    if not rows:
        return pd.DataFrame()
    df = pd.DataFrame(rows)
    return df.sort_values(["意向匹配度", "高价值"], ascending=[False, False]).reset_index(drop=True)


def password_hash(password: str, salt: bytes | None = None) -> str:
    salt = salt or secrets.token_bytes(16)
    digest = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt, PASSWORD_HASH_ITERATIONS)
    return "pbkdf2_sha256${}${}${}".format(
        PASSWORD_HASH_ITERATIONS,
        base64.b64encode(salt).decode("ascii"),
        base64.b64encode(digest).decode("ascii"),
    )


def verify_password(password: str, stored_hash: str) -> bool:
    try:
        algorithm, iterations, salt_b64, digest_b64 = stored_hash.split("$", 3)
        if algorithm != "pbkdf2_sha256":
            return False
        salt = base64.b64decode(salt_b64.encode("ascii"))
        expected = base64.b64decode(digest_b64.encode("ascii"))
        actual = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt, int(iterations))
        return hmac.compare_digest(actual, expected)
    except Exception:
        return False


def normalize_email(email: str) -> str:
    return normalize_text(email).strip().lower()


def current_user_id() -> int | None:
    user = st.session_state.get(AUTH_SESSION_KEY)
    if isinstance(user, dict) and user.get("id"):
        return int(user["id"])
    return None


def require_user_id() -> int:
    user_id = current_user_id()
    if not user_id:
        raise RuntimeError("用户未登录")
    return int(user_id)


def create_app_user(email: str, password: str, display_name: str = "") -> tuple[bool, str]:
    init_db()
    email = normalize_email(email)
    display_name = normalize_text(display_name) or email.split("@")[0]
    if not email or "@" not in email:
        return False, "请输入有效邮箱。"
    if len(password) < 8:
        return False, "密码至少 8 位。"
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    try:
        with sqlite3.connect(DB_PATH) as conn:
            cursor = conn.execute(
                """
                INSERT INTO app_users (email, password_hash, display_name, created_at, last_login_at)
                VALUES (?, ?, ?, ?, ?)
                """,
                (email, password_hash(password), display_name, now, now),
            )
            user_id = int(cursor.lastrowid)
        st.session_state[AUTH_SESSION_KEY] = {"id": user_id, "email": email, "display_name": display_name}
        clear_runtime_data_cache()
        return True, "注册成功。"
    except sqlite3.IntegrityError:
        return False, "该邮箱已注册，请直接登录。"


def authenticate_app_user(email: str, password: str) -> tuple[bool, str]:
    init_db()
    email = normalize_email(email)
    with sqlite3.connect(DB_PATH) as conn:
        row = conn.execute(
            "SELECT id, email, password_hash, display_name FROM app_users WHERE email = ?",
            (email,),
        ).fetchone()
        if not row or not verify_password(password, str(row[2])):
            return False, "邮箱或密码不正确。"
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        conn.execute("UPDATE app_users SET last_login_at = ? WHERE id = ?", (now, int(row[0])))
    st.session_state[AUTH_SESSION_KEY] = {"id": int(row[0]), "email": str(row[1]), "display_name": str(row[3] or row[1])}
    clear_runtime_data_cache()
    return True, "登录成功。"


def logout_app_user() -> None:
    st.session_state.pop(AUTH_SESSION_KEY, None)
    for key in ["active_profile_id", "active_resume_id"]:
        st.session_state.pop(key, None)
    clear_runtime_data_cache()


def init_db() -> None:
    try:
        if st.session_state.get("_careerpilot_db_initialized") == DB_SCHEMA_VERSION and DB_PATH.exists():
            return
    except Exception:
        pass
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS app_users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                email TEXT NOT NULL UNIQUE,
                password_hash TEXT NOT NULL,
                display_name TEXT,
                created_at TEXT,
                last_login_at TEXT
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS applications (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER,
                company TEXT,
                job_title TEXT,
                salary TEXT,
                location TEXT,
                category TEXT,
                match_score INTEGER,
                is_high_value INTEGER,
                is_generic_esg INTEGER,
                applied INTEGER DEFAULT 0,
                interview_status TEXT DEFAULT '未开始',
                offer_status TEXT DEFAULT '无',
                notes TEXT,
                created_at TEXT
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS user_profiles (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER,
                name TEXT NOT NULL,
                content TEXT NOT NULL,
                is_default INTEGER DEFAULT 0,
                created_at TEXT,
                updated_at TEXT
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS user_resumes (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER,
                name TEXT NOT NULL,
                content TEXT NOT NULL,
                is_default INTEGER DEFAULT 0,
                created_at TEXT,
                updated_at TEXT
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS recruitment_posts (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER,
                fingerprint TEXT NOT NULL,
                company TEXT,
                job_title TEXT,
                job_type TEXT,
                fresh_graduate TEXT,
                salary TEXT,
                location TEXT,
                standard_city TEXT,
                province TEXT,
                region TEXT,
                region_priority TEXT,
                education TEXT,
                experience TEXT,
                company_tier TEXT,
                category TEXT,
                high_value TEXT,
                generic_esg TEXT,
                skills TEXT,
                source TEXT,
                snippet TEXT,
                first_seen TEXT,
                last_seen TEXT
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS app_settings (
                user_id INTEGER,
                key TEXT NOT NULL,
                value TEXT NOT NULL,
                updated_at TEXT
            )
            """
        )
        ensure_user_columns(conn)
        ensure_recruitment_region_columns(conn)
        ensure_application_queue_columns(conn)
        ensure_profile_table_ready(conn)
        cleanup_legacy_database_state(conn)
    try:
        st.session_state["_careerpilot_db_initialized"] = DB_SCHEMA_VERSION
    except Exception:
        pass


def ensure_recruitment_region_columns(conn: sqlite3.Connection) -> None:
    existing = {row[1] for row in conn.execute("PRAGMA table_info(recruitment_posts)").fetchall()}
    for column in ["standard_city", "province", "region", "region_priority"]:
        if column not in existing:
            conn.execute(f"ALTER TABLE recruitment_posts ADD COLUMN {column} TEXT")


def ensure_user_columns(conn: sqlite3.Connection) -> None:
    for table in ["applications", "user_profiles", "user_resumes", "recruitment_posts", "app_settings"]:
        existing = {row[1] for row in conn.execute(f"PRAGMA table_info({table})").fetchall()}
        if "user_id" not in existing:
            conn.execute(f"ALTER TABLE {table} ADD COLUMN user_id INTEGER")


def ensure_application_queue_columns(conn: sqlite3.Connection) -> None:
    existing = {row[1] for row in conn.execute("PRAGMA table_info(applications)").fetchall()}
    for column in ["queue_date", "next_action"]:
        if column not in existing:
            conn.execute(f"ALTER TABLE applications ADD COLUMN {column} TEXT")


def seed_app_settings(conn: sqlite3.Connection) -> None:
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    conn.execute(
        """
        INSERT OR IGNORE INTO app_settings (key, value, updated_at)
        VALUES (?, ?, ?)
        """,
        ("target_preferences", json.dumps(DEFAULT_TARGET_PREFERENCES, ensure_ascii=False), now),
    )


def ensure_profile_table_ready(conn: sqlite3.Connection) -> None:
    repair_default_profiles(conn)


def cleanup_legacy_database_state(conn: sqlite3.Connection) -> None:
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    rows = conn.execute("SELECT id, name, content FROM user_resumes").fetchall()
    for resume_id, name, content in rows:
        if is_legacy_template_resume(repair_mojibake_text(str(name)), repair_mojibake_text(str(content))):
            conn.execute("DELETE FROM user_resumes WHERE id = ?", (resume_id,))

    raw_row = conn.execute("SELECT value FROM app_settings WHERE key = 'target_preferences'").fetchone()
    if raw_row:
        raw = repair_mojibake_text(str(raw_row[0]))
        try:
            data = json.loads(raw)
        except Exception:
            data = {}
        if not isinstance(data, dict):
            data = {}
        clean = dict(DEFAULT_TARGET_PREFERENCES)
        clean.update(data)
        clean["target_cities"] = split_preference_items(
            split_preference_items(clean.get("target_cities", [])) + split_preference_items(clean.get("extra_cities", ""))
        )
        clean["preferred_industries"] = split_preference_items(
            split_preference_items(clean.get("preferred_industries", [])) + split_preference_items(clean.get("extra_industries", ""))
        )
        if clean["target_cities"] == LEGACY_AUTO_TARGET_CITIES:
            clean["target_cities"] = []
        if clean["preferred_industries"] == LEGACY_AUTO_TARGET_INDUSTRIES:
            clean["preferred_industries"] = []
        if str(clean.get("avoid_keywords", "")).strip() == LEGACY_AUTO_AVOID_KEYWORDS:
            clean["avoid_keywords"] = ""
        if str(clean.get("notes", "")).strip() == LEGACY_AUTO_NOTES:
            clean["notes"] = ""
        clean["extra_cities"] = ""
        clean["extra_industries"] = ""
        conn.execute(
            """
            INSERT INTO app_settings (key, value, updated_at)
            VALUES ('target_preferences', ?, ?)
            ON CONFLICT(key) DO UPDATE SET
                value = excluded.value,
                updated_at = excluded.updated_at
            """,
            (json.dumps(clean, ensure_ascii=False), now),
        )


def repair_default_profiles(conn: sqlite3.Connection) -> None:
    rows = conn.execute("SELECT id, name, content, is_default FROM user_profiles ORDER BY id ASC").fetchall()
    if not rows:
        return
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    for profile_id, name, content, is_default in rows:
        if looks_mojibake(str(name)) or looks_mojibake(str(content)):
            conn.execute(
                """
                UPDATE user_profiles
                SET name = ?, content = ?, updated_at = ?
                WHERE id = ?
                """,
                (DEFAULT_TARGET_INTENTION_NAME, "", now, profile_id),
            )


def enforce_single_target_intention(conn: sqlite3.Connection) -> None:
    rows = conn.execute("SELECT id, is_default FROM user_profiles ORDER BY is_default DESC, id ASC").fetchall()
    if not rows:
        return
    keep_id = int(rows[0][0])
    conn.execute("UPDATE user_profiles SET is_default = CASE WHEN id = ? THEN 1 ELSE 0 END", (keep_id,))
    conn.execute("DELETE FROM user_profiles WHERE id <> ?", (keep_id,))


def clear_runtime_data_cache() -> None:
    try:
        st.cache_data.clear()
    except Exception:
        pass
    for key in ["report_excel_bytes", "report_pdf_bytes"]:
        st.session_state.pop(key, None)


def load_user_profiles() -> pd.DataFrame:
    init_db()
    user_id = require_user_id()
    with sqlite3.connect(DB_PATH) as conn:
        df = pd.read_sql_query("SELECT * FROM user_profiles WHERE user_id = ? ORDER BY is_default DESC, id ASC", conn, params=(user_id,))
    return repair_dataframe_text(df)


def get_active_profile() -> dict[str, Any]:
    profiles = load_user_profiles()
    if profiles.empty:
        return {"id": None, "name": "未设置", "content": "", "is_default": 1}

    active_id = st.session_state.get("active_profile_id")
    if active_id is None or active_id not in profiles["id"].tolist():
        default_rows = profiles[profiles["is_default"] == 1]
        active_row = default_rows.iloc[0] if not default_rows.empty else profiles.iloc[0]
        st.session_state.active_profile_id = int(active_row["id"])
    else:
        active_row = profiles[profiles["id"] == active_id].iloc[0]

    name = repair_mojibake_text(str(active_row["name"]))
    content = repair_mojibake_text(str(active_row["content"]))
    return {
        "id": int(active_row["id"]),
        "name": name,
        "content": content,
        "is_default": int(active_row["is_default"]),
    }


def save_user_profile(profile_id: int, name: str, content: str) -> None:
    user_id = require_user_id()
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute(
            """
            UPDATE user_profiles
            SET name = ?, content = ?, updated_at = ?
            WHERE id = ? AND user_id = ?
            """,
            (name.strip(), content.strip(), now, profile_id, user_id),
        )
    clear_runtime_data_cache()


def create_user_profile(name: str, content: str) -> int:
    user_id = require_user_id()
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    with sqlite3.connect(DB_PATH) as conn:
        if int(bool(conn.execute("SELECT COUNT(*) FROM user_profiles WHERE user_id = ?", (user_id,)).fetchone()[0])):
            is_default = 0
        else:
            is_default = 1
        cursor = conn.execute(
            """
            INSERT INTO user_profiles (user_id, name, content, is_default, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (user_id, name.strip(), content.strip(), is_default, now, now),
        )
        new_id = int(cursor.lastrowid)
    clear_runtime_data_cache()
    return new_id


def delete_user_profile(profile_id: int) -> bool:
    user_id = require_user_id()
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute("DELETE FROM user_profiles WHERE id = ? AND user_id = ?", (profile_id, user_id))
    if st.session_state.get("active_profile_id") == profile_id:
        st.session_state.pop("active_profile_id", None)
    clear_runtime_data_cache()
    return True


def profile_text_for_analysis() -> str:
    active = get_active_profile()
    return normalize_text(active.get("content", "") + "\n" + target_preferences_text())


def load_user_resumes() -> pd.DataFrame:
    init_db()
    user_id = require_user_id()
    with sqlite3.connect(DB_PATH) as conn:
        df = pd.read_sql_query("SELECT * FROM user_resumes WHERE user_id = ? ORDER BY is_default DESC, id ASC", conn, params=(user_id,))
    df = repair_dataframe_text(df)
    if df.empty:
        return df
    template_mask = df.apply(lambda row: is_legacy_template_resume(str(row.get("name", "")), str(row.get("content", ""))), axis=1)
    return df.loc[~template_mask].reset_index(drop=True)


def is_legacy_template_resume(name: str, content: str) -> bool:
    clean_name = normalize_text(name)
    clean_content = normalize_text(content)
    if not clean_content:
        return True
    marker_hits = sum(1 for marker in LEGACY_TEMPLATE_RESUME_MARKERS if marker in clean_name or marker in clean_content)
    if marker_hits >= 1 and len(clean_content) < 1600:
        return True
    placeholder_count = len(re.findall(r"X{2,}|xx|XX|【[^】]*(?:示例|模板|填写|替换)[^】]*】", clean_content))
    return placeholder_count >= 2


def create_user_resume(name: str, content: str, is_default: int = 0) -> int:
    user_id = require_user_id()
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    with sqlite3.connect(DB_PATH) as conn:
        if not conn.execute("SELECT COUNT(*) FROM user_resumes WHERE user_id = ?", (user_id,)).fetchone()[0]:
            is_default = 1
        cursor = conn.execute(
            """
            INSERT INTO user_resumes (user_id, name, content, is_default, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (user_id, name.strip(), content.strip(), int(is_default), now, now),
        )
        new_id = int(cursor.lastrowid)
    clear_runtime_data_cache()
    return new_id


def save_user_resume(resume_id: int, name: str, content: str) -> None:
    user_id = require_user_id()
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute(
            """
            UPDATE user_resumes
            SET name = ?, content = ?, updated_at = ?
            WHERE id = ? AND user_id = ?
            """,
            (name.strip(), content.strip(), now, resume_id, user_id),
        )
    clear_runtime_data_cache()


def delete_user_resume(resume_id: int) -> bool:
    user_id = require_user_id()
    with sqlite3.connect(DB_PATH) as conn:
        resume_count = conn.execute("SELECT COUNT(*) FROM user_resumes WHERE user_id = ?", (user_id,)).fetchone()[0]
        if resume_count <= 1:
            return False
        conn.execute("DELETE FROM user_resumes WHERE id = ? AND user_id = ?", (resume_id, user_id))
    if st.session_state.get("active_resume_id") == resume_id:
        st.session_state.pop("active_resume_id", None)
    clear_runtime_data_cache()
    return True


def load_app_setting(key: str, default: str = "") -> str:
    init_db()
    user_id = require_user_id()
    with sqlite3.connect(DB_PATH) as conn:
        row = conn.execute("SELECT value FROM app_settings WHERE user_id = ? AND key = ?", (user_id, key)).fetchone()
    return repair_mojibake_text(str(row[0])) if row else default


def save_app_setting(key: str, value: str) -> None:
    user_id = require_user_id()
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    with sqlite3.connect(DB_PATH) as conn:
        existing = conn.execute("SELECT 1 FROM app_settings WHERE user_id = ? AND key = ?", (user_id, key)).fetchone()
        if existing:
            conn.execute(
                "UPDATE app_settings SET value = ?, updated_at = ? WHERE user_id = ? AND key = ?",
                (value.strip(), now, user_id, key),
            )
        else:
            conn.execute(
                "INSERT INTO app_settings (user_id, key, value, updated_at) VALUES (?, ?, ?, ?)",
                (user_id, key, value.strip(), now),
            )
    clear_runtime_data_cache()


def load_target_preferences() -> dict[str, Any]:
    raw = load_app_setting("target_preferences", json.dumps(DEFAULT_TARGET_PREFERENCES, ensure_ascii=False))
    try:
        data = json.loads(raw)
        if not isinstance(data, dict):
            data = {}
    except Exception:
        data = {}
    merged = dict(DEFAULT_TARGET_PREFERENCES)
    merged.update(data)
    merged["target_roles"] = split_preference_items(merged.get("target_roles", []))
    merged["target_cities"] = split_preference_items(
        split_preference_items(merged.get("target_cities", [])) + split_preference_items(merged.get("extra_cities", ""))
    )
    merged["preferred_industries"] = split_preference_items(
        split_preference_items(merged.get("preferred_industries", [])) + split_preference_items(merged.get("extra_industries", ""))
    )
    merged["job_keywords"] = split_preference_items(merged.get("job_keywords", []))
    if merged["target_cities"] == LEGACY_AUTO_TARGET_CITIES:
        merged["target_cities"] = []
    if merged["preferred_industries"] == LEGACY_AUTO_TARGET_INDUSTRIES:
        merged["preferred_industries"] = []
    if str(merged.get("avoid_keywords", "")).strip() == LEGACY_AUTO_AVOID_KEYWORDS:
        merged["avoid_keywords"] = ""
    if str(merged.get("notes", "")).strip() == LEGACY_AUTO_NOTES:
        merged["notes"] = ""
    merged["extra_cities"] = ""
    merged["extra_industries"] = ""
    for key in ["min_monthly_salary", "min_daily_salary"]:
        try:
            merged[key] = int(merged.get(key) or DEFAULT_TARGET_PREFERENCES[key])
        except Exception:
            merged[key] = DEFAULT_TARGET_PREFERENCES[key]
    merged["accept_remote"] = bool(merged.get("accept_remote"))
    merged["accept_nationwide"] = bool(merged.get("accept_nationwide"))
    return merged


def save_target_preferences(preferences: dict[str, Any]) -> None:
    clean = dict(DEFAULT_TARGET_PREFERENCES)
    clean.update(preferences)
    clean["target_roles"] = split_preference_items(clean.get("target_roles", []))
    clean["target_cities"] = split_preference_items(clean.get("target_cities", []))
    clean["preferred_industries"] = split_preference_items(clean.get("preferred_industries", []))
    clean["job_keywords"] = split_preference_items(clean.get("job_keywords", []))
    clean["extra_cities"] = ""
    clean["extra_industries"] = ""
    clean["avoid_keywords"] = "、".join(split_preference_items(clean.get("avoid_keywords", "")))
    clean["notes"] = str(clean.get("notes", "")).strip()
    clean["min_monthly_salary"] = int(clean.get("min_monthly_salary") or DEFAULT_TARGET_PREFERENCES["min_monthly_salary"])
    clean["min_daily_salary"] = int(clean.get("min_daily_salary") or DEFAULT_TARGET_PREFERENCES["min_daily_salary"])
    clean["accept_remote"] = bool(clean.get("accept_remote"))
    clean["accept_nationwide"] = bool(clean.get("accept_nationwide"))
    save_app_setting("target_preferences", json.dumps(clean, ensure_ascii=False))


def target_preferences_text(preferences: dict[str, Any] | None = None) -> str:
    preferences = preferences or load_target_preferences()
    role_items = split_preference_items(preferences.get("target_roles", []))
    city_items = split_preference_items(preferences.get("target_cities", []))
    industry_items = split_preference_items(preferences.get("preferred_industries", []))
    keyword_items = split_preference_items(preferences.get("job_keywords", []))
    lines = [
        "求职方向：" + (" / ".join(role_items) if role_items else "不限"),
        "意向城市：" + (" / ".join(city_items) if city_items else "不限"),
        "目标行业：" + (" / ".join(industry_items) if industry_items else "不限"),
    ]
    if keyword_items:
        lines.append("岗位关键词：" + " / ".join(keyword_items))
    if preferences.get("avoid_keywords"):
        lines.append("排除关键词：" + str(preferences.get("avoid_keywords")).strip())
    return "\n".join(lines)


def compact_list_text(items: list[str], empty: str = "未设置", limit: int = 3) -> str:
    clean_items = split_preference_items(items)
    if not clean_items:
        return empty
    label = " / ".join(clean_items[:limit])
    if len(clean_items) > limit:
        label += f" 等 {len(clean_items)} 项"
    return label


def compact_profile_summary(content: str, max_length: int = 72) -> str:
    lines = normalize_resume_lines(content)
    if not lines:
        return "未设置"
    summary = lines[0]
    return summary[:max_length] + ("..." if len(summary) > max_length else "")


def get_active_resume() -> dict[str, Any]:
    resumes = load_user_resumes()
    if resumes.empty:
        return {"id": None, "name": "未设置简历", "content": "", "is_default": 1, "updated_at": ""}
    active_id = st.session_state.get("active_resume_id")
    if active_id is None or active_id not in resumes["id"].tolist():
        default_rows = resumes[resumes["is_default"] == 1]
        active_row = default_rows.iloc[0] if not default_rows.empty else resumes.iloc[0]
        st.session_state.active_resume_id = int(active_row["id"])
    else:
        active_row = resumes[resumes["id"] == active_id].iloc[0]
    return {
        "id": int(active_row["id"]),
        "name": str(active_row["name"]),
        "content": str(active_row["content"]),
        "is_default": int(active_row["is_default"]),
        "updated_at": str(active_row.get("updated_at", "")),
    }


def resume_text_for_analysis() -> str:
    active_resume = get_active_resume()
    return str(active_resume.get("content") or "")


def current_resume_fingerprint() -> str:
    return content_fingerprint(resume_text_for_analysis())


def add_application(record: dict[str, Any]) -> None:
    user_id = require_user_id()
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute(
            """
            INSERT INTO applications (
                user_id, company, job_title, salary, location, category, match_score,
                is_high_value, is_generic_esg, applied, interview_status,
                offer_status, notes, queue_date, next_action, created_at
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                user_id,
                record.get("company", ""),
                record.get("job_title", ""),
                record.get("salary", ""),
                record.get("location", ""),
                record.get("category", ""),
                int(record.get("match_score", 0) or 0),
                int(bool(record.get("is_high_value", False))),
                int(bool(record.get("is_generic_esg", False))),
                int(bool(record.get("applied", False))),
                record.get("interview_status", "未开始"),
                record.get("offer_status", "无"),
                record.get("notes", ""),
                record.get("queue_date", ""),
                record.get("next_action", ""),
                datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            ),
        )
    clear_runtime_data_cache()


def load_applications() -> pd.DataFrame:
    init_db()
    user_id = require_user_id()
    with sqlite3.connect(DB_PATH) as conn:
        df = pd.read_sql_query("SELECT * FROM applications WHERE user_id = ? ORDER BY id DESC", conn, params=(user_id,))
    df = repair_dataframe_text(df)
    if not df.empty:
        for col in ["is_high_value", "is_generic_esg", "applied"]:
            df[col] = df[col].astype(bool)
    return df


def save_application_edits(df: pd.DataFrame) -> None:
    if df.empty:
        return
    user_id = require_user_id()
    with sqlite3.connect(DB_PATH) as conn:
        for _, row in df.iterrows():
            conn.execute(
                """
                UPDATE applications
                SET company=?, job_title=?, salary=?, location=?, category=?,
                    match_score=?, is_high_value=?, is_generic_esg=?, applied=?,
                    interview_status=?, offer_status=?, notes=?, queue_date=?, next_action=?
                WHERE id=? AND user_id=?
                """,
                (
                    str(row.get("company", "")),
                    str(row.get("job_title", "")),
                    str(row.get("salary", "")),
                    str(row.get("location", "")),
                    str(row.get("category", "")),
                    int(row.get("match_score", 0) or 0),
                    int(bool(row.get("is_high_value", False))),
                    int(bool(row.get("is_generic_esg", False))),
                    int(bool(row.get("applied", False))),
                    str(row.get("interview_status", "未开始")),
                    str(row.get("offer_status", "无")),
                    str(row.get("notes", "")),
                    str(row.get("queue_date", "")),
                    str(row.get("next_action", "")),
                    int(row["id"]),
                    user_id,
                ),
            )
    clear_runtime_data_cache()


def delete_application(row_id: int) -> None:
    user_id = require_user_id()
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute("DELETE FROM applications WHERE id = ? AND user_id = ?", (row_id, user_id))
    clear_runtime_data_cache()


def today_label() -> str:
    return datetime.now().strftime("%Y-%m-%d")


def add_current_jd_to_today_queue(jd_analysis: dict[str, Any], resume_match: dict[str, Any] | None = None) -> None:
    basic = jd_analysis.get("basic", {})
    plan = build_job_action_plan(jd_analysis, resume_match)
    add_application(
        {
            "company": basic.get("公司名", ""),
            "job_title": basic.get("岗位名", "") or jd_analysis.get("category", ""),
            "salary": basic.get("薪资", ""),
            "location": basic.get("地点", ""),
            "category": jd_analysis.get("category", ""),
            "match_score": int(resume_match.get("score", 0) if resume_match else 0),
            "is_high_value": jd_analysis.get("value", {}).get("is_high_value", False),
            "is_generic_esg": jd_analysis.get("value", {}).get("is_generic_esg", False),
            "applied": False,
            "interview_status": "未开始",
            "offer_status": "无",
            "notes": "今日队列：来自单条JD分析",
            "queue_date": today_label(),
            "next_action": str(plan.get("actions", ["定制简历并确认投递渠道"])[0]),
        }
    )


def add_batch_rows_to_today_queue(rows: pd.DataFrame) -> int:
    if rows is None or rows.empty:
        return 0
    count = 0
    for _, row in rows.iterrows():
        add_application(
            {
                "company": row.get("公司", ""),
                "job_title": row.get("岗位", ""),
                "salary": row.get("薪资", ""),
                "location": row.get("地点", ""),
                "category": row.get("岗位分类", ""),
                "match_score": int(row.get("意向匹配度", 0) or 0),
                "is_high_value": str(row.get("高价值", "")) == "是",
                "is_generic_esg": str(row.get("低价值风险", "")) == "是",
                "applied": False,
                "interview_status": "未开始",
                "offer_status": "无",
                "notes": f"今日队列：{row.get('来源', '批量JD筛选')}",
                "queue_date": today_label(),
                "next_action": row.get("下一步动作", "定制简历并确认投递渠道"),
            }
        )
        count += 1
    return count


def build_report_frames() -> dict[str, pd.DataFrame]:
    jd_analysis = st.session_state.get("jd_analysis")
    resume_match = st.session_state.get("resume_match")
    gap_analysis = st.session_state.get("gap_analysis")
    interview_analysis = st.session_state.get("interview_analysis")
    internship_analysis = st.session_state.get("internship_analysis")
    custom_resume = st.session_state.get("custom_resume")
    recruitment_monitor = st.session_state.get("recruitment_monitor")
    offer_prediction = st.session_state.get("offer_prediction")
    batch_jd_analysis = st.session_state.get("batch_jd_analysis")

    frames: dict[str, pd.DataFrame] = {}
    active_profile = get_active_profile()
    active_resume = get_active_resume()
    target_meta = st.session_state.get("target_jd_meta", {})
    if jd_analysis:
        action_plan = build_job_action_plan(jd_analysis, resume_match)
        basic = jd_analysis.get("basic", {})
        frames["投递包摘要"] = pd.DataFrame(
            [
                {
                    "目标岗位": basic.get("岗位名", ""),
                    "目标公司": basic.get("公司名", ""),
                    "来源": target_meta.get("source", ""),
                    "岗位分类": jd_analysis.get("category", ""),
                    "岗位价值": "高" if jd_analysis.get("value", {}).get("is_high_value") else "待判",
                    "低价值风险": "高" if jd_analysis.get("value", {}).get("is_generic_esg") else "低",
                    "简历版本": active_resume.get("name", ""),
                    "简历匹配度": resume_match.get("score", "") if resume_match else "待分析",
                    "投递判断": action_plan["decision"],
                    "投递前动作": " / ".join(action_plan["actions"]),
                    "面试准备": " / ".join(action_plan["interview_focus"]),
                }
            ]
        )
    frames["目标意向"] = pd.DataFrame([{"目标意向名称": active_profile["name"], "目标意向内容": active_profile["content"]}])
    frames["目标偏好"] = pd.DataFrame(
        [{"偏好": line} for line in normalize_resume_lines(target_preferences_text())]
    )
    if active_resume.get("content"):
        parsed_resume = parse_resume_content(active_resume["content"])
        quality = resume_parse_quality(parsed_resume)
        frames["当前简历状态"] = pd.DataFrame(
            [
                {
                    "简历名称": active_resume.get("name", ""),
                    "更新时间": active_resume.get("updated_at", ""),
                    "有效行数": len(parsed_resume.get("lines", [])),
                    "识别板块": len(parsed_resume.get("sections", {})),
                    "技能命中": " / ".join(parsed_resume.get("skills", [])[:12]),
                    "质量": quality["label"],
                    "质量说明": quality["detail"],
                }
            ]
        )
    if jd_analysis:
        frames["岗位基本信息"] = pd.DataFrame([jd_analysis["basic"]])
        frames["技能关键词"] = jd_analysis["skills"]
        frames["岗位判断"] = pd.DataFrame(
            [
                {
                    "岗位分类": jd_analysis.get("category", ""),
                    "高价值词命中": jd_analysis.get("value", {}).get("high_score", ""),
                    "低价值风险词命中": jd_analysis.get("value", {}).get("low_score", ""),
                    "判断": jd_analysis.get("value", {}).get("label", ""),
                }
            ]
        )
    if resume_match:
        frames["简历匹配"] = pd.DataFrame(
            [
                {
                    "匹配度": resume_match["score"],
                    "匹配技能": " / ".join(resume_match["matched_skills"]),
                    "缺口技能": " / ".join(resume_match["missing_skills"]),
                    "优势": " / ".join(resume_match["strengths"]),
                    "缺口说明": " / ".join(resume_match["gap_examples"]),
                }
            ]
        )
    if batch_jd_analysis is not None and not batch_jd_analysis.empty:
        frames["批量JD分析"] = public_export_df(batch_jd_analysis)
    if custom_resume:
        frames["定制简历"] = pd.DataFrame(
            [
                {
                    "目标岗位": custom_resume["job_title"],
                    "岗位分类": custom_resume["category"],
                    "关键词": custom_resume["keyword_line"],
                    "直接可用版本": custom_resume.get("ready_resume_text", ""),
                    "摘要": custom_resume["summary"],
                    "核心能力": " / ".join(custom_resume.get("skills_section", [])),
                    "Bullet": " / ".join(custom_resume.get("experience_bullets", custom_resume["bullets"])),
                    "投递说明": custom_resume.get("application_pitch", ""),
                    "风险提醒": " / ".join(custom_resume["risk_notes"]),
                }
            ]
        )
    if recruitment_monitor is not None and not recruitment_monitor.empty:
        frames["行业招聘监测"] = public_export_df(recruitment_monitor)
    if offer_prediction:
        frames["Offer预测"] = pd.DataFrame(
            [
                {
                    "简历通过率": offer_prediction["简历通过率"],
                    "进入面试概率": offer_prediction["进入面试概率"],
                    "拿Offer概率": offer_prediction["拿 offer 概率"],
                    "短板数量": offer_prediction["shortcomings"],
                    "影响因素": " / ".join(offer_prediction["drivers"]),
                    "提升建议": " / ".join(offer_prediction["actions"]),
                    "投递步骤": " / ".join(offer_prediction.get("application_steps", [])),
                }
            ]
        )
    if gap_analysis:
        if gap_analysis.get("action_rows"):
            frames["不足行动包"] = pd.DataFrame(gap_analysis["action_rows"])
        if gap_analysis.get("weekly_plan"):
            frames["一周补强计划"] = pd.DataFrame(gap_analysis["weekly_plan"])
        gap_rows = []
        for category, items in gap_analysis["current_gaps"].items():
            for item in items:
                gap_rows.append({"分类": category, "不足项": item})
        frames["不足清单"] = pd.DataFrame(gap_rows)
        frames["努力方向"] = pd.DataFrame(gap_analysis["priorities"], columns=["优先级", "建议"])
    if interview_analysis:
        rows = []
        for answer in interview_analysis.get("answer_templates", []):
            rows.append({"分类": "回答框架", "面经问题": answer})
        for category, questions in interview_analysis["buckets"].items():
            for question in questions:
                rows.append({"分类": category, "面经问题": question})
        for category, questions in interview_analysis["generated_questions"].items():
            for question in questions:
                rows.append({"分类": category, "面经问题": question})
        frames["面试问题"] = pd.DataFrame(rows)
    if internship_analysis:
        frames["实习评估"] = pd.DataFrame(
            [
                {
                    "实习价值评分": internship_analysis["score"],
                    "结论": internship_analysis["verdict"],
                    "决策建议": internship_analysis["decision"],
                    "有利原因": " / ".join(internship_analysis["reasons"]),
                    "风险点": " / ".join(internship_analysis["risks"]),
                    "建议产出": " / ".join(internship_analysis["recommended_outputs"]),
                    "沟通话术": " / ".join(internship_analysis.get("negotiation_script", [])),
                    "第一周计划": " / ".join(internship_analysis.get("first_week_plan", [])),
                }
            ]
        )
        frames["实习维度评分"] = internship_analysis["dimension_scores"]
    return frames


def build_excel_report() -> bytes:
    output = io.BytesIO()
    frames = build_report_frames()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        if not frames:
            pd.DataFrame([{"提示": "暂无分析结果"}]).to_excel(writer, sheet_name="报告", index=False)
        for name, df in frames.items():
            safe_name = name[:31]
            df.to_excel(writer, sheet_name=safe_name, index=False)
    return output.getvalue()


def build_pdf_report() -> bytes | None:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=16 * mm, leftMargin=16 * mm, topMargin=16 * mm, bottomMargin=16 * mm)
        styles = getSampleStyleSheet()
        for style in styles.byName.values():
            style.fontName = "STSong-Light"
        story = [Paragraph("CareerPilot 投递包", styles["Title"]), Spacer(1, 8)]

        frames = build_report_frames()
        for name, df in frames.items():
            story.append(Paragraph(name, styles["Heading2"]))
            show_df = df.copy().astype(str).head(20)
            if show_df.empty:
                story.append(Paragraph("暂无数据", styles["BodyText"]))
                story.append(Spacer(1, 6))
                continue
            table_data = [show_df.columns.tolist()] + show_df.values.tolist()
            table = Table(table_data, repeatRows=1)
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EAF2F8")),
                        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                        ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"),
                        ("FONTSIZE", (0, 0), (-1, -1), 8),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 8))
        doc.build(story)
        return buffer.getvalue()
    except Exception:
        return None


def safe_html(value: Any) -> str:
    return html.escape(str(value), quote=True)


def render_main_workspace_nav() -> str:
    current = st.session_state.get("main_workspace", "jd")
    if current not in MAIN_WORKSPACE_LABELS:
        current = "jd"
        st.session_state.main_workspace = current

    return st.radio(
        "主工作区",
        list(MAIN_WORKSPACE_LABELS.keys()),
        format_func=lambda key: MAIN_WORKSPACE_LABELS.get(key, key),
        horizontal=True,
        label_visibility="collapsed",
        key="main_workspace",
    )


def render_app_styles() -> None:
    st.markdown(
        """
        <style>
        :root {
            --cp-bg: var(--background-color, #f6f7f4);
            --cp-panel: var(--secondary-background-color, #ffffff);
            --cp-panel-soft: color-mix(in srgb, var(--cp-panel) 92%, var(--cp-bg));
            --cp-sidebar: color-mix(in srgb, var(--cp-panel) 88%, var(--cp-bg));
            --cp-border: color-mix(in srgb, var(--text-color, #1d2520) 16%, transparent);
            --cp-text: var(--text-color, #1d2520);
            --cp-muted: color-mix(in srgb, var(--cp-text) 62%, transparent);
            --cp-teal: var(--primary-color, #0f766e);
            --cp-teal-dark: color-mix(in srgb, var(--cp-teal) 82%, var(--cp-text));
            --cp-accent-soft: color-mix(in srgb, var(--cp-teal) 13%, var(--cp-panel));
            --cp-gold: #b7791f;
            --cp-red: #b42318;
        }

        .stApp {
            background: var(--cp-bg);
            color: var(--cp-text);
        }

        header[data-testid="stHeader"] {
            background: transparent;
        }

        .block-container {
            max-width: 1280px;
            padding-top: 3.25rem;
            padding-bottom: 3rem;
        }

        section[data-testid="stSidebar"] {
            background: var(--cp-sidebar);
            border-right: 1px solid var(--cp-border);
        }

        section[data-testid="stSidebar"] [data-testid="stMarkdownContainer"] p,
        section[data-testid="stSidebar"] label {
            color: var(--cp-text);
        }

        h1, h2, h3 {
            color: var(--cp-text);
            letter-spacing: 0;
        }

        h2, h3 {
            margin-top: 1.25rem;
        }

        .cp-overview-grid {
            display: grid;
            grid-template-columns: repeat(4, minmax(0, 1fr));
            gap: 10px;
            margin: 0 0 12px;
        }

        .cp-overview-card {
            border: 1px solid var(--cp-border);
            border-radius: 8px;
            background: var(--cp-panel-soft);
            padding: 10px 12px;
        }

        .cp-overview-label {
            color: var(--cp-muted);
            font-size: 12px;
            font-weight: 700;
            margin-bottom: 5px;
        }

        .cp-overview-value {
            color: var(--cp-text);
            font-size: 16px;
            font-weight: 800;
            margin-bottom: 4px;
        }

        .cp-overview-copy {
            color: var(--cp-muted);
            font-size: 12px;
            line-height: 1.45;
        }

        .cp-workspace-head {
            display: flex;
            align-items: flex-end;
            justify-content: space-between;
            gap: 16px;
            margin: 4px 0 14px;
        }

        .cp-workspace-eyebrow {
            color: var(--cp-teal-dark);
            font-size: 12px;
            font-weight: 800;
            margin-bottom: 5px;
        }

        .cp-workspace-title {
            color: var(--cp-text);
            font-size: 23px;
            line-height: 1.25;
            font-weight: 850;
            margin: 0 0 5px;
        }

        .cp-workspace-copy {
            color: var(--cp-muted);
            font-size: 13px;
            line-height: 1.55;
            max-width: 760px;
        }

        .cp-mode-note {
            color: var(--cp-muted);
            font-size: 12px;
            line-height: 1.45;
            text-align: right;
            max-width: 250px;
        }

        .cp-panel-title {
            color: var(--cp-text);
            font-size: 17px;
            font-weight: 800;
            margin-bottom: 4px;
        }

        .cp-panel-copy {
            color: var(--cp-muted);
            font-size: 13px;
            line-height: 1.55;
            margin-bottom: 10px;
        }

        .cp-note {
            color: var(--cp-muted);
            font-size: 12px;
            line-height: 1.65;
            margin: 8px 0 12px;
        }

        .cp-note strong {
            color: var(--cp-text);
            font-weight: 800;
        }

        [data-testid="stDataFrame"] [role="columnheader"],
        [data-testid="stDataEditor"] [role="columnheader"] {
            justify-content: center !important;
            text-align: center !important;
        }

        [data-testid="stDataFrame"],
        [data-testid="stDataEditor"] {
            font-size: 11.5px;
        }

        .cp-empty-state {
            border: 1px dashed var(--cp-border);
            border-radius: 8px;
            background: var(--cp-panel-soft);
            padding: 18px 18px 16px;
            min-height: 238px;
        }

        .cp-empty-state-title {
            color: var(--cp-text);
            font-size: 18px;
            font-weight: 820;
            margin-bottom: 8px;
        }

        .cp-empty-state-copy {
            color: var(--cp-muted);
            font-size: 13px;
            line-height: 1.6;
            margin-bottom: 12px;
        }

        .cp-mini-steps {
            display: grid;
            gap: 8px;
        }

        .cp-mini-step {
            display: grid;
            grid-template-columns: 26px 1fr;
            gap: 8px;
            align-items: start;
            color: var(--cp-text);
            font-size: 13px;
            line-height: 1.45;
        }

        .cp-mini-step-index {
            width: 24px;
            height: 24px;
            border-radius: 999px;
            display: inline-flex;
            align-items: center;
            justify-content: center;
            background: var(--cp-accent-soft);
            color: var(--cp-teal-dark);
            font-size: 12px;
            font-weight: 800;
        }

        .cp-decision-card {
            border: 1px solid var(--cp-border);
            border-left: 4px solid var(--cp-teal);
            border-radius: 8px;
            background: var(--cp-panel);
            padding: 14px 16px;
            margin-bottom: 10px;
        }

        .cp-decision-label {
            color: var(--cp-muted);
            font-size: 12px;
            font-weight: 800;
            margin-bottom: 5px;
        }

        .cp-decision-value {
            color: var(--cp-text);
            font-size: 24px;
            line-height: 1.2;
            font-weight: 860;
            margin-bottom: 8px;
        }

        .cp-decision-copy {
            color: var(--cp-muted);
            font-size: 13px;
            line-height: 1.55;
        }

        .cp-fact-grid {
            display: grid;
            grid-template-columns: repeat(2, minmax(0, 1fr));
            gap: 8px;
            margin: 10px 0;
        }

        .cp-fact {
            border: 1px solid var(--cp-border);
            border-radius: 8px;
            background: var(--cp-panel-soft);
            padding: 9px 10px;
        }

        .cp-fact-label {
            color: var(--cp-muted);
            font-size: 12px;
            font-weight: 700;
            margin-bottom: 3px;
        }

        .cp-fact-value {
            color: var(--cp-text);
            font-size: 14px;
            line-height: 1.35;
            font-weight: 780;
        }

        div[data-testid="stMetric"] {
            background: var(--cp-panel);
            border: 1px solid var(--cp-border);
            border-radius: 8px;
            padding: 12px 14px;
        }

        div[data-testid="stMetricLabel"] {
            color: var(--cp-muted);
        }

        div[data-testid="stMetricValue"] {
            color: var(--cp-text);
            font-weight: 750;
        }

        .stTabs [data-baseweb="tab-list"] {
            gap: 6px;
            border-bottom: 1px solid var(--cp-border);
        }

        .stTabs [data-baseweb="tab"] {
            border-radius: 6px 6px 0 0;
            color: var(--cp-muted);
            font-size: 15px;
            padding-left: 14px;
            padding-right: 14px;
        }

        .stTabs [aria-selected="true"] {
            color: var(--cp-teal-dark);
            background: var(--cp-accent-soft);
        }

        div[data-testid="stExpander"] {
            background: var(--cp-panel);
            border: 1px solid var(--cp-border);
            border-radius: 8px;
        }

        div[data-testid="stDataFrame"],
        div[data-testid="stDataEditor"] {
            border: 1px solid var(--cp-border);
            border-radius: 8px;
            overflow: hidden;
        }

        .stButton > button,
        .stDownloadButton > button,
        button[kind="secondary"] {
            border-radius: 6px;
            border: 1px solid var(--cp-border);
            color: var(--cp-text);
            background: var(--cp-panel);
        }

        .stButton > button:hover,
        .stDownloadButton > button:hover {
            border-color: var(--cp-teal);
            color: var(--cp-teal-dark);
        }

        .stButton > button[kind="primary"] {
            background: var(--cp-teal);
            border-color: var(--cp-teal);
            color: #ffffff;
        }

        div[data-testid="stAlert"] {
            border-radius: 8px;
            border: 1px solid var(--cp-border);
        }

        textarea,
        input,
        div[data-baseweb="select"] > div,
        div[data-baseweb="base-input"] {
            border-radius: 6px;
        }

        @media (max-width: 760px) {
            .block-container {
                padding-left: 1rem;
                padding-right: 1rem;
            }
            .cp-overview-grid {
                grid-template-columns: 1fr;
            }
            .cp-workspace-head {
                display: block;
            }
            .cp-mode-note {
                text-align: left;
                margin-top: 8px;
                max-width: none;
            }
            .cp-fact-grid {
                grid-template-columns: 1fr;
            }
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def public_export_df(df: pd.DataFrame) -> pd.DataFrame:
    hidden_columns = {
        "JD原文",
        "原文片段",
        "fingerprint",
        "区域",
        "地域优先级",
        "地域修正",
        "泛ESG风险",
        "来源",
        "页面",
        "最佳意向",
        "最佳意向分",
        "读取质量",
        "技术匹配分",
        "语义相似分",
        "简历匹配分",
        "JD简历相似分",
        "岗位价值分",
        "风险分",
        "置信度",
        "高价值词命中",
        "低价值风险词命中",
    }
    output = df.drop(columns=[col for col in hidden_columns if col in df.columns], errors="ignore")
    output = output.rename(columns={"意向匹配度": "综合分"})
    if "链接" in output.columns:
        output = output[[col for col in output.columns if col != "链接"] + ["链接"]]
    return output


USER_HIDDEN_TABLE_COLUMNS = {
    "记录ID",
    "信息集ID",
    "id",
    "fingerprint",
    "JD原文",
    "原文片段",
    "raw_text",
    "text",
    "path",
    "页面",
    "来源路径",
    "技术匹配分",
    "语义相似分",
    "简历匹配分",
    "JD简历相似分",
    "岗位价值分",
    "风险分",
    "置信度",
    "高价值词命中",
    "低价值风险词命中",
    "最佳意向",
    "最佳意向分",
    "地域修正",
    "地域优先级",
    "读取质量",
}


USER_TABLE_RENAMES = {
    "意向匹配度": "综合分",
    "match_score": "综合分",
    "company": "公司",
    "job_title": "岗位",
    "url": "链接",
    "salary": "薪资",
    "location": "地点",
    "category": "分类",
    "is_high_value": "高价值",
    "is_generic_esg": "低价值风险",
    "applied": "已投递",
    "interview_status": "面试状态",
    "offer_status": "Offer状态",
    "queue_date": "队列日期",
    "next_action": "下一步动作",
    "notes": "备注",
    "created_at": "创建时间",
    "updated_at": "更新时间",
}


def user_table_df(df: pd.DataFrame, columns: list[str] | None = None, *, hide_scores: bool = True) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame()
    table = df.copy()
    if columns:
        table = table[[col for col in columns if col in table.columns]]
    hidden = set(USER_HIDDEN_TABLE_COLUMNS)
    if hide_scores:
        hidden.update(col for col in table.columns if str(col).endswith("分") and col not in {"意向匹配度", "综合分", "价值评分", "匹配度"})
        hidden.update(col for col in table.columns if "相似" in str(col) or "命中" in str(col))
    table = table.drop(columns=[col for col in hidden if col in table.columns], errors="ignore")
    table = table.rename(columns={key: value for key, value in USER_TABLE_RENAMES.items() if key in table.columns})
    if "链接" in table.columns:
        table = table[[col for col in table.columns if col != "链接"] + ["链接"]]
    return table


def user_table_row_height(df: pd.DataFrame) -> int:
    if df is None or df.empty:
        return 28
    visible = df.drop(columns=[col for col in ["链接"] if col in df.columns], errors="ignore")
    text_lengths = visible.astype(str).map(lambda value: len(normalize_text(value))).to_numpy().flatten()
    max_len = int(max(text_lengths)) if len(text_lengths) else 0
    if max_len >= 120:
        return 44
    if max_len >= 56:
        return 38
    if max_len >= 24:
        return 32
    return 28


def user_table_height(df: pd.DataFrame, row_height: int) -> int | str:
    if df is None or df.empty:
        return "auto"
    return min(420, 36 + max(1, len(df)) * row_height)


def user_table_column_width(series: pd.Series, column: str) -> str:
    if column == "链接":
        return "medium"
    max_len = int(series.astype(str).map(lambda value: len(normalize_text(value))).max()) if len(series) else 0
    if max_len >= 64:
        return "large"
    if max_len >= 18:
        return "medium"
    return "small"


def user_table_column_config(df: pd.DataFrame, extra: dict[str, Any] | None = None) -> dict[str, Any]:
    config: dict[str, Any] = {}
    short_center_columns = {
        "序号",
        "综合分",
        "匹配度",
        "投递建议",
        "类型",
        "是否招实习",
        "是否招应届生",
        "应届生",
        "薪资",
        "地点",
        "高价值",
        "高价值岗位",
        "低价值风险",
        "省份",
        "标准城市",
        "学历",
        "经验",
        "队列日期",
        "已投递",
        "面试状态",
        "Offer状态",
    }
    for column in df.columns:
        width = user_table_column_width(df[column], str(column))
        max_len = int(df[column].astype(str).map(lambda value: len(normalize_text(value))).max()) if len(df[column]) else 0
        alignment = "center" if str(column) in short_center_columns or max_len <= 10 else "left"
        if column == "链接":
            config[column] = st.column_config.LinkColumn(column, width=width, alignment="center")
        elif pd.api.types.is_numeric_dtype(df[column]):
            config[column] = st.column_config.NumberColumn(column, width=width, alignment="center")
        else:
            config[column] = st.column_config.TextColumn(column, width=width, alignment=alignment)
    if extra:
        config.update(extra)
    return config


def render_user_dataframe(
    df: pd.DataFrame,
    columns: list[str] | None = None,
    *,
    hide_index: bool = True,
    column_config: dict[str, Any] | None = None,
    key: str | None = None,
) -> None:
    table = user_table_df(df, columns)
    row_height = user_table_row_height(table)
    st.dataframe(
        table,
        width="stretch",
        height=user_table_height(table, row_height),
        hide_index=hide_index,
        row_height=row_height,
        column_config=user_table_column_config(table, column_config),
        key=key,
    )


def render_risk_logic_note() -> None:
    st.markdown(
        """
        <div class="cp-note">
            <div><strong>注：</strong></div>
            <div><strong>综合分：</strong>由目标岗位/行业/城市偏好、薪资与地域匹配、岗位价值、应届/实习友好度、当前简历匹配度共同加权，并扣除低价值、职责不清、经验门槛过高等风险。</div>
            <div><strong>投递建议：</strong>P1表示优先投递，P2表示值得投递但需补关键词或项目证据，P3表示可作为备选，谨慎表示低价值或门槛风险较高。</div>
            <div><strong>风险标签：</strong>来自JD中的职责范围、经验学历门槛、销售/行政杂务信号、产出是否明确等文本特征。</div>
            <div><strong>高价值：</strong>表示岗位更容易沉淀可复用项目、业务结果或技术/分析能力。</div>
            <div><strong>低价值风险：</strong>表示职责可能偏杂、偏执行或与目标方向弱相关。</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def target_jd_row_label(row: pd.Series) -> str:
    score = row.get("意向匹配度", row.get("匹配度", ""))
    company = str(row.get("公司", "") or "未识别公司")
    title = str(row.get("岗位", "") or row.get("岗位名", "") or "未识别岗位")
    prefix = f"{score}分｜" if str(score).strip() else ""
    return f"{prefix}{company}｜{title}"[:140]


BATCH_QUICK_VIEWS = ["全部", "强投优先", "目标城市", "目标行业", "排除低价值", "应届/实习友好", "缺口最小"]


def batch_row_text(row: pd.Series, columns: list[str]) -> str:
    return normalize_text(" ".join(str(row.get(column, "")) for column in columns))


def batch_gap_count(value: Any) -> int:
    text = normalize_text(str(value or ""))
    if not text or text in {"暂无", "无", "nan"}:
        return 0
    return len(split_preference_items(text))


def apply_batch_quick_view(df: pd.DataFrame, view_name: str, preferences: dict[str, Any] | None = None) -> pd.DataFrame:
    if df is None or df.empty or view_name == "全部":
        return df
    preferences = preferences or load_target_preferences()
    view_df = df.copy()
    risk_column = "低价值风险" if "低价值风险" in view_df.columns else "泛ESG风险"

    if view_name == "强投优先" and "投递建议" in view_df.columns:
        view_df = view_df[view_df["投递建议"].astype(str).str.contains("P1|P2", regex=True, na=False)]
    elif view_name == "目标城市":
        target_cities = split_preference_items(preferences.get("target_cities", []))
        if target_cities:
            mask = view_df.apply(
                lambda row: bool(city_match_label(batch_row_text(row, ["地点", "标准城市", "省份", "匹配原因"]), target_cities)),
                axis=1,
            )
            view_df = view_df[mask]
    elif view_name == "目标行业":
        target_industries = split_preference_items(preferences.get("preferred_industries", []))
        if target_industries:
            mask = view_df.apply(
                lambda row: bool(industry_preference_hits(batch_row_text(row, ["公司", "岗位", "岗位分类", "技能关键词", "匹配原因", "JD原文"]), target_industries)),
                axis=1,
            )
            view_df = view_df[mask]
    elif view_name == "排除低价值" and risk_column in view_df.columns:
        view_df = view_df[view_df[risk_column].astype(str) != "是"]
    elif view_name == "应届/实习友好":
        masks = []
        if "是否招实习" in view_df.columns:
            masks.append(view_df["是否招实习"].astype(str) == "是")
        if "应届生" in view_df.columns:
            masks.append(view_df["应届生"].astype(str).str.contains("是|应届|校招|经验不限|在校生|实习", regex=True, na=False))
        if masks:
            combined = masks[0]
            for mask in masks[1:]:
                combined = combined | mask
            view_df = view_df[combined]
    elif view_name == "缺口最小" and "主要缺口" in view_df.columns:
        view_df = view_df.assign(_gap_count=view_df["主要缺口"].apply(batch_gap_count))
        sort_cols = ["_gap_count"]
        ascending = [True]
        if "意向匹配度" in view_df.columns:
            sort_cols.append("意向匹配度")
            ascending.append(False)
        view_df = view_df.sort_values(sort_cols, ascending=ascending).drop(columns=["_gap_count"])

    if "意向匹配度" in view_df.columns and view_name in {"强投优先", "目标城市", "目标行业", "排除低价值", "应届/实习友好"}:
        view_df = view_df.sort_values("意向匹配度", ascending=False)
    return view_df


def render_target_jd_picker(view_df: pd.DataFrame, key_prefix: str, source_label: str) -> None:
    if view_df is None or view_df.empty:
        return
    candidates = view_df.reset_index(drop=True).head(100)
    if "JD原文" not in candidates.columns and "原文片段" not in candidates.columns:
        return
    with st.expander("从当前结果设为目标 JD", expanded=True):
        st.markdown("#### 设为当前目标 JD")
        selected_index = st.selectbox(
            "从当前结果中选择一个岗位做深度分析",
            list(range(len(candidates))),
            format_func=lambda index: target_jd_row_label(candidates.iloc[index]),
            key=f"{key_prefix}_target_jd_select",
        )
        selected_row = candidates.iloc[selected_index]
        if st.button("设为当前目标 JD 并分析", type="primary", key=f"{key_prefix}_set_target_jd"):
            jd_text = str(selected_row.get("JD原文") or selected_row.get("原文片段") or "")
            if len(normalize_text(jd_text)) < 20:
                st.warning("这条记录缺少足够的 JD 原文，建议打开详情页或重新用插件保存。")
            else:
                set_current_target_jd(jd_text, source_label, target_jd_row_label(selected_row))
                st.success("已设为当前目标 JD。现在可以进入“简历工作台”做匹配、定制简历和不足分析。")


def render_state_alerts() -> None:
    jd_fingerprint = current_jd_fingerprint()
    resume_fingerprint = current_resume_fingerprint()
    target_meta = st.session_state.get("target_jd_meta")
    if target_meta and jd_fingerprint:
        st.caption(f"当前目标 JD：{target_meta.get('title', '目标JD')}｜来源：{target_meta.get('source', '')}｜更新时间：{target_meta.get('updated_at', '')}")

    if st.session_state.get("resume_match"):
        stale_parts = []
        if st.session_state.get("resume_match_jd_fingerprint") != jd_fingerprint:
            stale_parts.append("目标 JD 已变化")
        if st.session_state.get("resume_match_resume_fingerprint") != resume_fingerprint:
            stale_parts.append("当前简历已变化")
        if stale_parts:
            st.warning("当前简历匹配结果可能已过期：" + "、".join(stale_parts) + "。请重新运行简历匹配。")

    if st.session_state.get("custom_resume"):
        stale_parts = []
        if st.session_state.get("custom_resume_jd_fingerprint") != jd_fingerprint:
            stale_parts.append("目标 JD 已变化")
        if st.session_state.get("custom_resume_resume_fingerprint") != resume_fingerprint:
            stale_parts.append("当前简历已变化")
        if stale_parts:
            st.info("定制简历片段可能已过期：" + "、".join(stale_parts) + "。建议重新生成。")


def build_job_action_plan(jd_analysis: dict[str, Any] | None, resume_match: dict[str, Any] | None = None) -> dict[str, list[str] | str]:
    if not jd_analysis:
        return {"decision": "先完成目标 JD 分析。", "actions": [], "bullets": [], "evidence_gaps": [], "interview_focus": []}

    value = jd_analysis.get("value", {})
    skills = jd_skill_list(jd_analysis)
    missing = resume_match.get("missing_skills", []) if resume_match else skills[:5]
    score = int(resume_match.get("score", 0) if resume_match else 0)

    if resume_match:
        if score >= 82 and value.get("is_high_value"):
            decision = "优先投递，投前重点打磨简历证据。"
        elif score >= 72:
            decision = "值得投递，先补齐关键证据再投。"
        elif score >= 60:
            decision = "可作为备选，投递前需要明显改简历。"
        else:
            decision = "暂不优先，除非公司或机会本身特别值得。"
    elif value.get("is_generic_esg"):
        decision = "先谨慎核实职责，避免纯披露/运营型岗位。"
    elif value.get("is_high_value"):
        decision = "值得深看，下一步应立刻用当前简历做匹配。"
    else:
        decision = "先补充 JD 详情，再判断是否投入简历定制时间。"

    actions = []
    if value.get("is_high_value"):
        actions.append("打开 JD 原文确认职责里是否有真实项目、核心任务、可量化指标和明确交付物；如果只有模糊支持性描述，降低优先级。")
    if value.get("is_generic_esg"):
        actions.append("先向招聘方确认是否能接触数据、需求、客户、代码、报告或项目核心环节；如果只是杂务、排版或宣传，不作为主投。")
    if missing:
        actions.append("投递前在简历中为这些词各补一句证据：" + " / ".join(missing[:5]))
    if skills:
        actions.append("把简历摘要第一句和最近一段项目经历改到这些关键词上：" + " / ".join(skills[:5]))
    actions.append("投递前生成定制简历片段；投递后在投递库记录投递渠道、简历版本和下次跟进日期。")

    bullets = generate_resume_bullets(jd_analysis, resume_match)[:3] if resume_match else []
    evidence_gaps = missing[:5] or ["需要补充可量化的项目产出、工具使用和交付场景。"]
    interview_focus = (skills[:5] or [jd_analysis.get("category", "岗位方向")]) + ["项目复盘", "动机匹配"]
    return {
        "decision": decision,
        "actions": actions[:5],
        "bullets": bullets,
        "evidence_gaps": evidence_gaps,
        "interview_focus": interview_focus[:6],
    }


def job_decision_label(jd_analysis: dict[str, Any], resume_match: dict[str, Any] | None = None) -> tuple[str, str]:
    value = jd_analysis.get("value", {})
    score = int(resume_match.get("score", 0) if resume_match else 0)
    if value.get("is_generic_esg"):
        return "谨慎核实", "warning"
    if resume_match:
        if score >= 82 and value.get("is_high_value"):
            return "强投", "success"
        if score >= 72 or value.get("is_high_value"):
            return "可投", "success"
        if score >= 60:
            return "备选", "info"
        return "暂缓", "warning"
    if value.get("is_high_value"):
        return "值得深看", "success"
    return "待判断", "info"


def job_decision_reasons(jd_analysis: dict[str, Any], resume_match: dict[str, Any] | None = None) -> list[str]:
    basic = jd_analysis.get("basic", {})
    value = jd_analysis.get("value", {})
    reasons = [
        f"岗位分类：{jd_analysis.get('category', '待判断')}",
        str(value.get("label", "需要结合 JD 详情判断")),
    ]
    location = basic.get("地点")
    if location and location != "未识别":
        reasons.append(f"工作地点：{location}")
    salary = basic.get("薪资")
    if salary and salary != "未识别":
        reasons.append(f"薪资信息：{salary}")
    preference_score, preference_reasons = target_preference_adjustment(
        jd_analysis.get("raw_text", ""),
        jd_analysis,
        "",
        load_target_preferences(),
    )
    if preference_score or preference_reasons:
        reasons.extend(preference_reasons[:3])
    if resume_match:
        reasons.append(f"当前简历匹配度：{int(resume_match.get('score', 0))}")
        missing = resume_match.get("missing_skills", [])
        if missing:
            reasons.append("待补证据：" + " / ".join(missing[:4]))
    return list(dict.fromkeys(reason for reason in reasons if reason))[:6]


def build_resume_focus_plan(jd_analysis: dict[str, Any], resume_match: dict[str, Any] | None = None) -> dict[str, list[str]]:
    skills = jd_skill_list(jd_analysis)
    category = jd_analysis.get("category", "目标岗位")
    matched = resume_match.get("matched_skills", []) if resume_match else []
    missing = resume_match.get("missing_skills", []) if resume_match else skills[:6]

    category_focus = {
        "数据/商业分析岗": ["数据分析项目", "指标体系或业务复盘", "SQL / Excel / Python 工具使用"],
        "产品岗": ["需求分析或产品优化项目", "用户研究/竞品分析", "方案落地和指标验证"],
        "运营岗": ["运营活动或增长复盘", "用户分层与转化指标", "内容/社群/活动执行结果"],
        "咨询/项目岗": ["项目推进案例", "行业研究和报告输出", "客户沟通与交付物沉淀"],
        "研发/工程岗": ["工程项目或代码作品", "模块实现和测试验证", "性能/稳定性/协作交付"],
        "财务/金融岗": ["财务或经营分析案例", "模型/报表/预算相关产出", "风险判断和业务解释"],
        "法务/合规岗": ["合规风险案例", "合同/政策/监管分析", "跨部门沟通记录"],
        "供应链/采购岗": ["供应链优化或采购分析", "供应商评估", "库存/成本/交付指标"],
        "LCA技术岗": ["LCA建模案例", "清单数据收集与质量判断", "方法学和边界设定说明"],
        "产品碳足迹岗": ["产品碳足迹项目", "ISO14067 / PEF 方法理解", "数据质量和报告输出"],
        "碳核算岗": ["碳盘查或排放核算案例", "活动数据和排放因子处理", "GHG / ISO14064 理解"],
        "CBAM/出海合规岗": ["CBAM数据清单案例", "欧盟合规政策研究", "客户材料和申报逻辑整理"],
    }
    highlight = category_focus.get(category, [f"{category}相关项目", "可量化成果", "跨部门协作和交付"])
    if matched:
        highlight = list(dict.fromkeys(highlight + [f"{skill}已有证据" for skill in matched[:4]]))

    add_keywords = list(dict.fromkeys(missing[:6] + [skill for skill in skills[:6] if skill not in matched]))
    if not add_keywords:
        add_keywords = ["项目目标", "个人动作", "量化结果", "复盘改进"]

    downplay = ["课程罗列", "泛泛写学习能力", "只写协助/支持但没有个人产出"]
    if jd_analysis.get("value", {}).get("is_generic_esg"):
        downplay.append("宣传、排版、会议组织等杂务经历")
    if category not in ["市场/销售岗"]:
        downplay.append("与岗位无关的销售指标或纯拉新表述")

    summary_seed = " / ".join((matched or skills)[:4]) or category
    summary = [
        f"摘要第一句对齐 {category}，突出 {summary_seed}。",
        "第二句写清楚你能交付的结果：分析表、方案文档、项目推进、报告或可展示作品。",
    ]
    bullets = generate_resume_bullets(jd_analysis, resume_match)[:3]
    if not bullets:
        bullets = [
            "围绕目标问题完成信息收集、数据/材料整理、分析判断和结果复盘，输出可用于业务决策的交付物。",
            "将岗位关键词对应到具体经历中，补充工具、方法、协作对象和量化结果。",
        ]
    return {
        "highlight": highlight[:6],
        "add_keywords": add_keywords[:8],
        "downplay": list(dict.fromkeys(downplay))[:5],
        "summary": summary,
        "bullets": bullets,
    }


def render_resume_focus_plan(jd_analysis: dict[str, Any], resume_match: dict[str, Any] | None = None) -> None:
    plan = build_resume_focus_plan(jd_analysis, resume_match)
    st.markdown("#### 投前简历改法")
    focus_cols = st.columns(3)
    with focus_cols[0]:
        st.markdown("##### 突出")
        for item in plan["highlight"]:
            st.write(f"- {item}")
    with focus_cols[1]:
        st.markdown("##### 补关键词")
        for item in plan["add_keywords"]:
            st.write(f"- {item}")
    with focus_cols[2]:
        st.markdown("##### 少写")
        for item in plan["downplay"]:
            st.write(f"- {item}")
    with st.expander("可直接改的摘要和 bullet", expanded=False):
        st.markdown("##### 摘要")
        for item in plan["summary"]:
            st.write(f"- {item}")
        st.markdown("##### 经历 bullet")
        for item in plan["bullets"]:
            st.write(f"- {item}")


def render_score_breakdown(jd_analysis: dict[str, Any], resume_match: dict[str, Any] | None = None) -> None:
    st.markdown("#### 最终判断")
    value = jd_analysis.get("value", {}) if jd_analysis else {}
    if resume_match:
        st.metric("最终匹配度", f"{int(resume_match.get('score', 0))} / 100")
    else:
        st.info("完成简历匹配后，这里会显示最终匹配度。")
    render_risk_logic_note()
    notes = []
    if value.get("is_high_value"):
        notes.append("岗位价值较高：职责里出现项目、数据、交付物、咨询、分析或可沉淀成果等信号。")
    if value.get("is_generic_esg"):
        notes.append("低价值风险较高：JD 中出现职责不清、杂务、销售伪装或成果不可量化等信号。")
    risk_tags = value.get("risk_tags") or []
    if risk_tags:
        notes.append("风险标签：" + " / ".join(risk_tags[:6]))
    for item in notes or ["暂无明显风险标签；仍建议结合公司、岗位正文和面试信息确认。"]:
        st.write(f"- {item}")


def render_job_action_plan(jd_analysis: dict[str, Any] | None, resume_match: dict[str, Any] | None = None) -> None:
    plan = build_job_action_plan(jd_analysis, resume_match)
    st.markdown("#### 下一步行动")
    st.info(str(plan["decision"]))
    action_tab, evidence_tab, resume_tab = st.tabs(["投递前动作", "证据与面试", "简历改写"])
    with action_tab:
        for item in plan["actions"]:
            st.write(f"- {item}")
    with evidence_tab:
        for item in plan["evidence_gaps"][:3]:
            st.write(f"- {item}")
        focus = " / ".join(plan["interview_focus"])
        if focus:
            st.caption("面试准备：" + focus)
    with resume_tab:
        if plan["bullets"]:
            for item in plan["bullets"]:
                st.write(f"- {item}")
        else:
            st.caption("完成简历匹配后，这里会给出可直接替换的 bullet。")


def render_resume_workspace_heading() -> None:
    st.markdown(
        """
        <div class="cp-workspace-head">
            <div>
                <div class="cp-workspace-eyebrow">简历工作台</div>
                <div class="cp-workspace-title">围绕当前目标岗位，把简历改到能投递</div>
                <div class="cp-workspace-copy">
                    先用你保存的当前简历跑匹配，再看优势、缺口和可直接替换的句子；定制简历和不足清单都基于同一个目标 JD。
                </div>
            </div>
            <div class="cp-mode-note">建议路径：匹配分析 → 定制简历 → 不足清单</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_resume_empty_state(jd_analysis: dict[str, Any] | None, active_resume: dict[str, Any], has_match: bool) -> None:
    if not jd_analysis:
        title = "还没有目标岗位"
        copy = "先在岗位工作台导入并分析一条 JD，简历工作台才能判断这份简历该突出什么。"
        steps = ["进入岗位工作台粘贴 JD。", "完成 JD 分析并设为当前目标。", "回到这里运行简历匹配。"]
    elif not active_resume.get("content", "").strip():
        title = "还没有当前简历"
        copy = "左侧当前简历为空；保存真实简历后，这里才会用于匹配和定制。"
        steps = ["在左侧当前简历上传或粘贴内容。", "保存为当前简历。", "回到这里点击匹配分析。"]
    elif not has_match:
        title = "等待运行简历匹配"
        copy = "点击左侧按钮后，这里会集中显示匹配度、已覆盖证据、必须补齐的关键词和投递前动作。"
        steps = ["确认左侧当前简历是最新版本。", "点击使用当前简历分析匹配。", "根据右侧结论进入定制简历。"]
    else:
        return
    step_html = "".join(
        f'<div class="cp-mini-step"><span class="cp-mini-step-index">{idx}</span><span>{safe_html(item)}</span></div>'
        for idx, item in enumerate(steps, start=1)
    )
    st.markdown(
        '<div class="cp-empty-state">'
        f'<div class="cp-empty-state-title">{safe_html(title)}</div>'
        f'<div class="cp-empty-state-copy">{safe_html(copy)}</div>'
        f'<div class="cp-mini-steps">{step_html}</div>'
        "</div>",
        unsafe_allow_html=True,
    )


def render_resume_match_snapshot(resume_match: dict[str, Any]) -> None:
    score = int(resume_match.get("score", 0))
    if score >= 78:
        verdict = "可以进入定制简历"
    elif score >= 60:
        verdict = "需要补齐关键证据"
    else:
        verdict = "先补基础匹配"
    facts = [
        ("已覆盖", f"{len(resume_match.get('matched_skills', []))} 项"),
        ("待补证据", f"{len(resume_match.get('missing_skills', []))} 项"),
    ]
    fact_cards = "".join(
        '<div class="cp-fact">'
        f'<div class="cp-fact-label">{safe_html(name)}</div>'
        f'<div class="cp-fact-value">{safe_html(value)}</div>'
        "</div>"
        for name, value in facts
    )
    st.markdown(
        '<div class="cp-decision-card">'
        '<div class="cp-decision-label">当前匹配度</div>'
        f'<div class="cp-decision-value">{score} / 100</div>'
        f'<div class="cp-decision-copy">{safe_html(verdict)}</div>'
        "</div>"
        f'<div class="cp-fact-grid">{fact_cards}</div>',
        unsafe_allow_html=True,
    )
    st.markdown("##### 可主打的优势")
    for item in resume_match.get("strengths", [])[:4]:
        st.write(f"- {item}")
    st.markdown("##### 投前必须处理")
    for item in resume_match.get("gap_examples", [])[:4]:
        st.write(f"- {item}")


def render_decision_workspace_heading() -> None:
    st.markdown(
        """
        <div class="cp-workspace-head">
            <div>
                <div class="cp-workspace-eyebrow">求职决策</div>
                <div class="cp-workspace-title">把机会判断、行动队列和投递状态放在一起</div>
                <div class="cp-workspace-copy">
                    这里不再重复分析 JD，而是回答三个问题：这个机会推进概率如何、今天先处理哪几个、投递状态有没有跟上。
                </div>
            </div>
            <div class="cp-mode-note">建议路径：Offer 预测 → 今日队列 → 投递管理</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_decision_empty_state(message: str, steps: list[str]) -> None:
    step_html = "".join(
        f'<div class="cp-mini-step"><span class="cp-mini-step-index">{idx}</span><span>{safe_html(item)}</span></div>'
        for idx, item in enumerate(steps, start=1)
    )
    st.markdown(
        '<div class="cp-empty-state">'
        '<div class="cp-empty-state-title">还不能生成决策</div>'
        f'<div class="cp-empty-state-copy">{safe_html(message)}</div>'
        f'<div class="cp-mini-steps">{step_html}</div>'
        "</div>",
        unsafe_allow_html=True,
    )


def render_offer_prediction_snapshot(result: dict[str, Any] | None) -> None:
    if not result:
        render_decision_empty_state(
            "完成左侧参数后，右侧会显示简历通过率、面试概率、Offer 概率和下一步动作。",
            ["确认目标 JD。", "选择公司层级、学历匹配和英文能力。", "点击预测通过率。"],
        )
        return
    pass_rate = int(result.get("简历通过率", 0))
    interview_rate = int(result.get("进入面试概率", 0))
    offer_rate = int(result.get("拿 offer 概率", 0))
    if offer_rate >= 35:
        verdict = "优先推进"
    elif interview_rate >= 35:
        verdict = "值得投递，先补证据"
    else:
        verdict = "备选或降优先级"
    facts = [
        ("简历通过", f"{pass_rate}%"),
        ("进入面试", f"{interview_rate}%"),
        ("拿 Offer", f"{offer_rate}%"),
        ("短板数量", str(result.get("shortcomings", 0))),
    ]
    fact_cards = "".join(
        '<div class="cp-fact">'
        f'<div class="cp-fact-label">{safe_html(name)}</div>'
        f'<div class="cp-fact-value">{safe_html(value)}</div>'
        "</div>"
        for name, value in facts
    )
    st.markdown(
        '<div class="cp-decision-card">'
        '<div class="cp-decision-label">推进判断</div>'
        f'<div class="cp-decision-value">{safe_html(verdict)}</div>'
        '<div class="cp-decision-copy">把概率当作排序工具，不当作绝对结果；真正影响结果的是投递版简历和岗位证据。</div>'
        "</div>"
        f'<div class="cp-fact-grid">{fact_cards}</div>',
        unsafe_allow_html=True,
    )
    st.markdown("##### 下一步动作")
    for item in result.get("actions", [])[:4]:
        st.write(f"- {item}")


def render_internship_snapshot(analysis: dict[str, Any] | None) -> None:
    if not analysis:
        render_decision_empty_state(
            "填入实习信息后，这里会判断它是否值得去，以及入职前该问清哪些条件。",
            ["填写公司、岗位、周期和薪资。", "粘贴实习 JD 或导师描述。", "重点看真实项目、导师反馈和可写进简历的产出。"],
        )
        return
    score = int(analysis.get("score", 0))
    facts = [
        ("价值评分", f"{score} / 100"),
        ("建议产出", f"{len(analysis.get('recommended_outputs', []))} 项"),
        ("入职前问题", f"{len(analysis.get('questions_to_ask', []))} 个"),
        ("第一周动作", f"{len(analysis.get('first_week_plan', []))} 项"),
    ]
    fact_cards = "".join(
        '<div class="cp-fact">'
        f'<div class="cp-fact-label">{safe_html(name)}</div>'
        f'<div class="cp-fact-value">{safe_html(value)}</div>'
        "</div>"
        for name, value in facts
    )
    st.markdown(
        '<div class="cp-decision-card">'
        '<div class="cp-decision-label">实习判断</div>'
        f'<div class="cp-decision-value">{safe_html(analysis.get("verdict", "待判断"))}</div>'
        f'<div class="cp-decision-copy">{safe_html(analysis.get("decision", ""))}</div>'
        "</div>"
        f'<div class="cp-fact-grid">{fact_cards}</div>',
        unsafe_allow_html=True,
    )
    st.markdown("##### 接 Offer 前先问")
    for item in analysis.get("questions_to_ask", [])[:4]:
        st.write(f"- {item}")


def render_application_summary(df: pd.DataFrame) -> None:
    today = today_label()
    today_count = int((df["queue_date"].astype(str) == today).sum()) if "queue_date" in df.columns and not df.empty else 0
    applied_count = int(df["applied"].astype(bool).sum()) if "applied" in df.columns and not df.empty else 0
    interview_count = 0
    if "interview_status" in df.columns and not df.empty:
        interview_count = int(df["interview_status"].astype(str).str.contains("笔试|一面|二面|HR面", regex=True).sum())
    offer_count = int(df["offer_status"].astype(str).str.contains("Offer", regex=False).sum()) if "offer_status" in df.columns and not df.empty else 0
    cards = [
        ("今日队列", f"{today_count} 个", "今天需要推进的岗位"),
        ("已投递", f"{applied_count} 个", "投递库中已提交记录"),
        ("面试中", f"{interview_count} 个", "笔试、一面、二面或 HR 面"),
        ("Offer", f"{offer_count} 个", "已拿到或标记 Offer"),
    ]
    html_cards = "".join(
        '<div class="cp-overview-card">'
        f'<div class="cp-overview-label">{safe_html(label)}</div>'
        f'<div class="cp-overview-value">{safe_html(value)}</div>'
        f'<div class="cp-overview-copy">{safe_html(copy)}</div>'
        "</div>"
        for label, value, copy in cards
    )
    st.markdown(f'<div class="cp-overview-grid">{html_cards}</div>', unsafe_allow_html=True)


def render_interview_report_workspace_heading() -> None:
    st.markdown(
        """
        <div class="cp-workspace-head">
            <div>
                <div class="cp-workspace-eyebrow">面试与报告</div>
                <div class="cp-workspace-title">把面试准备和求职材料导出收在最后一步</div>
                <div class="cp-workspace-copy">
                    面试页负责把 JD、简历和面经整理成可回答的问题；报告页负责把岗位、简历、缺口和投递记录打包输出。
                </div>
            </div>
            <div class="cp-mode-note">建议路径：提炼问题 → 练回答 → 导出投递包</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_interview_snapshot(interview_analysis: dict[str, Any] | None) -> None:
    if not interview_analysis:
        render_decision_empty_state(
            "粘贴面经或上传文档后，这里会集中显示可直接背的回答框架、按 JD 生成的问题和原始面经命中。",
            ["粘贴面经原文，或上传 PDF/Word/TXT。", "最好先完成目标 JD 分析。", "点击提炼面试问题。"],
        )
        return
    answer_count = len(interview_analysis.get("answer_templates", []))
    generated_count = sum(len(items) for items in interview_analysis.get("generated_questions", {}).values())
    history_count = sum(len(items) for items in interview_analysis.get("buckets", {}).values())
    facts = [
        ("回答框架", f"{answer_count} 条"),
        ("按 JD 生成题", f"{generated_count} 个"),
        ("面经命中", f"{history_count} 条"),
        ("准备优先级", "先项目后行为"),
    ]
    fact_cards = "".join(
        '<div class="cp-fact">'
        f'<div class="cp-fact-label">{safe_html(name)}</div>'
        f'<div class="cp-fact-value">{safe_html(value)}</div>'
        "</div>"
        for name, value in facts
    )
    st.markdown(
        '<div class="cp-decision-card">'
        '<div class="cp-decision-label">面试准备</div>'
        '<div class="cp-decision-value">已生成问题清单</div>'
        '<div class="cp-decision-copy">先准备目标岗位项目题，再补行为面和英文表达；不要背空答案。</div>'
        "</div>"
        f'<div class="cp-fact-grid">{fact_cards}</div>',
        unsafe_allow_html=True,
    )
    st.markdown("##### 先背这几条")
    for item in interview_analysis.get("answer_templates", [])[:4]:
        st.write(f"- {item}")


def render_report_readiness_panel() -> None:
    jd_ready = bool(st.session_state.get("jd_analysis"))
    resume_ready = bool(st.session_state.get("resume_match"))
    gap_ready = bool(st.session_state.get("gap_analysis"))
    applications_ready = not load_applications().empty
    facts = [
        ("目标 JD", "已就绪" if jd_ready else "缺失"),
        ("简历匹配", "已就绪" if resume_ready else "缺失"),
        ("不足清单", "已就绪" if gap_ready else "缺失"),
        ("投递记录", "已就绪" if applications_ready else "缺失"),
    ]
    fact_cards = "".join(
        '<div class="cp-fact">'
        f'<div class="cp-fact-label">{safe_html(name)}</div>'
        f'<div class="cp-fact-value">{safe_html(value)}</div>'
        "</div>"
        for name, value in facts
    )
    ready_count = sum([jd_ready, resume_ready, gap_ready, applications_ready])
    st.markdown(
        '<div class="cp-decision-card">'
        '<div class="cp-decision-label">报告完整度</div>'
        f'<div class="cp-decision-value">{ready_count} / 4</div>'
        '<div class="cp-decision-copy">完整投递包建议包含 JD 判断、简历匹配、不足行动和投递记录。</div>'
        "</div>"
        f'<div class="cp-fact-grid">{fact_cards}</div>',
        unsafe_allow_html=True,
    )


def render_browser_extension_import(key_prefix: str, default: bool = False, show_toggle: bool = True) -> str:
    use_plugin_import = True
    if show_toggle:
        use_plugin_import = st.checkbox(
            "使用浏览器插件导入",
            value=default,
            key=f"{key_prefix}_use_plugin_import",
            help="适合已登录招聘页；不勾选时隐藏插件导入面板。",
        )
    if not use_plugin_import:
        return ""
    with st.container(border=True):
        st.markdown("##### 浏览器插件导入")
        st.write("在 Edge/Chrome 已登录页面里点击 JD Saver 插件，文件会保存到：")
        st.code("\n".join(str(path) for path in JD_EXPORT_DIRS))
        if st.button("扫描插件导出文件", key=f"{key_prefix}_scan_exports"):
            text, table = scan_browser_export_folder()
            st.session_state[f"{key_prefix}_exported_text"] = text
            st.session_state[f"{key_prefix}_exported_table"] = table
            if table.empty:
                st.warning("没有扫描到导出文件。请先安装插件并在招聘页面点击保存。")
            else:
                st.success(f"已导入 {len(table)} 个导出文件。")
        table = st.session_state.get(f"{key_prefix}_exported_table")
        if table is not None and not table.empty:
            render_user_dataframe(table)
        st.caption("这条路线不读取 Cookie、不接管登录，也不绕过网站验证；只是读取你当前页面可见的正文并保存到本地。")
    return st.session_state.get(f"{key_prefix}_exported_text", "")


def render_jd_workspace_heading() -> None:
    st.markdown(
        """
        <div class="cp-workspace-head">
            <div>
                <div class="cp-workspace-eyebrow">岗位工作台</div>
                <div class="cp-workspace-title">先判断值不值得投，再决定怎么改简历</div>
                <div class="cp-workspace-copy">
                    单条 JD 适合深看目标岗位；批量筛选适合从招聘网站导入后排序；行业监测适合观察一批公司或岗位的新机会。
                </div>
            </div>
            <div class="cp-mode-note">建议路径：单条判断 → 批量筛选 → 加入今日队列</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_jd_empty_state() -> None:
    st.markdown(
        """
        <div class="cp-empty-state">
            <div class="cp-empty-state-title">等待导入一个岗位 JD</div>
            <div class="cp-empty-state-copy">
                右侧会在分析后集中显示投递建议、岗位价值、风险标签和下一步动作，不再让结果散落在页面下方。
            </div>
            <div class="cp-mini-steps">
                <div class="cp-mini-step"><span class="cp-mini-step-index">1</span><span>粘贴 JD 正文，或从文件/插件/公开链接导入。</span></div>
                <div class="cp-mini-step"><span class="cp-mini-step-index">2</span><span>点击分析 JD，系统会自动设为当前目标岗位。</span></div>
                <div class="cp-mini-step"><span class="cp-mini-step-index">3</span><span>根据判断结果进入简历工作台补证据、改摘要。</span></div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_job_decision_snapshot(jd_analysis: dict[str, Any], resume_match: dict[str, Any] | None = None) -> None:
    plan = build_job_action_plan(jd_analysis, resume_match)
    label, _tone = job_decision_label(jd_analysis, resume_match)
    reasons = job_decision_reasons(jd_analysis, resume_match)
    basic = jd_analysis.get("basic", {})
    value = jd_analysis.get("value", {})
    facts = [
        ("岗位分类", jd_analysis.get("category", "待判断")),
        ("岗位价值", "高" if value.get("is_high_value") else "待判断"),
        ("低价值风险", "高" if value.get("is_generic_esg") else "低"),
        ("地点", basic.get("地点") or "未识别"),
    ]
    fact_cards = "".join(
        '<div class="cp-fact">'
        f'<div class="cp-fact-label">{safe_html(name)}</div>'
        f'<div class="cp-fact-value">{safe_html(value_text)}</div>'
        "</div>"
        for name, value_text in facts
    )
    st.markdown(
        '<div class="cp-decision-card">'
        '<div class="cp-decision-label">投递建议</div>'
        f'<div class="cp-decision-value">{safe_html(label)}</div>'
        f'<div class="cp-decision-copy">{safe_html(plan["decision"])}</div>'
        "</div>"
        f'<div class="cp-fact-grid">{fact_cards}</div>',
        unsafe_allow_html=True,
    )
    st.markdown("##### 关键理由")
    for reason in reasons[:4]:
        st.write(f"- {reason}")
    risk_tags = value.get("risk_tags") or []
    if risk_tags:
        st.caption("风险标签：" + " / ".join(risk_tags[:5]))
    st.markdown("##### 下一步动作")
    for item in list(plan["actions"])[:3]:
        st.write(f"- {item}")
    if st.button("加入今日队列", key="add_current_jd_today_queue_snapshot"):
        add_current_jd_to_today_queue(jd_analysis, resume_match)
        st.success("已加入今日求职队列。")


def render_jd_tab() -> None:
    left_col, right_col = st.columns([1.08, 0.92], gap="large")
    with left_col:
        with st.container(border=True):
            st.markdown('<div class="cp-panel-title">导入 JD</div>', unsafe_allow_html=True)
            st.markdown(
                '<div class="cp-panel-copy">优先粘贴完整岗位描述；链接、插件和文件导入收进高级入口，减少首屏干扰。</div>',
                unsafe_allow_html=True,
            )
            text = st.text_area("粘贴 JD 文本", height=260, placeholder="粘贴招聘 JD、岗位描述或网页复制内容...")
            exported_jd_text = ""
            upload_text = ""
            crawled_jd_text = ""
            use_extra_imports = st.checkbox("使用其他导入方式", value=False, key="jd_use_extra_imports")
            if use_extra_imports:
                import_modes = st.columns(3)
                use_url_import = import_modes[0].checkbox("公开链接抓取", value=False, key="jd_use_url_import")
                use_plugin_import = import_modes[1].checkbox("浏览器插件", value=False, key="jd_use_plugin_import")
                use_file_import = import_modes[2].checkbox("上传文件", value=False, key="jd_use_file_import")

                if use_url_import:
                    with st.container(border=True):
                        st.markdown("##### 公开链接抓取")
                        url_block = st.text_area(
                            "JD 链接（一行一个 URL）",
                            height=100,
                            placeholder="https://...\nhttps://...",
                            help="会并发抓取多个公开网页。需要登录、验证码或强动态渲染的招聘页可能抓不到，建议改用复制文本或截图 OCR。",
                        )
                        cols = st.columns([1, 3])
                        max_workers = cols[0].slider("并发爬虫数", min_value=1, max_value=8, value=4)
                        use_dynamic_crawl = cols[1].checkbox("公开网页启用动态抓取", value=False, help="使用 Playwright 打开公开页面，适合前端渲染页面；已登录页面仍建议用浏览器插件。")
                        if cols[1].button("并发抓取 JD 链接"):
                            urls = extract_urls(url_block)
                            if not urls:
                                st.warning("请先粘贴至少一个 JD 链接。")
                            else:
                                with st.spinner(f"正在并发抓取 {len(urls)} 个链接..."):
                                    crawl_results = crawl_jd_urls(
                                        urls,
                                        max_workers=max_workers,
                                        use_dynamic=use_dynamic_crawl,
                                    )
                                st.session_state.jd_crawl_results = crawl_results
                                st.session_state.crawled_jd_text = "\n\n".join([item["text"] for item in crawl_results if item["ok"]])
                                ok_count = sum(1 for item in crawl_results if item["ok"])
                                if ok_count:
                                    st.success(f"成功抓取 {ok_count}/{len(crawl_results)} 个链接。")
                                else:
                                    st.error("没有成功抓到可用 JD 文本，建议复制网页正文或上传截图。")
                        crawled_jd_text = st.session_state.get("crawled_jd_text", "")
                        crawl_results = st.session_state.get("jd_crawl_results", [])
                        if crawl_results:
                            with st.expander("查看爬虫结果"):
                                for item in crawl_results:
                                    if item["ok"]:
                                        st.success(item["url"])
                                        st.text_area(f"抓取文本 - {item['url']}", value=item["text"][:5000], height=160)
                                    else:
                                        st.error(f"{item['url']}：{item['error']}")

                if use_plugin_import:
                    with st.container(border=True):
                        st.markdown("##### 浏览器插件导入")
                        st.caption("在 Edge/Chrome 已登录页面里点击 JD Saver 插件，文件会保存到本地目录；这里仅扫描本地导出文本。")
                        if st.button("扫描插件导出文件", key="jd_scan_exports_inline"):
                            text_from_export, table_from_export = scan_browser_export_folder()
                            st.session_state["jd_exported_text"] = text_from_export
                            st.session_state["jd_exported_table"] = table_from_export
                            if table_from_export.empty:
                                st.warning("没有扫描到导出文件。请先安装插件并在招聘页面点击保存。")
                            else:
                                st.success(f"已导入 {len(table_from_export)} 个导出文件。")
                        exported_table = st.session_state.get("jd_exported_table")
                        if exported_table is not None and not exported_table.empty:
                            render_user_dataframe(exported_table)
                        exported_jd_text = st.session_state.get("jd_exported_text", "")

                if use_file_import:
                    with st.container(border=True):
                        st.markdown("##### 上传文件")
                        uploaded = st.file_uploader(
                            "上传 JD 文件",
                            type=["pdf", "docx", "txt", "md", "html", "htm", "xlsx", "xls", "csv", "png", "jpg", "jpeg", "webp"],
                            key="jd_upload",
                        )
                        upload_text = extract_text_from_upload(uploaded) if uploaded else ""
                        if upload_text:
                            with st.expander("查看上传文件提取文本"):
                                st.text_area("提取结果", value=upload_text, height=180)
            if st.button("分析 JD", type="primary", width="stretch"):
                full_text = "\n".join([text, exported_jd_text, crawled_jd_text, upload_text]).strip()
                if not full_text:
                    st.warning("请先粘贴或上传 JD。")
                else:
                    set_current_target_jd(full_text, "单条JD分析", "手动导入JD")
                    st.success("JD 分析完成，并已设为当前目标 JD。")

    jd_analysis = st.session_state.get("jd_analysis")
    with right_col:
        with st.container(border=True):
            st.markdown('<div class="cp-panel-title">判断结果</div>', unsafe_allow_html=True)
            st.markdown(
                '<div class="cp-panel-copy">这里优先回答“投不投、为什么、下一步做什么”。</div>',
                unsafe_allow_html=True,
            )
            if jd_analysis:
                render_job_decision_snapshot(jd_analysis, st.session_state.get("resume_match"))
            else:
                render_jd_empty_state()

    if not jd_analysis:
        return

    basic = jd_analysis["basic"]
    value = jd_analysis["value"]
    score_tab, resume_tab, detail_tab = st.tabs(["评分拆解", "简历改法", "原始字段"])
    with score_tab:
        render_score_breakdown(jd_analysis, st.session_state.get("resume_match"))
    with resume_tab:
        render_resume_focus_plan(jd_analysis, st.session_state.get("resume_match"))
    with detail_tab:
        detail_cols = st.columns(4)
        detail_cols[0].metric("岗位分类", jd_analysis["category"])
        detail_cols[1].metric("岗位价值", "高" if value["is_high_value"] else "待判断")
        detail_cols[2].metric("低价值风险", "高" if value["is_generic_esg"] else "低")
        detail_cols[3].metric("地点", basic.get("地点", "未识别"))

        basic_df = pd.DataFrame([basic]).drop(columns=["区域", "地域优先级"], errors="ignore")
        render_user_dataframe(basic_df)

        skills = jd_skill_list(jd_analysis)
        if skills:
            st.write("岗位关键词：" + " / ".join(skills[:12]))
        else:
            st.caption("暂未识别到明确岗位关键词，可补充更完整的 JD 文本。")

        if value["is_high_value"]:
            st.success(value["label"])
        elif value["is_generic_esg"]:
            st.warning(value["label"])
        else:
            st.info(value["label"])


def render_batch_jd_tab() -> None:
    st.subheader("批量 JD 分析")
    active_profile = get_active_profile()
    active_resume = get_active_resume()
    st.caption("把插件批量导出的岗位、粘贴的多条 JD、Excel/CSV/TXT 文件和公开链接统一拆成多条岗位，逐条计算目标意向匹配度。")

    with st.expander(f"当前目标意向：{active_profile['name']}"):
        st.write(profile_text_for_analysis())
    with st.expander(f"当前对比简历：{active_resume['name']}"):
        if active_resume["content"]:
            st.write(active_resume["content"])
        else:
            st.info("还没有设置当前简历。可在左侧边栏的“当前简历”里新增。")

    imported_records: list[dict[str, str]] = []
    scan_cols = st.columns([1, 3])
    export_dates = available_export_dates()
    selected_export_date = scan_cols[1].selectbox(
        "导出日期",
        ["全部"] + export_dates,
        index=0,
        help="插件会按日期保存到 Downloads/CareerPilot_JD/YYYY-MM-DD。",
    )
    enrich_export_details = st.checkbox(
        "分析前用详情页补全公司/岗位",
        value=False,
        key="batch_enrich_export_details",
        help="有详情页链接时，会按本次实际选择分析的岗位数量自动补全，不再手动选择数量。",
    )
    if scan_cols[0].button("扫描并选择导入文件", type="primary"):
        with st.spinner("正在扫描插件导出文件..."):
            records, table = scan_exported_jd_records(
                date_filter=selected_export_date,
                enrich_detail_pages=False,
                detail_limit=0,
            )
        st.session_state.batch_export_records = records
        st.session_state.batch_export_table = table
        st.session_state.batch_export_deleted_ids = set()
        st.session_state.batch_export_selected_ids = {jd_record_id(record) for record in records}
        st.session_state.batch_export_info_set_selected_ids = (
            set(table["信息集ID"].astype(str).tolist())
            if table is not None and not table.empty and "信息集ID" in table.columns
            else set()
        )
        st.session_state.pop("batch_jd_analysis", None)
        if records:
            render_export_info_set_dialog(records, table, "batch_export")
        else:
            st.warning("没有读到插件导出的岗位。请在浏览器列表页优先点击 Collect current + next pages。")
    if st.checkbox("查看插件导出目录", value=False, key="batch_show_export_dirs"):
        st.code("\n".join(str(path) for path in JD_EXPORT_DIRS))

    export_table = st.session_state.get("batch_export_table")
    if export_table is not None and not export_table.empty:
        export_records = st.session_state.get("batch_export_records", [])
        selected_info_set_records = selected_export_info_set_records(export_records, export_table, "batch_export")
        selected_info_set_ids = set(st.session_state.get("batch_export_info_set_selected_ids", set()))
        file_count = len(selected_info_set_ids)
        st.caption(f"已选择 {file_count} 个文件，包含 {len(selected_info_set_records)} 条岗位。需要调整时点击上方“扫描并选择导入文件”。")
        selected_export_records = render_import_record_selector(
            selected_info_set_records,
            "batch_export",
            allow_import_toggle=False,
            title="按岗位删除/排除",
            caption="导入粒度以上方信息集为准；这里仅用于按岗位排除不想分析的记录，不会删除硬盘里的原始导出文件。",
        )
        imported_records.extend(selected_export_records)

    pasted_text = ""
    uploaded_files = []
    url_block = ""
    use_dynamic_crawl = False
    selected_url_records: list[dict[str, str]] = []
    use_extra_batch_imports = st.checkbox("使用其他导入方式", value=False, key="batch_use_extra_imports")
    if use_extra_batch_imports:
        st.markdown("#### 其他导入方式")
        import_modes = st.columns(3)
        use_paste_import = import_modes[0].checkbox("粘贴多条 JD", value=False, key="batch_use_paste_import")
        use_file_import = import_modes[1].checkbox("上传表格/文件", value=False, key="batch_use_file_import")
        use_url_import = import_modes[2].checkbox("公开链接抓取", value=False, key="batch_use_url_import")

        if use_paste_import:
            with st.container(border=True):
                st.markdown("##### 粘贴多条 JD")
                pasted_text = st.text_area("每条之间空一行", height=180, key="batch_pasted_text")
        if use_file_import:
            with st.container(border=True):
                st.markdown("##### 上传表格/文件")
                uploaded_files = st.file_uploader(
                    "支持 Excel / CSV / TXT / HTML / PDF / DOCX",
                    type=["xlsx", "xls", "csv", "txt", "md", "html", "htm", "pdf", "docx"],
                    accept_multiple_files=True,
                    key="batch_jd_upload",
                )
        if use_url_import:
            with st.container(border=True):
                st.markdown("##### 公开链接抓取")
                url_block = st.text_area("一行一个 URL", height=90, key="batch_url_block")
                use_dynamic_crawl = st.checkbox(
                    "公开链接动态抓取",
                    value=False,
                    help="只对公开链接生效；登录态岗位请用浏览器插件。",
                    key="batch_dynamic_crawl",
                )
                if st.button("抓取公开链接并预览", key="batch_fetch_urls_preview"):
                    urls = extract_urls(url_block)
                    if not urls:
                        st.warning("请先输入至少一个 URL。")
                    else:
                        with st.spinner(f"正在抓取 {len(urls)} 个公开链接..."):
                            crawl_results = crawl_jd_urls(urls, max_workers=5, use_dynamic=use_dynamic_crawl)
                        url_records: list[dict[str, str]] = []
                        for item in crawl_results:
                            url_records.extend(records_from_crawl_result(item))
                        st.session_state.batch_url_records = dedupe_jd_records(url_records)
                        st.session_state.batch_url_deleted_ids = set()
                        st.session_state.batch_url_selected_ids = {jd_record_id(record) for record in st.session_state.batch_url_records}
                        st.session_state.batch_url_crawl_summary = f"公开链接抓取完成：{len(urls)} 个链接，拆出 {len(st.session_state.batch_url_records)} 条岗位。"
                if st.session_state.get("batch_url_crawl_summary"):
                    st.caption(st.session_state.batch_url_crawl_summary)
                if st.session_state.get("batch_url_records"):
                    selected_url_records = render_import_record_selector(
                        st.session_state.get("batch_url_records", []),
                        "batch_url",
                    )

    options = st.columns(2)
    include_resume = options[0].checkbox("叠加当前简历匹配", value=bool(active_resume["content"]))
    min_score = options[1].slider("最低显示匹配度", min_value=0, max_value=100, value=0)

    resume_text = active_resume["content"].strip() if include_resume and active_resume["content"].strip() else ""
    if include_resume and resume_text:
        st.caption(f"批量分析将直接调用当前简历：{active_resume['name']}")
    elif include_resume:
        st.warning("当前没有简历内容，请先在左侧“当前简历”上传或粘贴后再叠加匹配。")

    if st.button("开始批量分析", type="primary"):
        records = list(imported_records)
        records.extend(selected_url_records)
        records.extend(split_batch_jd_text(pasted_text))

        for uploaded in uploaded_files or []:
            try:
                records.extend(records_from_uploaded_file(uploaded))
            except Exception as exc:
                st.warning(f"{uploaded.name} 解析失败：{exc}")

        urls = extract_urls(url_block)
        if urls and not selected_url_records and not st.session_state.get("batch_url_records"):
            with st.spinner(f"正在抓取 {len(urls)} 个公开链接..."):
                crawl_results = crawl_jd_urls(urls, max_workers=5, use_dynamic=use_dynamic_crawl)
            crawl_record_count = 0
            for item in crawl_results:
                link_records = records_from_crawl_result(item)
                crawl_record_count += len(link_records)
                records.extend(link_records)
            st.caption(f"公开链接抓取完成：{len(urls)} 个链接，拆出 {crawl_record_count} 条岗位。")

        records = dedupe_jd_records(records)
        if not records:
            st.warning("没有可分析的 JD。建议先用浏览器插件在列表页批量保存岗位。")
        else:
            if enrich_export_details:
                detail_limit = len(records)
                with st.spinner(f"正在按本次选择的 {detail_limit} 条岗位补全详情页..."):
                    records = enrich_records_with_detail_pages(records, detail_limit)
                    records = dedupe_jd_records(records)
            with st.spinner(f"正在逐条分析 {len(records)} 条 JD..."):
                df = analyze_batch_jd_records(
                    records,
                    profile_text_for_analysis(),
                    resume_text=resume_text,
                )
            st.session_state.batch_jd_analysis = df
            st.success(f"批量分析完成，共 {len(df)} 条结果。")

    df = st.session_state.get("batch_jd_analysis")
    if df is None or df.empty:
        st.info("推荐流程：招聘网站搜索结果页 -> 插件 Collect current + next pages -> 回到这里按日期扫描并分析。")
        return

    quick_view = st.radio("一键视图", BATCH_QUICK_VIEWS, horizontal=True, key="batch_quick_view")
    filter_cols = st.columns(6)
    only_p1 = filter_cols[0].checkbox("只看P1/P2", value=False)
    only_high_value = filter_cols[1].checkbox("只看高价值", value=False)
    exclude_generic = filter_cols[2].checkbox("排除低价值风险", value=True)
    only_internship = filter_cols[3].checkbox("只看实习", value=False)
    province_filter = filter_cols[4].selectbox("省份筛选", province_filter_options(df), index=0)
    keyword_filter = filter_cols[5].text_input("关键词筛选", placeholder="公司/岗位/技能/城市/省份")

    view_df = apply_batch_quick_view(df.copy(), quick_view, load_target_preferences())
    if only_p1 and "投递建议" in view_df.columns:
        view_df = view_df[view_df["投递建议"].astype(str).str.contains("P1|P2", regex=True)]
    if only_high_value and "高价值" in view_df.columns:
        view_df = view_df[view_df["高价值"].astype(str) == "是"]
    risk_column = "低价值风险" if "低价值风险" in view_df.columns else "泛ESG风险"
    if exclude_generic and risk_column in view_df.columns:
        view_df = view_df[view_df[risk_column].astype(str) != "是"]
    if only_internship and "是否招实习" in view_df.columns:
        view_df = view_df[view_df["是否招实习"].astype(str) == "是"]
    view_df = filter_by_province(view_df, province_filter)
    if keyword_filter:
        haystack = view_df.astype(str).agg(" ".join, axis=1)
        view_df = view_df[haystack.str.contains(keyword_filter, case=False, na=False, regex=False)]
    if min_score and "意向匹配度" in view_df.columns:
        view_df = view_df[view_df["意向匹配度"] >= min_score]
    view_df = view_df.reset_index(drop=True)

    st.markdown("#### 批量分析结果")
    top_cols = [
        "意向匹配度",
        "投递建议",
        "下一步动作",
        "公司",
        "岗位",
        "类型",
        "是否招实习",
        "应届生",
        "薪资",
        "地点",
        "岗位分类",
        "高价值",
        "低价值风险",
        "风险标签",
        "技能关键词",
        "主要缺口",
        "薪资判断",
        "链接",
    ]
    show_cols = [col for col in top_cols if col in df.columns]

    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("筛选后岗位数", len(view_df))
    c1.caption(f"视图：{quick_view}")
    c2.metric("P1优先", int(view_df["投递建议"].astype(str).str.contains("P1").sum()) if "投递建议" in view_df.columns else 0)
    c3.metric("高价值", int((view_df["高价值"] == "是").sum()) if "高价值" in view_df.columns else 0)
    c4.metric("已识别省份", recognized_province_total(view_df))
    c5.metric("实习机会", int((view_df["是否招实习"] == "是").sum()) if "是否招实习" in view_df.columns else 0)
    table_tab, target_tab, insight_tab, export_tab = st.tabs(["结果表", "设为目标JD", "分布与优先岗位", "导出"])
    with table_tab:
        render_user_dataframe(view_df, show_cols)
        render_risk_logic_note()
        queue_count = min(5, len(view_df))
        if queue_count and st.button(f"把前 {queue_count} 个加入今日队列", key="batch_add_today_queue"):
            added = add_batch_rows_to_today_queue(view_df.head(queue_count))
            st.success(f"已加入今日求职队列：{added} 个岗位。")
    with target_tab:
        render_target_jd_picker(view_df, "batch_jd", "批量JD筛选")
    with insight_tab:
        if not view_df.empty:
            st.markdown("#### 最该优先看的岗位")
            for _, row in view_df.head(5).iterrows():
                st.write(f"- {row.get('意向匹配度', '')}分｜{row.get('公司', '')}｜{row.get('岗位', '')}｜{row.get('下一步动作', '')}")
            st.markdown("#### 省份热力地图")
            render_china_province_map(view_df, "批量JD省份分布")
            if "岗位分类" in view_df.columns:
                dist = view_df["岗位分类"].value_counts().reset_index()
                dist.columns = ["岗位分类", "数量"]
                fig = px.bar(dist, x="岗位分类", y="数量", color="岗位分类")
                st.plotly_chart(fig, width="stretch")
    with export_tab:
        output = io.BytesIO()
        public_export_df(view_df).to_excel(output, index=False)
        st.download_button(
            "导出批量JD分析 Excel",
            output.getvalue(),
            "batch_jd_analysis.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )


def render_resume_tab() -> None:
    active_resume = get_active_resume()
    jd_analysis = st.session_state.get("jd_analysis")
    resume_match = st.session_state.get("resume_match")
    can_match = bool(jd_analysis and active_resume.get("content", "").strip())

    left_col, right_col = st.columns([1.05, 0.95], gap="large")
    with left_col:
        with st.container(border=True):
            st.markdown('<div class="cp-panel-title">当前简历</div>', unsafe_allow_html=True)
            st.markdown(
                '<div class="cp-panel-copy">这里直接调用左侧保存的真实简历；需要换版本时，先在左侧更新再回来匹配。</div>',
                unsafe_allow_html=True,
            )
            if active_resume.get("content", "").strip():
                st.text_area(
                    f"当前简历：{active_resume.get('name') or '简历'}",
                    value=active_resume["content"],
                    height=300,
                    disabled=True,
                    key=f"resume_global_preview_{active_resume.get('id', 'none')}",
                )
            else:
                render_resume_empty_state(jd_analysis, active_resume, False)

            if st.button("使用当前简历分析匹配", type="primary", width="stretch", disabled=not can_match):
                st.session_state.resume_text = active_resume["content"].strip()
                st.session_state.resume_match = match_resume_to_jd(
                    jd_analysis,
                    st.session_state.resume_text,
                    profile_text_for_analysis(),
                )
                st.session_state.resume_match_jd_fingerprint = current_jd_fingerprint()
                st.session_state.resume_match_resume_fingerprint = current_resume_fingerprint()
                st.success("简历匹配分析完成。")
                resume_match = st.session_state.get("resume_match")

            if not jd_analysis:
                st.caption("需要先完成目标 JD 分析。")
            elif not active_resume.get("content", "").strip():
                st.caption("需要先在左侧保存一份真实简历。")

    with right_col:
        with st.container(border=True):
            st.markdown('<div class="cp-panel-title">匹配结论</div>', unsafe_allow_html=True)
            st.markdown(
                '<div class="cp-panel-copy">优先看匹配度、已覆盖证据和投前必须补齐的内容。</div>',
                unsafe_allow_html=True,
            )
            if resume_match and can_match:
                render_resume_match_snapshot(resume_match)
            else:
                render_resume_empty_state(jd_analysis, active_resume, bool(resume_match))

    if not resume_match or not can_match:
        return

    summary_tab, rewrite_tab, evidence_tab = st.tabs(["结论", "可直接改的简历句子", "证据与缺口"])
    with summary_tab:
        render_job_action_plan(st.session_state.get("jd_analysis"), resume_match)
    with rewrite_tab:
        for bullet in generate_resume_bullets(st.session_state.get("jd_analysis"), resume_match):
            st.write(f"- {bullet}")
    with evidence_tab:
        if resume_match["matched_skills"]:
            st.write("已覆盖：" + " / ".join(resume_match["matched_skills"]))
        if resume_match["missing_skills"]:
            st.write("待补证据：" + " / ".join(resume_match["missing_skills"]))
        if resume_match.get("evidence"):
            st.markdown("#### 已识别证据")
            for item in resume_match["evidence"]:
                st.write(f"- {item}")


def render_custom_resume_tab() -> None:
    st.subheader("定制简历")
    jd_analysis = st.session_state.get("jd_analysis")
    active_profile = get_active_profile()
    active_resume = get_active_resume()
    if not jd_analysis:
        st.warning("请先在岗位工作台完成目标 JD 分析。")
        return
    if not active_resume["content"].strip():
        st.warning("请先在左侧“当前简历”上传或粘贴真实简历。")
        return

    st.caption("定制简历直接调用左侧“当前简历”；如果要换版本，请先在左侧更新后再生成。")
    with st.expander(f"当前简历：{active_resume['name']}", expanded=False):
        st.text_area(
            "当前简历内容",
            height=240,
            value=active_resume["content"],
            disabled=True,
            key=f"custom_global_resume_preview_{active_resume.get('id', 'none')}",
        )

    if st.button("生成定制版简历内容", type="primary"):
        full_resume = active_resume["content"].strip()
        if not full_resume:
            st.warning("请先在左侧维护当前简历。")
        else:
            st.session_state.custom_resume = build_custom_resume(full_resume, jd_analysis, profile_text_for_analysis())
            st.session_state.custom_resume_jd_fingerprint = current_jd_fingerprint()
            st.session_state.custom_resume_resume_fingerprint = current_resume_fingerprint()
            st.success("定制简历内容已生成。")

    result = st.session_state.get("custom_resume")
    if not result:
        return

    st.markdown("#### 目标定位")
    st.write(f"目标岗位：{result['job_title']}；岗位分类：{result['category']}")
    if result["company"]:
        st.write(f"目标公司：{result['company']}")
    st.write("关键词布局：" + result["keyword_line"])

    ready_tab, parts_tab, pitch_tab, risk_tab = st.tabs(["直接可用", "分块查看", "投递说明", "证据与风险"])
    with ready_tab:
        st.text_area("可复制到简历里再按真实经历微调", value=result.get("ready_resume_text", ""), height=430)
    with parts_tab:
        st.markdown("#### 求职摘要")
        st.text_area("放到简历开头", value=result["summary"], height=120)
        st.markdown("#### 核心能力栏")
        for item in result.get("skills_section", []):
            st.write(f"- {item}")
        st.markdown("#### 经历 bullet 成稿")
        for bullet in result.get("experience_bullets", result["bullets"]):
            st.write(f"- {bullet}")
        st.markdown("#### 项目经历可替换版本")
        for item in result["project_rewrite"]:
            st.write(f"- {item}")
    with pitch_tab:
        st.text_area("可用于邮件/私信/网申自我介绍", value=result.get("application_pitch", ""), height=140)
    with risk_tab:
        cols = st.columns(2)
        with cols[0]:
            st.markdown("#### 已覆盖关键词")
            st.write(" / ".join(result["matched_skills"]) or "暂无明显覆盖")
        with cols[1]:
            st.markdown("#### 需要补证据的关键词")
            st.write(" / ".join(result["missing_skills"]) or "暂无明显缺口")
        st.markdown("#### 改写风险提醒")
        for item in result["risk_notes"]:
            st.write(f"- {item}")
        if result.get("evidence_lines"):
            st.markdown("#### 本次改写参考的原简历证据")
            for item in result["evidence_lines"]:
                st.write(f"- {item}")

    export_text = "\n".join(
        [
            f"目标岗位：{result['job_title']}",
            f"关键词：{result['keyword_line']}",
            "",
            "直接可用版本：",
            result.get("ready_resume_text", ""),
            "",
            "投递说明：",
            result.get("application_pitch", ""),
            "",
            "简历摘要：",
            result["summary"],
            "",
            "经历 bullet：",
            *[f"- {item}" for item in result.get("experience_bullets", result["bullets"])],
            "",
            "项目经历可替换版本：",
            *[f"- {item}" for item in result["project_rewrite"]],
            "",
            "风险提醒：",
            *[f"- {item}" for item in result["risk_notes"]],
        ]
    )
    st.download_button("下载定制简历片段 TXT", export_text.encode("utf-8-sig"), "custom_resume.txt", mime="text/plain")


def render_recruitment_monitor_tab() -> None:
    st.subheader("行业招聘监测")
    st.caption("输入一批 JD 链接或招聘文本，系统会标注新发现岗位、公司、是否招应届生、薪资、实习/正式工和岗位价值。")

    url_block = ""
    exported_recruitment_text = ""
    pasted_text = ""
    upload_text = ""
    max_workers = 5
    use_dynamic_crawl = False

    input_modes = st.columns(4)
    use_paste_input = input_modes[0].checkbox("粘贴招聘文本", value=True, key="recruitment_use_paste")
    use_link_input = input_modes[1].checkbox("抓取公开链接", value=False, key="recruitment_use_links")
    use_plugin_input = input_modes[2].checkbox("浏览器插件", value=False, key="recruitment_use_plugin")
    use_file_input = input_modes[3].checkbox("上传文件", value=False, key="recruitment_use_file")

    if use_paste_input:
        with st.container(border=True):
            st.markdown("##### 粘贴招聘文本")
            pasted_text = st.text_area("多条 JD 之间空一行", height=220, placeholder="可以一次粘贴多条 JD；建议每条之间空一行。")
    if use_link_input:
        with st.container(border=True):
            st.markdown("##### 招聘链接抓取")
            url_block = st.text_area("一行一个 URL", height=100, placeholder="https://...\nhttps://...")
            crawl_cols = st.columns(2)
            max_workers = crawl_cols[0].slider("并发爬虫数", min_value=1, max_value=8, value=5, key="recruitment_workers")
            use_dynamic_crawl = crawl_cols[1].checkbox("公开网页启用动态抓取", value=False, key="recruitment_dynamic_crawl")
    if use_plugin_input:
        exported_recruitment_text = render_browser_extension_import("recruitment", default=True, show_toggle=False)
    if use_file_input:
        with st.container(border=True):
            st.markdown("##### 上传招聘信息文件")
            uploaded = st.file_uploader("支持 TXT / MD / CSV / Excel / HTML", type=["txt", "md", "csv", "xlsx", "xls", "html", "htm"], key="recruitment_upload")
            upload_text = extract_text_from_upload(uploaded) if uploaded else ""

    if st.button("开始监测并解析", type="primary"):
        urls = extract_urls(url_block)
        crawl_results = []
        if urls:
            with st.spinner(f"正在并发抓取 {len(urls)} 个招聘链接..."):
                crawl_results = crawl_jd_urls(
                    urls,
                    max_workers=max_workers,
                    use_dynamic=use_dynamic_crawl,
                )
            crawl_record_count = sum(len(records_from_crawl_result(item)) for item in crawl_results)
            st.caption(f"公开链接抓取完成：{len(urls)} 个链接，拆出 {crawl_record_count} 条岗位。")
        source_text = "\n\n".join([pasted_text, exported_recruitment_text, upload_text]).strip()
        df = analyze_recruitment_sources(source_text, crawl_results)
        if df.empty:
            st.warning("没有解析出招聘岗位，请补充 JD 文本或换公开链接。")
        else:
            st.session_state.recruitment_monitor = register_recruitment_records(df)
            st.success("招聘监测完成。")

    df = st.session_state.get("recruitment_monitor")
    if df is not None and not df.empty:
        monitor_cols = st.columns(3)
        province_filter = monitor_cols[0].selectbox("省份筛选", province_filter_options(df), index=0, key="recruitment_province_filter")
        only_new = monitor_cols[1].checkbox("只看新发现", value=False)
        monitor_keyword = monitor_cols[2].text_input("关键词筛选", placeholder="公司/岗位/城市/省份", key="recruitment_keyword_filter")
        view_df = df.copy()
        view_df = filter_by_province(view_df, province_filter)
        if only_new and "是否新发现" in view_df.columns:
            view_df = view_df[view_df["是否新发现"].astype(str) == "是"]
        if monitor_keyword:
            haystack = view_df.astype(str).agg(" ".join, axis=1)
            view_df = view_df[haystack.str.contains(monitor_keyword, case=False, na=False, regex=False)]
        view_df = view_df.reset_index(drop=True)

        display_cols = [
            "是否新发现",
            "公司",
            "岗位",
            "类型",
            "是否招实习",
            "是否招应届生",
            "薪资",
            "地点",
            "公司层级",
            "岗位分类",
            "高价值岗位",
            "低价值风险",
            "技能关键词",
            "链接",
        ]
        table_tab, target_tab, map_tab, export_tab = st.tabs(["岗位列表", "设为目标JD", "省份地图", "导出"])
        with table_tab:
            render_user_dataframe(view_df, display_cols)
            render_risk_logic_note()
        with target_tab:
            render_target_jd_picker(view_df, "recruitment", "行业招聘监测")
        with map_tab:
            render_china_province_map(view_df, "行业招聘省份分布")
        with export_tab:
            output = io.BytesIO()
            public_export_df(view_df).to_excel(output, index=False)
            st.download_button("导出本次监测结果 Excel", output.getvalue(), "recruitment_monitor.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    history = load_recruitment_posts()
    if not history.empty:
        with st.expander("查看历史监测库"):
            render_user_dataframe(history)


def render_offer_prediction_tab() -> None:
    jd_analysis = st.session_state.get("jd_analysis")
    resume_match = st.session_state.get("resume_match")
    gap_analysis = st.session_state.get("gap_analysis")

    if not jd_analysis:
        render_decision_empty_state(
            "Offer 预测需要先有一个明确目标 JD。",
            ["去岗位工作台导入 JD。", "完成单条 JD 分析。", "回到这里调整公司层级和个人匹配参数。"],
        )
        return

    inferred_company = jd_basic_value(jd_analysis, "公司名")
    inferred_tier = infer_company_tier(inferred_company, jd_analysis.get("raw_text", ""))
    tier_options = ["头部/知名机构", "中高层级", "普通公司", "小公司/冷门岗位", "未知"]
    left_col, right_col = st.columns([0.95, 1.05], gap="large")
    with left_col:
        with st.container(border=True):
            st.markdown('<div class="cp-panel-title">预测参数</div>', unsafe_allow_html=True)
            st.markdown(
                '<div class="cp-panel-copy">参数越贴近真实投递场景，预测越适合作为排序依据。</div>',
                unsafe_allow_html=True,
            )
            if not resume_match:
                st.warning("建议先完成简历匹配；未完成时会用默认匹配度估算。")
            company_tier = st.selectbox("公司层级", tier_options, index=tier_options.index(inferred_tier) if inferred_tier in tier_options else 4)
            education_fit = st.selectbox("学历匹配", ["高于或满足要求", "勉强满足", "低于要求"])
            english_level = st.selectbox("英文能力", ["强", "一般", "弱"])
            if st.button("预测通过率 / 面试率 / Offer率", type="primary", width="stretch"):
                st.session_state.offer_prediction = predict_offer_probabilities(
                    jd_analysis,
                    resume_match,
                    gap_analysis,
                    company_tier,
                    education_fit,
                    english_level,
                )
                st.success("Offer 预测完成。")

    result = st.session_state.get("offer_prediction")
    with right_col:
        with st.container(border=True):
            st.markdown('<div class="cp-panel-title">推进结论</div>', unsafe_allow_html=True)
            st.markdown(
                '<div class="cp-panel-copy">优先看是否值得推进，以及今天该先做哪一步。</div>',
                unsafe_allow_html=True,
            )
            render_offer_prediction_snapshot(result)
    if not result:
        return

    action_tab, factor_tab, funnel_tab = st.tabs(["执行清单", "影响因素", "漏斗图"])
    with action_tab:
        st.markdown("#### 直接执行")
        for item in result.get("actions", []):
            st.write(f"- {item}")
        for item in result.get("application_steps", []):
            st.write(f"- {item}")
    with factor_tab:
        for item in result["drivers"]:
            st.write(f"- {item}")
    with funnel_tab:
        chart_df = pd.DataFrame(
            [
                {"阶段": "简历通过", "概率": result["简历通过率"]},
                {"阶段": "进入面试", "概率": result["进入面试概率"]},
                {"阶段": "拿 Offer", "概率": result["拿 offer 概率"]},
            ]
        )
        fig = px.bar(chart_df, x="阶段", y="概率", color="阶段", range_y=[0, 100], title="Offer 漏斗预测")
        st.plotly_chart(fig, width="stretch")


def render_interview_tab() -> None:
    left_col, right_col = st.columns([1.03, 0.97], gap="large")
    with left_col:
        with st.container(border=True):
            st.markdown('<div class="cp-panel-title">导入面经</div>', unsafe_allow_html=True)
            st.markdown(
                '<div class="cp-panel-copy">可粘贴牛客、社群、公众号或个人复盘内容；已有目标 JD 时会自动生成更贴近岗位的问题。</div>',
                unsafe_allow_html=True,
            )
            text = st.text_area("粘贴面经文本", height=260, placeholder="粘贴牛客、公众号、社群或个人复盘中的面经内容...")
            uploaded = st.file_uploader("上传面经文档", type=["pdf", "docx", "txt", "md"], key="interview_upload")
            upload_text = extract_text_from_upload(uploaded) if uploaded else ""
            if upload_text:
                with st.expander("查看面经提取文本"):
                    st.text_area("面经提取结果", value=upload_text, height=180)

            if st.button("提炼面试问题", type="primary", width="stretch"):
                full_text = "\n".join([text, upload_text]).strip()
                jd_analysis = st.session_state.get("jd_analysis")
                st.session_state.interview_analysis = analyze_interview(full_text, jd_analysis)
                st.success("面试问题提炼完成。")

    interview_analysis = st.session_state.get("interview_analysis")
    with right_col:
        with st.container(border=True):
            st.markdown('<div class="cp-panel-title">准备结论</div>', unsafe_allow_html=True)
            st.markdown(
                '<div class="cp-panel-copy">先看要背什么、要练什么，再去下方分块查看完整问题。</div>',
                unsafe_allow_html=True,
            )
            render_interview_snapshot(interview_analysis)
    if not interview_analysis:
        return

    answer_tab, question_tab, history_tab = st.tabs(["回答框架", "按 JD 准备题", "面经原题"])
    with answer_tab:
        for item in interview_analysis.get("answer_templates", []):
            st.write(f"- {item}")
    with question_tab:
        cols = st.columns(2)
        for idx, (category, questions) in enumerate(interview_analysis["generated_questions"].items()):
            with cols[idx % 2]:
                st.markdown(f"#### {category}")
                for question in questions:
                    st.write(f"- {question}")
    with history_tab:
        cols = st.columns(2)
        for idx, (category, questions) in enumerate(interview_analysis["buckets"].items()):
            with cols[idx % 2]:
                st.markdown(f"#### {category}")
                if questions:
                    for question in questions[:8]:
                        st.write(f"- {question}")
                else:
                    st.caption("暂无历史面经命中。")


def render_internship_tab() -> None:
    active_profile = get_active_profile()
    left_col, right_col = st.columns([1.04, 0.96], gap="large")
    with left_col:
        with st.container(border=True):
            st.markdown('<div class="cp-panel-title">实习信息</div>', unsafe_allow_html=True)
            st.markdown(
                '<div class="cp-panel-copy">评估重点不是名气，而是这段实习能不能变成全职求职证据。</div>',
                unsafe_allow_html=True,
            )
            with st.expander(f"当前目标意向：{active_profile['name']}", expanded=False):
                st.write(profile_text_for_analysis())
            cols = st.columns(2)
            company = cols[0].text_input("实习公司", placeholder="例如：SGS / TÜV / 制造业低碳部门 / 咨询公司")
            role = cols[1].text_input("实习岗位", placeholder="例如：数据分析实习生 / 产品运营实习生 / 咨询实习生 / 研发实习生")
            cols = st.columns(2)
            industry = cols[0].text_input("行业/方向", placeholder="例如：TIC / 制造业 / 咨询 / 出海合规")
            duration = cols[1].text_input("周期/地点/薪资", placeholder="例如：上海，3个月，每周4天")

            internship_text = st.text_area(
                "粘贴实习 JD 或工作内容",
                height=220,
                placeholder="粘贴实习招聘描述、导师说的工作内容、项目方向，或你已经拿到的 offer 信息...",
            )
            uploaded = st.file_uploader("上传实习 JD 文件", type=["pdf", "docx", "txt", "md", "html", "htm"], key="internship_upload")
            upload_text = extract_text_from_upload(uploaded) if uploaded else ""
            if upload_text:
                with st.expander("查看实习文件提取文本"):
                    st.text_area("实习文本提取结果", value=upload_text, height=160)

            if st.button("评估这个实习是否值得去", type="primary", width="stretch"):
                full_text = "\n".join([internship_text, upload_text]).strip()
                if not any([company, role, industry, duration, full_text]):
                    st.warning("请先填写或粘贴实习信息。")
                else:
                    st.session_state.internship_analysis = analyze_internship(
                        company,
                        role,
                        industry,
                        duration,
                        full_text,
                        profile_text_for_analysis(),
                    )
                    st.success("实习评估完成。")

    analysis = st.session_state.get("internship_analysis")
    with right_col:
        with st.container(border=True):
            st.markdown('<div class="cp-panel-title">是否值得去</div>', unsafe_allow_html=True)
            st.markdown(
                '<div class="cp-panel-copy">先看能不能沉淀项目、产出和导师反馈，再决定是否推进。</div>',
                unsafe_allow_html=True,
            )
            render_internship_snapshot(analysis)
    if not analysis:
        return

    action_tab, proof_tab, score_tab = st.tabs(["入职前怎么谈", "入职后怎么做", "评分细节"])
    with action_tab:
        st.markdown("#### 可直接问 HR/导师")
        for item in analysis.get("negotiation_script", []):
            st.write(f"- {item}")
        st.markdown("#### 接 offer 前问清楚")
        for question in analysis["questions_to_ask"]:
            st.write(f"- {question}")
    with proof_tab:
        st.markdown("#### 第一周行动")
        for item in analysis.get("first_week_plan", []):
            st.write(f"- {item}")
        st.markdown("#### 简历可写成")
        for item in analysis.get("resume_bullets", []):
            st.write(f"- {item}")
        st.markdown("#### 要争取的产出")
        for item in analysis["recommended_outputs"]:
            st.write(f"- {item}")
    with score_tab:
        cols = st.columns(2)
        with cols[0]:
            st.markdown("#### 为什么有利")
            for reason in analysis["reasons"]:
                st.write(f"- {reason}")
        with cols[1]:
            st.markdown("#### 风险点")
            for risk in analysis["risks"]:
                st.write(f"- {risk}")
        st.markdown("#### 维度判断")
        dimension_view = analysis["dimension_scores"].drop(columns=[col for col in ["得分", "权重"] if col in analysis["dimension_scores"].columns], errors="ignore")
        render_user_dataframe(dimension_view)


def render_gap_tab() -> None:
    st.subheader("不足分析 / 努力方向")
    if not st.session_state.get("jd_analysis"):
        st.warning("请先在“岗位工作台 > 单条JD分析”完成目标 JD 分析。")
        return
    if not st.session_state.get("resume_match"):
        st.info("请先在“简历解析与匹配”中使用当前简历完成匹配分析。")
        return
    if st.button("生成不足清单与努力方向", type="primary"):
        st.session_state.gap_analysis = build_gap_analysis(
            st.session_state.get("jd_analysis"),
            st.session_state.get("resume_match"),
            st.session_state.get("interview_analysis"),
        )
        st.success("不足分析完成。")

    gap_analysis = st.session_state.get("gap_analysis")
    if not gap_analysis:
        st.info("完成 JD 和简历分析后，点击按钮生成当前不足与优先级建议。")
        return

    st.info(gap_analysis["summary"])
    st.caption("下面的【】必须替换成你的真实项目、数字、工具或交付物；没有证据就不要写进投递版。")
    action_tab, resume_tab, plan_tab, detail_tab = st.tabs(["马上做什么", "简历改写句", "一周计划", "缺口明细"])
    with action_tab:
        action_rows = pd.DataFrame(gap_analysis.get("action_rows", []))
        if not action_rows.empty:
            preferred_cols = ["优先级", "短板", "今天就做", "交付物", "不要写"]
            show_cols = [col for col in preferred_cols if col in action_rows.columns]
            render_user_dataframe(action_rows, show_cols)
        else:
            st.caption("暂无明确短板，请先补充 JD 或完成简历匹配。")
    with resume_tab:
        st.markdown("#### 可替换进简历的表达")
        for item in gap_analysis.get("resume_patches", []):
            st.write(f"- {item}")
        st.markdown("#### 面试回答骨架")
        for item in gap_analysis.get("interview_scripts", []):
            st.write(f"- {item}")
    with plan_tab:
        weekly_df = pd.DataFrame(gap_analysis.get("weekly_plan", []))
        if not weekly_df.empty:
            render_user_dataframe(weekly_df)
        st.markdown("#### 面试准备重点")
        for item in gap_analysis.get("interview_focus", [])[:8]:
            st.write(f"- {item}")
    with detail_tab:
        cols = st.columns(2)
        for idx, (category, items) in enumerate(gap_analysis["current_gaps"].items()):
            with cols[idx % 2]:
                st.markdown(f"#### {category}")
                for item in items:
                    st.write(f"- {item}")
        st.markdown("#### 原优先级判断")
        priority_df = pd.DataFrame(gap_analysis["priorities"], columns=["优先级", "建议"])
        render_user_dataframe(priority_df)


def render_applications_tab() -> None:
    jd_analysis = st.session_state.get("jd_analysis")
    resume_match = st.session_state.get("resume_match")
    today = today_label()

    df = load_applications()
    if df.empty:
        render_decision_empty_state(
            "暂无投递记录。可以从岗位工作台的判断结果加入今日队列，也可以在下方手动新增。",
            ["先分析目标 JD。", "点击加入今日队列，或展开手动新增。", "投递后及时更新面试状态和下一步动作。"],
        )
    else:
        render_application_summary(df)

        today_df = df[df.get("queue_date", "").astype(str) == today] if "queue_date" in df.columns else pd.DataFrame()
        st.markdown("#### 今日求职队列")
        if today_df.empty:
            st.info("今天还没有队列岗位。可从单条 JD 决策台或批量分析结果加入。")
        else:
            queue_cols = ["company", "job_title", "match_score", "next_action", "applied", "interview_status", "notes"]
            render_user_dataframe(today_df, queue_cols)
            st.caption(f"今日队列共 {len(today_df)} 个岗位。处理顺序建议：先定制简历，再确认投递渠道，最后更新已投递状态。")

    with st.expander("手动新增投递记录", expanded=bool(df.empty)):
        with st.form("add_application_form"):
            basic = jd_analysis["basic"] if jd_analysis else {}
            cols = st.columns(3)
            company = cols[0].text_input("公司", value=basic.get("公司名", "") if basic else "")
            job_title = cols[1].text_input("岗位", value=basic.get("岗位名", "") if basic else "")
            salary = cols[2].text_input("薪资", value=basic.get("薪资", "") if basic else "")
            cols = st.columns(3)
            location = cols[0].text_input("地点", value=basic.get("地点", "") if basic else "")
            category = cols[1].text_input("分类", value=jd_analysis.get("category", "") if jd_analysis else "")
            match_score = cols[2].number_input("匹配度", min_value=0, max_value=100, value=int(resume_match.get("score", 0) if resume_match else 0))
            cols = st.columns(3)
            applied = cols[0].checkbox("已投递")
            interview_status = cols[1].selectbox("面试状态", ["未开始", "已投递", "笔试", "一面", "二面", "HR面", "已结束"])
            offer_status = cols[2].selectbox("Offer 状态", ["无", "等待中", "Offer", "拒绝", "放弃"])
            queue_cols = st.columns(2)
            queue_date = queue_cols[0].text_input("队列日期", value="")
            next_action = queue_cols[1].text_input("下一步动作", value="")
            notes = st.text_area("备注", height=80)
            submitted = st.form_submit_button("加入投递库")
            if submitted:
                add_application(
                    {
                        "company": company,
                        "job_title": job_title,
                        "salary": salary,
                        "location": location,
                        "category": category,
                        "match_score": match_score,
                        "is_high_value": jd_analysis["value"]["is_high_value"] if jd_analysis else False,
                        "is_generic_esg": jd_analysis["value"]["is_generic_esg"] if jd_analysis else False,
                        "applied": applied,
                        "interview_status": interview_status,
                        "offer_status": offer_status,
                        "queue_date": queue_date,
                        "next_action": next_action,
                        "notes": notes,
                    }
                )
                st.success("已加入投递库。")

    if df.empty:
        return

    st.markdown("#### 投递记录")
    edited = st.data_editor(
        df,
        width="stretch",
        hide_index=True,
        column_config={
            "id": None,
            "match_score": st.column_config.ProgressColumn("综合分", min_value=0, max_value=100),
            "is_high_value": st.column_config.CheckboxColumn("高价值"),
            "is_generic_esg": st.column_config.CheckboxColumn("低价值风险"),
            "applied": st.column_config.CheckboxColumn("已投递"),
            "queue_date": st.column_config.TextColumn("队列日期"),
            "next_action": st.column_config.TextColumn("下一步动作"),
            "company": st.column_config.TextColumn("公司"),
            "job_title": st.column_config.TextColumn("岗位"),
            "salary": st.column_config.TextColumn("薪资"),
            "location": st.column_config.TextColumn("地点"),
            "category": st.column_config.TextColumn("分类"),
            "offer_status": st.column_config.TextColumn("Offer状态"),
            "notes": st.column_config.TextColumn("备注"),
            "created_at": None,
            "updated_at": None,
            "source": None,
            "jd_text": None,
        },
        disabled=["id", "created_at"],
        key="application_editor",
    )
    cols = st.columns([1, 1, 2])
    if cols[0].button("保存表格修改"):
        save_application_edits(edited)
        st.success("修改已保存。")
    delete_options = {
        f"{row.get('company', '未识别公司')}｜{row.get('job_title', '未识别岗位')}｜{row.get('interview_status', '')}": int(row["id"])
        for _, row in df.iterrows()
        if "id" in row and pd.notna(row["id"])
    }
    delete_label = cols[1].selectbox("删除记录", ["不删除"] + list(delete_options.keys()))
    if cols[2].button("删除选中记录"):
        if delete_label != "不删除":
            delete_application(delete_options[delete_label])
            st.success("记录已删除，请刷新或重新进入页面查看。")
        else:
            st.warning("请先选择要删除的记录。")

    excel = io.BytesIO()
    edited.to_excel(excel, index=False)
    st.download_button("导出投递表 Excel", excel.getvalue(), "applications.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


def render_dashboard_tab() -> None:
    jd_analysis = st.session_state.get("jd_analysis")
    resume_match = st.session_state.get("resume_match")
    gap_analysis = st.session_state.get("gap_analysis")

    top_cols = st.columns([1.05, 0.95], gap="large")
    with top_cols[0]:
        with st.container(border=True):
            st.markdown('<div class="cp-panel-title">报告完整度</div>', unsafe_allow_html=True)
            st.markdown(
                '<div class="cp-panel-copy">导出前先确认材料是否齐全；缺项不会阻止导出，但会影响报告可用性。</div>',
                unsafe_allow_html=True,
            )
            render_report_readiness_panel()
    with top_cols[1]:
        with st.container(border=True):
            st.markdown('<div class="cp-panel-title">导出投递包</div>', unsafe_allow_html=True)
            st.markdown(
                '<div class="cp-panel-copy">Excel 适合继续编辑和筛选；PDF 适合归档或发送前快速查看。</div>',
                unsafe_allow_html=True,
            )
            export_cols = st.columns(2)
            if export_cols[0].button("生成 Excel", key="build_excel_report_btn"):
                st.session_state.report_excel_bytes = build_excel_report()
            if st.session_state.get("report_excel_bytes"):
                st.download_button(
                    "下载 Excel",
                    st.session_state.report_excel_bytes,
                    "careerpilot_application_pack.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    width="stretch",
                )
            if export_cols[1].button("生成 PDF", key="build_pdf_report_btn"):
                st.session_state.report_pdf_bytes = build_pdf_report()
                if not st.session_state.report_pdf_bytes:
                    st.warning("PDF 导出需要安装 reportlab。")
            if st.session_state.get("report_pdf_bytes"):
                st.download_button(
                    "下载 PDF",
                    st.session_state.report_pdf_bytes,
                    "careerpilot_application_pack.pdf",
                    mime="application/pdf",
                    width="stretch",
                )

    skill_tab, match_tab, gap_tab, application_tab = st.tabs(["技能关键词", "匹配雷达", "能力缺口", "投递进度"])
    with skill_tab:
        st.markdown("#### 技能关键词")
        if jd_analysis and not jd_analysis["skills"].empty:
            fig = px.bar(jd_analysis["skills"], x="技能", y="命中次数", color="技能")
            st.plotly_chart(fig, width="stretch")
        else:
            st.caption("暂无 JD 技能数据。")

    with match_tab:
        st.markdown("#### 匹配度雷达图")
        if resume_match:
            matched_count = len(resume_match["matched_skills"])
            missing_count = len(resume_match["missing_skills"])
            hard_skill = max(20, min(100, 50 + matched_count * 8 - missing_count * 3))
            tool_skill = 75 if any(skill in resume_match["matched_skills"] for skill in ["Python", "数据分析"]) else 45
            english = 75 if "英文能力" in resume_match["matched_skills"] else 45
            project = 80 if any("项目" in strength or "交付" in strength for strength in resume_match["strengths"]) else 55
            values = [hard_skill, tool_skill, english, project, resume_match["score"]]
            labels = ["硬技能", "工具/数据", "英文", "项目经历", "岗位匹配"]
            fig = go.Figure(
                data=go.Scatterpolar(r=values + [values[0]], theta=labels + [labels[0]], fill="toself")
            )
            fig.update_layout(polar=dict(radialaxis=dict(visible=True, range=[0, 100])), showlegend=False)
            st.plotly_chart(fig, width="stretch")
        else:
            st.caption("暂无简历匹配数据。")

    with gap_tab:
        st.markdown("#### 能力缺口柱状图")
        if gap_analysis:
            gap_counts = pd.DataFrame(
                [{"分类": key, "数量": len(items)} for key, items in gap_analysis["current_gaps"].items()]
            )
            fig = px.bar(gap_counts, x="分类", y="数量", color="分类")
            st.plotly_chart(fig, width="stretch")
        else:
            st.caption("暂无不足分析数据。")

    with application_tab:
        st.markdown("#### 投递进度")
        df = load_applications()
        if not df.empty:
            progress = df["interview_status"].value_counts().reset_index()
            progress.columns = ["状态", "数量"]
            fig = px.pie(progress, names="状态", values="数量", hole=0.45)
            st.plotly_chart(fig, width="stretch")
        else:
            st.caption("暂无投递数据。")


def render_jd_workspace_tab() -> None:
    render_jd_workspace_heading()
    mode = st.radio(
        "选择岗位工作模式",
        APP_NAVIGATION["jd_modes"],
        horizontal=True,
        label_visibility="collapsed",
    )
    if mode == "单条JD分析":
        render_jd_tab()
    elif mode == "批量JD筛选":
        render_batch_jd_tab()
    else:
        render_recruitment_monitor_tab()


def render_resume_workspace_tab() -> None:
    render_resume_workspace_heading()
    mode = st.radio(
        "选择简历工作模式",
        APP_NAVIGATION["resume_modes"],
        horizontal=True,
        label_visibility="collapsed",
    )
    if mode == "简历解析与匹配":
        render_resume_tab()
    elif mode == "定制简历":
        render_custom_resume_tab()
    else:
        render_gap_tab()


def render_decision_workspace_tab() -> None:
    render_decision_workspace_heading()
    mode = st.radio(
        "选择决策工作模式",
        APP_NAVIGATION["decision_modes"],
        horizontal=True,
        label_visibility="collapsed",
    )
    if mode == "Offer预测":
        render_offer_prediction_tab()
    elif mode == "实习评估":
        render_internship_tab()
    else:
        render_applications_tab()


def render_interview_report_workspace_tab() -> None:
    render_interview_report_workspace_heading()
    mode = st.radio(
        "选择准备与报告模式",
        APP_NAVIGATION["report_modes"],
        horizontal=True,
        label_visibility="collapsed",
    )
    if mode == "面经分析":
        render_interview_tab()
    else:
        render_dashboard_tab()


def render_sidebar() -> None:
    st.sidebar.title(APP_TITLE)
    st.sidebar.caption("本地化全职业岗位分析、简历匹配与投递决策工具")
    user = st.session_state.get(AUTH_SESSION_KEY) or {}
    if user:
        st.sidebar.caption(f"当前用户：{user.get('display_name') or user.get('email')}")
        if st.sidebar.button("退出登录", key="logout_user_btn"):
            logout_app_user()
            st.rerun()
    st.sidebar.markdown("#### 求职目标")
    profiles = load_user_profiles()
    active_profile = get_active_profile()
    prefs = load_target_preferences()
    profile_name_label = active_profile.get("name") or "未设置"
    st.sidebar.caption(profile_name_label)
    if active_profile.get("content"):
        st.sidebar.write(compact_profile_summary(active_profile.get("content", "")))
    st.sidebar.caption("方向：" + compact_list_text(prefs.get("target_roles", [])))
    st.sidebar.caption("城市：" + compact_list_text(prefs.get("target_cities", [])))
    st.sidebar.caption("行业：" + compact_list_text(prefs.get("preferred_industries", [])))

    with st.sidebar.expander("编辑求职目标", expanded=profiles.empty):
        profile_name = st.text_input(
            "目标名称",
            value=active_profile.get("name", ""),
            key=f"target_goal_name_{active_profile.get('id') or 'new'}",
        )
        selected_roles = free_input_multiselect(
            "求职方向",
            TARGET_ROLE_OPTIONS,
            split_preference_items(prefs.get("target_roles", [])),
            key="target_role_selector",
            help_text="选择常见岗位方向；也可以直接输入新的方向后回车确认。",
        )
        selected_cities = free_input_multiselect(
            "意向城市",
            CHINA_CITY_OPTIONS,
            split_preference_items(prefs.get("target_cities", [])),
            key="target_city_selector",
            help_text=f"可搜索 {len(CHINA_CITY_OPTIONS)} 个中国城市；也可以直接输入新城市后回车确认。",
        )
        selected_industries = free_input_multiselect(
            "目标行业",
            RECRUITMENT_INDUSTRY_OPTIONS,
            split_preference_items(prefs.get("preferred_industries", [])),
            key="target_industry_selector",
            help_text="按主流招聘网站行业口径细分；也可以直接输入新的行业标签后回车确认。",
        )
        job_keywords = st.text_input(
            "岗位关键词",
            value="、".join(split_preference_items(prefs.get("job_keywords", []))),
            help="例如：SQL、用户研究、项目交付；多个关键词可用顿号、逗号或空格分隔。",
        )
        profile_content = st.text_area(
            "目标说明",
            value=active_profile.get("content", ""),
            height=160,
            key=f"target_goal_content_{active_profile.get('id') or 'new'}",
        )
        show_advanced_goal = st.checkbox("显示高级偏好", value=False, key="show_advanced_goal_preferences")
        if show_advanced_goal:
            avoid_keywords = st.text_area("排除关键词", value=str(prefs.get("avoid_keywords", "")), height=72)
            notes = st.text_area("补充偏好", value=str(prefs.get("notes", "")), height=72)
        else:
            avoid_keywords = str(prefs.get("avoid_keywords", ""))
            notes = str(prefs.get("notes", ""))
        target_cols = st.columns(2)
        if target_cols[0].button("保存求职目标"):
            if not profile_name.strip() or not profile_content.strip():
                st.warning("目标名称和目标说明不能为空。")
            else:
                try:
                    if profiles.empty or active_profile.get("id") is None:
                        new_id = create_user_profile(profile_name, profile_content)
                        st.session_state.active_profile_id = new_id
                    else:
                        save_user_profile(int(active_profile["id"]), profile_name, profile_content)
                    save_target_preferences(
                        {
                            "target_roles": selected_roles,
                            "target_cities": selected_cities,
                            "extra_cities": "",
                            "accept_remote": bool(prefs.get("accept_remote")),
                            "accept_nationwide": bool(prefs.get("accept_nationwide")),
                            "min_monthly_salary": int(prefs.get("min_monthly_salary") or DEFAULT_TARGET_PREFERENCES["min_monthly_salary"]),
                            "min_daily_salary": int(prefs.get("min_daily_salary") or DEFAULT_TARGET_PREFERENCES["min_daily_salary"]),
                            "preferred_industries": selected_industries,
                            "extra_industries": "",
                            "job_keywords": split_preference_items(job_keywords),
                            "avoid_keywords": avoid_keywords,
                            "notes": notes,
                        }
                    )
                    st.success("求职目标已保存。")
                except sqlite3.IntegrityError:
                    st.error("目标名称已存在，请换一个名称。")
        if target_cols[1].button("清空结构化项"):
            save_target_preferences(DEFAULT_TARGET_PREFERENCES)
            st.success("已清空求职方向、城市、行业和关键词。")
        if not profiles.empty and st.button("删除目标"):
            delete_user_profile(int(active_profile["id"]))
            st.success("目标意向已删除，可重新新建。")

    st.sidebar.markdown("#### 当前简历")
    resumes = load_user_resumes()
    if resumes.empty:
        st.sidebar.caption("尚未设置当前简历。")
        with st.sidebar.expander("新增简历", expanded=False):
            new_resume_name = st.text_input("简历名称", value="", key="sidebar_new_resume_name")
            new_resume_upload = st.file_uploader("上传简历 PDF / Word / TXT", type=["pdf", "docx", "txt", "md"], key="sidebar_new_resume_upload")
            new_resume_upload_parsed = parsed_resume_from_upload(new_resume_upload) if new_resume_upload else None
            new_resume_content_key = "sidebar_new_resume_content"
            ensure_widget_text(new_resume_content_key)
            new_resume_upload_text = sync_uploaded_resume_to_widget(
                new_resume_upload,
                new_resume_upload_parsed,
                new_resume_content_key,
                "sidebar_new_resume_upload_sig",
            )
            if new_resume_upload_parsed and new_resume_upload_text:
                st.caption(f"已结构化读取 {len(new_resume_upload_parsed['sections'])} 个板块、{len(new_resume_upload_parsed['skills'])} 个技能关键词，并自动填入简历内容。")
                upload_signature = uploaded_file_signature(new_resume_upload)
                if upload_signature and st.session_state.get("sidebar_new_resume_saved_sig") != upload_signature:
                    try:
                        auto_name = unique_resume_name(new_resume_name or Path(new_resume_upload.name).stem)
                        new_id = create_user_resume(auto_name, new_resume_upload_text, is_default=1)
                        st.session_state.active_resume_id = new_id
                        st.session_state.sidebar_new_resume_saved_sig = upload_signature
                        clear_resume_dependent_results()
                        st.success("已自动保存为当前简历，其他功能会直接调用这份内容。")
                    except sqlite3.IntegrityError:
                        st.error("简历名称已存在，请换一个名称后保存。")
            elif new_resume_upload:
                st.warning("没有从上传文件中读取到有效文本。若 PDF 是扫描件，请先转成可复制文字。")
            new_resume_content = st.text_area("简历内容", height=180, key=new_resume_content_key)
            if st.button("保存为当前简历"):
                if new_resume_content.strip():
                    new_id = create_user_resume(unique_resume_name(new_resume_name), new_resume_content, is_default=1)
                    st.session_state.active_resume_id = new_id
                    clear_resume_dependent_results()
                    st.success("简历已保存。")
                else:
                    st.warning("请先粘贴简历内容。")
    else:
        active_resume = get_active_resume()
        resume_names = resumes["name"].tolist()
        resume_index = resume_names.index(active_resume["name"]) if active_resume["name"] in resume_names else 0
        st.sidebar.caption(f"当前：{active_resume['name']}")
        selected_resume_name = st.sidebar.selectbox("选择简历", resume_names, index=resume_index, key="sidebar_resume_select")
        selected_resume = resumes[resumes["name"] == selected_resume_name].iloc[0]
        if int(selected_resume["id"]) != st.session_state.get("active_resume_id"):
            st.session_state.active_resume_id = int(selected_resume["id"])
            active_resume = get_active_resume()

        with st.sidebar.expander("简历状态", expanded=False):
            render_global_resume_status(active_resume)

        with st.sidebar.expander("编辑当前简历", expanded=False):
            resume_name = st.text_input("简历名称", value=active_resume["name"], key=f"resume_name_{active_resume['id']}")
            resume_content_key = f"resume_content_{active_resume['id']}"
            ensure_widget_text(resume_content_key, active_resume["content"])
            resume_upload = st.file_uploader("上传 PDF / Word / TXT 更新当前简历", type=["pdf", "docx", "txt", "md"], key=f"sidebar_resume_upload_{active_resume['id']}")
            resume_upload_parsed = parsed_resume_from_upload(resume_upload) if resume_upload else None
            resume_upload_text = sync_uploaded_resume_to_widget(
                resume_upload,
                resume_upload_parsed,
                resume_content_key,
                f"sidebar_resume_upload_sig_{active_resume['id']}",
            )
            if resume_upload_text:
                st.caption(f"已结构化读取上传简历：{len(resume_upload_parsed['sections'])} 个板块、{len(resume_upload_parsed['skills'])} 个技能关键词，并自动填入简历内容。")
                upload_signature = uploaded_file_signature(resume_upload)
                autosave_key = f"sidebar_resume_upload_saved_sig_{active_resume['id']}"
                if upload_signature and st.session_state.get(autosave_key) != upload_signature:
                    save_user_resume(int(active_resume["id"]), resume_name, resume_upload_text)
                    st.session_state[autosave_key] = upload_signature
                    clear_resume_dependent_results()
                    st.success("上传文件已自动保存为当前简历。")
            elif resume_upload:
                st.warning("没有从上传文件中读取到有效文本。若 PDF 是扫描件，请先转成可复制文字。")
            resume_content = st.text_area("简历内容", height=220, key=resume_content_key)
            cols = st.columns(2)
            if cols[0].button("保存简历"):
                try:
                    old_fingerprint = content_fingerprint(active_resume["content"])
                    save_user_resume(int(active_resume["id"]), resume_name, resume_content)
                    if content_fingerprint(resume_content) != old_fingerprint:
                        clear_resume_dependent_results()
                    st.success("简历已保存。")
                except sqlite3.IntegrityError:
                    st.error("简历名称已存在。")
            if cols[1].button("复制简历"):
                try:
                    new_id = create_user_resume(unique_resume_name(f"{resume_name} - 副本"), resume_content)
                    st.session_state.active_resume_id = new_id
                    clear_resume_dependent_results()
                    st.success("已复制为新简历。")
                except sqlite3.IntegrityError:
                    st.error("副本名称已存在。")
            if st.button("删除当前简历"):
                if delete_user_resume(int(active_resume["id"])):
                    st.success("简历已删除。")
                else:
                    st.warning("至少保留一个简历，或先新建另一个。")

        with st.sidebar.expander("新增简历", expanded=False):
            new_resume_name = st.text_input("新简历名称", key="sidebar_add_resume_name")
            new_resume_upload = st.file_uploader("上传新简历 PDF / Word / TXT", type=["pdf", "docx", "txt", "md"], key="sidebar_add_resume_upload")
            new_resume_upload_parsed = parsed_resume_from_upload(new_resume_upload) if new_resume_upload else None
            new_resume_content_key = "sidebar_add_resume_content"
            ensure_widget_text(new_resume_content_key)
            new_resume_upload_text = sync_uploaded_resume_to_widget(
                new_resume_upload,
                new_resume_upload_parsed,
                new_resume_content_key,
                "sidebar_add_resume_upload_sig",
            )
            if new_resume_upload_parsed and new_resume_upload_text:
                st.caption(f"已结构化读取 {len(new_resume_upload_parsed['sections'])} 个板块、{len(new_resume_upload_parsed['skills'])} 个技能关键词，并自动填入新简历内容。")
                upload_signature = uploaded_file_signature(new_resume_upload)
                if upload_signature and st.session_state.get("sidebar_add_resume_saved_sig") != upload_signature:
                    try:
                        auto_name = unique_resume_name(new_resume_name or Path(new_resume_upload.name).stem)
                        new_id = create_user_resume(auto_name, new_resume_upload_text)
                        st.session_state.active_resume_id = new_id
                        st.session_state.sidebar_add_resume_saved_sig = upload_signature
                        clear_resume_dependent_results()
                        st.success("新简历已自动保存并设为当前简历。")
                    except sqlite3.IntegrityError:
                        st.error("简历名称已存在，请换一个名称后保存。")
            elif new_resume_upload:
                st.warning("没有从上传文件中读取到有效文本。若 PDF 是扫描件，请先转成可复制文字。")
            new_resume_content = st.text_area("新简历内容", height=140, key=new_resume_content_key)
            if st.button("新增并设为当前简历"):
                if not new_resume_name.strip() or not new_resume_content.strip():
                    st.warning("请填写名称和内容。")
                else:
                    try:
                        new_id = create_user_resume(unique_resume_name(new_resume_name), new_resume_content)
                        st.session_state.active_resume_id = new_id
                        clear_resume_dependent_results()
                        st.success("新简历已创建。")
                    except sqlite3.IntegrityError:
                        st.error("简历名称已存在。")

    with st.sidebar.expander("高级：本地数据位置", expanded=False):
        st.code(str(DB_PATH))


def render_auth_screen() -> None:
    st.title(APP_TITLE)
    st.caption("请先登录。每个账号的简历、目标意向、偏好和投递记录都会独立保存。")
    login_tab, register_tab = st.tabs(["登录", "注册"])
    with login_tab:
        with st.form("login_form"):
            email = st.text_input("邮箱", key="login_email")
            password = st.text_input("密码", type="password", key="login_password")
            submitted = st.form_submit_button("登录", type="primary")
        if submitted:
            ok, message = authenticate_app_user(email, password)
            if ok:
                st.success(message)
                st.rerun()
            else:
                st.error(message)
    with register_tab:
        with st.form("register_form"):
            display_name = st.text_input("昵称", key="register_display_name")
            email = st.text_input("邮箱", key="register_email")
            password = st.text_input("密码", type="password", key="register_password")
            password_confirm = st.text_input("确认密码", type="password", key="register_password_confirm")
            submitted = st.form_submit_button("注册并登录", type="primary")
        if submitted:
            if password != password_confirm:
                st.error("两次输入的密码不一致。")
            else:
                ok, message = create_app_user(email, password, display_name)
                if ok:
                    st.success(message)
                    st.rerun()
                else:
                    st.error(message)


def main() -> None:
    st.set_page_config(page_title=APP_TITLE, page_icon="🧭", layout="wide")
    render_app_styles()
    clear_legacy_runtime_state()
    init_db()
    if not current_user_id():
        render_auth_screen()
        return
    render_sidebar()
    render_state_alerts()

    workspace = render_main_workspace_nav()
    if workspace == "jd":
        render_jd_workspace_tab()
    elif workspace == "resume":
        render_resume_workspace_tab()
    elif workspace == "decision":
        render_decision_workspace_tab()
    else:
        render_interview_report_workspace_tab()


if __name__ == "__main__":
    main()
